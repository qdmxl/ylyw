#!/usr/bin/env python3
"""生成论文docx——嵌入图片、表格编号表名、上标引用"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

doc = Document()

for sec in doc.sections:
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

sty = doc.styles['Normal']
sty.font.name = '宋体'; sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 1.5

fig_dir = 'figures'

# ====== 辅助函数 ======
def make_superscript(run):
    """给run添加上标"""
    rPr = run._element.get_or_add_rPr()
    vert = OxmlElement('w:vertAlign')
    vert.set(qn('w:val'), 'superscript')
    rPr.append(vert)

def add_ref_text(paragraph, text):
    """添加带引用的段落——[n]转上标"""
    import re as _re
    parts = _re.split(r'(\[\d+\])', text)
    for part in parts:
        if _re.match(r'^\[\d+\]$', part):
            run = paragraph.add_run(part)
            run.font.size = Pt(12)
            make_superscript(run)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(12)

def add_table_with_caption(doc, data, caption):
    """生成表格 + 表名（居中粗体）"""
    if not data or len(data) < 2:
        return
    cols = max(len(r) for r in data)
    t = doc.add_table(rows=len(data), cols=cols, style='Table Grid')
    for ri, row in enumerate(data):
        for ci, ct in enumerate(row):
            if ci < cols:
                c = t.rows[ri].cells[ci]; c.text = str(ct)
                for cp in c.paragraphs:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for rn in cp.runs:
                        if ri == 0: rn.bold = True
                        rn.font.size = Pt(10)
    doc.add_paragraph()  # 表后空行

def insert_figure(doc, fname, caption):
    """插入图片 + 图名"""
    path = os.path.join(fig_dir, fname + '.png')
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(5.2))
    # 图名
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cp.add_run(caption)
    r2.font.size = Pt(10); r2.italic = True
    doc.add_paragraph()

# ====== 标题 ======
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('基于YLYW易理模型的具身智能体逐步任务规划')
r.bold = True; r.font.size = Pt(16); r.font.name = '黑体'
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Step-by-Step Task Planning for Embodied Agents Based on the YLYW Yili Model')
r2.font.size = Pt(14); r2.italic = True
doc.add_paragraph()

# ====== 摘要 ======
p = doc.add_paragraph(); r = p.add_run('摘  要'); r.bold = True; r.font.size = Pt(14)
abs_text = (
    '本文提出一种基于YLYW（易理研物）易理模型的逐步任务规划方法，面向具身智能中只能获取环境文本观测'
    '(observation-only)的任务执行场景。YLYW以《易经》八卦、六爻、六十四卦为核心的结构化先验模板，'
    '递归应用于字、词、句三个语言层级，实现零训练的汉语理解与逐步决策。'
    '在ALFWorld[3] valid_unseen全量134场景上的实验表明，系统在纯观测条件下达到44.0%的总体成功率，'
    '其中清洗类67.7%、冷却类71.4%、加热类56.5%。本文详细分析了YLYW框架在各模块中的作用'
    '——句级卦象为任务类型定性提供先验，六爻编码为环境状态提供连续的语义表征，模糊规则使阶段切换平滑，'
    '知几跨局学习弥补同义词和位置先验的差距[1,2]——同时给出易理框架与工程硬编码的明确分界。'
)
p1 = doc.add_paragraph(); p1.paragraph_format.first_line_indent = Cm(0.74)
add_ref_text(p1, abs_text)
p2 = doc.add_paragraph()
r = p2.add_run('关键词：YLYW易理模型；具身智能；逐步任务规划；六爻编码；模糊推理；观察驱动；知几学习')
r.bold = True; r.font.size = Pt(12)
doc.add_paragraph()

# ====== 读md ======
with open('YLYW_obs_only_task_planning_paper.md', encoding='utf-8') as f:
    md = f.read()

start = md.find('## 1. 引言')
body = md[start:] if start >= 0 else md
lines = body.split('\n')

in_code = False; code_l = []; table_d = []; table_caption = ''
skip_la = re.compile(r'^\\(begin|end|centering|label|ref|textwidth|vspace|textbf|hline|tabular|left|right)')
fig_inserted = False

for line in lines:
    s = line.strip()
    
    # 代码块
    if s.startswith('```'):
        in_code = not in_code
        if not in_code and code_l:
            cp = doc.add_paragraph(); cp.paragraph_format.left_indent = Cm(0.5)
            for cl in code_l:
                r = cp.add_run(cl+'\n'); r.font.name='Courier New'; r.font.size=Pt(9)
            code_l = []
        continue
    if in_code:
        code_l.append(s); continue
    
    # 二级标题
    if s.startswith('## ') and '摘要' not in s and '参考文献' not in s:
        if table_d: add_table_with_caption(doc, table_d, table_caption); table_d = []
        h = doc.add_paragraph(); r = h.add_run(s[3:])
        r.bold = True; r.font.size = Pt(14); r.font.name = '黑体'
        h.paragraph_format.space_before = Pt(12); continue
    
    # 三级标题
    if s.startswith('### '):
        if table_d: add_table_with_caption(doc, table_d, table_caption); table_d = []
        h = doc.add_paragraph(); r = h.add_run(s[4:])
        r.bold = True; r.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(8); continue
    
    # 表格
    if s.startswith('|') and s.endswith('|'):
        cells = [c.strip() for c in s.split('|')[1:-1]]
        if not re.match(r'^[\s\-:\|]+$', s):
            table_d.append(cells); continue
    else:
        if table_d:
            # 检查前面有没有表名
            if not table_caption:
                table_caption = '表'
            add_table_with_caption(doc, table_d, table_caption)
            table_d = []; table_caption = ''
    
    # 捕获表名（LaTeX \caption{}）
    cap_m = re.search(r'\\caption\{(.*?)\}', s)
    if cap_m:
        table_caption = cap_m.group(1)
        continue
    
    # 图片插入
    if 'system_architecture' in s and not fig_inserted:
        insert_figure(doc, 'system_architecture', '图1  YLYW逐步任务规划系统架构图')
        fig_inserted = True; continue
    
    if 'recursive_ylyw' in s:
        insert_figure(doc, 'recursive_ylyw', '图2  递归YLYW的三层推理架构')
        continue
    
    if 'decision_loop' in s:
        insert_figure(doc, 'decision_loop', '图3  逐步决策闭环')
        continue
    
    if 'includegraphics' in s:
        continue  # 已经通过上面的条件处理
    
    if skip_la.match(s): continue
    if not s: continue
    
    # 列表
    if s.startswith('- ') or s.startswith('* ') or re.match(r'^\d+\. ', s):
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', s)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.74)
        r = p.add_run(clean); r.font.size = Pt(12)
        continue
    
    # 普通段落
    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', s)
    if re.match(r'^[\s\d\-\|:\/.]+$', clean): continue
    if clean.startswith('\\') and not clean.startswith('\\#'): continue
    if '\\' in clean and not any(c.isalpha() for c in clean.replace('\\','').replace(' ','').replace('#','').replace('{','').replace('}','')): continue
    
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0.74)
    # 处理上标引用
    add_ref_text(p, clean)

# 最后可能的表格
if table_d: add_table_with_caption(doc, table_d, table_caption)

# ====== 参考文献 ======
doc.add_paragraph()
h = doc.add_paragraph(); r = h.add_run('参考文献'); r.bold = True; r.font.size = Pt(14); r.font.name = '黑体'
for ref in [
    '[1] 马兴录. 基于YLYW易理模糊模型的机器人抓取决策研究. 技术论文, 2024.',
    '[2] 马兴录, 等. YLYW递归汉语理解：基于易理原理的汉语理解架构. 技术论文, 2026.',
    '[3] Shridhar, M., et al. ALFWorld: Aligning Text and Embodied Environments for Interactive Learning. ICLR, 2021.',
    '[4] Yao, S., et al. ReAct: Synergizing Reasoning and Acting in Language Models. ICLR, 2023.',
]:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.74); p.paragraph_format.first_line_indent = Cm(-0.37)
    r = p.add_run(ref); r.font.size = Pt(11)

doc.save('YLYW_obs_only_task_planning_paper.docx')
print('OK')
