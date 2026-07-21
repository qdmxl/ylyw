#!/usr/bin/env python3
"""将论文Markdown转换为规范docx格式"""

import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = os.path.join(os.path.dirname(__file__), 'YLYW_递归汉语理解_技术论文.md')
DOCX_PATH = os.path.join(os.path.dirname(__file__), 'YLYW_递归汉语理解_技术论文.docx')
IMG_DIR = os.path.dirname(__file__)

with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# 样式
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.7)

def set_run_font(run, east_asia=None):
    ea = east_asia or run.font.name or '宋体'
    rpr = run.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from lxml import etree
        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
        rfonts.set(qn('w:ascii'), run.font.name or 'Times New Roman')
        rfonts.set(qn('w:hAnsi'), run.font.name or 'Times New Roman')
    rfonts.set(qn('w:eastAsia'), ea)

def add_run(paragraph, text, bold=False, italic=False, size=None, font=None, east_asia=None, superscript=False, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = size
    if font: run.font.name = font
    set_run_font(run, east_asia)
    if superscript:
        run.font.superscript = True
    if color:
        run.font.color.rgb = color
    return run

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def parse_inline(text):
    """解析行内格式：**bold** `code` [N](ref)"""
    parts = []
    i = 0
    while i < len(text):
        if text[i:i+2] == '**':
            j = text.find('**', i+2)
            if j != -1:
                parts.append(('bold', text[i+2:j]))
                i = j + 2
                continue
        elif text[i] == '`':
            j = text.find('`', i+1)
            if j != -1:
                parts.append(('code', text[i+1:j]))
                i = j + 1
                continue
        elif text[i:i+1] == '[' and i+2 < len(text):
            # 可能的上标引用 [N]
            j = text.find(']', i)
            if j != -1 and j - i <= 4 and text[i+1:j].isdigit():
                parts.append(('sup', text[i:j+1]))
                i = j + 1
                continue
        # 普通字符
        j = i + 1
        while j < len(text) and text[j] not in '*`[':
            j += 1
        if j > i:
            parts.append(('normal', text[i:j]))
        i = j
    return parts

def add_formatted_paragraph(doc, text, style_name='Normal', alignment=None, 
                           first_indent=True, font_size=None, bold=False):
    if not text.strip():
        return doc.add_paragraph()
    p = doc.add_paragraph(style=style_name)
    if alignment:
        p.alignment = alignment
    if not first_indent and alignment != WD_ALIGN_PARAGRAPH.CENTER:
        p.paragraph_format.first_line_indent = Cm(0)
    
    parts = parse_inline(text)
    for typ, content in parts:
        if typ == 'normal':
            add_run(p, content, bold=bold, east_asia='宋体')
        elif typ == 'bold':
            add_run(p, content, bold=True, east_asia='宋体')
        elif typ == 'code':
            add_run(p, content, font='Consolas', size=Pt(9), east_asia='Consolas')
        elif typ == 'sup':
            add_run(p, content, superscript=True, size=Pt(9), east_asia='宋体')
    return p

def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    
    if level == 1:
        size = Pt(16)
        b = True
    elif level == 2:
        size = Pt(14)
        b = True
    else:
        size = Pt(12)
        b = True
    add_run(p, text, bold=b, size=size, east_asia='黑体')
    return p

def add_table_from_data(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D6E4F0')
    
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if j < len(headers):
                cell = table.rows[i+1].cells[j]
                cell.text = ''
                # 解析加粗
                if val.startswith('**') and val.endswith('**'):
                    run = cell.paragraphs[0].add_run(val[2:-2])
                    run.bold = True
                else:
                    run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    return table

# ============================================================
# 解析Markdown内容
# ============================================================

state = 'normal'
code_block = []
table_data = []
in_table = False

def render_table_data(doc, headers, rows):
    add_table_from_data(doc, headers, rows)
    doc.add_paragraph()  # 表后空行

for line in lines:
    raw = line.rstrip()
    
    if raw.startswith('```'):
        if state == 'code':
            # 渲染代码块
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for cl in code_block:
                run = p.add_run(cl + '\n')
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            code_block = []
            state = 'normal'
        else:
            state = 'code'
        continue
    
    if state == 'code':
        code_block.append(raw)
        continue
    
    # 空行
    if not raw:
        if in_table and table_data:
            # 表结束
            headers = [c.strip() for c in table_data[0].split('|')[1:-1]]
            rows = []
            for tl in table_data[2:]:  # 跳过表头和分隔行
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if cells:
                    rows.append(cells)
            if rows and headers:
                render_table_data(doc, headers, rows)
            table_data = []
            in_table = False
        continue
    
    # 表行
    if '|' in raw and '---' in raw:
        if not in_table:
            in_table = True
        continue
    if in_table:
        table_data.append(raw)
        continue
    
    # 标题
    if raw.startswith('# '):
        t = raw[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        add_run(p, t, bold=True, size=Pt(18), east_asia='黑体')
        continue
    if raw.startswith('## '):
        t = raw[3:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if 'A Recursive' in t else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        if 'A Recursive' in t:
            add_run(p, t, italic=True, size=Pt(14), east_asia='宋体')
        else:
            # 提取章节号
            add_run(p, t, bold=True, size=Pt(15), east_asia='黑体')
        continue
    if raw.startswith('### '):
        add_heading_custom(doc, raw[4:].strip(), level=3)
        continue
    if raw.startswith('#### '):
        add_heading_custom(doc, raw[5:].strip(), level=4)
        continue
    
    # 表头（**表N:** 开头）
    if raw.startswith('**表') and '**' in raw[4:]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        parts = parse_inline(raw)
        for typ, content in parts:
            if typ == 'normal':
                add_run(p, content, bold=True, east_asia='宋体')
            elif typ == 'bold':
                add_run(p, content, bold=True, east_asia='宋体')
        continue
    
    # 图（![...](...)）
    if raw.startswith('!['):
        m = re.match(r'!\[.*\]\((.+)\)', raw)
        if m:
            img_file = m.group(1)
            if img_file.endswith('.svg'):
                img_file = img_file[:-4] + '.png'
            img_path = os.path.join(IMG_DIR, img_file)
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.add_run()
                run.add_picture(img_path, width=Cm(14))
        continue
    
    # 图说明（**图N:**）
    if raw.startswith('**图') and '**' in raw[4:] and '：' in raw:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        add_run(p, raw, size=Pt(10), east_asia='宋体')
        continue
    
    # 分隔线
    if raw == '---':
        continue
    
    # 普通段落
    add_formatted_paragraph(doc, raw)

# 最后还可能有表
if in_table and table_data:
    headers = [c.strip() for c in table_data[0].split('|')[1:-1]]
    rows = []
    for tl in table_data[2:]:
        cells = [c.strip() for c in tl.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    if rows and headers:
        render_table_data(doc, headers, rows)

doc.save(DOCX_PATH)
print(f'✅ docx已生成: {DOCX_PATH}')
print(f'  文件大小: {os.path.getsize(DOCX_PATH)/1024:.0f} KB')
