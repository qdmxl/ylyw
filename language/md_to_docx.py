#!/usr/bin/env python3
"""将YLYW技术论文Markdown转换为docx格式"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

MD_PATH = os.path.join(os.path.dirname(__file__), 'YLYW_递归汉语理解_技术论文.md')
DOCX_PATH = os.path.join(os.path.dirname(__file__), 'YLYW_递归汉语理解_技术论文.docx')

with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# 样式设置
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def add_formatted_text(paragraph, text, bold=False, italic=False, font_size=None, font_name=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = font_size
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return run

def parse_inline(text):
    """解析行内格式：加粗、斜体、代码"""
    parts = []
    i = 0
    while i < len(text):
        if text[i:i+2] == '**':
            j = text.find('**', i+2)
            if j != -1:
                parts.append(('bold', text[i+2:j]))
                i = j + 2
                continue
        elif text[i:i+1] == '*':
            j = text.find('*', i+1)
            if j != -1 and text[j:j+2] != '**':
                parts.append(('italic', text[i+1:j]))
                i = j + 1
                continue
        elif text[i:i+1] == '`':
            j = text.find('`', i+1)
            if j != -1:
                parts.append(('code', text[i+1:j]))
                i = j + 1
                continue
        # 正常文字
        j = i + 1
        while j < len(text) and text[j] not in '*`':
            j += 1
        if j > i:
            parts.append(('normal', text[i:j]))
        i = j
    return parts

def add_paragraph_with_format(doc, text, style_name='Normal', alignment=None, font_size=None, bold=False):
    p = doc.add_paragraph(style=style_name)
    if alignment:
        p.alignment = alignment
        
    parts = parse_inline(text)
    for typ, content in parts:
        run = p.add_run(content)
        if typ == 'bold' or bold:
            run.bold = True
        if typ == 'italic':
            run.italic = True
        if typ == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        if font_size:
            run.font.size = font_size
    return p

# 状态跟踪
in_code_block = False
code_lines = []
in_table = False
table_lines = []

def render_code_block(doc, lines):
    if not lines:
        return
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    for line in lines:
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def render_table(doc, lines):
    """渲染markdown表格"""
    if len(lines) < 2:
        return
    # 解析表头
    header = [c.strip() for c in lines[0].split('|')[1:-1]]
    # 跳过分隔行（第2行）
    data = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            data.append(cells)
    
    if not data:
        return
    
    rows = len(data) + 1
    cols = len(header)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # 表头
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    
    # 数据行
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            if j < cols:
                cell = table.rows[i+1].cells[j]
                cell.text = ''
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(10)
    
    doc.add_paragraph()  # 表后空行

# 标题中英对照
title_cn = None
title_en = None

for line in lines:
    stripped = line.rstrip()
    
    # 代码块处理
    if stripped.startswith('```'):
        if in_code_block:
            render_code_block(doc, code_lines)
            code_lines = []
            in_code_block = False
        else:
            in_code_block = True
        continue
    
    if in_code_block:
        code_lines.append(stripped)
        continue
    
    # 空行
    if not stripped:
        if in_table:
            render_table(doc, table_lines)
            table_lines = []
            in_table = False
        continue
    
    # 表处理
    if '|' in stripped and '---' not in stripped:
        if not in_table and table_lines:
            render_table(doc, table_lines)
            table_lines = []
        in_table = True
        table_lines.append(stripped)
        continue
    else:
        if in_table:
            render_table(doc, table_lines)
            table_lines = []
            in_table = False
        if '---' in stripped:
            continue
    
    # 标题
    if stripped.startswith('# '):
        title_cn = stripped[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_cn)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        continue
    
    if stripped.startswith('## '):
        text = stripped[3:].strip()
        p = doc.add_paragraph()
        if 'A Recursive' in text:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(14)
        else:
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(15)
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        continue
    
    if stripped.startswith('### '):
        text = stripped[4:].strip()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        continue
    
    if stripped.startswith('#### '):
        text = stripped[5:].strip()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        continue
    
    if stripped.startswith('---'):
        continue
    
    # 普通段落
    add_paragraph_with_format(doc, stripped)

# 如果最后还有表
if table_lines:
    render_table(doc, table_lines)

doc.save(DOCX_PATH)
print(f'✅ 已生成: {DOCX_PATH}')
