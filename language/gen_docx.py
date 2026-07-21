#!/usr/bin/env python3
"""生成论文docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

for sec in doc.sections:
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

sty = doc.styles['Normal']
sty.font.name = '宋体'
sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 1.5

# 标题
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('基于YLYW易理模型的具身智能体逐步任务规划')
r.bold = True; r.font.size = Pt(16); r.font.name = '黑体'
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Step-by-Step Task Planning for Embodied Agents Based on the YLYW Yili Model')
r2.font.size = Pt(14); r2.italic = True
doc.add_paragraph()

# 摘要
p = doc.add_paragraph(); r = p.add_run('摘  要'); r.bold = True; r.font.size = Pt(14)
abs_t = ('本文提出一种基于YLYW（易理研物）易理模型的逐步任务规划方法，面向具身智能中只能获取环境文本观测'
    '(observation-only)的任务执行场景。YLYW以《易经》八卦、六爻、六十四卦为核心的结构化先验模板，'
    '递归应用于字、词、句三个语言层级，实现零训练的汉语理解与逐步决策。'
    '在ALFWorld valid_unseen全量134场景上的实验表明，系统在纯观测条件下达到44.0%的总体成功率，'
    '其中清洗类67.7%、冷却类71.4%、加热类56.5%。'
    '本文详细分析了YLYW框架在各模块中的作用——句级卦象为任务类型定性提供先验，六爻编码为环境状态提供'
    '连续的语义表征，模糊规则使阶段切换平滑，知几跨局学习弥补同义词和位置先验的差距——同时给出易理框架与'
    '工程硬编码的明确分界。')
p1 = doc.add_paragraph(); p1.paragraph_format.first_line_indent = Cm(0.74)
r = p1.add_run(abs_t); r.font.size = Pt(12)
p2 = doc.add_paragraph()
r = p2.add_run('关键词：YLYW易理模型；具身智能；逐步任务规划；六爻编码；模糊推理；观察驱动；知几学习')
r.bold = True; r.font.size = Pt(12)
doc.add_paragraph()

# 正文
with open('YLYW_obs_only_task_planning_paper.md', encoding='utf-8') as f:
    content = f.read()

start = content.find('## 1. 引言')
body = content[start:] if start >= 0 else content
lines = body.split('\n')

in_code = False; code_l = []; table_d = []
skip_pat = re.compile(r'^\\(begin|end|centering|includegraphics|label|ref|textwidth|vspace|textbf|hline|tabular)')

def tbl_flush():
    global table_d
    if not table_d or len(table_d) < 2: return
    cols = max(len(r) for r in table_d)
    t = doc.add_table(rows=len(table_d), cols=cols, style='Table Grid')
    for ri, row in enumerate(table_d):
        for ci, ct in enumerate(row):
            if ci < cols:
                c = t.rows[ri].cells[ci]; c.text = str(ct)
                for cp in c.paragraphs:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for rn in cp.runs:
                        if ri == 0: rn.bold = True
                        rn.font.size = Pt(10)
    doc.add_paragraph()
    table_d = []

def fig_ph(cap):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\u3010\u6b64\u4f4d\u7f6e\u9884\u7559\u56fe\uff1a' + cap + '\u3011')
    r.font.size = Pt(10); r.italic = True; r.font.color.rgb = RGBColor(0x99,0x99,0x99)

for line in lines:
    s = line.strip()
    if s.startswith('```'):
        in_code = not in_code
        if not in_code and code_l:
            cp = doc.add_paragraph(); cp.paragraph_format.left_indent = Cm(0.5)
            for cl in code_l:
                r = cp.add_run(cl+'\n'); r.font.name='Courier New'; r.font.size=Pt(9)
            code_l = []
        continue
    if in_code:
        code_l.append(s); continue
    
    if s.startswith('## ') and '\u6458\u8981' not in s and '\u53c2\u8003\u6587\u732e' not in s:
        tbl_flush()
        h = doc.add_paragraph(); r = h.add_run(s[3:])
        r.bold = True; r.font.size = Pt(14); r.font.name = '\u9ed1\u4f53'
        h.paragraph_format.space_before = Pt(12); continue
    
    if s.startswith('### '):
        tbl_flush()
        h = doc.add_paragraph(); r = h.add_run(s[4:])
        r.bold = True; r.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(8); continue
    
    if s.startswith('|') and s.endswith('|'):
        cells = [c.strip() for c in s.split('|')[1:-1]]
        if not re.match(r'^[\s\-:\|]+$', s):
            table_d.append(cells); continue
    else:
        tbl_flush()
    
    if 'includegraphics' in s or 'bagua_radicals' in s or 'recursive_ylyw' in s or 'decision_loop' in s:
        m = re.search(r'\\caption\{(.*?)\}', s)
        cap = m.group(1) if m else '图'
        fig_ph(cap); continue
    
    if skip_pat.match(s): continue
    if not s: continue
    
    if s.startswith('- ') or s.startswith('* ') or re.match(r'^\d+\. ', s):
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', s)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.74)
        r = p.add_run(clean); r.font.size = Pt(12); continue
    
    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', s)
    if re.match(r'^[\s\d\-\|:\/.]+$', clean): continue
    if clean.startswith('\\') and not clean.startswith('\\#'): continue
    if '\\' in clean and not any(c.isalpha() for c in clean.replace('\\','').replace(' ','').replace('#','').replace('{','').replace('}','')): continue
    
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(clean); r.font.size = Pt(12)

# 参考文献
doc.add_paragraph()
h = doc.add_paragraph(); r = h.add_run('参考文献'); r.bold = True; r.font.size = Pt(14); r.font.name = '黑体'
for ref in [
    '[1] 马兴录. 基于YLYW易理模糊模型的机器人抓取决策研究. 技术论文, 2024. (未发表)',
    '[2] 马兴录, 等. YLYW递归汉语理解：基于易理原理的汉语理解架构. 技术论文, 2026. (未发表)',
    '[3] Shridhar, M., et al. ALFWorld: Aligning Text and Embodied Environments for Interactive Learning. ICLR, 2021.',
    '[4] Yao, S., et al. ReAct: Synergizing Reasoning and Acting in Language Models. ICLR, 2023.',
]:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.74); p.paragraph_format.first_line_indent = Cm(-0.37)
    r = p.add_run(ref); r.font.size = Pt(11)

doc.save('YLYW_obs_only_task_planning_paper.docx')
print('OK')
