#!/usr/bin/env python3
"""Markdown论文 → DOCX转换器（读取paper/技术论文_v0.1.md）"""
import sys, os, re

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
md_path = os.path.join(os.path.dirname(__file__), '技术论文_v0.1.md')

with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

def add_para(text, bold=False, italic=False, size=None, align=None, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if align is not None: para.alignment = align
    if color: run.font.color.rgb = RGBColor(*color)

def add_heading_styled(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

def add_table_from_md(lines):
    """Parse markdown table from lines and add to docx"""
    headers = [c.strip() for c in lines[0].split('|')[1:-1]]
    rows = []
    for line in lines[2:]:  # skip header and separator
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = hd
        for pp in cell.paragraphs:
            for run in pp.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = val
            for pp in cell.paragraphs:
                for run in pp.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

def process_text(text):
    """Process inline formatting: **bold**, *italic*"""
    # Handle **bold**
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Handle *italic*
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Handle `code`
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

# Parse markdown line by line
lines = content.split('\n')
i = 0
table_lines = []
in_table = False
in_code = False
in_refs = False

while i < len(lines):
    line = lines[i]
    
    # Skip empty lines
    if not line.strip():
        if in_table:
            add_table_from_md(table_lines)
            table_lines = []
            in_table = False
        i += 1
        continue
    
    # Reference section: just add as paragraph
    if line.startswith('## 参考文献'):
        in_refs = True
        add_heading_styled('参考文献', level=1)
        i += 1
        continue
    
    if in_refs and line.startswith('['):
        add_para(process_text(line.strip()), size=9)
        i += 1
        continue
    
    if in_refs:
        i += 1
        continue
    
    # Skip horizontal rules
    if line.strip() == '---':
        doc.add_paragraph()
        i += 1
        continue
    
    # Title
    if line.startswith('# ') and not line.startswith('## '):
        add_para(process_text(line[2:].strip()), bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Section headers
    if line.startswith('## '):
        add_heading_styled(process_text(line[3:].strip()), level=1)
        i += 1
        continue
    
    if line.startswith('### '):
        add_heading_styled(process_text(line[4:].strip()), level=2)
        i += 1
        continue
    
    # Author line
    if line.startswith('**') and '课题组' in line:
        add_para(process_text(line.strip()), size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Affiliation
    if '青岛科技大学' in line and '266061' in line:
        add_para(process_text(line.strip()), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Table detection
    if line.strip().startswith('|'):
        table_lines.append(line)
        in_table = True
        i += 1
        continue
    
    # Code blocks
    if line.strip().startswith('```'):
        in_code = not in_code
        i += 1
        continue
    
    if in_code:
        # Add as monospace paragraph
        p_code = doc.add_paragraph()
        run = p_code.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        i += 1
        continue
    
    # Bold text marker (like "**表1.**" or "**创新点1：**")
    is_bold = line.strip().startswith('**') and ('：' in line or ':**' in line)
    
    # Regular paragraph
    text = process_text(line.strip())
    
    # Handle bullet points
    if text.startswith('- ') or text.startswith('* ') or text.startswith('• '):
        text = text[2:]
    
    # Handle numbered items
    if re.match(r'^\d+\.\s', text):
        pass  # keep numbering
    
    add_para(text, bold=is_bold)
    
    i += 1

# Handle remaining table
if table_lines:
    add_table_from_md(table_lines)

# Save
out_path = os.path.join(os.path.dirname(__file__), 'YLYW技术论文_v0.1.docx')
doc.save(out_path)
print(f'✅ DOCX saved: {out_path}')
print(f'   Size: {os.path.getsize(out_path)} bytes')
