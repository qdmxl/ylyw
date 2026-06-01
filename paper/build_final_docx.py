#!/usr/bin/env python3
"""
生成学术排版Word文档
- 完整表格格式（边框、对齐、字体）
- 架构图嵌入
- 中文宋体/黑体 + 英文Times New Roman
- 标准学术论文排版
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re, os

# ============================================================
#  样式系统
# ============================================================
doc = Document()

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# 创建/修改样式
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(11)
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.space_after = Pt(6)

# 标题样式
for i, (size, color) in enumerate([(22, '1A1A2E'), (16, '1A5276'), (13, '2C3E50'), (12, '34495E')], 1):
    style_name = f'Heading {i}'
    if style_name in [s.name for s in doc.styles]:
        style = doc.styles[style_name]
    else:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

def add_body(text):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    return p

def add_heading_custom(text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_table(headers, rows, col_widths=None, caption=None):
    """添加格式化表格"""
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        # 表头背景色
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A5276"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 交替行背景
            if r % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    
    doc.add_paragraph()  # 表后间距
    return table

def add_code_block(code_text):
    """添加代码块（等宽字体段落）"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ============================================================
#  论文正文
# ============================================================

# === 封面标题 ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('YLYW：一种基于《易经》先验符号知识的\n联邦式神经符号具身决策框架')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.color.rgb = RGBColor.from_string('1A1A2E')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n马老师课题组')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('青岛科技大学 信息科学技术学院，山东 青岛 266061')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# === 摘要 ===
add_heading_custom('摘要', 1)
add_body('当前具身智能决策方法主要依赖数据驱动的深度学习范式——层次化决策依赖大模型预训练，端到端VLA模型需要数十万条机器人轨迹。这些方法在数据效率、可解释性和物理合规性方面存在根本性瓶颈。本文提出YLYW（易理研物），一种知识驱动的第三条道路：将《易经》的六十四卦符号系统形式化为联邦式神经符号架构中的先验知识手册。')

add_body('系统的核心创新在于：（1）以八卦作为连续模糊隶属度的符号基元，解决了传统符号系统的"符号接地"困境；（2）以六十四卦的结构化模板替代通用谓词逻辑，提供强归纳偏置；（3）提出基于物体质心的爻模板系统性优化方法，实现卦象模板与物理语义的精准对齐；（4）首次将《周易》"乘承比应当位得中"五种爻位关系形式化为可计算算子，实现卦象决定策略类型、爻位关系决定执行参数的分层决策。')

add_body('在42/64卦规则库下，300物体的零样本基线测试中，系统达到90.0%的策略合理率和0%的严重错误率，纯先验推理仅需1.7ms/物体。从初始48.0%到90.0%的性能跃迁，其提升过程本身即为易理模型可解释性的生动验证。进一步的小样本微调实验表明，使用仅24个专家标注演示即可将策略合理率从17.0%显著提升至51.5%。实验结果验证了YLYW在零样本和小样本条件下高效决策的潜力，为数据高效具身决策提供了一条可解释的新路径。')

p = doc.add_paragraph()
run = p.add_run('关键词：')
run.font.bold = True
run.font.size = Pt(11)
run = p.add_run('神经符号系统；《易经》；先验知识；零样本学习；机器人抓取；联邦式架构；爻位关系运算')
run.font.size = Pt(11)

doc.add_page_break()

# === 1 引言（精简版） ===
add_heading_custom('1 引言', 1)
add_heading_custom('1.1 问题背景', 2)
add_body('具身智能的目标是创造能在物理世界中自主行动的智能体。当前主流范式以深度学习为核心，本质是数据驱动的统计拟合。这导致了三个根本性挑战：数据饥渴（依赖海量交互轨迹）、黑箱决策（不可解释）和物理失配（缺乏物理常识）。这些困境共同指向一个核心问题：智能体是否应该拥有一套关于世界变化与运行的先验知识？深度学习范式的答案是"否"，而我们基于《易经》哲学的回答是"是"。')

add_heading_custom('1.2 《易经》作为一个被忽视的先验知识体系', 2)
add_body('《易经》是中国古代最重要的哲学经典，被誉为"群经之首"。从现代信息科学的视角审视，《易经》呈现出令人惊讶的结构化特征：以阴阳为原子符号，三爻成八卦（2³=8）、六爻成六十四卦（2⁶=64）；定义了爻位间的"乘、承、比、应、当位、得中"等结构化关系网络；通过"变卦"描述状态转移。这些特征使其天然适合作为计算先验。需要特别说明的是：本文将《易经》视为中国古代对自然变化规律的符号化总结与形式化建模，不涉及任何超自然解读。')

add_heading_custom('1.3 本文贡献', 2)
contributions = [
    '以八卦作为连续模糊隶属度的符号基元，解决传统符号系统的"符号接地"困境。',
    '设计联邦式三层架构（L1八卦基元→L2六爻编码→L3六十四卦规则），保持先验手册的计算独立性。',
    '提出基于物体爻向量质心的系统性模板优化方法，实现从"人工直觉赋值"到"数据驱动的先验精化"。',
    '首次将《周易》"乘承比应当位得中"五种爻位关系形式化为可计算算子，作为卦象匹配之后的质量评估与策略修正层。',
    '完整实现约2200行Python代码库，完成300物体零样本基线（90.0%合理率）、小样本微调、消融实验等系统性验证。',
]
for i, c in enumerate(contributions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f'{i}. {c}')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# === 3.1 总体架构 ===
add_heading_custom('2 系统架构', 1)
add_heading_custom('2.1 总体架构', 2)

add_body('YLYW采用联邦式神经符号架构，核心设计原则有三：（1）先验知识独立——易理手册作为独立的符号模块，保持计算完整性和可追溯性；（2）连续模糊表示——所有物理特征保留连续值[0,1]，通过模糊隶属度实现符号接地；（3）结构化先验偏置——利用64卦的内在关系网络提供强归纳偏置。系统完整架构如图1所示。')

# === 插入架构图 ===
img_path = '/home/lijinhan/MXL/科研/ylyw/paper/architecture_diagram.png'
if os.path.exists(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('图1. YLYW系统完整架构：物理世界→L1八卦→L2六爻→L3卦象→L3+爻位→决策→物理约束')
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

add_body('架构各层功能如下：L1八卦基元将13维连续物理特征模糊化为8个卦的隶属度向量，解决符号接地问题。L2六爻编码通过6个预定义加权公式将物理特征聚合为6维爻值向量。L3卦象匹配通过余弦相似度在42个卦象模板中搜索最佳匹配，决定策略类型（做什么）。L3+爻位关系通过乘承比应当位得中分析，决定策略执行参数（怎么做）。决策层整合卦象预设与爻位修正，输出可追溯的完整推理链。')

# === 表1: 八卦物理映射 ===
add_heading_custom('2.2 八卦-物理原型映射', 2)
add_body('八个卦象各自对应一组"理想物理属性"，构成先验知识的核心。表1展示了八卦的符号、卦德及其物理原型。')

add_table(
    ['卦名', '符号', '卦德', '自然象征', '力需求', '稳定性', '变形性', '滚动性', '可见性', '脆弱性'],
    [
        ['乾', '☰', '健', '天', '0.90', '0.80', '0.10', '0.20', '0.70', '0.20'],
        ['坤', '☷', '顺', '地', '0.30', '0.60', '0.80', '0.30', '0.50', '0.30'],
        ['震', '☳', '动', '雷', '0.50', '0.20', '0.40', '0.90', '0.60', '0.60'],
        ['艮', '☶', '止', '山', '0.40', '0.90', '0.20', '0.10', '0.50', '0.40'],
        ['离', '☲', '明/附丽', '火', '0.40', '0.50', '0.40', '0.40', '0.90', '0.50'],
        ['坎', '☵', '陷/险', '水', '0.60', '0.30', '0.50', '0.50', '0.40', '0.70'],
        ['兑', '☱', '悦', '泽', '0.30', '0.50', '0.60', '0.30', '0.60', '0.40'],
        ['巽', '☴', '入', '风', '0.40', '0.40', '0.50', '0.40', '0.50', '0.50'],
    ],
    caption='表1. 八卦-物理原型映射（每个卦的"理想物理属性"先验值）'
)

# === 表2: 抓取策略 ===
add_heading_custom('2.3 卦象-策略映射', 2)
add_body('六十四卦关联着预定义的抓取策略类型。表2列出了核心卦象及其对应的策略语义。')

add_table(
    ['卦名', '符号', '卦辞摘要', '策略类型', '力预设', '速度', '适用物体'],
    [
        ['乾为天', '☰☰', '刚健中正', 'power_grasp', '0.85', 'fast', '坚硬/重物'],
        ['坤为地', '☷☷', '柔顺包容', 'precision_grasp', '0.25', 'slow', '易碎/柔性'],
        ['震为雷', '☳☳', '震动不安', 'dynamic_grasp', '0.50', 'fast', '球体/圆柱'],
        ['艮为山', '☶☶', '静止如山', 'stable_grasp', '0.40', 'slow', '静止物体'],
        ['离为火', '☲☲', '光明附丽', 'adhesion_grasp', '0.40', 'medium', '光滑平面'],
        ['坎为水', '☵☵', '重险陷也', 'cautious_grasp', '0.50', 'slow', '带孔/凹陷'],
        ['兑为泽', '☱☱', '和悦相待', 'soft_grasp', '0.30', 'slow', '软质/食品'],
        ['巽为风', '☴☴', '随风而动', 'compliant_grasp', '0.35', 'medium', '不规则物体'],
        ['泽雷随', '☱☳', '随时而动', 'following_grasp', '0.50', 'medium', '滚动/滑动'],
        ['雷地豫', '☳☷', '预判而动', 'predictive_grasp', '0.45', 'fast', '移动中物体'],
        ['火泽睽', '☲☱', '求同存异', 'adaptive_irregular', '0.45', 'slow', '不规则形状'],
        ['天雷无妄', '☰☳', '真实不虚', 'direct_grasp', '0.55', 'fast', '常规形状'],
    ],
    caption='表2. 核心卦象-抓取策略映射（节选12/42卦）'
)

doc.add_page_break()

# === 3 爻位关系运算 ===
add_heading_custom('3 爻位关系运算', 1)
add_body('以上三层确定了"匹配哪个卦"——即策略的类型。但《易经》的推理不止于此：在确定卦象之后，还须分析六爻之间的结构关系，以判断策略执行的质量与可信度。本节将《周易》中五种经典的爻位关系——当位、得中、乘承、亲比、呼应——形式化为可计算的算子，作为卦象匹配之后的质量评估与策略修正层（L3+）。')

add_heading_custom('3.1 算法：爻位关系运算', 2)
add_body('算法1展示了完整的爻位关系运算伪代码。输入为6维爻向量，输出为综合爻位质量评分和策略修正系数。')

# 插入伪代码
code = '''Algorithm 1: Yao Relations Analysis
──────────────────────────────────────────────
Input:  yao_vector ∈ [0,1]⁶
Output: S_yao (quality score), modifier (force correction)

1. 当位分析:
   yang_positions ← {0, 2, 4}  // 阳位: 初(0),三(2),五(4)
   dangwei ← count(i: (yao[i]≥0.5) == (i∈yang_positions))
   S_dw ← dangwei / 6.0

2. 得中分析:
   er_is_yin ← yao[1] < 0.5       // 二爻(1)为阴 = 六二
   wu_is_yang ← yao[4] ≥ 0.5      // 五爻(4)为阳 = 九五
   if er_is_yin ∧ wu_is_yang: S_dz ← 1.0
   elif wu_is_yang:           S_dz ← 0.75
   elif er_is_yin:            S_dz ← 0.50
   else:                      S_dz ← 0.25

3. 乘承分析 (i=0..4, 5对相邻爻):
   cheng ← 0; cheng_hao ← 0
   for i in 0..4:
     if yao[i+1]<0.5 ∧ yao[i]≥0.5:   cheng += 1      // 乘(逆)
     elif yao[i]<0.5 ∧ yao[i+1]≥0.5: cheng_hao += 1  // 承(顺)
   S_cc ← max(0, 1.0 - cheng×0.3 + cheng_hao×0.15)

4. 亲比分析:
   harmony ← count(i in 0..4: (yao[i]≥0.5)==(yao[i+1]≥0.5))
   S_bi ← harmony / 5.0

5. 呼应分析:
   ying ← count((a,b)∈{(0,3),(1,4),(2,5)}:
                (yao[a]≥0.5)≠(yao[b]≥0.5))
   S_ying ← ying / 3.0

6. 综合评分:
   S_yao ← 0.40·S_dw + 0.20·S_dz + 0.15·S_cc
           + 0.10·S_bi + 0.15·S_ying

7. 策略修正:
   modifier ← 1.0
   if S_dw ≤ 0.33: modifier -= 0.10
   if cheng ≥ 2:   modifier -= 0.10
   if ying == 0:   modifier -= 0.05
   elif ying == 3: modifier += 0.05

return S_yao, clamp(modifier, 0.75, 1.05)'''
add_code_block(code)

add_heading_custom('3.2 五种爻位关系的易学基础', 2)

# 表3: 爻位关系定义
add_table(
    ['关系', '周易定义', '程序化判据', '策略影响'],
    [
        ['当位', '阳爻居初/三/五位，阴爻居二/四/上位', '(yao[i]≥0.5) == (i∈{0,2,4})', '影响稳定性基础'],
        ['得中', '二爻和五爻为"中位"，六二、九五最佳', '六二阴(yao[1]<0.5) + 九五阳(yao[4]≥0.5)', '影响决策优先级'],
        ['乘(逆)', '阴爻在上压制下方阳爻', 'yao[i+1]<0.5 and yao[i]≥0.5', '降低期望力 (-0.10/处)'],
        ['承(顺)', '阴爻在下承载上方阳爻', 'yao[i]<0.5 and yao[i+1]≥0.5', '提升策略流畅度'],
        ['亲比', '相邻两爻同阴阳为亲比', '(yao[i]≥0.5)==(yao[i+1]≥0.5)', '影响动作连贯性'],
        ['呼应', '初-四/二-五/三-上，阴阳相反为"有应"', '(yao[a]≥0.5)≠(yao[b]≥0.5)', '影响策略确信度'],
    ],
    caption='表3. 五种爻位关系的形式化定义'
)

add_body('五种关系的评分按权重融合为综合爻位质量：S_yao = 0.40·S_dw + 0.20·S_dz + 0.15·S_cc + 0.10·S_bi + 0.15·S_ying。权重分配反映了易学传统中对各关系的重视程度：当位最重（"天地设位"），得中次之（"中正之道"），乘承与应再次，亲比最末。')

# 表4: 谨慎级别
add_table(
    ['爻位质量', '谨慎级别', '力修正系数', '含义'],
    [
        ['≥ 0.70', 'relaxed', '1.00 – 1.05', '爻位优良，可略增力'],
        ['0.50 – 0.70', 'normal', '0.95 – 1.00', '爻位正常，标准执行'],
        ['0.30 – 0.50', 'cautious', '0.85 – 0.95', '爻位偏差，降低力预设'],
        ['< 0.30', 'very_cautious', '0.75 – 0.85', '爻位严重失序，大幅降力'],
    ],
    caption='表4. 爻位质量→谨慎级别→力修正系数映射'
)

doc.add_page_break()

# === 4 实验 ===
add_heading_custom('4 实验', 1)
add_heading_custom('4.1 实验设置', 2)
add_body('定义了8种基本物体类型（球体、立方体、圆柱体、碗、瓶子、盘子、不规则石块、花瓶），每种有预设物理参数模板。使用合成场景生成器（SimulationScene）按物理模板生成带±5%噪声的特征向量。所有实验不使用任何训练数据，纯先验知识推理。评估采用启发式策略合理性评估：每种物体类型预定义"合理策略集"与"明显错误策略集"。')

add_heading_custom('4.2 爻模板优化前后对比', 2)

# 表5: 优化前基线
add_table(
    ['指标', '优化前(20卦)', '优化后(42卦+质心)', '提升'],
    [
        ['合理匹配率', '48.0%', '90.0%', '+42.0%'],
        ['明显错误率', '2.0%', '0.0%', '−2.0%'],
        ['可用卦象', '20/64', '42/64', '+22卦'],
        ['策略类型覆盖', '18', '39', '+21'],
        ['推理速度', '0.5ms', '1.7ms', '—'],
        ['模型参数量', '0', '0', '—'],
    ],
    caption='表5. 爻模板优化前后零样本基线对比'
)

# 表6: 按物体类型
add_table(
    ['物体类型', '优化前合理率', '优化后合理率', '提升', '典型策略', '典型卦象'],
    [
        ['盘子', '100.0%', '100.0%', '—', 'low_visibility_grasp', '地火明夷'],
        ['花瓶', '50.0%', '97.1%', '+47.1%', 'progressive_grasp', '火地晋'],
        ['立方体', '57.1%', '95.0%', '+37.9%', 'top_down_grasp', '地泽临'],
        ['碗', '33.3%', '94.7%', '+61.4%', 'progressive_grasp', '风天小畜'],
        ['球体', '28.6%', '93.8%', '+65.2%', 'following_grasp', '泽雷随'],
        ['瓶子', '100.0%', '91.7%', '−8.3%', 'coordinated_grasp', '风火家人'],
        ['圆柱体', '0.0%', '80.0%', '+80.0%', 'endurance_grasp', '雷风恒'],
        ['石块', '16.7%', '77.4%', '+60.7%', 'corrective_grasp', '山风蛊'],
    ],
    caption='表6. 按物体类型的优化前后结果对比（300物体大样本）'
)

# 表7: 爻位关系统计
add_table(
    ['指标', '数值', '说明'],
    [
        ['力修正系数均值', '0.94', '多数情境力预设小幅下调'],
        ['爻位质量均值', '0.50', '自然物体处于中等爻位质量'],
        ['cautious占比', '52%', '超半数情境被识别为需谨慎'],
        ['normal占比', '46%', '正常执行情境'],
        ['relaxed占比', '1%', '极少优化情境'],
        ['very_cautious', '0%', '300物体中无触发'],
    ],
    caption='表7. 爻位关系运算全局统计（300物体）'
)

# 表8: 爻位关系案例分析
add_table(
    ['案例', '物体', '当位', '乘(逆)', '应', '爻位质量', '力修正', '谨慎级别'],
    [
        ['最佳', '盘子', '3/6', '1处', '2/3', '0.64', '1.00', 'normal'],
        ['典型', '球体', '1/6', '1处', '2/3', '0.47', '0.90', 'cautious'],
        ['最差', '石块', '0/6', '2处', '3/3', '0.38', '0.85', 'cautious'],
    ],
    caption='表8. 爻位关系典型案例分析'
)

# 表9: 与VLA对比
add_table(
    ['维度', 'VLA典型值', 'YLYW'],
    [
        ['训练数据', '10⁴ – 10⁶ 轨迹', '0'],
        ['推理硬件', 'GPU集群', 'CPU'],
        ['代码量', '数万行', '约2200行'],
        ['模型大小', '100MB – 10GB', '约80KB'],
        ['推理延迟', '10 – 1000ms', '1.7ms'],
        ['可解释性', '黑箱', '卦→爻→辞全链追溯'],
        ['物理约束', '隐式学习', '显式嵌入'],
        ['零样本合理率', '—', '90.0%'],
    ],
    caption='表9. YLYW与代表性VLA模型的间接比较'
)

doc.add_page_break()

# === 5 结论 ===
add_heading_custom('5 结论与未来工作', 1)
add_body('本文提出并实现了YLYW——一种基于《易经》先验符号知识的联邦式神经符号具身决策框架。系统将八卦的连续模糊隶属度、六爻编码、六十四卦结构模板和爻位关系运算整合为统一的先验推理引擎。在42/64卦规则库下，300物体零样本基线达到90.0%合理率、0%严重错误率，验证了知识驱动方法在数据效率和可解释性方面的根本性优势。')

add_body('爻模板质心优化方法展示了可解释性作为工程工具的独特价值——从48%到90%的性能跃迁中，每一步提升都有可追溯的物理语义支撑。爻位关系运算层的引入进一步完善了易理模型，使卦象决定"做什么"而爻位关系决定"怎么做"的分层决策体系得以建立。')

add_body('未来工作将集中在：（1）补齐剩余22卦的工程转译；（2）在真实机器人上验证抓取策略；（3）引入物理约束层（逆动力学硬性映射）实现零物理违规保证；（4）构建与LLM的混合架构原型，使YLYW作为可靠的"系统1"直觉核心与大模型的"系统2"长程规划能力形成互补。')

doc.add_page_break()

# === 参考文献 ===
add_heading_custom('参考文献', 1)

refs = [
    '[1] Driess D, et al. PaLM-E: An Embodied Multimodal Language Model. ICML, 2023.',
    '[2] Ahn M, et al. Do As I Can, Not As I Say: Grounding Language in Robotic Affordances. CoRL, 2022.',
    '[3] Brohan A, et al. RT-2: Vision-Language-Action Models. arXiv:2307.15818, 2023.',
    '[4] Kim M J, et al. OpenVLA: An Open-Source Vision-Language-Action Model. arXiv:2406.09246, 2024.',
    '[5] Octo Model Team. Octo: An Open-Source Generalist Robot Policy. RSS, 2024.',
    '[6] 汪玉等. 具身智能中VLA模型的现状与挑战. 中国科学: 信息科学, 2025.',
    '[7] Black K, et al. π₀: A Vision-Language-Action Flow Model. arXiv:2410.24164, 2024.',
    '[8] Garcez A, Lamb L C. Neurosymbolic AI: The 3rd Wave. Artificial Intelligence Review, 2023.',
    '[9] Manhaeve R, et al. DeepProbLog: Neural Probabilistic Logic Programming. NeurIPS, 2018.',
    '[10] Badreddine S, et al. Logic Tensor Networks. Artificial Intelligence, 2022.',
    '[11] van Krieken E, et al. Analyzing Differentiable Fuzzy Logic Operators. AIJ, 2022.',
    '[12] Morrison D, et al. Closing the Loop for Robotic Grasping. RSS, 2018.',
    '[13] Fang H S, et al. GraspNet-1Billion. CVPR, 2020.',
    '[14] Fang H S, et al. AnyGrasp: Robust Grasp Perception. IEEE TRO, 2023.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(ref)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('附：易学原典参考文献')
run.font.bold = True
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

classics = [
    '[15] 王弼（魏）. 周易注.',
    '[16] 孔颖达（唐）. 周易正义.',
    '[17] 程颐（宋）. 程氏易传.',
    '[18] 朱熹（宋）. 周易本义.',
    '[19] 黄寿祺, 张善文. 周易译注. 上海古籍出版社, 2007.',
    '[20] 高亨. 周易古经今注. 中华书局, 1984.',
    '[21] 李学勤. 周易溯源. 巴蜀书社, 2005.',
]
for ref in classics:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(ref)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
#  保存
# ============================================================
out_path = '/home/lijinhan/MXL/科研/ylyw/paper/YLYW技术论文_v0.4.docx'
doc.save(out_path)
print(f'✅ 论文已保存: {out_path}')
