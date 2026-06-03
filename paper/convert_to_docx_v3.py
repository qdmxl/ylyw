#!/usr/bin/env python3
"""Markdown论文 v0.6 → DOCX转换器（图片嵌入 + 上标引用版）"""
import sys, os, re

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
md_path = os.path.join(os.path.dirname(__file__), '技术论文_v0.6.md')

with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# Normal style
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

PAPER_DIR = os.path.dirname(os.path.abspath(__file__))

def add_para_superscript(text, bold=False, italic=False, size=None, align=None, indent=True, font_name=None):
    """Add paragraph with superscript citations and proper formatting"""
    if not text.strip():
        return None
    
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    else:
        para.paragraph_format.first_line_indent = Cm(0)
    
    if align is not None:
        para.alignment = align
    
    # Split text by citation patterns [number] or [number,number] or [number-number]
    # Pattern: [...] with digits, commas, hyphens
    pattern = r'(\[\d+(?:[-,]\d+)*\])'
    parts = re.split(pattern, text)
    
    fname = font_name or '宋体'
    fsize = size or 12
    
    for part in parts:
        if not part:
            continue
        if re.match(r'\[\d+(?:[-,]\d+)*\]$', part):
            # Citation - superscript
            run = para.add_run(part)
            run.font.superscript = True
            run.font.size = Pt(fsize - 2)  # slightly smaller for superscript
            run.font.name = fname
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
            run.bold = bold
        else:
            run = para.add_run(part)
            run.font.size = Pt(fsize)
            run.font.name = fname
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
            run.bold = bold
            run.italic = italic
    
    return para

def add_heading_styled(text, level):
    """Add heading"""
    h = doc.add_heading(text, level=min(level, 3))
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(15)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(13)
    h.paragraph_format.first_line_indent = Cm(0)
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(3)
    return h

def add_image(img_name, width_inches=5.5):
    """Insert an image centered"""
    img_path = os.path.join(PAPER_DIR, img_name)
    if not os.path.exists(img_path):
        print(f"  ⚠️  Image not found: {img_name}")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

def add_table_from_md(lines):
    """Parse markdown table and add to docx"""
    if len(lines) < 2:
        return
    parts = lines[0].strip().split('|')
    if len(parts) < 3:
        return
    headers = [p.strip() for p in parts[1:-1]]
    if not headers:
        return
    
    rows = []
    for line in lines[2:]:
        parts = line.strip().split('|')
        if len(parts) >= 3:
            inner = [p.strip() for p in parts[1:-1]]
            if inner and any(c for c in inner):
                rows.append(inner)
    if not rows:
        return
    
    ncols = len(headers)
    for row in rows:
        while len(row) < ncols:
            row.append('')
    
    table = doc.add_table(rows=1+len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, hd in enumerate(headers):
        if i >= ncols: break
        cell = table.rows[0].cells[i]
        hd = re.sub(r'\*\*(.+?)\*\*', r'\1', hd)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(hd)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml import OxmlElement
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)
    
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            if c >= ncols: break
            cell = table.rows[r+1].cells[c]
            val = re.sub(r'\*\*(.+?)\*\*', r'\1', val)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(3)
    spacer.paragraph_format.space_after = Pt(3)
    spacer.paragraph_format.first_line_indent = Cm(0)

def process_text(text):
    """Remove markdown formatting"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

# Define figure triggers: (trigger text pattern, image filename)
FIGURE_MAP = [
    ('图1. YLYW系统完整架构', 'architecture_diagram.png'),
    ('图2. 各物体类型零样本合理率优化前后柱状对比', 'fig_before_after_comparison.png'),
    ('图3. 卦象命中分布', 'fig_hexagram_hit_distribution.png'),
    ('图4. 爻位质量分布直方图', 'fig_yao_quality_distribution.png'),
    ('图5. 双八卦安全评估下按物体类型的力修正系数对比', 'fig_force_modifier_by_type.png'),
]

# Parse state
lines = content.split('\n')
i = 0
table_lines = []
in_table = False
in_code = False
in_refs = False
in_appendix = False

while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    
    # Empty lines
    if not stripped and not in_code:
        if in_table:
            add_table_from_md(table_lines)
            table_lines = []
            in_table = False
        i += 1
        continue
    
    # Reference section
    if stripped.startswith('## 参考文献'):
        in_refs = True
        in_appendix = False
        add_heading_styled('参考文献', level=1)
        i += 1
        continue
    
    if stripped.startswith('### 附：'):
        in_refs = True
        in_appendix = False
        add_heading_styled(stripped.replace('### ', '').strip(), level=2)
        i += 1
        continue
    
    # Reference items (no superscript in reference list, no indent)
    if in_refs and (stripped.startswith('[') or stripped.startswith('*论文版本') or stripped.startswith('*完成日期')):
        text = process_text(stripped)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        i += 1
        continue
    
    if in_refs:
        i += 1
        continue
    
    # Appendix
    if stripped.startswith('## 附录'):
        in_appendix = True
        in_refs = False
        add_heading_styled(stripped.replace('## ', '').strip(), level=1)
        i += 1
        continue
    
    if in_appendix and stripped.startswith('表'):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(process_text(stripped))
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        i += 1
        continue
    
    # Horizontal rules
    if stripped == '---':
        i += 1
        continue
    
    # Title
    if line.startswith('# ') and not line.startswith('## '):
        title_text = process_text(line[2:].strip())
        while i+1 < len(lines) and lines[i+1].strip() and not lines[i+1].strip().startswith('#') and not lines[i+1].strip().startswith('**'):
            i += 1
            title_text += '\n' + process_text(lines[i].strip())
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        i += 1
        continue
    
    # Author / affiliation
    if ('马兴录' in stripped or '课题组' in stripped) and not stripped.startswith('#'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(process_text(stripped))
        run.font.size = Pt(14)
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        i += 1
        continue
    
    if ('青岛科技大学' in stripped or '266061' in stripped or '通讯作者' in stripped) and not stripped.startswith('#'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(process_text(stripped))
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        i += 1
        continue
    
    # H2
    m_h2 = re.match(r'^## (.+)', line)
    if m_h2:
        text = process_text(m_h2.group(1).strip())
        if len(text) < 60:
            add_heading_styled(text, level=1)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        i += 1
        continue
    
    # H3
    m_h3 = re.match(r'^### (.+)', line)
    if m_h3:
        text = process_text(m_h3.group(1).strip())
        if len(text) < 60:
            add_heading_styled(text, level=2)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        i += 1
        continue
    
    # H4
    m_h4 = re.match(r'^#### (.+)', line)
    if m_h4:
        text = process_text(m_h4.group(1).strip())
        if len(text) < 60:
            add_heading_styled(text, level=3)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        i += 1
        continue
    
    # Blockquote
    if stripped.startswith('> '):
        text = process_text(stripped[2:])
        add_para_superscript(text, italic=True, size=10.5)
        i += 1
        continue
    
    # Table detection
    if stripped.startswith('|'):
        table_lines.append(line)
        in_table = True
        i += 1
        continue
    
    # Code blocks
    if stripped.startswith('```'):
        in_code = not in_code
        i += 1
        continue
    
    if in_code:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        i += 1
        continue
    
    # Check for figure trigger
    is_figure = False
    for trigger, img_name in FIGURE_MAP:
        if trigger in stripped:
            # Add the caption paragraph
            text = process_text(stripped)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # Insert the image
            add_image(img_name)
            is_figure = True
            break
    
    if is_figure:
        i += 1
        continue
    
    # Figure/Table caption (bold, no indent, centered)
    is_caption = bool(re.match(r'^(图\d+|表\d+)', stripped))
    
    # Regular paragraph
    text = process_text(stripped)
    if not text:
        i += 1
        continue
    
    # Footer
    if text.startswith('*论文版本') or text.startswith('*完成日期') or text.startswith('*青岛'):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        i += 1
        continue
    
    # Bold header like "**范式一：**"
    is_bold_header = bool(re.match(r'\*\*.+：\*\*', stripped))
    
    if is_caption:
        # Table captions: bold, no indent
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    else:
        add_para_superscript(text, bold=is_bold_header)
    
    i += 1

# Handle remaining table
if table_lines:
    add_table_from_md(table_lines)

# Save
out_path = os.path.join(PAPER_DIR, 'YLYW技术论文_v0.6.docx')
doc.save(out_path)
print(f'✅ Saved: {out_path}')
print(f'   Size: {os.path.getsize(out_path)} bytes')
print(f'   Paragraphs: {len(doc.paragraphs)}')
print(f'   Tables: {len(doc.tables)}')
print(f'   Images embedded: 5')
