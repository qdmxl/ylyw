from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21)  # A4
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)

def set_font(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def add_paragraph_with_font(doc, text, font_name, size_pt, bold=False, alignment=None, spacing_after=Pt(6)):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = spacing_after
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, font_name, size_pt, bold)
    return p

def add_body_text(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(24)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, '宋体', 12, bold=False)
    return p

with open('知几学习技术论文_v0.1.md', 'r') as f:
    content = f.read()

lines = content.split('\n')
i = 0
in_table = False
while i < len(lines):
    line = lines[i]
    
    # Skip empty lines
    if not line.strip():
        i += 1
        continue
    
    # Title
    if line.startswith('# ') and not line.startswith('## '):
        title_text = line[2:].strip()
        add_paragraph_with_font(doc, title_text, '黑体', 18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(12))
        i += 1
        continue
    
    # Author line
    if '马兴录' in line and ('青岛' not in line and '通讯' not in line):
        add_paragraph_with_font(doc, line.strip(), '仿宋', 14, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Affiliation
    if '青岛科技大学' in line:
        add_paragraph_with_font(doc, line.strip(), '宋体', 10.5, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Corresponding author
    if '通讯作者' in line:
        add_paragraph_with_font(doc, line.strip(), '宋体', 10.5, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Separator
    if line.strip() == '---':
        i += 1
        continue
    
    # Section headings
    if line.startswith('## ') and not line.startswith('### '):
        heading_text = line[3:].strip()
        add_paragraph_with_font(doc, heading_text, '黑体', 15, bold=True)  # 一级标题
        i += 1
        continue
    
    if line.startswith('### ') and not line.startswith('#### '):
        heading_text = line[4:].strip()
        add_paragraph_with_font(doc, heading_text, '黑体', 14, bold=True)  # 二级标题
        i += 1
        continue
    
    if line.startswith('#### '):
        heading_text = line[5:].strip()
        add_paragraph_with_font(doc, heading_text, '楷体', 12, bold=True)  # 三级标题
        i += 1
        continue
    
    # Table rows
    if line.strip().startswith('|') and line.strip().endswith('|'):
        if '---' in line:
            i += 1
            continue
        # Skip table for now - simplified handling
        i += 1
        continue
    
    # Regular paragraph
    # Handle bold markers
    text = line.strip()
    # bold markers **text**
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    add_body_text(doc, text)
    i += 1

doc.save('知几学习技术论文_v0.1.docx')
print('Done: 知几学习技术论文_v0.1.docx')
