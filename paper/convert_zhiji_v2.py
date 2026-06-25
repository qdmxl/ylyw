from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import copy

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)

def set_font(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def add_superscript(run, text):
    """Add superscript text (for references)"""
    sup_run = run._element.makeelement(qn('w:r'), {})
    rPr = OxmlElement('w:rPr')
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)
    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:hAnsi'), '宋体')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')  # 9pt = 18 half-points
    rPr.append(sz)
    sup_run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    sup_run.append(t)
    run._element.addnext(sup_run)

def add_paragraph_with_font(doc, text, font_name, size_pt, bold=False, alignment=None, spacing_after=Pt(6), first_line_indent=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = spacing_after
    pf.line_spacing = 1.5
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_font(run, font_name, size_pt, bold)
    return p, run

def add_body_with_refs(doc, text):
    """Add body paragraph with [N] superscript references"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(24)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    
    # Split text by reference markers like [1], [2], [7,8] etc.
    parts = re.split(r'(\[\d+(?:,\d+)*\])', text)
    for part in parts:
        if re.match(r'^\[\d+(?:,\d+)*\]$', part):
            # Reference - add as superscript
            run = p.add_run(part)
            set_font(run, '宋体', 12, bold=False)
            # Make it superscript
            rPr = run._element.get_or_add_rPr()
            vertAlign = OxmlElement('w:vertAlign')
            vertAlign.set(qn('w:val'), 'superscript')
            rPr.append(vertAlign)
            # slightly smaller
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '18')
            rPr.append(sz)
        else:
            run = p.add_run(part)
            set_font(run, '宋体', 12, bold=False)
    return p

def add_heading_text(doc, text, level=1):
    """Add heading with proper style, handling refs"""
    if level == 1:
        font, size = '黑体', 15
    elif level == 2:
        font, size = '黑体', 14
    else:
        font, size = '楷体', 12
    
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    
    parts = re.split(r'(\[\d+(?:,\d+)*\])', text)
    for part in parts:
        if re.match(r'^\[\d+(?:,\d+)*\]$', part):
            run = p.add_run(part)
            set_font(run, font, size, bold=True)
            rPr = run._element.get_or_add_rPr()
            vertAlign = OxmlElement('w:vertAlign')
            vertAlign.set(qn('w:val'), 'superscript')
            rPr.append(vertAlign)
        else:
            run = p.add_run(part)
            set_font(run, font, size, bold=True)
    return p

with open('知几学习技术论文_v0.2.md', 'r') as f:
    content = f.read()

lines = content.split('\n')
i = 0

while i < len(lines):
    line = lines[i]
    
    if not line.strip():
        i += 1
        continue
    
    # Title
    if line.startswith('# ') and not line.startswith('## '):
        title_text = line[2:].strip()
        add_paragraph_with_font(doc, title_text, '黑体', 18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(12))
        i += 1
        continue
    
    # Author line
    if '马兴录' in line and '青岛' not in line and '通讯' not in line and '摘要' not in line and '关键词' not in line:
        p, run = add_paragraph_with_font(doc, line.strip(), '仿宋', 14, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Affiliation
    if '青岛科技大学' in line:
        add_paragraph_with_font(doc, line.strip(), '宋体', 10.5, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Corresponding author
    if '通讯作者' in line:
        add_paragraph_with_font(doc, line.strip(), '宋体', 10.5, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Separator
    if line.strip() == '---':
        i += 1
        continue
    
    # Section headings
    if line.startswith('## ') and not line.startswith('### '):
        heading_text = line[3:].strip()
        add_heading_text(doc, heading_text, level=1)
        i += 1
        continue
    
    if line.startswith('### ') and not line.startswith('#### '):
        heading_text = line[4:].strip()
        add_heading_text(doc, heading_text, level=2)
        i += 1
        continue
    
    if line.startswith('#### '):
        heading_text = line[5:].strip()
        add_heading_text(doc, heading_text, level=3)
        i += 1
        continue
    
    # Blockquote lines
    if line.strip().startswith('>'):
        text = line.strip()[1:].strip()
        add_body_with_refs(doc, text)
        i += 1
        continue
    
    # Table rows - skip
    if line.strip().startswith('|'):
        i += 1
        continue
    
    # Math lines ($$) - add as is
    if line.strip().startswith('$$') and line.strip().endswith('$$'):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_after = Pt(6)
        pf.line_spacing = 1.5
        run = p.add_run(line.strip())
        set_font(run, 'Times New Roman', 11, bold=False)
        i += 1
        continue
    
    # Regular paragraph
    text = line.strip()
    # Remove bold markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    add_body_with_refs(doc, text)
    i += 1

doc.save('知几学习技术论文_v0.2.docx')
print('Done: 知几学习技术论文_v0.2.docx')
