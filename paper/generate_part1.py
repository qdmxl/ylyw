#!/usr/bin/env python3
"""Generate YLYW paper Part 1 (Chapters 1-3) as .docx"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ============================================================
# Font & Style Setup
# ============================================================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.first_line_indent = Pt(21)
style.paragraph_format.line_spacing = 1.5

# Heading styles
for level, size in [(1, 14), (2, 12), (3, 10.5)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.paragraph_format.first_line_indent = Pt(0)
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)


def add_paragraph(text, style_name='Normal', alignment=None, first_indent=True):
    """Add a paragraph with proper formatting."""
    p = doc.add_paragraph(style=style_name)
    if alignment:
        p.alignment = alignment
    if not first_indent:
        p.paragraph_format.first_line_indent = Pt(0)
    # Parse text for references [n] → superscript
    import re
    parts = re.split(r'(\[\d+\])', text)
    for part in parts:
        if re.match(r'\[\d+\]', part):
            run = p.add_run(part)
            run.font.superscript = True
            run.font.size = Pt(8)
        else:
            run = p.add_run(part)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_formula(text, label=None):
    """Add centered formula in Cambria Math."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    if label:
        run2 = p.add_run(f'    {label}')
        run2.font.name = 'Cambria Math'
        run2.font.size = Pt(11)
    return p


def add_heading(text, level=1):
    """Add heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


# ============================================================
# Title Page
# ============================================================
# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(60)
run = p.add_run('YLYW：一种基于《易经》先验符号知识的可学习具身决策框架')
run.font.name = '黑体'
run.font.size = Pt(16)
run.font.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# Authors
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(24)
run = p.add_run('马兴录，马圣洁，张国安，李金函，于敬涛，李望')
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Affiliation
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('青岛科技大学 信息科学技术学院，山东 青岛 266061')
run.font.name = '宋体'
run.font.size = Pt(10.5)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Corresponding author
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('通讯作者：马兴录 E-mail: maxinglu@qust.edu.cn')
run.font.name = '宋体'
run.font.size = Pt(10.5)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Classification
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(12)
run = p.add_run('中图分类号：TP18')
run.font.name = '宋体'
run.font.size = Pt(10.5)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# Chinese Abstract
# ============================================================
add_heading('摘要', level=1)

abstract_zh = (
    '当前具身智能决策方法主要依赖数据驱动的深度学习范式。大模型预训练方法（如PaLM-E）和端到端视觉-语言-动作（VLA）模型（如RT-2）'
    '需要海量训练数据，且决策过程不可解释，成为制约具身智能走向实际应用的关键瓶颈。本文提出YLYW（易理研物），一种基于《易经》六十四卦'
    '先验符号知识的可学习具身决策框架。YLYW以八卦作为连续模糊隶属度的符号基元，解决了传统符号系统的接地困境；以六十四卦结构化模板替代'
    '通用谓词逻辑，提供强归纳偏置；以"乘承比应当位得中"五种爻位关系实现策略参数的精细修正，构成"卦定策略类型、爻定执行参数"的联邦式'
    '三层推理架构（L1八卦隶属度、L2六爻编码、L3卦象匹配）。在此基础上，本文进一步提出知几学习——一种作用于先验参数的对称校准机制。'
    '知几学习源自"知几其神乎"的哲学洞察，从微小征兆中察觉成功模式（吉之几，正向强化）和失败模式（凶之几，负向抑制），对位置先验矩阵'
    'P(obj,loc)、物体匹配置信度M(target,entity)和同义词映射S(word→set)三类参数做精确对称更新。实验在两个域上验证：物理域300物体零样本'
    '决策达92.7%合理率（消融实验验证易理规则贡献+33.6%、三层架构贡献+12.7%、模糊隶属度贡献+23.0%的独立贡献）；ALFWorld具身导航域134局'
    '测试中，静态先验达67.2%，引入知几学习后逐步提升至73.9%并一轮收敛（标注一致子集达90.0%，接近270亿参数大语言模型的93.28%）。本文'
    '首次证明了基于《易经》先验知识的可解释系统，在配合知几学习的精确校准后，可在多个具身智能域实现接近大模型的决策能力，为具身智能提供了'
    '一条数据高效、可解释、可增长的新路径。'
)
add_paragraph(abstract_zh)

# Keywords
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(12)
run = p.add_run('关键词：')
run.font.name = '黑体'
run.font.size = Pt(10.5)
run.font.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run = p.add_run('YLYW；易理先验知识；神经符号系统；模糊隶属度；知几学习；对称校准；零样本决策；具身智能；ALFWorld；可解释AI')
run.font.name = '宋体'
run.font.size = Pt(10.5)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# English Title & Abstract
# ============================================================
add_heading('YLYW: A Learnable Embodied Decision Framework Based on I Ching Prior Symbolic Knowledge', level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('MA Xinglu, MA Shengjie, ZHANG Guoan, LI Jinhan, YU Jingtao, LI Wang')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('School of Information Science and Technology, Qingdao University of Science and Technology, Qingdao 266061, China')
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)

# English Abstract
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Abstract: ')
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)
run.font.bold = True

abstract_en = (
    'Current embodied intelligence decision methods primarily rely on data-driven deep learning paradigms. '
    'Large-scale pretrained models (e.g., PaLM-E) and end-to-end Vision-Language-Action (VLA) models (e.g., RT-2) '
    'demand massive training data while producing opaque, uninterpretable decisions. This paper proposes YLYW '
    '(Yi-Li-Yan-Wu, literally "Studying Things through I Ching Principles"), a learnable embodied decision framework '
    'grounded in the prior symbolic knowledge of the 64 hexagrams from the I Ching (Book of Changes). YLYW employs '
    'the Eight Trigrams as continuous fuzzy membership primitives to resolve the symbol grounding problem, uses the '
    '64 hexagram structured templates as strong inductive biases replacing generic predicate logic, and applies five '
    'yao-position relations (cheng, cheng, bi, ying, dang-wei-de-zhong) for fine-grained parameter correction. This '
    'constitutes a federated three-layer reasoning architecture: L1 trigram membership, L2 six-yao encoding, and '
    'L3 hexagram matching, following the principle that "hexagrams determine strategy type; yao lines determine '
    'execution parameters." Furthermore, we introduce Zhi-Ji Learning (Incipience Learning), a symmetric calibration '
    'mechanism operating on prior parameters. Inspired by the philosophical insight "How divine is the knowledge of '
    'incipience," Zhi-Ji Learning detects subtle patterns of success (auspicious incipience, positive reinforcement) '
    'and failure (inauspicious incipience, negative suppression), performing precise symmetric updates on three parameter '
    'classes: position prior matrix P(obj,loc), object matching confidence M(target,entity), and synonym mapping '
    'S(word→set). Experiments validate the framework across two domains: in the physical domain, 300-object zero-shot '
    'decision-making achieves 92.7% rationality (ablation studies confirm independent contributions of I Ching rules '
    '+33.6%, three-layer architecture +12.7%, and fuzzy membership +23.0%); in the ALFWorld embodied navigation domain, '
    '134-episode testing shows static priors reaching 67.2%, improving to 73.9% with Zhi-Ji Learning and converging '
    'within one round (reaching 90.0% on annotation-consistent subsets, approaching the 93.28% of a 27-billion-parameter '
    'LLM). This work demonstrates for the first time that an interpretable system based on I Ching prior knowledge, '
    'when combined with precise Zhi-Ji calibration, can achieve decision capabilities approaching large models across '
    'multiple embodied intelligence domains.'
)
run2 = p.add_run(abstract_en)
run2.font.name = 'Times New Roman'
run2.font.size = Pt(10.5)

# English Keywords
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Keywords: ')
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)
run.font.bold = True
run = p.add_run('YLYW; I Ching prior knowledge; neuro-symbolic system; fuzzy membership; Zhi-Ji Learning; symmetric calibration; zero-shot decision; embodied intelligence; ALFWorld; explainable AI')
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)

doc.add_page_break()

# ============================================================
# Chapter 1: Introduction
# ============================================================
add_heading('1 引言', level=1)

add_heading('1.1 问题背景：数据驱动范式的三重瓶颈', level=2)

add_paragraph(
    '具身智能（Embodied Intelligence）是人工智能的核心前沿方向，要求智能体在物理世界中感知环境、理解任务并执行操作。'
    '近年来，以深度学习为代表的数据驱动范式取得了显著进展，但同时暴露出三重根本性瓶颈。'
)

add_paragraph(
    '第一，数据饥渴问题。端到端视觉-语言-动作（VLA）模型如RT-2需要超过130万条机器人操作轨迹进行训练[1]，'
    'AlphaGo需要数百万局自对弈才能掌握围棋策略[2]。这种对海量数据的依赖使得模型难以在资源受限的现实场景中快速部署。'
    '每增加一个新任务或新环境，都可能需要重新采集大量标注数据，导致边际成本居高不下。'
)

add_paragraph(
    '第二，不可解释性问题。深度强化学习策略本质上是高维数值矩阵的非线性组合[3]，'
    '其决策过程缺乏人类可理解的语义解释。当机器人在执行任务时做出错误决策，研究者无法追溯错误的根因，'
    '也无法通过局部修正来消除特定失败模式。这种"黑箱"特性严重制约了系统在安全关键场景中的应用。'
)

add_paragraph(
    '第三，物理不合规问题。基于随机探索的强化学习方法在训练初期会产生大量不符合物理常识的行为[4]，'
    '例如试图将易碎物品用力投掷、将液体倒置等。这些行为在仿真环境中可以容忍，但在真实物理环境中会导致不可逆的损坏和安全风险。'
    '现有方法通常依赖奖励工程或安全约束来缓解这一问题，但根本原因在于系统缺乏物理世界的先验常识。'
)

add_paragraph(
    '上述三重瓶颈的共同根源在于：现有方法试图从零开始、仅通过数据统计来学习世界的运行规律，'
    '而忽视了人类文明数千年积累的结构化知识体系可以作为强先验来引导决策。'
)

add_heading('1.2 《易经》作为先验知识体系', level=2)

add_paragraph(
    '《易经》是中华文明最古老的哲学经典之一，其核心思想是通过阴阳变化的组合来描述万事万物的运行规律。'
    '六十四卦体系以6个二值爻位的全排列构成64种基本卦象，每卦6爻共计384种基本情境，'
    '"弥纶天地之道"[5]，形成了一套完备的情境分类与策略关联体系。'
)

add_paragraph(
    '从计算智能的视角审视，《易经》体系具有三个关键优势。第一，八卦作为连续模糊基元可以解决符号接地问题。'
    '传统符号AI将概念处理为离散的布尔谓词，导致"脆性"问题——边界情况无法妥善处理。'
    '而八卦的物理原型（如乾=强健高稳定、坤=柔顺低力需求、坎=含水凹陷等）天然对应连续的物理属性空间，'
    '每个物体可以同时以不同程度关联多个卦象，完美契合模糊逻辑的框架。'
)

add_paragraph(
    '第二，六十四卦作为结构化先验模板提供强归纳偏置。与通用谓词逻辑需要从零构建规则不同，'
    '64卦×6爻的组合已经预定义了一套完整的情境-策略映射关系。这些映射经过数千年的经验验证和哲学提炼，'
    '蕴含了丰富的物理直觉和操作智慧，可直接作为机器人决策的先验知识库。'
)

add_paragraph(
    '第三，爻位关系（乘、承、比、应、当位得中）提供了参数精细调节的理论框架。'
    '不同爻位之间的关系类型直接对应物理操作中的约束关系，例如"承"关系对应支撑关系、'
    '"比"关系对应相邻物体的交互、"应"关系对应远程关联。这些关系可量化为修正系数，'
    '作用于执行参数的精细调整。'
)

add_heading('1.3 从静态框架到可学习系统', level=2)

add_paragraph(
    '基于上述《易经》先验知识构建的静态决策框架已展现出令人鼓舞的性能：在物理域300个物体的零样本决策中达到92.7%的合理率，'
    '在ALFWorld具身导航域达到67.2%的任务成功率。然而，静态先验存在固有局限——它无法从经验中学习和进化。'
)

add_paragraph(
    '为弥合这一差距，本文提出知几学习机制。"知几"源自《易经·系辞》"知几其神乎"[5]，'
    '意为对事物变化最微小征兆的洞察能力。在YLYW框架中，知几学习是一种对称校准机制：'
    '从每次任务执行的成败中提取微小征兆（几），将成功模式（吉之几）通过正向强化增强相关参数，'
    '将失败模式（凶之几）通过负向抑制削弱错误关联。这种机制直接作用于先验参数表，'
    '无需梯度计算，单条目精确更新，一轮即可收敛。'
)

add_paragraph(
    '引入知几学习后，ALFWorld域的成功率从67.2%逐步提升至73.9%，'
    '在标注一致子集上达到90.0%，接近270亿参数大语言模型的93.28%。这一结果表明，'
    '将古代哲学智慧与现代计算方法结合，可以实现数据高效、可解释且性能优异的具身决策系统。'
)

add_heading('1.4 本文贡献', level=2)

contributions = [
    '提出YLYW框架，首次将《易经》六十四卦体系形式化为可计算的具身决策系统，建立"卦定策略类型、爻定执行参数"的联邦式三层推理架构。',
    '提出八卦连续模糊隶属度机制，以高斯核函数将物理属性映射为八卦原型的连续隶属度值，解决了传统符号系统的接地困境。',
    '提出知几学习——一种基于"吉凶之几"对称校准的参数学习机制，仅通过单条目精确更新即可从经验中持续增长先验知识，无需梯度计算。',
    '在物理域300物体零样本决策实验中达到92.7%合理率，消融实验量化了易理规则（+33.6%）、三层架构（+12.7%）和模糊隶属度（+23.0%）的独立贡献。',
    '在ALFWorld具身导航域134局测试中，知几学习将静态先验的67.2%提升至73.9%并一轮收敛，标注一致子集达90.0%，接近270亿参数LLM的93.28%。',
    '首次证明基于古代哲学先验的可解释系统在配合精确校准后可达到接近数据驱动大模型的决策能力，为具身智能提供了数据高效、可解释、可增长的新范式。',
]

for i, c in enumerate(contributions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(f'（{i}）{c}')
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# Chapter 2: YLYW Architecture
# ============================================================
add_heading('2 YLYW架构', level=1)

add_heading('2.1 总体架构与设计哲学', level=2)

add_paragraph(
    'YLYW的核心设计哲学是"联邦式"架构——各推理层保持独立性，不溶解进统一的神经网络权重中。'
    '这一设计选择带来三个关键优势：其一，每层的输入输出具有明确的物理语义，整条推理链完全可追溯；'
    '其二，各层可独立调试、替换和升级，不影响其他层的功能；其三，先验知识以显式参数表形式存在，可直接人工审查和修正。'
)

add_paragraph(
    'YLYW的三层推理流程如下：L1层接收物体的多模态特征向量，计算其对八种卦象原型的模糊隶属度；'
    'L2层将8维隶属度向量通过语义锚定映射为6维爻值向量，每个爻位对应一个独立的物理语义维度；'
    'L3层通过余弦相似度在64个卦象理想模板中找到最佳匹配，并根据匹配卦象的预定义策略生成决策输出。'
    '三层之间信息单向流动，无反馈回路，确保推理的确定性和可复现性。'
)

add_paragraph(
    '这种联邦式架构与现有的端到端范式形成鲜明对比。端到端模型将所有知识压缩在不可分解的权重矩阵中，'
    '而YLYW将知识显式分布在三个语义清晰的层次：感知层（L1）对应"观物取象"，编码层（L2）对应"立象以尽意"，'
    '决策层（L3）对应"极数知来"。每一层都可以独立解释其功能和输出。'
)

add_heading('2.2 L1层：八卦基元——连续模糊隶属度', level=2)

add_paragraph(
    'L1层的核心创新在于将《易经》八卦重新诠释为连续模糊隶属度的物理原型。'
    '八种卦象原型及其物理语义定义如下：乾（☰）对应强健、高稳定性物体；坤（☷）对应柔顺、低力需求物体；'
    '坎（☵）对应含水、凹陷形态；离（☲）对应热源、高温物体；震（☳）对应动荡、不稳定状态；'
    '艮（☶）对应静止、封闭结构；巽（☴）对应轻柔、低密度物体；兑（☱）对应开口、容器形态。'
)

add_paragraph(
    '给定物体的特征向量f和卦象原型的特征中心p，隶属度计算采用高斯径向基核函数：'
)

add_formula('μ(f, p) = exp(−‖f − p‖² / 2σ²)', '(1)')

add_paragraph(
    '其中σ为带宽参数，控制隶属度的衰减速率。该函数确保隶属度μ∈[0,1]连续取值，'
    '一个物体可以同时以不同程度关联多个卦象原型。例如，一个装有热水的陶瓷杯可能同时具有'
    '较高的兑（开口容器）隶属度、中等的离（含热）隶属度和一定的艮（稳定结构）隶属度。'
)

add_paragraph(
    '这种连续模糊表示解决了传统符号系统的二值化困境。在经典符号AI中，"杯子"要么"是"容器要么"不是"容器，'
    '无法表达"杯子同时也是易碎的、可盛装液体的、有把手的"等多重属性的程度差异。'
    'YLYW的模糊隶属度机制允许每个物体在八维卦象空间中拥有独特的"卦象指纹"，'
    '为后续的精细化决策提供了丰富的信息基础。'
)

add_heading('2.3 L2层：六爻编码', level=2)

add_paragraph(
    'L2层将L1输出的8维隶属度向量映射为6维爻值向量y∈[0,1]⁶。'
    '这一映射并非简单的降维，而是基于语义锚定的重新编码——每个爻位被赋予独立且正交的物理语义：'
)

add_paragraph(
    '初爻（第1爻）：基础稳定性，表征物体的底部支撑强度和放置稳定度；'
    '二爻（第2爻）：力需求与可达性，表征操作该物体所需的力度以及末端执行器的可达程度；'
    '三爻（第3爻）：接触面与操作难度，表征物体表面的摩擦特性和抓取几何复杂度；'
    '四爻（第4爻）：脆弱性与变形风险，表征物体在受力时的损坏概率和形变程度；'
    '五爻（第5爻）：优先级与重要性，表征该物体在当前任务语境中的操作优先等级；'
    '上爻（第6爻）：环境约束，表征周围环境对该物体操作施加的限制条件。'
)

add_paragraph(
    '每个爻值为[0,1]连续值，以阈值0.5区分阳爻（≥0.5，对应"—"）和阴爻（<0.5，对应"--"）。'
    '六个爻值组合形成一个完整的卦象编码，既保留了连续数值的精细信息（用于参数调节），'
    '又可离散化为卦象符号（用于规则匹配）。这种双重表示是YLYW实现"模糊感知+符号推理"融合的关键设计。'
)

add_heading('2.4 L3层：六十四卦规则匹配', level=2)

add_paragraph(
    'L3层是决策的最终环节。它维护一个64×6的理想卦象模板矩阵T，每行对应一个标准卦象的理想爻值模式。'
    '给定L2输出的爻值向量y，通过余弦相似度计算与所有模板的匹配程度：'
)

add_formula('sim(y, T_k) = (y · T_k) / (‖y‖ · ‖T_k‖),  k = 1, 2, ..., 64', '(2)')

add_paragraph(
    '取相似度最高的卦象作为匹配结果。每个卦象预先关联一套完整的策略描述，包括三个层次的信息：'
    '策略类型（如"轻取慎放"、"双手协作"、"避让等待"等）、执行参数范围（如力度上限、速度范围、接近角度等）、'
    '以及注意事项（如"注意液面倾斜"、"保持竖直方向"等）。'
)

add_paragraph(
    '这种"卦定策略类型、爻定执行参数"的分工机制是YLYW的核心创新之一。'
    '卦象层面提供策略的质性选择（做什么），而爻值层面的连续数值则通过线性插值确定量化参数（怎么做）。'
    '例如，匹配到"谦卦"（表示需要谦虚谨慎对待的操作）后，具体的接近速度由二爻的连续值线性映射到[0.1, 0.5]m/s区间，'
    '抓取力度由三爻值映射到[2, 8]N区间。'
)

add_heading('2.5 爻位关系运算', level=2)

add_paragraph(
    '在完成基本的卦象匹配后，YLYW进一步通过五种传统爻位关系对执行参数进行精细修正。'
    '这五种关系源自《易经》传统义理学，本文将其形式化为可计算的修正系数：'
)

add_paragraph(
    '乘（下压上）：当阴爻位于阳爻之上时构成"乘"关系，表示上方对下方施加压力。'
    '在操作语境中对应"需要考虑物体自重对下方支撑面的影响"，修正系数作用于力度参数，使其增加10%-20%。'
)

add_paragraph(
    '承（上承下）：当阳爻位于阴爻之上时构成"承"关系，表示下方对上方的支撑。'
    '对应"下方基础稳固，操作可以更大胆"，修正系数作用于速度参数，允许提高15%-25%。'
)

add_paragraph(
    '比（相邻同性）：相邻两个同性爻（均阳或均阴）构成"比"关系，表示协同或竞争。'
    '对应"相邻属性具有一致性，参数可保持稳定"，修正系数接近1.0，起稳定作用。'
)

add_paragraph(
    '应（上下对应）：上卦与下卦对应位置的爻构成"应"关系（一阴一阳为正应）。'
    '对应"远程属性之间存在呼应关系"，修正系数用于调和表面属性与深层属性的冲突。'
)

add_paragraph(
    '当位得中：爻值是否处于其"正位"（阳爻在奇数位、阴爻在偶数位）及是否居中（二、五为中位）。'
    '当位得中表示参数处于最优区间，修正系数为1.0；失位则需向正位方向调整。'
)

add_paragraph(
    '实验表明，在运动控制域中爻位关系修正系数的均值为0.94，相比静态视觉域更为显著。'
    '这是因为运动控制涉及力度、速度、加速度等连续参数，爻位关系的精细调节能够产生更明显的效果。'
)

add_heading('2.6 可解释性分析', level=2)

add_paragraph(
    'YLYW架构的一个核心优势在于决策链的每一步都具有明确的物理语义，可以被人类理解和审查。'
    '以一个具体案例说明：当系统需要操作一个"装满水的玻璃杯"时，推理链如下：'
)

add_paragraph(
    'L1层：计算八卦隶属度，得到兑=0.82（开口容器）、坎=0.71（含水）、离=0.12（非热源）、乾=0.15（非强健）...；'
    'L2层：映射为爻值向量[0.35, 0.28, 0.62, 0.85, 0.73, 0.41]，解读为"基础稳定性中等、力需求低、操作难度中高、脆弱性高、优先级高、环境约束中等"；'
    'L3层：匹配到"坎卦"（水在上，险也），策略为"轻取慎放、保持水平、缓速移动"。'
)

add_paragraph(
    '这种全链条可解释性使得系统的性能提升过程完全可追溯。消融实验显示，'
    '从最简基线（48.0%）到完整系统（92.7%）的每一步改进都可以归因到具体的卦爻关系和规则贡献，'
    '研究者可以精确定位系统的弱点并针对性优化。'
)

doc.add_page_break()

# ============================================================
# Chapter 3: Zhi-Ji Learning
# ============================================================
add_heading('3 知几学习', level=1)

add_heading('3.1 哲学基础', level=2)

add_paragraph(
    '"知几其神乎！几者，动之微，吉之先见者也"[5]。《易经·系辞》中的这段论述揭示了一种深刻的认知能力：'
    '在事物变化的最初征兆阶段就能洞察其发展方向。"几"是变化的最微小起点，是吉凶未判时的临界状态。'
    '能够"知几"——洞察这些微小征兆——被视为最高层次的智慧。'
)

add_paragraph(
    '进一步，《易经》指出"知微知彰，知柔知刚"，强调对吉凶两种征兆的对称认知。'
    '吉之几——成功的微小征兆，如某个位置频繁出现目标物体；凶之几——失败的微小征兆，'
    '如某个物体名称反复导致错误匹配。两种征兆同属"几"的范畴，需要对称的响应机制。'
)

add_paragraph(
    '本文将这一哲学洞察形式化为知几学习机制：一种作用于先验参数的对称校准方法。'
    '其核心理念是——不需要改变整个系统的架构或权重，仅需从每次交互中提取微小征兆，'
    '对特定参数条目进行精确的正向强化或负向抑制，即可实现先验知识的持续增长。'
)

add_heading('3.2 问题形式化', level=2)

add_paragraph(
    '设系统的知识状态为K，初始状态为先验知识Ω。经过一系列交互后，知识状态更新为：'
)

add_formula('K = Ω ⊕ ΔK(trajectory, reward)', '(3)')

add_paragraph(
    '其中⊕表示知识叠加操作（对不同类型的参数分别为加法更新或集合并集），ΔK为从轨迹和奖励中提取的知识增量。'
    '知识增量的计算遵循：'
)

add_formula('ΔK = Σᵢ δ(observedᵢ) × rewardᵢ', '(4)')

add_paragraph(
    '其中δ(observedᵢ)为第i个观察到的征兆的定位函数（指向具体的参数条目），'
    'rewardᵢ∈{+1, -1}为该征兆的性质判断（+1对应吉之几，-1对应凶之几）。'
    '这一形式化的关键特征是：更新是精确定位的（只影响被观察到的特定条目），而非全局的（不影响无关参数）。'
)

add_heading('3.3 三类校准参数', level=2)

add_paragraph(
    '知几学习作用于三类先验参数，每类参数具有独立的更新规则：'
)

add_paragraph(
    '第一类：位置先验矩阵P(obj, loc)。该矩阵记录"物体obj出现在位置loc"的先验置信度。'
    '当在位置loc成功找到物体obj时执行正向更新P[obj][loc] += 1.0（吉之几：该位置确实存放该物体）；'
    '当在位置loc未找到预期物体时执行负向更新P[obj][loc] -= 0.5（凶之几：该位置先验不可靠）。'
    '非对称的更新步长（+1.0/-0.5）体现了"确认比否定更有信息量"的认知原则。'
)

add_paragraph(
    '第二类：物体匹配置信度M(target, entity)。该矩阵记录"任务目标target与环境实体entity"的匹配可能性。'
    '正向更新：成功交互后M[target][entity] += 1.0；负向更新：错误拿取后M[target][entity] -= 3.0。'
    '此处负向步长为正向的3倍，体现了"错拿的代价远高于确认的收益"——一次错误操作可能导致任务彻底失败，'
    '而一次正确确认仅是渐进式的信心积累。'
)

add_paragraph(
    '第三类：同义词映射S(word → set)。该映射记录"语言表达word对应的物体集合"。'
    '当观察到新的word-entity对应关系时，执行集合扩展S[word].add(entity)。'
    '同义词映射只增不减（单调扩展），因为语言的多义性意味着新的有效映射不应否定已有映射。'
)

add_heading('3.4 对称校准机制', level=2)

add_paragraph(
    '知几学习通过统一接口observe(trajectory, won)实现对称校准。'
    '无论任务成功（won=True）还是失败（won=False），系统都从同一条轨迹中提取征兆，'
    '区别仅在于更新方向：成功时正向强化（+α），失败时负向抑制（-β）。'
)

add_paragraph(
    '这种对称机制与神经科学中多巴胺系统的工作原理高度类似[6]。'
    '多巴胺神经元在奖励预测误差为正时（获得超预期奖励）增强相关突触连接，'
    '在预测误差为负时（未获得预期奖励）削弱相关突触连接——同一套突触权重，双向调制。'
    '知几学习的P/M/S参数表对应"突触权重"，+α/-β对应"多巴胺的双向调制"，'
    '征兆定位对应"资格迹（eligibility trace）"指向最近激活的突触。'
)

add_paragraph(
    '这种类比并非牵强附会。从信息论角度，两者都解决同一个问题：'
    '在信用分配（credit assignment）困难的环境中，如何将全局奖励信号精确归因到局部参数。'
    '深度RL通过反向传播在所有参数上分配梯度（O(N)更新），而知几学习通过征兆定位仅更新被激活的条目（O(1)更新），'
    '这解释了其极高的样本效率。'
)

add_heading('3.5 五种征兆定位策略', level=2)

add_paragraph(
    '知几学习的效果取决于征兆定位的精确性。本文定义五种征兆类型，每种对应一个卦象隐喻和特定的参数更新操作：'
)

add_paragraph(
    '错拿之几（睽卦——背离）：系统拿取了错误物体。定位操作：将M[target][wrong_entity]大幅降低（-3.0），'
    '同时将正确物体的匹配度小幅提升（+1.0）。睽卦象征背离和误解，对应物体匹配的语义偏差。'
)

add_paragraph(
    '空位之几（困卦——困顿）：目标位置为空，物体不在预期位置。定位操作：降低P[obj][expected_loc]（-0.5），'
    '后续在其他位置找到时提升P[obj][actual_loc]（+1.0）。困卦象征陷入困境，对应空间先验的失效。'
)

add_paragraph(
    '瓶颈之几（蹇卦——跛行困难）：行动受阻，如路径不通或容器已满。定位操作：记录阻塞条件并在后续规划中规避。'
    '蹇卦象征行路艰难，对应环境约束的动态变化。'
)

add_paragraph(
    '超时之几（节卦——节制）：行动超过时间预算，效率过低。定位操作：降低低效策略路径的优先级。'
    '节卦象征节制和限度，对应资源约束下的效率优化。'
)

add_paragraph(
    '模式之几（明夷卦——光明隐没）：重复出现相同失败模式但原因不明。定位操作：触发更保守的策略选择。'
    '明夷卦象征光明被遮蔽，对应系统对自身盲区的认知。'
)

add_heading('3.6 经验持久化', level=2)

add_paragraph(
    '知几学习的另一关键设计是经验的持久化存储。每轮交互结束后，系统将更新后的P/M/S参数以JSON格式保存到磁盘。'
    '下一轮交互开始时加载已保存的参数作为新的先验起点。这意味着系统的知识状态是单调递增的——每一轮交互都在之前的基础上积累，'
    '不会因为系统重启而丢失已学习的经验。'
)

add_paragraph(
    '实验表明，这种持久化机制使得知几学习在ALFWorld域中一轮即收敛。'
    '第一轮（134局）的成功率从67.2%提升至73.9%，第二轮加载第一轮经验后直接达到稳定水平，'
    '后续轮次不再有显著变化。这说明134局的经验足以校准该域中的关键参数条目，'
    '体现了知几学习"精确归因"策略的高效性。'
)

add_heading('3.7 与强化学习的本质区别', level=2)

add_paragraph(
    '知几学习与传统强化学习（RL）虽然都从奖励信号中学习，但在机制层面存在本质区别：'
)

add_paragraph(
    '更新范围：RL通过反向传播计算全参数梯度θ←θ-η∇L(θ)，每次更新影响所有网络权重；'
    '知几学习仅更新被征兆定位的特定条目P[obj][loc]+=reward×α，其余参数完全不变。'
)

add_paragraph(
    '样本复杂度：RL的策略网络通常有10⁶-10⁹个参数，收敛需要O(ε⁻ᴺ)样本（N为参数量相关）；'
    '知几学习的有效参数空间K远小于N（K为环境中实际出现的obj-loc/target-entity组合数），'
    '收敛仅需O(ε⁻ᴷ)样本，K通常为10¹-10²量级。'
)

add_formula('RL: θ ← θ − η∇L(θ),  全参数梯度更新', '(5)')
add_formula('知几: P[obj][loc] += reward × α,  单条目精确更新', '(6)')

add_paragraph(
    '核心差异可归结为一句话：1次精确归因的价值≥10⁴次统计平均。'
    '知几学习之所以能用极少样本收敛，是因为它将信用分配问题从"在所有参数上统计分摊"'
    '简化为"在被激活的条目上精确归因"。这得益于YLYW架构的显式参数表设计——每个参数条目都有明确的物理语义，'
    '使得征兆定位可以做到单步、确定、无歧义。'
)

add_paragraph(
    '从计算复杂度角度，单次知几更新的时间复杂度为O(1)（字典查找+标量加法），'
    '而单次RL反向传播为O(N)（N为网络参数量）。在ALFWorld域的134局实验中，'
    '知几学习的总参数更新次数约为500次（平均每局3-4次征兆定位），总计算量可忽略不计；'
    '而同等规模的RL训练通常需要数万次梯度更新，每次涉及数百万参数的矩阵运算。'
)

doc.add_page_break()

# ============================================================
# References (for Part 1)
# ============================================================
add_heading('参考文献', level=1)

references = [
    'Brohan A, Brown N, Carbajal J, et al. RT-2: Vision-Language-Action Models Transfer Web Knowledge to Robotic Control[J]. arXiv preprint arXiv:2307.15818, 2023.',
    'Silver D, Huang A, Maddison C J, et al. Mastering the game of Go with deep neural networks and tree search[J]. Nature, 2016, 529(7587): 484-489.',
    'Mnih V, Kavukcuoglu K, Silver D, et al. Human-level control through deep reinforcement learning[J]. Nature, 2015, 518(7540): 529-533.',
    'García J, Fernández F. A comprehensive survey on safe reinforcement learning[J]. Journal of Machine Learning Research, 2015, 16(1): 1437-1480.',
    '黄寿祺, 张善文. 周易译注[M]. 上海: 上海古籍出版社, 2007.',
    'Schultz W. Predictive reward signal of dopamine neurons[J]. Journal of Neurophysiology, 1998, 80(1): 1-27.',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'[{i}] {ref}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)


# ============================================================
# Save
# ============================================================
output_path = '/home/lijinhan/MXL/科研/ylyw/paper/_part1.docx'
doc.save(output_path)
print(f'Successfully saved to: {output_path}')
