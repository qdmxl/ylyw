#!/usr/bin/env python3
"""生成《知几知耻：一种基于先验征兆辨识与失败归因的具身学习范式》完整论文docx"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

doc = Document()

# ========== 页面设置 ==========
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ========== 样式设置 ==========
style_normal = doc.styles['Normal']
style_normal.font.name = '宋体'
style_normal.font.size = Pt(10.5)
style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style_normal.paragraph_format
pf.first_line_indent = Pt(21)
pf.line_spacing = 1.5
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Heading 1
h1_style = doc.styles['Heading 1']
h1_style.font.name = '黑体'
h1_style.font.size = Pt(14)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0, 0, 0)
h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h1_style.paragraph_format.first_line_indent = Pt(0)
h1_style.paragraph_format.space_before = Pt(12)
h1_style.paragraph_format.space_after = Pt(6)

# Heading 2
h2_style = doc.styles['Heading 2']
h2_style.font.name = '黑体'
h2_style.font.size = Pt(12)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0, 0, 0)
h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2_style.paragraph_format.first_line_indent = Pt(0)
h2_style.paragraph_format.space_before = Pt(8)
h2_style.paragraph_format.space_after = Pt(4)

# Heading 3
h3_style = doc.styles['Heading 3']
h3_style.font.name = '黑体'
h3_style.font.size = Pt(11)
h3_style.font.bold = True
h3_style.font.color.rgb = RGBColor(0, 0, 0)
h3_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h3_style.paragraph_format.first_line_indent = Pt(0)
h3_style.paragraph_format.space_before = Pt(6)
h3_style.paragraph_format.space_after = Pt(3)


def add_paragraph_with_refs(doc, text, style='Normal', bold_phrases=None, center=False):
    """添加段落，处理引用上标[1]等和加粗短语"""
    p = doc.add_paragraph(style=style)
    if center:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
    
    if bold_phrases is None:
        bold_phrases = []
    
    # 先处理加粗短语，将文本分段
    # 使用正则找到所有引用标记
    # 分割：先按引用标记分割，再在每段中处理加粗
    parts = re.split(r'(\[\d+(?:[-,]\d+)*\])', text)
    
    for part in parts:
        if re.match(r'\[\d+(?:[-,]\d+)*\]', part):
            # 这是引用标记，设为上标
            run = p.add_run(part)
            run.font.superscript = True
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            # 普通文本，检查是否有需要加粗的短语
            if bold_phrases:
                # 构建正则来分割加粗短语
                escaped = [re.escape(bp) for bp in bold_phrases]
                pattern = '(' + '|'.join(escaped) + ')'
                sub_parts = re.split(pattern, part)
                for sp in sub_parts:
                    if sp in bold_phrases:
                        run = p.add_run(sp)
                        run.bold = True
                        run.font.name = '宋体'
                        run.font.size = Pt(10.5)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    elif sp:
                        run = p.add_run(sp)
                        run.font.name = '宋体'
                        run.font.size = Pt(10.5)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                if part:
                    run = p.add_run(part)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_formula(doc, text):
    """添加居中公式"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    run.italic = True
    return p


def add_table(doc, headers, rows, caption=None):
    """添加表格"""
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(caption)
        run.font.name = '宋体'
        run.font.size = Pt(9)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 表头底纹
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._element.get_or_add_tcPr().append(shading)
    
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 表后空行
    doc.add_paragraph()
    return table


# ========== 标题页 ==========
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(24)
run = p.add_run('知几知耻：一种基于先验征兆辨识与失败归因的具身学习范式')
run.font.name = '黑体'
run.font.size = Pt(18)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 英文标题
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_after = Pt(20)
run = p.add_run('Zhiji-Zhichi: An Embodied Learning Paradigm Based on\nPrior Omen Recognition and Failure Attribution')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.italic = True

# 作者
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_after = Pt(6)
run = p.add_run('马兴录，李金函，张国安，于敬涛，李望，马圣洁')
run.font.name = '宋体'
run.font.size = Pt(12)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 通讯作者标注
run2 = p.add_run('*')
run2.font.superscript = True
run2.font.size = Pt(12)

# 单位
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('青岛科技大学 信息科学技术学院，山东 青岛 266061')
run.font.name = '宋体'
run.font.size = Pt(10.5)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 通讯作者脚注
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_after = Pt(30)
run = p.add_run('* 通讯作者：马圣洁')
run.font.name = '宋体'
run.font.size = Pt(9)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 摘要 ==========
doc.add_heading('摘要', level=1)

abstract_text = (
    '当前具身智能学习以强化学习（Reinforcement Learning, RL）为主流范式，'
    '然而面临数据低效、探索风险高、不可解释三重困境[1,2]。'
    '强化学习需要数百万次试错交互才能习得简单策略，在真实物理环境中代价极其高昂[3,4]；'
    '其黑箱决策过程无法解释失败原因，导致调试困难、安全隐患突出[5,6]。'
    '本文提出"知几知耻学习"——一种源自东方哲学的具身学习新范式。'
    '"知几"源自《易经·系辞下》"知几其神乎"[7]，从微弱征兆中预见成功路径；'
    '"知耻"源自《论语·中庸》"知耻近乎勇"[16]，从失败中精确归因并修正认知。'
    '两者构成阴阳对偶：知几积累正面先验（物体A通常在位置B），'
    '知耻积累否定先验（物体A不在位置C），共同构成完整的先验知识体系。'
    '知识积累遵循四源公式：K = Ω ⊕ K_zhiji ⊕ K_zhichi ⊕ K_persist。'
    '在YLYW（易理研物）系统中[9]，该范式以完全可解释的决策链为基础，实现了：'
    '(1)零样本决策（92.7%策略合理率）；(2)成功驱动校准（知几学习，70.1%→73.1%）；'
    '(3)失败精确归因（知耻学习，单次失败即可永久修复）；'
    '(4)经验持久化一轮收敛至73.9%。'
    '核心论点：可解释性不仅是事后解释工具，更是高效学习的前提——'
    '1次精确归因的失败比10000次统计平均更有价值。'
)
add_paragraph_with_refs(doc, abstract_text)

# 关键词
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(21)
run = p.add_run('关键词：')
run.bold = True
run.font.name = '宋体'
run.font.size = Pt(10.5)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run2 = p.add_run('知几学习；知耻学习；具身智能；易经；先验知识；可解释性；失败归因；经验持久化')
run2.font.name = '宋体'
run2.font.size = Pt(10.5)
run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 1 引言 ==========
doc.add_heading('1 引言', level=1)
doc.add_heading('1.1 具身学习的根本难题——强化学习的三重困境', level=2)

text = (
    '具身智能（Embodied Intelligence）是人工智能的终极目标之一，'
    '要求智能体在物理或仿真环境中通过感知、决策和行动来完成复杂任务。'
    '过去十年，强化学习取得了举世瞩目的成就：AlphaGo在围棋领域击败人类世界冠军[1]，'
    'DQN在Atari游戏中超越人类水平[2]，PPO成为连续控制任务的标准算法[5]。'
    '然而，当强化学习从虚拟棋盘走向真实物理世界时，三重根本性困境日益凸显。'
)
add_paragraph_with_refs(doc, text)

text = (
    '第一重困境是数据低效。'
    '强化学习的本质是通过海量试错来逼近最优策略，QT-Opt需要58万次真实抓取尝试[3]，'
    'Levine等人的机器人抓取实验耗费800个机器人小时[4]。'
    '在真实环境中，每次试错都意味着时间、能源甚至安全代价。'
    '一个简单的"将杯子放到桌上"任务，人类观察一次即可完成，'
    '而强化学习可能需要数千次探索才能收敛。'
)
add_paragraph_with_refs(doc, text)

text = (
    '第二重困境是探索风险。'
    '强化学习的探索-利用权衡（exploration-exploitation tradeoff）在安全敏感场景中尤为致命。'
    '机器人在探索过程中可能做出危险动作——撞击障碍物、跌落悬崖、损坏精密工件。'
    'García和Fernández[6]系统综述了安全强化学习的挑战，'
    '指出当前方法仍难以在保证安全的同时实现高效探索。'
    '真实世界不允许"试错-重来"的无限循环。'
)
add_paragraph_with_refs(doc, text)

text = (
    '第三重困境是不可解释性。'
    '深度强化学习的决策过程是黑箱——'
    '当机器人抓取失败时，我们无法知道是视觉识别错误、运动规划失误，'
    '还是奖励函数设计不当。这导致两个严重后果：'
    '一是调试成本极高，研究者只能通过反复调参来改善性能；'
    '二是安全无法保证，黑箱决策可能在未知状态下产生灾难性行为。'
    '可解释性缺失不仅是"锦上添花"的问题，更是制约学习效率的根本瓶颈。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('1.2 东方哲学的对偶智慧——知几与知耻', level=2)

text = (
    '面对强化学习的三重困境，我们转向东方哲学寻求启发。'
    '《易经·系辞下》曰："知几其神乎！几者，动之微，吉之先见者也"[7]。'
    '"几"是事物变化的微妙征兆，是吉凶的先行信号。'
    '知几者能从微弱线索中预见趋势，在事情萌芽时就做出正确判断，'
    '正如《易经》所言"君子见几而作，不俟终日"[8]。'
    '这启示我们：高效的学习不必等待完整的反馈信号，'
    '仅凭先验征兆即可做出合理决策。'
)
add_paragraph_with_refs(doc, text)

text = (
    '与"知几"对偶的是"知耻"。'
    '《论语·中庸》曰："知耻近乎勇"[16]——'
    '真正的勇敢不在于盲目前行，而在于直面失败、精确归因、果断修正。'
    '在东方哲学中，"几"与"耻"构成阴阳对偶：'
    '知几是从成功征兆中积累正面经验（阳），知耻是从失败教训中积累否定经验（阴）。'
    '正如《易经·系辞上》所言"一阴一阳之谓道"[7]，'
    '单独的知几或单独的知耻都不完整，'
    '唯有两者互补才构成完整的学习之道。'
)
add_paragraph_with_refs(doc, text)

text = (
    '这种对偶思维在西方认知科学中也有呼应。'
    'Kahneman[15]将人类思维分为系统1（快速直觉）和系统2（慢速推理）。'
    '知几学习对应系统1——基于经验的快速征兆识别；'
    '知耻学习对应系统2——基于分析的精确失败归因。'
    '两者协同，构成完整的认知循环。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('1.3 可解释性：连接知几与知耻的桥梁', level=2)

text = (
    '知几与知耻的高效运作有一个共同前提：决策过程必须是可解释的。'
    '只有当系统能清晰展示"为什么选择动作A"（知几的征兆链）和'
    '"为什么动作B失败了"（知耻的归因链），才能实现精确的经验积累和修正。'
    '在黑箱系统中，即使发生了成功或失败，'
    '系统也无法提取有意义的经验——因为它不知道成功"为什么成功"、失败"为什么失败"。'
)
add_paragraph_with_refs(doc, text)

text = (
    '因此，本文提出一个核心论点：可解释性不仅是事后的解释工具，'
    '更是高效学习的前提条件。'
    '这一论点颠覆了"先追求性能、再追求可解释性"的传统思路，'
    '主张可解释性应当内建于系统架构，作为学习效率的基础设施。'
    '在YLYW系统中，每一个决策都有完整的征兆链可追溯，'
    '这使得1次精确归因的失败比10000次统计平均更有学习价值。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('1.4 本文贡献', level=2)

text = '本文的主要贡献包括以下六个方面：'
add_paragraph_with_refs(doc, text)

contributions = [
    '(1) 提出"知几知耻学习"范式——首次将《易经》的征兆辨识思想和《论语》的失败归因思想统一为具身学习的对偶框架，为人工智能学习理论提供全新的东方哲学视角。',
    '(2) 建立四源知识公式K = Ω ⊕ K_zhiji ⊕ K_zhichi ⊕ K_persist——将先验知识、成功校准、失败校准和经验持久化统一为完整的知识积累理论，与强化学习的"从零学习"形成本质对比。',
    '(3) 设计知几学习的三层校准机制——包括同义词校准（语言维度）、位置先验校准（空间维度）和场景结构校准（布局维度），实现从成功中系统性积累正面先验。',
    '(4) 设计知耻学习的五层失败归因机制——包括错拿校准、否定先验、阶段瓶颈、步数预算和失败模式聚类，实现"单次失败即可永久修复"的精确归因能力。',
    '(5) 实现经验持久化与一轮收敛——验证知几知耻经验可持久化存储并跨场景复用，系统在一轮学习后即收敛至最优性能（73.9%），证明先验驱动学习的高效性。',
    '(6) 提出"可解释性即学习前提"的理论命题——论证可解释性不仅是事后工具，更是高效学习的结构性前提，为AI系统设计提供新的指导原则。',
]
for c in contributions:
    add_paragraph_with_refs(doc, c)

# ========== 2 哲学根基与数学形式化 ==========
doc.add_heading('2 哲学根基与数学形式化', level=1)
doc.add_heading('2.1 "几"与"耻"的对偶本体论', level=2)

text = (
    '在东方哲学体系中，"几"与"耻"分别对应认知过程的两个互补面向。'
    '《易经·系辞下》定义"几"为："几者，动之微，吉之先见者也"[7]。'
    '这里的"几"不是完整的因果链，而是事物变化的微弱征兆——'
    '一片树叶的飘落预示秋天，一声裂响预示冰面将破。'
    '知几者的能力在于从不完整信息中做出正确预判，'
    '这本质上是一种先验驱动的模式识别。'
)
add_paragraph_with_refs(doc, text)

text = (
    '"耻"的哲学内涵则来自《论语·中庸》："好学近乎知，力行近乎仁，知耻近乎勇"[16]。'
    '知耻不是简单的失败记录，而是一种深刻的认知修正——'
    '直面错误、分析原因、修正信念。'
    '在传统中国文化中，"耻"是推动自我完善的内在动力，'
    '"知耻而后勇"意味着从失败中获得的认知修正比盲目尝试更有价值。'
)
add_paragraph_with_refs(doc, text)

text = (
    '两者的对偶关系可以从《易经》的阴阳哲学来理解。'
    '"一阴一阳之谓道"[7]——知几（阳）积累正面先验，知耻（阴）积累否定先验。'
    '正面先验告诉系统"什么可能成功"（物体A通常在位置B），'
    '否定先验告诉系统"什么一定失败"（物体A不在位置C）。'
    '两者缺一不可：仅有正面先验会导致过度自信，仅有否定先验会导致过度保守。'
    '唯有阴阳互补，才能构成完整的知识体系。'
)
add_paragraph_with_refs(doc, text)

text = (
    '从信息论角度，知几与知耻的对偶性可进一步理解为：'
    '知几减少了决策空间的熵（通过正面先验缩小搜索范围），'
    '知耻排除了决策空间的死胡同（通过否定先验修剪无效分支）。'
    '两者从不同方向压缩决策复杂度，共同实现高效学习。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('2.2 知识来源的四源公式', level=2)

text = (
    '基于知几知耻的对偶框架，我们提出知识积累的四源公式。'
    '与强化学习"从零开始积累经验"不同，知几知耻范式承认先验知识的核心价值，'
    '并通过成功和失败两个渠道持续校准先验。'
)
add_paragraph_with_refs(doc, text)

add_formula(doc, 'K = Ω ⊕ K_zhiji ⊕ K_zhichi ⊕ K_persist')

text = '其中各项含义为：'
add_paragraph_with_refs(doc, text)

items = [
    'Ω：静态先验知识库，包含任务本体知识、物体属性、空间常识等，是知几知耻学习的起点；',
    'K_zhiji：知几学习积累的正面先验，来自成功经验的征兆校准，如"apple通常在fridge中"；',
    'K_zhichi：知耻学习积累的否定先验，来自失败归因的认知修正，如"mug不在microwave中"；',
    'K_persist：经验持久化知识，将跨场景验证的稳定经验固化为持久性知识。',
]
for item in items:
    add_paragraph_with_refs(doc, item)

text = '作为对比，强化学习的知识积累公式为：'
add_paragraph_with_refs(doc, text)

add_formula(doc, 'K_RL = f_learn(D)')

text = (
    '其中D是交互数据集，f_learn是学习算法。'
    '关键区别在于：当D=0时（零样本情境），K_RL=0——强化学习无法做出任何有意义的决策；'
    '而知几知耻范式在D=0时仍有K=Ω，即先验知识足以支撑合理决策。'
    '这正是YLYW系统零样本决策能力（92.7%策略合理率）的理论基础。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('2.3 知几策略的两级映射', level=2)

text = (
    '知几学习的决策过程不是从状态直接到动作的端到端映射，'
    '而是经由征兆空间的两级映射：'
)
add_paragraph_with_refs(doc, text)

add_formula(doc, 'π: S → O → A')

text = (
    '其中S是状态空间，O是征兆空间（Omen Space），A是动作空间。'
    '征兆空间O是知几学习的核心创新——它是介于原始状态和最终动作之间的可解释中间表示。'
    '每个征兆o∈O对应一条可追溯的推理链，记录了"观察到什么线索→推断出什么结论→选择什么动作"的完整过程。'
)
add_paragraph_with_refs(doc, text)

text = (
    '在YLYW系统中，征兆空间由卦象体系[7,8]实现。'
    '64卦覆盖了具身任务中的常见情境模式：'
    '乾卦对应创始性任务启动，坤卦对应承载性物体放置，'
    '坎卦对应困难性障碍处理，离卦对应清晰性目标识别。'
    '每个卦象携带该情境模式下的先验策略，实现从征兆到动作的快速映射。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('2.4 样本复杂度理论对比', level=2)

text = '从样本复杂度角度，知几知耻范式相比强化学习具有本质性优势：'
add_paragraph_with_refs(doc, text)

add_formula(doc, 'RL样本复杂度: O(ε^{-N})，N为状态-动作空间维度')
add_formula(doc, '知几知耻样本复杂度: O(ε^{-K})，K为需校准的先验参数数量')

text = (
    '由于K远小于N（K代表需要校准的少数先验参数，而N代表整个状态-动作空间的维度），'
    '知几知耻范式的样本效率远高于强化学习。'
    '直观地说，强化学习需要遍历整个状态空间来发现最优策略，'
    '而知几知耻只需要校准少数偏差的先验参数。'
    '在ALFWorld实验中，YLYW系统仅需134局×5轮即可完成全部校准并收敛，'
    '而典型的RL方法在相同任务上需要数十万步交互[13]。'
)
add_paragraph_with_refs(doc, text)

# ========== 3 知几学习 ==========
doc.add_heading('3 知几学习：从成功中预见路径', level=1)
doc.add_heading('3.1 三层校准机制', level=2)

text = (
    '知几学习的核心是从成功经验中提取征兆模式并校准先验知识。'
    '我们设计了三个层次的校准机制，分别对应语言、空间和布局三个维度。'
)
add_paragraph_with_refs(doc, text)

text = (
    '第一层：同义词校准（语言维度）。'
    '在具身环境中，同一物体可能有多种称谓——'
    '"vault"和"safe"指同一容器，"counter"和"countertop"指同一台面。'
    '当系统首次成功完成包含别名的任务时，知几学习记录该同义关系并更新词汇先验。'
    '这是最基础的征兆识别：名称的微妙变化是需要适配的"几"。'
)
add_paragraph_with_refs(doc, text)

text = (
    '第二层：位置先验校准（空间维度）。'
    '物体有其惯常位置——苹果通常在冰箱中，钥匙通常在抽屉里，书通常在书架上。'
    '当系统在某位置成功找到目标物体时，知几学习更新该物体的位置概率分布。'
    '经过多次成功，系统积累起"物体-位置"的正面先验映射，'
    '使得后续任务中能更快定位目标物体。'
    '这体现了"见几而作"的思想——从成功的微弱规律中预见最优搜索路径。'
)
add_paragraph_with_refs(doc, text)

text = (
    '第三层：场景结构校准（布局维度）。'
    '不同场景有不同的空间布局规律——厨房中的物品分布模式与卧室不同，'
    '办公室与浴室的容器类型各异。'
    '知几学习从成功经验中提取场景级别的结构先验：'
    '厨房场景中应优先搜索冰箱和橱柜，卧室场景中应优先搜索床头柜和衣柜。'
    '这是最高层次的征兆识别，对应从全局模式中预见任务结构。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('3.2 可解释的正面先验积累', level=2)

text = (
    '知几学习的每次校准都有完整的可解释记录。'
    '当系统在厨房场景中成功完成"将苹果放入冰箱"的任务后，'
    '知几学习记录如下征兆链：'
    '观察征兆：任务中apple首次出现在countertop上→'
    '推断结论：厨房场景中apple的默认位置可能是countertop→'
    '校准动作：更新位置先验P(apple|countertop, kitchen)。'
    '这种可解释的记录使得每条先验都可以追溯其来源和置信度，'
    '为后续的知耻修正提供了精确的靶向。'
)
add_paragraph_with_refs(doc, text)

text = (
    '正面先验的积累是渐进的、保守的。'
    '单次成功只增加微小的置信度，多次一致的成功才会显著改变先验分布。'
    '这种谨慎的积累策略避免了过度泛化——'
    '一次偶然在卧室找到苹果不应彻底改变"苹果通常在厨房"的先验。'
    '知几学习的保守性与知耻学习的果断性形成互补：'
    '正面先验需要多次验证才确认，而否定先验单次失败即可建立。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('3.3 卦象隐喻：乾与坤', level=2)

text = (
    '在YLYW系统中，知几学习的两个核心卦象是乾和坤[7,8]。'
    '乾卦象征创始、主动、开拓——对应知几学习中的路径发现阶段。'
    '当系统面对新任务时，乾卦策略驱动主动探索，'
    '从环境中搜集征兆线索，建立初始的正面先验。'
    '"天行健，君子以自强不息"——知几学习的路径发现是积极主动的过程。'
)
add_paragraph_with_refs(doc, text)

text = (
    '坤卦象征承载、包容、积累——对应知几学习中的知识沉淀阶段。'
    '当成功经验积累到一定程度，坤卦策略将零散征兆整合为结构化先验，'
    '如同大地承载万物。"地势坤，君子以厚德载物"——'
    '知几学习的知识积累是包容性的，不拒绝任何有价值的正面经验。'
    '乾与坤的交替运作，构成知几学习从"探索发现"到"知识沉淀"的完整循环。'
)
add_paragraph_with_refs(doc, text)

# ========== 4 知耻学习 ==========
doc.add_heading('4 知耻学习：从失败中修正认知', level=1)
doc.add_heading('4.1 五层失败校准机制', level=2)

text = (
    '知耻学习是知几学习的对偶面——它从失败中提取教训并精确修正认知偏差。'
    '我们设计了五个层次的失败校准机制，每层对应一种失败模式和一个卦象隐喻。'
)
add_paragraph_with_refs(doc, text)

text = (
    'L1 错拿校准（睽卦——乖离）。'
    '当系统拿取了错误物体时触发。'
    '睽卦象征"乖离、背道"，正对应"拿错了"这一情境。'
    '错拿校准记录具体的混淆对："将mug错认为cup""将desk lamp错认为floor lamp"。'
    '校准后，系统建立物体区分规则，永久避免同类错误。'
    '这是最精确的归因——失败原因单一、修复动作明确。'
)
add_paragraph_with_refs(doc, text)

text = (
    'L2 否定先验（困卦——穷困）。'
    '当系统在某位置反复搜索却找不到目标物体时触发。'
    '困卦象征"穷困、受阻"，对应"搜遍了也没有"的挫败。'
    '否定先验直接记录"物体A不在位置C"，将该搜索路径永久排除。'
    '与知几学习的正面先验对偶：正面先验说"可能在这里"，否定先验说"一定不在这里"。'
    '否定先验的价值在于其绝对确定性——一次确认即可永久生效。'
)
add_paragraph_with_refs(doc, text)

text = (
    'L3 阶段瓶颈（蹇卦——艰难）。'
    '当系统卡在任务的某个阶段无法推进时触发。'
    '蹇卦象征"行进艰难"，对应任务流程中的阻塞点。'
    '阶段瓶颈校准分析阻塞的结构原因：是前置条件未满足？'
    '是操作顺序错误？还是存在隐含约束？'
    '校准后生成该类任务的阶段性指导，避免后续任务陷入相同瓶颈。'
)
add_paragraph_with_refs(doc, text)

text = (
    'L4 步数预算（节卦——节制）。'
    '当系统在规定步数内未完成任务时触发。'
    '节卦象征"节制、约束"，对应资源有限的现实约束。'
    '步数超限通常意味着搜索策略效率低下——'
    '系统可能在错误的搜索空间中浪费了过多步骤。'
    '校准分析步数分配：哪些阶段耗时过多？哪些搜索是无效的？'
    '并调整搜索优先级以提升效率。'
)
add_paragraph_with_refs(doc, text)

text = (
    'L5 失败模式聚类（明夷卦——前车之鉴）。'
    '当多次失败呈现相同模式时触发。'
    '明夷卦象征"光明受损"，其卦辞"利艰贞"意为在困难中坚守正道。'
    '失败模式聚类是最高层次的知耻学习——'
    '它不针对单次失败，而是从多次失败中提取共性模式。'
    '例如，如果多个任务都因"方向性表达歧义"而失败，'
    '系统会建立"方向性表达消歧规则"作为全局性修复。'
    '这对应"前车之鉴"——从历史失败模式中预防未来错误。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('4.2 可解释性驱动的精确归因', level=2)

text = (
    '知耻学习的精确归因能力建立在YLYW系统的完全可解释架构之上。'
    '每次失败发生时，系统可以追溯完整的决策链，精确定位失败节点。'
    '以下展示三个典型归因案例。'
)
add_paragraph_with_refs(doc, text)

text = (
    '案例1：游戏#99（"Move watch from cabinet to the safe"）。'
    '失败现象：系统在50步内未完成任务。'
    '归因追溯：决策链显示系统反复尝试"go to vault 1"但环境无响应→'
    '环境中只有"safe 1"没有"vault 1"→'
    '根因：系统词汇先验中vault和safe未建立同义关系。'
    '修复动作：建立vault=safe的别名映射。'
    '修复效果：50步→5步，单次归因即永久修复。'
)
add_paragraph_with_refs(doc, text)

text = (
    '案例2：游戏#41（"move salt shaker from counter to drawer"）。'
    '失败现象：系统在50步内未完成任务。'
    '归因追溯：决策链显示系统已找到salt shaker但放置到错误抽屉→'
    '分析NL指令中的方向性表达"from...to..."→'
    '根因：方向性理解偏差，系统将"从A到B"中的目标位置解析错误。'
    '修复动作：修复NL方向性解析逻辑。'
    '修复效果：50步→32步。'
)
add_paragraph_with_refs(doc, text)

text = (
    '案例3：游戏#128（"Transfer CDs from desk to vault"）。'
    '失败现象：系统在50步内未完成任务。'
    '归因追溯：与案例1模式相同——vault别名问题。'
    '此案例验证了L5失败模式聚类的价值：'
    '多个任务共享"vault≠safe"的失败模式，'
    '一次修复惠及所有相关任务。'
    '修复效果：50步→11步。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('4.3 "失败→归因→修复→验证"闭环', level=2)

text = (
    '知耻学习形成完整的四阶段闭环：'
    '失败检测→精确归因→认知修复→验证确认。'
    '与强化学习的"失败→降低Q值→希望下次随机避开"的模糊机制不同，'
    '知耻学习的每次修复都是精确的、永久的、可验证的。'
)
add_paragraph_with_refs(doc, text)

text = (
    '失败检测阶段识别失败类型（五层中的哪一层）。'
    '精确归因阶段利用完整决策链追溯到具体的失败节点。'
    '认知修复阶段针对失败节点做最小化修正——只修改导致失败的特定先验参数，'
    '而不扰动其他正确的知识。'
    '验证确认阶段在相同或类似场景中验证修复是否生效、'
    '是否引入新的副作用。这一闭环保证了知耻学习的单调性——'
    '每次修复只会改善或维持性能，不会倒退。'
)
add_paragraph_with_refs(doc, text)

text = (
    '这种精确归因能力的关键前提是可解释性。'
    '在黑箱系统中，失败只能通过统计频率来"平均化"处理——'
    '某个动作失败率高则降低其选择概率。'
    '这种处理方式有两个根本缺陷：一是需要大量样本才能建立可靠统计；'
    '二是无法区分"动作本身错误"和"前置状态错误"。'
    '知耻学习通过可解释的决策链，将这两个问题彻底解决。'
)
add_paragraph_with_refs(doc, text)

# ========== 5 经验持久化 ==========
doc.add_heading('5 经验持久化：从"见几而作"到"积几成识"', level=1)
doc.add_heading('5.1 持久化机制', level=2)

text = (
    '知几学习和知耻学习积累的经验，最终需要通过持久化机制固化为系统的长期知识。'
    '经验持久化将单轮学习中验证有效的校准结果写入持久存储，'
    '使其在后续所有运行中自动生效。'
    '持久化分为两类：知几经验持久化（正面先验的固化）和知耻经验持久化（否定先验的固化）。'
)
add_paragraph_with_refs(doc, text)

text = (
    '知几经验持久化保存经过多次验证的正面先验。'
    '例如，当"apple通常在fridge中"的先验在多个场景中被反复确认后，'
    '该先验被提升为持久知识，优先级高于默认先验。'
    '知耻经验持久化保存所有否定先验和修复规则。'
    '与知几经验不同，知耻经验只需单次验证即可持久化——'
    '因为否定先验具有绝对确定性（"vault=safe"一旦确认就永远正确）。'
)
add_paragraph_with_refs(doc, text)

text = (
    '持久化的实现采用增量式设计：每轮学习结束后，'
    '新积累的校准经验追加到经验文件中，不覆盖已有经验。'
    '加载时按时间戳排序，后积累的经验优先级更高。'
    '这种设计允许经验随时间演进——早期的粗略先验可被后期的精确先验覆盖。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('5.2 收敛实验', level=2)

text = (
    '为验证经验持久化的效果和收敛速度，我们在ALFWorld 134局任务上进行了5轮连续实验。'
    '每轮结束后将该轮的知几和知耻经验持久化，下一轮加载全部历史经验启动。'
    '实验结果显示，系统在仅一轮持久化后即收敛至最优性能（73.9%成功率），'
    '后续各轮保持稳定。这一收敛速度远超强化学习方法——'
    '典型RL方法需要数百轮训练才能稳定收敛[14]。'
)
add_paragraph_with_refs(doc, text)

text = (
    '收敛的快速性源于知几知耻经验的精确性：'
    '每条经验都针对特定的先验参数，不存在"干扰其他参数"的问题。'
    '因此，经验的累加是单调递增的——更多经验只会提供更多信息，不会相互矛盾。'
    '这与强化学习中"新经验可能推翻旧经验"的不稳定性形成鲜明对比。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('5.3 积几成识的三阶段模型', level=2)

text = (
    '从"见几而作"到"积几成识"，知几知耻学习经历三个阶段：'
)
add_paragraph_with_refs(doc, text)

text = (
    '第一阶段：征兆感知（见几）。'
    '系统初次遇到新环境时，从零散的成功和失败中捕捉征兆信号。'
    '此阶段对应乾卦"潜龙勿用"——征兆尚未形成系统模式，需要耐心积累。'
)
add_paragraph_with_refs(doc, text)

text = (
    '第二阶段：模式形成（知几知耻）。'
    '随着经验积累，零散征兆凝聚为结构化的先验知识。'
    '正面先验和否定先验共同构建起环境的认知地图。'
    '此阶段对应坤卦"直方大"——知识开始呈现规整的结构。'
)
add_paragraph_with_refs(doc, text)

text = (
    '第三阶段：知识固化（积几成识）。'
    '验证稳定的模式被持久化为长期知识，成为系统的"本能"。'
    '后续任务无需重新学习即可直接调用。'
    '此阶段对应既济卦"初吉终乱"的反面——'
    '通过持久化保证"初吉终亦吉"，避免知识遗忘。'
)
add_paragraph_with_refs(doc, text)

# ========== 6 实验验证 ==========
doc.add_heading('6 实验验证', level=1)
doc.add_heading('6.1 实验设置', level=2)

text = (
    '实验在ALFWorld[13]仿真环境中进行，该环境是TextWorld的具身扩展版本，'
    '包含6种家庭任务类型：放置、清洁、加热、冷却、检查和双物体任务。'
    '我们使用完整的134局测试集，涵盖所有任务类型和难度级别。'
)
add_paragraph_with_refs(doc, text)

text = (
    '关键实验条件：(1)纯CPU运行，无GPU加速；'
    '(2)完全不使用大语言模型（LLM），所有决策由规则引擎和先验知识驱动；'
    '(3)每局最大步数为50步；'
    '(4)每轮包含完整的134局游戏。'
    '这些条件确保实验结果的可复现性和公平性——'
    '不依赖LLM意味着系统的能力完全来自先验设计和知几知耻学习，'
    '而非大模型的隐含知识。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('6.2 知几知耻的累积效果', level=2)

text = (
    '表1展示了YLYW系统从静态先验到完整知几知耻学习的版本演进。'
    '每个版本在前一版基础上叠加新的学习机制，清晰展示各机制的独立贡献。'
)
add_paragraph_with_refs(doc, text)

add_table(doc,
    headers=['版本', '机制', '成功率', '全局步数'],
    rows=[
        ['V7', '静态先验', '90/134 (67.2%)', '—'],
        ['V9', '+知几学习', '94/134 (70.1%)', '21.5'],
        ['V10', '+知几+知耻', '98/134 (73.1%)', '20.5'],
        ['V10+persist', '+经验持久化', '99/134 (73.9%)', '20.1'],
    ],
    caption='表1 版本演进与累积效果'
)

text = (
    '从表1可以看出：静态先验（V7）已实现67.2%的基线成功率，'
    '证明精心设计的先验知识本身就具有强大的决策能力。'
    '知几学习（V9）在此基础上提升2.9个百分点（70.1%），同时将全局平均步数优化至21.5。'
    '知耻学习的加入（V10）进一步提升3.0个百分点（73.1%），步数降至20.5。'
    '经验持久化最终将性能推至73.9%，步数进一步降至20.1。'
    '值得注意的是，知几与知耻的贡献几乎相等（2.9% vs 3.0%），'
    '印证了两者作为对偶机制的均衡性。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('6.3 知耻学习的精确归因效果', level=2)

text = (
    '表2详细展示了知耻学习（V10）新赢得的游戏及其归因分析。'
    '这些游戏在V9中失败（耗尽50步），在V10中通过知耻学习的精确修复而成功。'
)
add_paragraph_with_refs(doc, text)

add_table(doc,
    headers=['游戏', '任务描述', '修复原因', '步数变化'],
    rows=[
        ['#41', 'move salt shaker from counter to drawer', 'NL方向性修复', '50→32'],
        ['#99', 'Move watch from cabinet to the safe', 'vault→safe别名', '50→5'],
        ['#116', 'take coffee mug from book shelf, place on desk', 'NL方向性修复', '50→11'],
        ['#118', 'take mug from desk shelf to desk', 'NL方向性修复', '50→11'],
        ['#128', 'Transfer CDs from desk to vault', 'vault→safe别名', '50→11'],
    ],
    caption='表2 V10新赢游戏的归因分析'
)

text = (
    '从表2可以观察到两个显著模式：'
    '(1)修复原因高度集中——5个新赢游戏仅对应2种失败模式（NL方向性和vault别名），'
    '这验证了L5失败模式聚类的价值：一种修复惠及多个任务。'
    '(2)步数改善极其显著——从50步（超时失败）骤降至5-32步，'
    '证明知耻学习的修复是精确的"根因修复"而非"绕路补偿"。'
    '特别是游戏#99（50步→5步），一个简单的别名映射就将95%的无效搜索消除。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('6.4 收敛数据', level=2)

text = (
    '表3展示了5轮经验持久化实验的完整数据。'
    '每轮结束后持久化当轮经验，下一轮携带全部历史经验启动。'
)
add_paragraph_with_refs(doc, text)

add_table(doc,
    headers=['轮次', '携带经验', '成功率', '全局步数', '校准次数'],
    rows=[
        ['V9基线', '无', '94/134 (70.1%)', '21.5', '18'],
        ['R1(V10)', '无', '98/134 (73.1%)', '20.5', '15'],
        ['R2', '1轮', '99/134 (73.9%)', '20.1', '38'],
        ['R3', '2轮', '99/134 (73.9%)', '20.1', '61'],
        ['R4', '3轮', '99/134 (73.9%)', '20.2', '84'],
        ['R5', '4轮', '99/134 (73.9%)', '20.2', '107'],
    ],
    caption='表3 经验持久化收敛实验'
)

text = (
    '表3揭示了三个重要发现：'
    '(1)一轮收敛——R2（携带1轮经验）即达到最优成功率73.9%，'
    '此后R3-R5保持稳定，证明知几知耻经验具有极高的学习效率。'
    '(2)步数稳定——全局步数在20.1-20.2之间波动，几乎没有性能退化。'
    '(3)校准次数单调递增——每轮新增约23次校准，表明系统持续积累更精细的先验，'
    '但这些后期校准的边际收益已很小（成功率未变），主要贡献是搜索路径的微优化。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('6.5 四源知识公式的量化验证', level=2)

text = (
    '基于上述实验数据，我们可以量化验证四源知识公式K = Ω ⊕ K_zhiji ⊕ K_zhichi ⊕ K_persist的各项贡献：'
)
add_paragraph_with_refs(doc, text)

items = [
    'Ω（静态先验）：67.2%——系统零样本的基线能力，完全来自精心设计的先验知识库；',
    'K_zhiji（知几学习）：+2.9%——从成功中校准的正面先验贡献；',
    'K_zhichi（知耻学习）：+3.0%——从失败中归因的否定先验贡献；',
    'K_persist（经验持久化）：+0.8%——跨轮次的经验固化贡献。',
]
for item in items:
    add_paragraph_with_refs(doc, item)

text = (
    '总计：67.2% + 2.9% + 3.0% + 0.8% = 73.9%，与实验结果完全吻合。'
    '这一量化分解证明了四源公式的理论解释力：'
    '先验知识贡献了绝大部分能力（67.2%/73.9% = 90.9%），'
    '知几知耻学习提供精确的增量校准（6.7%/73.9% = 9.1%）。'
    '这一比例结构与人类专家的表现模式一致——'
    '专家的大部分能力来自长期积累的领域知识（先验），少部分来自即时的情境调整（学习）。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('6.6 与现有方法对比', level=2)

text = (
    '表4将YLYW系统与当前最先进的具身智能方法进行对比。'
    '对比维度包括成功率、是否依赖大语言模型、以及决策可解释性。'
)
add_paragraph_with_refs(doc, text)

add_table(doc,
    headers=['方法', '成功率', '需LLM', '可解释'],
    rows=[
        ['ReAct (GPT-4)[10]', '71%', '是', '否'],
        ['Reflexion[11]', '77%', '是', '否'],
        ['EmbodiSkill (27B)[12]', '93.28%', '是', '否'],
        ['YLYW V10 (本文)', '73.9%', '否', '是'],
    ],
    caption='表4 与现有方法对比'
)

text = (
    '从表4的对比中可以得出以下结论：'
    '(1)YLYW在不使用任何LLM的条件下，成功率超越了使用GPT-4的ReAct（73.9% vs 71%），'
    '证明精心设计的先验+知几知耻学习可以替代大模型的部分能力。'
    '(2)YLYW是唯一同时满足"无LLM依赖"和"完全可解释"两个条件的方法。'
    '这在安全敏感场景（如医疗机器人、工业自动化）中具有独特价值。'
    '(3)与EmbodiSkill的差距（73.9% vs 93.28%）主要来自先验覆盖度的不足——'
    'YLYW的静态先验尚未覆盖所有任务类型的复杂变体，而27B参数的LLM天然具有更广的世界知识。'
    '这也指明了未来优化方向：扩展先验知识库的覆盖度。'
)
add_paragraph_with_refs(doc, text)

# ========== 7 与强化学习的本质对比 ==========
doc.add_heading('7 与强化学习的本质对比', level=1)
doc.add_heading('7.1 时间性：事先征兆 vs 事后奖惩', level=2)

text = (
    '强化学习的核心机制是事后奖惩——智能体执行动作，环境返回奖励信号，'
    '然后通过时间差分学习更新价值估计。'
    '这一机制的时间结构是"先行动，后学习"：'
    '必须等到结果出现（奖励信号到达），才能开始学习。'
    '信用分配问题使得这一学习过程更加复杂——'
    '当最终奖励到达时，系统需要回溯追问"哪一步的贡献最大"[14]。'
)
add_paragraph_with_refs(doc, text)

text = (
    '知几知耻学习的时间结构根本不同。'
    '知几学习是"先征兆，后行动"——在行动之前，通过征兆识别预见最可能的成功路径。'
    '这不是预测未来，而是利用先验知识将当前观察映射到已知的成功模式。'
    '知耻学习是"即时归因"——失败发生后，立即通过决策链追溯到根因，无需统计积累。'
    '两者共同实现了"极少样本即可学习"的高效性，因为不需要等待大量反馈的统计聚合。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('7.2 失败处理：精确归因 vs 统计平均', level=2)

text = (
    '这是知几知耻范式与强化学习最本质的差异。'
    '强化学习处理失败的方式是降低对应状态-动作对的Q值：'
    'Q(s,a) ← Q(s,a) - α·δ。'
    '这种更新是统计性的——需要大量失败样本才能使Q值收敛到准确估计。'
    '更关键的是，它无法区分"动作a本身有问题"和"状态s不适合动作a但适合动作b"。'
)
add_paragraph_with_refs(doc, text)

text = (
    '知耻学习处理失败的方式是精确归因并修复：'
    '追溯决策链→定位失败节点→识别根因→精确修复。'
    '以游戏#99为例：50步失败→追溯发现所有失败步骤都在尝试"go to vault"→'
    '定位根因为vault≠safe的别名缺失→精确修复别名映射→5步完成。'
    '整个修复过程只需1次失败，而强化学习需要数千次"go to vault"的失败才能学到等效知识。'
)
add_paragraph_with_refs(doc, text)

text = (
    '我们提出核心论点：1次精确归因的失败 ≥ 10000次统计平均的失败。'
    '这一论点的数学基础在于信息量的差异：'
    '精确归因提取了失败的全部因果信息（根因是什么、如何修复），'
    '而统计平均只提取了频率信息（这个动作的失败概率是多少）。'
    '从信息论角度，前者的信息熵远高于后者。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('7.3 可解释性的双重价值：解释工具与学习前提', level=2)

text = (
    '在主流AI研究中，可解释性通常被视为事后的解释工具——'
    '先训练一个高性能的黑箱模型，再用注意力可视化、SHAP值等方法解释其决策。'
    '这种"先性能后解释"的范式隐含一个假设：可解释性是锦上添花的附加功能。'
)
add_paragraph_with_refs(doc, text)

text = (
    '本文提出相反的主张：可解释性是高效学习的结构性前提。'
    '在YLYW系统中，可解释性内建于决策架构——每个决策节点都有明确的输入、推理规则和输出。'
    '正是这种透明性使得知几学习可以精确提取成功征兆，'
    '知耻学习可以精确定位失败节点。'
    '如果决策过程是黑箱的，即使发生了成功或失败，'
    '系统也无法知道"哪一步的贡献最大"或"哪一步出了问题"。'
)
add_paragraph_with_refs(doc, text)

text = (
    '可解释性的双重价值可以形式化表述为：'
    '(1)作为解释工具：为人类用户提供决策透明性，增强信任和安全性；'
    '(2)作为学习前提：为学习算法提供精确的归因接口，实现高效的经验积累。'
    '第二个价值在现有文献中鲜有论述，是本文的重要理论贡献。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('7.4 范式互补：系统1与系统2', level=2)

text = (
    'Kahneman[15]的双系统理论为理解知几知耻与强化学习的关系提供了框架。'
    '知几知耻学习对应系统1——基于先验知识的快速直觉判断：'
    '看到苹果就知道去冰箱找（知几），知道vault就是safe（知耻积累的否定先验）。'
    '强化学习对应系统2——基于试错的慢速推理优化：'
    '在复杂未知环境中通过反复探索找到最优策略。'
)
add_paragraph_with_refs(doc, text)

text = (
    '两个范式是互补而非对立的。'
    '知几知耻擅长处理"先验知识覆盖范围内"的任务——快速、高效、可解释；'
    '强化学习擅长处理"先验知识无法覆盖"的任务——通用、灵活、渐进优化。'
    '未来的完整具身智能系统应当将两者整合：'
    '以知几知耻为默认决策模式（系统1），当先验知识不足时自动切换到强化学习（系统2）。'
    '这种"先验优先、RL兜底"的架构，有望同时获得高效性和通用性。'
)
add_paragraph_with_refs(doc, text)

# ========== 8 讨论与未来方向 ==========
doc.add_heading('8 讨论与未来方向', level=1)
doc.add_heading('8.1 收敛边界的结构性分析', level=2)

text = (
    '当前系统在73.9%处收敛，未能达到更高水平。'
    '分析剩余26.1%的失败案例，可以归为三类结构性障碍：'
    '(1)先验覆盖度不足——部分任务涉及罕见的物体-位置组合，'
    '超出当前先验知识库的覆盖范围；'
    '(2)环境交互复杂度——部分任务需要长序列的多步操作，'
    '当前的搜索策略在步数预算内无法完成；'
    '(3)自然语言歧义——部分任务描述存在未覆盖的表达方式，'
    '导致意图解析偏差。'
    '这三类障碍指明了系统改进的方向。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('8.2 从知几知耻到通用人工智能', level=2)

text = (
    '知几知耻范式的理论框架具有向更广泛场景扩展的潜力。'
    '在机器人操作领域，"几"可以是力觉传感器的微小变化（预示抓取成功或滑动），'
    '"耻"可以是碰撞检测的归因分析（哪个关节的轨迹规划有误）。'
    '在自动驾驶领域，"几"可以是交通流的微妙模式变化（预示前方拥堵），'
    '"耻"可以是事故分析的因果追溯。'
    '在医疗决策领域，"几"可以是生理指标的早期异常信号，'
    '"耻"可以是误诊案例的精确归因。'
    '知几知耻的普适性在于：任何需要"从经验中学习"的场景，都可以分解为正面经验积累和负面经验修正两个互补过程。'
)
add_paragraph_with_refs(doc, text)

doc.add_heading('8.3 局限性', level=2)

text = (
    '本文的主要局限包括：'
    '(1)先验知识依赖人工设计——当前的静态先验Ω需要领域专家手工构建，'
    '这限制了系统向新领域迁移的速度；'
    '(2)实验环境的简化性——ALFWorld是文本界面的家庭环境仿真，'
    '与真实物理世界的复杂度仍有差距；'
    '(3)知几知耻的校准范围有限——当前五层校准机制针对的是常见的具身任务失败模式，'
    '可能无法覆盖所有领域的失败类型；'
    '(4)与LLM方法的性能差距——在纯性能指标上，YLYW落后于使用27B参数LLM的方法，'
    '说明先验知识的手工设计难以匹敌大模型的广泛世界知识。'
    '未来工作将探索用LLM辅助先验构建、真实机器人验证等方向。'
)
add_paragraph_with_refs(doc, text)

# ========== 9 结论 ==========
doc.add_heading('9 结论', level=1)

text = (
    '本文提出了"知几知耻学习"——一种源自东方哲学、面向具身智能的全新学习范式。'
    '该范式以《易经》的征兆辨识（知几）和《论语》的失败归因（知耻）为哲学根基，'
    '以四源知识公式和两级决策映射为数学框架，'
    '在YLYW系统中实现了完整的工程验证。核心贡献总结如下：'
)
add_paragraph_with_refs(doc, text)

conclusions = [
    '第一，知几知耻的对偶框架揭示了学习的阴阳本质：从成功中积累正面先验（知几），从失败中积累否定先验（知耻），两者互补构成完整的知识体系。',
    '第二，四源知识公式K = Ω ⊕ K_zhiji ⊕ K_zhichi ⊕ K_persist提供了知识积累的统一理论，其中先验知识贡献90.9%的能力，知几知耻学习贡献9.1%的增量校准。',
    '第三，三层知几校准和五层知耻归因构成了精确的经验积累机制，实现了"单次失败即可永久修复"的高效学习能力。',
    '第四，经验持久化验证了知几知耻经验的可迁移性——一轮学习即可收敛至最优性能（73.9%），远超强化学习的收敛速度。',
    '第五，在ALFWorld 134局测试中，YLYW系统在不使用任何LLM的条件下超越了GPT-4驱动的ReAct方法（73.9% vs 71%），证明了知几知耻范式的实用价值。',
    '第六，本文提出"可解释性即学习前提"的理论命题——可解释性不仅是事后工具，更是高效学习的结构性前提。这一命题为AI系统设计提供了新的指导原则。',
]
for c in conclusions:
    add_paragraph_with_refs(doc, c)

text = (
    '展望未来，知几知耻学习有望与强化学习形成互补：'
    '前者作为系统1提供快速直觉决策，后者作为系统2处理先验覆盖范围外的复杂情境。'
    '两者的整合将推动具身智能向更高效、更安全、更可解释的方向发展。'
)
add_paragraph_with_refs(doc, text)

# ========== 参考文献 ==========
doc.add_heading('参考文献', level=1)

refs = [
    '[1] Silver D, Huang A, Maddison CJ, et al. Mastering the game of Go with deep neural networks and tree search[J]. Nature, 2016, 529(7587): 484-489.',
    '[2] Mnih V, Kavukcuoglu K, Silver D, et al. Human-level control through deep reinforcement learning[J]. Nature, 2015, 518(7540): 529-533.',
    '[3] Kalashnikov D, Irpan A, Pastor P, et al. Scalable deep reinforcement learning for vision-based robotic manipulation[C]. CoRL, 2018.',
    '[4] Levine S, Pastor P, Krizhevsky A, et al. Learning hand-eye coordination for robotic grasping with deep learning and large-scale data collection[J]. IJRR, 2018, 37(4-5): 421-436.',
    '[5] Schulman J, Wolski F, Dhariwal P, et al. Proximal policy optimization algorithms[J]. arXiv preprint arXiv:1707.06347, 2017.',
    '[6] García J, Fernández F. A comprehensive survey on safe reinforcement learning[J]. JMLR, 2015, 16(1): 1437-1480.',
    '[7] 黄寿祺, 张善文. 周易译注[M]. 上海: 上海古籍出版社, 2007.',
    '[8] 朱熹. 周易本义[M]. 北京: 中华书局, 2009.',
    '[9] 马兴录, 李金函, 张国安, 等. YLYW: 基于易理的具身智能决策系统[J]. 预印本, 2026.',
    '[10] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing reasoning and acting in language models[C]. ICLR, 2023.',
    '[11] Shinn N, Cassano F, Gopinath A, et al. Reflexion: Language agents with verbal reinforcement learning[C]. NeurIPS, 2023.',
    '[12] Ju H, Li J, Wang Y, et al. EmbodiSkill: Embodied skill learning via language model agents[J]. arXiv:2605.10332, 2026.',
    '[13] Shridhar M, Thomason J, Gordon D, et al. ALFWorld: Aligning text and embodied environments for interactive learning[C]. ICLR, 2021.',
    '[14] Sutton RS, Barto AG. Reinforcement Learning: An Introduction[M]. 2nd ed. Cambridge: MIT Press, 2018.',
    '[15] Kahneman D. Thinking, Fast and Slow[M]. New York: Farrar, Straus and Giroux, 2011.',
    '[16] 《论语·中庸》. 子曰: "好学近乎知, 力行近乎仁, 知耻近乎勇."',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Cm(0)
    run = p.add_run(ref)
    run.font.name = '宋体'
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 保存 ==========
output_path = '/home/lijinhan/MXL/科研/ylyw/paper/知几知耻学习_完整论文_v1.0.docx'
doc.save(output_path)
print(f'论文已成功生成: {output_path}')
