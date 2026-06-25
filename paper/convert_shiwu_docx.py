#!/usr/bin/env python3
"""实物验证_学灵通机器人.md → DOCX 转换器"""
import sys, os, re

PAPER_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(PAPER_DIR, '实物验证_学灵通机器人.md')

with open(MD_PATH, 'r', encoding='utf-8') as f:
    content = f.read()

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_para(text, size=12, bold=False, italic=False, indent=True, align=None, font_name='宋体'):
    """Add a formatted paragraph"""
    if not text.strip():
        return None
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    if align is not None:
        p.alignment = align
    
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    run.italic = italic
    return p


def add_heading_styled(text, level):
    h = doc.add_heading(text, level=min(level, 3))
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(13)
    h.paragraph_format.first_line_indent = Cm(0)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_table_from_md(lines, indent_first=True):
    """Parse and add a markdown table"""
    if len(lines) < 2:
        return
    parts = lines[0].strip().split('|')
    if len(parts) < 3:
        return
    headers = [p.strip() for p in parts[1:-1]]
    if not headers:
        return
    
    rows = []
    for line in lines[2:]:
        parts = line.strip().split('|')
        if len(parts) >= 3:
            inner = [p.strip() for p in parts[1:-1]]
            if inner and any(c for c in inner):
                rows.append(inner)
    if not rows:
        return
    
    ncols = len(headers)
    for row in rows:
        while len(row) < ncols:
            row.append('')
    
    table = doc.add_table(rows=1+len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, hd in enumerate(headers):
        if i >= ncols: break
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(hd)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)
    
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            if c >= ncols: break
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(3)
    spacer.paragraph_format.space_after = Pt(3)
    spacer.paragraph_format.first_line_indent = Cm(0)


def process_inline(text):
    """Handle **bold** and *italic* inline marks, inline code"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


# Parse
lines = content.split('\n')
i = 0
in_code = False
table_lines = []

while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    
    # Skip empty
    if not stripped and not in_code:
        if table_lines:
            add_table_from_md(table_lines)
            table_lines = []
        i += 1
        continue
    
    if not stripped and in_code:
        add_para('', indent=False)
        i += 1
        continue
    
    # Title (# )
    if line.startswith('# ') and not line.startswith('## '):
        text = process_inline(line[2:].strip())
        add_para(text, size=18, bold=True, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体')
        i += 1
        continue
    
    # Metadata
    if stripped.startswith('**版本') or stripped.startswith('**日期'):
        text = process_inline(stripped)
        add_para(text, size=10.5, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # H2
    if line.startswith('## '):
        text = process_inline(line[3:].strip())
        add_heading_styled(text, level=1)
        i += 1
        continue
    
    # H3
    if line.startswith('### '):
        text = process_inline(line[4:].strip())
        add_heading_styled(text, level=2)
        i += 1
        continue
    
    # H4
    if line.startswith('#### '):
        text = process_inline(line[5:].strip())
        add_heading_styled(text, level=3)
        i += 1
        continue
    
    # Horizontal rule
    if stripped == '---':
        add_para('─' * 60, size=8, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Code blocks
    if stripped.startswith('```'):
        in_code = not in_code
        i += 1
        continue
    
    if in_code:
        text = line  # preserve indentation in code
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text if text else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        i += 1
        continue
    
    # Tables
    if stripped.startswith('|'):
        table_lines.append(line)
        i += 1
        continue
    
    # Flowchart / ASCII diagrams - treat as code
    if any(c in stripped for c in ['┌', '└', '├', '─', '│', '▶', '▼', '▲', '◀']):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(7)
        i += 1
        continue
    
    # Bullet list items
    if stripped.startswith('- ') or stripped.startswith('  - '):
        text = process_inline(stripped.lstrip('- '))
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run('• ' + text)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        i += 1
        continue
    
    # Numbered list
    if re.match(r'^\d+\.\s', stripped):
        text = process_inline(stripped)
        add_para(text, indent=True)
        i += 1
        continue
    
    # Math formulas (inline)
    if stripped.startswith('$$'):
        text = process_inline(stripped.strip('$').strip())
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
        i += 1
        continue
    
    # Regular text
    text = process_inline(stripped)
    if text:
        add_para(text)
    i += 1

# Remaining table
if table_lines:
    add_table_from_md(table_lines)

# Save
out_path = os.path.join(PAPER_DIR, '实物验证_学灵通机器人.docx')
doc.save(out_path)
print(f'✅ Saved: {out_path}')
print(f'   Size: {os.path.getsize(out_path):,} bytes')
print(f'   Paragraphs: {len(doc.paragraphs)}')
print(f'   Tables: {len(doc.tables)}')
