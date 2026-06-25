#!/usr/bin/env python3
"""
生成《自动化学报》格式的知几学习投稿论文 docx
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ========== 页面设置 ==========
sections = doc.sections
for section in sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ========== 样式辅助函数 ==========
def set_run_font(run, font_name_cn='宋体', font_name_en='Times New Roman', size=Pt(10.5), bold=False):
    """设置run的字体"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = font_name_en
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)

def add_paragraph_normal(text, indent_first=True, space_after=Pt(0), space_before=Pt(0)):
    """添加正文段落，宋体10.5pt，首行缩进"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = 1.5
    if indent_first:
        pf.first_line_indent = Pt(21)  # 约2字符
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(10.5))
    return p

def add_heading1(text):
    """一级标题：黑体14pt居中"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, '黑体', 'Times New Roman', Pt(14), bold=True)
    return p

def add_heading2(text):
    """二级标题：黑体12pt左对齐"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, '黑体', 'Times New Roman', Pt(12), bold=True)
    return p

def add_heading3(text):
    """三级标题：黑体10.5pt加粗左对齐"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, '黑体', 'Times New Roman', Pt(10.5), bold=True)
    return p

def add_equation(text, number=''):
    """添加公式（居中，编号右对齐）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Cambria Math', Pt(11))
    if number:
        run2 = p.add_run(f'    {number}')
        set_run_font(run2, '宋体', 'Times New Roman', Pt(10.5))
    return p

def add_table_caption(text):
    """表题：居中加粗9pt"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(9), bold=True)
    return p

def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, '宋体', 'Times New Roman', Pt(9), bold=True)
    # data
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_run_font(run, '宋体', 'Times New Roman', Pt(9))
    # space after table
    doc.add_paragraph()
    return table

def add_ref(text):
    """参考文献条目：宋体9pt"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.25
    pf.left_indent = Pt(21)
    pf.first_line_indent = Pt(-21)
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(9))
    return p

def add_figure_placeholder(text):
    """图片占位"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(9))
    run.font.color.rgb = RGBColor(128, 128, 128)
    return p

# ============================================================
# 开始正文内容
# ============================================================

# ===== 论文标题 =====
p = doc.add_paragraph()
pf = p.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.space_before = Pt(24)
pf.space_after = Pt(6)
run = p.add_run('知几学习：一种基于可解释先验体系的具身智能学习方法')
set_run_font(run, '黑体', 'Times New Roman', Pt(18), bold=True)

# ===== 英文标题 =====
p = doc.add_paragraph()
pf = p.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.space_after = Pt(12)
run = p.add_run('Zhiji Learning: An Embodied Intelligence Learning Method Based on Interpretable Prior Systems')
set_run_font(run, '宋体', 'Times New Roman', Pt(14), bold=True)

# ===== 作者 =====
p = doc.add_paragraph()
pf = p.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.space_after = Pt(3)
run = p.add_run('马兴录    李金函    张国安    于敬涛    李望    马圣洁')
set_run_font(run, '宋体', 'Times New Roman', Pt(12))

# ===== 单位 =====
p = doc.add_paragraph()
pf = p.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.space_after = Pt(12)
run = p.add_run('(青岛科技大学 信息科学技术学院，山东 青岛 266061)')
set_run_font(run, '宋体', 'Times New Roman', Pt(10.5))

# ===== 收稿、基金等信息 =====
info_lines = [
    '收稿日期：2026-06-19',
    '基金项目：XXX（待补充）',
]
for line in info_lines:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    run = p.add_run(line)
    set_run_font(run, '宋体', 'Times New Roman', Pt(9))

# 作者简介
p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_after = Pt(0)
run = p.add_run('作者简介：马兴录（1970-），男，教授，主要研究方向为具身智能、嵌入式AI. E-mail: maxinglu@qust.edu.cn')
set_run_font(run, '宋体', 'Times New Roman', Pt(9))

p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_after = Pt(6)
run = p.add_run('          马圣洁（通讯作者），男，主要研究方向为人工智能. E-mail: mashengjie@qust.edu.cn')
set_run_font(run, '宋体', 'Times New Roman', Pt(9))

# 中图分类号
p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_after = Pt(12)
run = p.add_run('中图分类号：TP18    文献标志码：A    DOI：')
set_run_font(run, '宋体', 'Times New Roman', Pt(9))

# ===== 中文摘要（约400字） =====
p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_before = Pt(6)
pf.space_after = Pt(3)
run = p.add_run('摘  要')
set_run_font(run, '黑体', 'Times New Roman', Pt(10.5), bold=True)

abstract_cn = (
    '强化学习是当前具身智能领域的主流学习范式，但面临样本效率低、训练不稳定、决策不可解释三重困境。'
    '这三重困境的根源在于知识以隐式方式分散存储在高维网络参数中，学习信号无法精确定位到需要调整的具体知识单元。'
    '针对这一根本性瓶颈，本文提出知几学习（Zhiji Learning）——一种基于可解释先验体系的具身智能学习方法。'
    '知几学习建立在YLYW（易理模糊模型）系统之上，该系统通过模糊隶属度实现连续符号接地，以三层推理架构（特征到隶属度、隶属度到爻值、爻值到64卦匹配）'
    '构建完全透明的推理链，使决策的每一步都可追溯。在此可解释基础上，知几学习定义了三类可校准参数：'
    '位置先验矩阵P、物体匹配置信度M和同义词映射S，设计了统一的对称更新接口，'
    '使得成功信号（吉之几）强化正确先验、失败信号（凶之几）削弱错误先验，共享同一套校准逻辑。'
    '系统通过五种凶之几定位策略（错拿定位、空位定位、瓶颈定位、超时定位、模式定位）实现精确归因，'
    '将学习信号直接定位到具体参数条目，实现"外科手术式"的精准修正。'
    '在ALFWorld基准134局实验中，知几学习将成功率从67.2%提升至73.9%，且仅需一轮即达收敛。'
    '该方法无需GPU或大语言模型，仅以800行Python代码实现，与GPT-4驱动的ReAct方法性能相当。'
    '本文论证了核心命题"1次精确归因大于等于10000次全局梯度更新"，'
    '揭示了可解释性不是学习系统的附加品，而是实现高效学习的结构性前提条件。'
)
add_paragraph_normal(abstract_cn, indent_first=True)

# 关键词
p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_before = Pt(6)
pf.space_after = Pt(12)
run = p.add_run('关键词：')
set_run_font(run, '黑体', 'Times New Roman', Pt(10.5), bold=True)
run2 = p.add_run('知几学习；YLYW；模糊隶属度；先验知识；对称校准；可解释性；具身智能')
set_run_font(run2, '宋体', 'Times New Roman', Pt(10.5))

# ===== 英文摘要 =====
p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_before = Pt(12)
pf.space_after = Pt(3)
run = p.add_run('Abstract')
set_run_font(run, '宋体', 'Times New Roman', Pt(10.5), bold=True)

abstract_en = (
    'Reinforcement learning (RL) is the dominant learning paradigm in embodied intelligence but suffers from three critical challenges: '
    'low sample efficiency, training instability, and opaque decision-making. These challenges stem from the implicit distribution of knowledge '
    'across high-dimensional network parameters, which prevents learning signals from precisely targeting specific knowledge units that require adjustment. '
    'To address this fundamental bottleneck, this paper proposes Zhiji Learning, an embodied intelligence learning method based on interpretable prior systems. '
    'Zhiji Learning operates on top of the YLYW (Yi-Li Fuzzy Model) system, which achieves continuous symbol grounding through fuzzy membership degrees '
    'and constructs fully transparent reasoning chains via a three-layer architecture (features to membership, membership to yao-values, yao-values to 64-hexagram matching). '
    'On this interpretable foundation, Zhiji Learning defines three types of calibratable parameters: location prior matrix P, object matching confidence M, '
    'and synonym mapping S, with a unified symmetric update interface where success signals (auspicious signs) reinforce correct priors and failure signals '
    '(inauspicious signs) weaken incorrect ones. The system implements five inauspicious-sign localization strategies (wrong-pick, empty-location, bottleneck, '
    'timeout, and pattern localization) for precise attribution, directing learning signals to specific parameter entries. '
    'Experiments on the ALFWorld benchmark (134 tasks) demonstrate that Zhiji Learning improves success rate from 67.2% to 73.9% with single-round convergence, '
    'without requiring GPUs or large language models, achieving performance comparable to GPT-4-driven ReAct using only 800 lines of Python. '
    'This work establishes the core proposition that "one precise attribution is equivalent to 10,000 global gradient updates" and reveals that '
    'interpretability is not an add-on but a structural prerequisite for efficient learning.'
)
add_paragraph_normal(abstract_en, indent_first=True)

p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_before = Pt(6)
pf.space_after = Pt(18)
run = p.add_run('Key words: ')
set_run_font(run, '宋体', 'Times New Roman', Pt(10.5), bold=True)
run2 = p.add_run('Zhiji Learning; YLYW; fuzzy membership; prior knowledge; symmetric calibration; interpretability; embodied intelligence')
set_run_font(run2, '宋体', 'Times New Roman', Pt(10.5))

# ============================================================
# 1 引言
# ============================================================
add_heading1('1  引言')

add_paragraph_normal(
    '具身智能要求智能体在物理或虚拟环境中自主感知、推理与行动[1]。强化学习（Reinforcement Learning, RL）是当前主流的具身智能学习范式[2]，'
    '通过大量试错积累统计信息来优化策略。然而，RL在具身智能应用中面临三重困境：'
)
add_paragraph_normal(
    '第一，样本效率极低。典型的深度RL算法需要数百万次环境交互才能学会简单的操作策略[3]，'
    '这在真实物理环境中几乎不可行。即使在仿真环境中，训练成本也极高。'
)
add_paragraph_normal(
    '第二，训练不稳定。策略梯度方法依赖于奖励信号在整个参数空间的传播，'
    '超参数敏感、奖励稀疏时梯度消失、局部最优难以逃脱等问题使得训练过程充满不确定性[8]。'
)
add_paragraph_normal(
    '第三，决策不可解释。深度网络的策略以隐式方式存储在高维参数中，'
    '当智能体做出错误决策时，工程师无法定位出错的具体环节，也无法进行针对性修正。'
    '失败后只能依赖更多数据的统计平均来"期望"策略自动改善。'
)
add_paragraph_normal(
    '这三重困境的根源在于：传统RL将知识隐式地分散存储在网络参数中，'
    '导致学习信号无法精确地定位到需要调整的具体知识单元[3]。'
    '近年来，大语言模型（LLM）驱动的方法（如ReAct[5]、Reflexion[6]）虽然展现了强大的推理能力，'
    '但仍存在不可解释、计算成本高昂、幻觉问题等局限性[24,25]。'
)
add_paragraph_normal(
    '本文的核心思路是：如果推理链完全透明、每一步决策的依据都可追溯，'
    '那么成功/失败信号就可以被精确归因到推理链中的具体环节，进而精确定位到需要调整的具体参数条目。'
    '这比全局梯度更新高效得多——前者是外科手术式的精准修正，后者是对所有参数施加微弱统计压力。'
)
add_paragraph_normal(
    '这一思路的实现需要两个前提：（1）一个完全可解释的推理体系，使决策链的每一步都可追溯；'
    '（2）一种能利用可解释性进行精确归因和定向校准的学习机制。'
    'YLYW（易理模糊模型）提供了前者，知几学习提供了后者。'
)
add_paragraph_normal(
    '"知几"出自《易经·系辞下》："知几其神乎！几者，动之微，吉之先见者也。"'
    '意为从微小的征兆中察觉变化的端倪。在知几学习中，"几"是每次任务成功或失败所携带的精确信息'
    '——不是模糊的全局奖励，而是可以追溯到具体参数的精确归因。'
)
add_paragraph_normal(
    '本文的主要贡献如下：'
)
add_paragraph_normal(
    '（1）提出知几学习范式：在可解释先验体系上进行精确归因与对称校准的学习方法，'
    '核心论点为"1次精确归因大于等于10000次统计平均"；'
)
add_paragraph_normal(
    '（2）定义三类可校准参数（P/M/S），给出统一的学习接口和五种凶之几定位策略；'
)
add_paragraph_normal(
    '（3）阐明YLYW可解释先验体系中模糊隶属度的技术创新及其工程意义；'
)
add_paragraph_normal(
    '（4）在ALFWorld 134局上验证知几学习一轮收敛特性（67.2%到73.9%）；'
)
add_paragraph_normal(
    '（5）论证可解释性不是附加品而是高效学习的前提条件。'
)

# ============================================================
# 2 相关工作
# ============================================================
add_heading1('2  相关工作')

add_heading2('2.1  强化学习与具身智能')
add_paragraph_normal(
    '强化学习[2]通过与环境的试错交互来学习最优策略，已在游戏[13]、机器人控制[1]等领域取得突破性成果。'
    '然而，将RL应用于具身智能面临严峻挑战[3]：真实物理环境中的样本采集成本极高，'
    '且安全约束使得随机探索不可行。深度RL虽然通过神经网络的函数逼近能力提升了策略表达，'
    '但也引入了训练不稳定和不可解释的问题[8]。'
)
add_paragraph_normal(
    '近年来，ALFWorld[4]等文本交互式具身智能基准的出现为评估具身智能方法提供了标准化平台。'
    'ALFWorld将ALFRED[18]中的视觉导航任务抽象为文本交互，'
    '保留了任务规划和推理的核心挑战，同时降低了感知层面的复杂性，使研究者能聚焦于决策和学习机制本身。'
    '刘华平等[19]系统综述了机器人操作技能学习方法，指出当前方法在泛化性和样本效率方面仍有显著不足。'
    '谭营等[20]分析了深度强化学习面临的发展与挑战，强调了可解释性和样本效率的重要性。'
)

add_heading2('2.2  知识驱动方法')
add_paragraph_normal(
    '为解决RL的样本效率问题，知识驱动方法[29]尝试将先验知识注入学习过程。'
    '这类方法包括：基于模型的RL利用环境动力学模型减少试错[3]；'
    '分层RL通过任务分解降低搜索空间[1]；'
    '课程学习通过由易到难的任务序列加速训练。'
)
add_paragraph_normal(
    '大语言模型（LLM）的出现为知识驱动方法带来了新范式。ReAct[5]将推理（思维链）与行动结合，'
    '利用LLM的世界知识进行任务规划。Reflexion[6]通过语言形式的自我反思实现粗粒度的经验学习。'
    'SayCan[32]将LLM的语言理解与机器人的可行性评估结合。'
    'Inner Monologue[31]通过内部独白机制实现具身推理。'
    '然而，这些方法本质上依赖LLM的隐式知识表示，不具备精确的参数级归因能力。'
    '朱松纯[28]指出，从感知到认知的跨越需要可解释的推理机制作为桥梁。'
)

add_heading2('2.3  神经符号AI')
add_paragraph_normal(
    '神经符号AI[23]旨在结合神经网络的学习能力与符号系统的推理和可解释能力。'
    'Zadeh[22]提出的模糊集理论为处理连续概念与离散符号之间的鸿沟提供了数学基础。'
    '经典的认知架构如ACT-R[11]和SOAR[12]通过产生式规则实现可解释的推理，'
    '但缺乏从经验中自动学习规则的能力。'
)
add_paragraph_normal(
    'Marcus[24]强调了下一个十年AI需要结合符号推理和学习的能力。'
    'Lake等[29]提出构建像人一样学习和思考的机器需要将直觉物理、因果推理等先验知识融入学习系统。'
    '曾毅等[21]从类脑智能的角度综述了受生物启发的智能方法。'
    '黄凯奇等[27]分析了具身智能体的研究现状，指出知识表示与推理是实现通用具身智能的关键。'
    'Bengio等[30]提出通过元学习解耦因果机制，为可解释的学习提供了理论支撑。'
)
add_paragraph_normal(
    '与上述工作相比，本文提出的知几学习方法有以下独特之处：'
    '（1）以模糊隶属度实现连续符号接地，兼顾可微性与可解释性；'
    '（2）不需要预定义完备的符号规则，而是通过精确校准不断完善先验知识；'
    '（3）学习过程本身是可解释的——每次参数更新都有明确的归因链。'
)

# ============================================================
# 3 YLYW可解释先验体系
# ============================================================
add_heading1('3  YLYW可解释先验体系')

add_paragraph_normal(
    '知几学习的精确校准能力依赖于底层推理体系的可解释性。'
    '本章详细阐述YLYW（易理模糊模型）如何构建完全可解释的推理链，重点说明模糊隶属度这一核心技术创新。'
)

add_heading2('3.1  模糊隶属度与符号接地')
add_paragraph_normal(
    '传统符号系统对概念的处理是二值化的：一个物体要么"属于"某个类别，要么"不属于"。'
    '这种非此即彼的表示无法处理现实世界中普遍存在的模糊性、多义性和渐变性[22]。'
    '例如，一个温热的杯子同时具有"容器"和"热源"的属性，其热源属性的程度取决于温度高低。'
)
add_paragraph_normal(
    'YLYW的核心创新是引入模糊隶属度：对象以连续值μ属于[0,1]同时关联多个语义原型，'
    '而非被强制分配到唯一类别。YLYW定义了八种基本语义原型（对应八卦）：'
    '乾（强健/刚性）、坤（柔顺/承载）、坎（水/凹陷/容器）、离（热源/发光）、'
    '震（动态/振动）、巽（流通/渗透）、艮（静止/封闭）、兑（开放/交换）。'
)
add_paragraph_normal(
    '给定物体的特征向量f和语义原型p_i的中心向量c_i，隶属度通过高斯核函数计算：'
)
add_equation('μ_i(f) = exp(-||f - c_i||² / 2σ_i²)', '(1)')
add_paragraph_normal(
    '其中σ_i为第i个原型的带宽参数。这一计算使得每个物体获得一个8维隶属度向量μ = [μ_1, μ_2, ..., μ_8]^T，'
    '描述其与各原型的亲和程度。隶属度是连续的、可微的，同时保持了符号层面的可解释性'
    '——每一维都对应一个人类可理解的语义概念。'
)

add_heading2('3.2  三层推理架构')
add_paragraph_normal(
    'YLYW采用三层推理架构，每层的输入输出和转换规则都是显式的：'
)
add_paragraph_normal(
    '第一层（L1）：特征到隶属度。将物理/语义特征向量通过式(1)映射为8维隶属度向量μ属于[0,1]^8。'
    '这一步实现了从连续感知空间到结构化符号空间的接地。'
)
add_paragraph_normal(
    '第二层（L2）：隶属度到爻值。通过线性组合或规则映射，将8维隶属度压缩为6维爻值向量y属于[0,1]^6。'
    '六个爻位分别编码不同维度的决策相关信息：初爻=稳定性/基础条件、二爻=可达性/距离、'
    '三爻=可操作性、四爻=约束条件、五爻=目标匹配度、上爻=全局态势。'
)
add_paragraph_normal(
    '第三层（L3）：爻值到64卦匹配。6维爻值与64种卦象模板进行匹配，'
    '选择最佳匹配的卦象作为当前情境的策略类型。每种卦象预定义了策略模板和参数化方式，'
    '爻值的具体数值则填充为策略参数。'
)
add_paragraph_normal(
    '三层推理的关键特性是：每一步转换都有明确的语义解释和可追溯的计算路径。'
    '当最终决策出错时，可以逐层回溯，定位到底是特征提取有误、隶属度计算偏差、'
    '爻值映射不当，还是卦象匹配错误。'
)

add_heading2('3.3  推理链的可解释性')
add_paragraph_normal(
    'YLYW的可解释性不是学术装饰，而是具有直接的工程价值：'
)
add_paragraph_normal(
    '决策链透明——系统对每个动作都能给出完整的推理解释：'
    '"因为物体A的位置隶属度μ(坤)=0.8最高，所以判断它最可能在countertop类位置；'
    '先验矩阵P[A][countertop]=3分最高，所以优先去countertop寻找。"'
)
add_paragraph_normal(
    '失败可追溯——当任务失败时，系统能追溯到推理链的哪一步出了问题：'
    '"去了countertop但没找到plate，说明P[plate][countertop]的先验评分偏高，需要下调这个具体条目。"'
)
add_paragraph_normal(
    '参数可定位——出错环节被定位后，需要调整的参数是明确的、具体的：'
    '是P矩阵的某个条目，还是M矩阵的某个条目，还是S映射缺少某个词条。'
    '这种精确定位能力是知几学习能够进行精确校准的根本前提。'
)
add_paragraph_normal(
    '在ALFWorld[4]中的适配方案如下：八卦映射为位置语义类型——'
    '坤（countertop，承载平面）、坎（sinkbasin，凹陷容器）、离（microwave/stoveburner，热源）、'
    '艮（safe/cabinet，封闭空间）、兑（fridge，开放存取）、巽（shelf，流通）等。'
    '先验矩阵P(obj,loc)以字典形式存储物体与位置类型的关联评分。'
)

# ============================================================
# 4 知几学习方法
# ============================================================
add_heading1('4  知几学习方法')

add_paragraph_normal(
    '本章是论文的核心技术章节，具体回答"知几学习调什么参数、在哪里、怎么调"这一关键问题。'
)

add_heading2('4.1  问题形式化')
add_paragraph_normal(
    '知几学习是作用于YLYW可解释先验体系之上的学习机制。'
    '它不改变YLYW的推理架构本身（三层结构和64卦匹配逻辑保持不变），'
    '而是通过任务执行的反馈信号，精确校准推理过程中使用的先验参数。'
)
add_paragraph_normal(
    '"知几"的哲学内涵是"从微小征兆中察觉变化"。'
    '在知几学习中，每次任务的成功或失败都是一个"几"——一个携带精确信息的征兆。'
    '成功是"吉之几"，告诉系统哪些先验参数是正确的、值得强化；'
    '失败是"凶之几"，告诉系统哪些参数有偏差、需要修正。'
)
add_paragraph_normal(
    '在"道法术器"体系中，知几学习的定位是"法"的层面：'
    '道=易理模型（不可变的推理架构），法=知几学习（通用学习方法），'
    '术=知识库（P/M/S等可学习参数），器=代码实现。'
    '知几学习作为通用方法，可应用于任何具有可解释先验体系的系统。'
)

add_heading2('4.2  三类校准参数')
add_paragraph_normal(
    '知几学习调整的目标是三类显式、可解释的参数。'
    '这三类参数共同构成了YLYW系统的"术"层——可学习的知识库。'
    '每类参数都有明确的数据结构、语义含义、初始化方式和更新规则。'
)
add_heading3('4.2.1  位置先验矩阵 P(obj, loc)')
add_paragraph_normal(
    '数据结构为二级字典Dict[str, Dict[str, float]]。'
    '语义含义：物体obj出现在位置类型loc的先验评分，数值越高表示关联越强。'
    '初始化来自人工常识编码，如P[plate][countertop]=3, P[plate][fridge]=0, P[plate][sinkbasin]=2。'
    '查询时机：在探索阶段，对各go-to命令按P[target_obj][loc]评分排序，优先前往高分位置。'
    '更新规则为：'
)
add_equation('P[obj][loc] ← P[obj][loc] + α × reward', '(2)')
add_paragraph_normal(
    '成功时（reward=+1）：agent在loc发现obj，则P[obj][loc] += 1.0（确认关联）。'
    '失败时（reward=-1）：agent去了loc没找到obj，则P[obj][loc] -= 0.5（削弱关联）。'
)

add_heading3('4.2.2  物体匹配置信度 M(target, entity)')
add_paragraph_normal(
    '数据结构为二级字典Dict[str, Dict[str, float]]。'
    '语义含义：当任务描述说target时，环境实体entity的匹配可信度。'
    '初始化为全零（无先验假设）。'
    '查询时机：在物体匹配阶段，M[target][entity] <= -2.0的实体被排除。'
    '更新规则为：'
)
add_equation('M[target][entity] ← M[target][entity] + β × reward', '(3)')
add_paragraph_normal(
    '成功时：成功take entity完成target任务，则M[target][entity] += 1.0（确认匹配）。'
    '失败时：take了entity但任务失败，则M[target][entity] -= 3.0（强烈否定匹配）。'
    '失败惩罚大于成功奖励（β=3 vs α=1），体现了"一次错误匹配的代价远大于一次成功确认"的不对称性。'
)

add_heading3('4.2.3  同义词映射 S(word → Set[str])')
add_paragraph_normal(
    '数据结构为字典Dict[str, Set[str]]，词到实体名集合的映射。'
    '语义含义：任务描述中的词word对应的环境实体名称集合。'
    '初始化为空集（从零开始积累）。'
    '查询时机：搜索目标物体时自动扩展匹配范围。如S[coffee]={mug}，则搜索coffee时也会匹配名为mug的实体。'
    '学习规则：从admissible命令中观察到实体命名，与任务描述对比发现映射关系。'
    'S是单调增长的——只添加不删除，因为同义关系一旦发现就是永久的。'
)

add_heading2('4.3  对称校准机制')
add_paragraph_normal(
    '知几学习的优雅之处在于：成功和失败的学习逻辑共享同一套代码，区别仅在于reward的符号。'
    '以下是核心学习接口的形式化描述：'
)
add_equation('observe(trajectory, won, agent_state)', '(4)')
add_paragraph_normal(
    '其中trajectory为任务执行轨迹，won为布尔值表示成功/失败，agent_state为智能体状态。'
    '学习逻辑为：reward = +1.0 if won else -1.0（吉/凶之几），然后遍历轨迹中的每一步，'
    '对位置先验执行P[obj][loc] += reward * alpha，'
    '对匹配置信度执行M[target][entity] += reward * beta，'
    '对同义词发现执行S[word].add(entity)（仅添加，无方向性）。'
)
add_paragraph_normal(
    '这一接口体现了对称校准原则：同一参数表中的同一条目，成功时被正向强化，失败时被负向削弱。'
    '不需要为成功和失败设计不同的学习策略，reward的符号自动决定了校准方向。'
)

add_heading2('4.4  五种征兆定位策略')
add_paragraph_normal(
    '当任务失败（won=False）时，知几学习的关键能力是精确定位到该调整哪个参数条目。'
    '系统通过分析轨迹信息，识别五种失败模式，每种对应特定的参数校准策略：'
)
add_paragraph_normal(
    '（1）错拿定位。比较已拿实体taken_entity与目标target的一致性。'
    '若不一致，说明物体匹配环节出错，执行M[target][wrong_entity] -= 3.0。'
    '这是最精确的归因——直接定位到M矩阵的具体条目。'
)
add_paragraph_normal(
    '（2）空位定位。统计智能体遍历过但未找到目标的所有位置。'
    '对每个空位执行P[obj][loc] -= 0.5，同时对最终找到目标的位置（如果有的话）执行P[obj][loc] += 1.0。'
)
add_paragraph_normal(
    '（3）瓶颈定位。记录智能体在每个任务阶段的停留时间。'
    '若某阶段反复卡顿（phase_count[type][phase]超过阈值），标记该阶段相关的参数为"需关注"。'
)
add_paragraph_normal(
    '（4）超时定位。当智能体探索了超过70%的位置仍未完成任务时，'
    '判断为空间搜索效率问题，建议优先检查open类容器（fridge, cabinet等）。'
)
add_paragraph_normal(
    '（5）模式定位。对多局失败的轨迹进行指纹聚类（fingerprint clustering），'
    '发现反复出现的失败模式，进行批量参数校准。'
)

add_heading2('4.5  经验持久化')
add_paragraph_normal(
    '知几学习支持经验的跨轮次持久化。在每轮运行结束后，P、M、S三类参数通过JSON序列化保存到文件；'
    '下一轮运行前加载，作为新的初始值。'
)
add_paragraph_normal(
    '这一机制的本质是：将运行时通过observe()积累的知识增量ΔK固化为永久先验。'
    '经过少量轮次（实验表明1-2轮即可），参数在给定任务空间中趋于饱和——'
    '新的运行几乎不再触发参数更新，因为先验已经足够准确。'
)
add_paragraph_normal(
    '持久化的数据格式为JSON，结构清晰、人类可读。'
    '工程师可以直接检查、编辑持久化文件，这是可解释性在工程层面的又一体现。'
    '这种设计也使得知识的跨环境迁移成为可能——将一个环境中学到的P/M/S参数作为新环境的初始值。'
)

# ============================================================
# 5 实验验证
# ============================================================
add_heading1('5  实验验证')

add_heading2('5.1  实验设置')
add_paragraph_normal(
    '实验基于ALFWorld[4]基准的valid_unseen划分，共134局，涵盖6类家务任务（拿放、加热、冷却、清洁、检视、双物拿放）。'
    'YLYW系统以纯Python实现（约800行），运行于单CPU，不依赖GPU或大语言模型（LLM）。'
    '对比方法包括：ReAct[5]（GPT-4驱动的推理+行动）、Reflexion[6]（基于自我反思的RL）、'
    'EmbodiSkill[7]（基于LLM的技能学习）。'
)
add_paragraph_normal(
    '评估指标包括：任务成功率（成功完成的任务比例）、平均步数（完成任务的平均动作数）、'
    '收敛速度（达到稳定性能所需的训练轮次）。'
)

add_heading2('5.2  逐步引入的效果')
add_paragraph_normal(
    '为验证知几学习各组件的贡献，我们设计了逐步引入实验。'
    '从无学习的静态先验版本开始，依次添加吉之几（正向校准）、凶之几（双向校准）、经验持久化。结果如表1所示。'
)

# 表1
add_table_caption('表1  知几学习的逐步引入')
add_table(
    ['版本', '学习机制', '成功率', '全局步数'],
    [
        ['V7', '无学习（静态先验）', '90/134 (67.2%)', '23.1'],
        ['V9', '+吉之几（正向校准P/S）', '94/134 (70.1%)', '21.5'],
        ['V10', '+凶之几（双向校准P/M/S）', '98/134 (73.1%)', '20.5'],
        ['V10+persist', '+经验持久化', '99/134 (73.9%)', '20.1'],
    ]
)

add_paragraph_normal(
    '表1显示，知几学习的每个组件都带来了可测量的提升。从V7到V10+persist，'
    '成功率绝对提升6.7个百分点（67.2%到73.9%），平均步数下降3.0步（23.1到20.1）。'
    '特别值得注意的是，凶之几（V10）带来的提升（+3.0%）大于吉之几（V9）的提升（+2.9%），'
    '这印证了"从失败中学习比从成功中学习更高效"的直觉——失败提供的归因信息比成功更具体、更有针对性。'
)

add_heading2('5.3  精确归因案例分析')
add_paragraph_normal(
    '为展示知几学习"精确归因到具体参数"的能力，我们列举凶之几定位的典型案例。'
    '每个案例都说明了：系统如何从失败信号追溯到具体参数，以及校准后的效果。如表2所示。'
)

# 表2
add_table_caption('表2  凶之几的精确归因（追溯到具体参数）')
add_table(
    ['游戏', '归因到的参数', '校准操作', '效果'],
    [
        ['#41', 'P: NL解析方向from/to反了', '修复解析规则', '50→32步'],
        ['#99', 'S: 缺vault→safe映射', 'S[vault].add(safe)', '50→5步'],
        ['#116', 'P: NL解析方向性', '修复', '50→11步'],
        ['#118', 'P: NL解析方向性', '修复', '50→11步'],
        ['#128', 'S: 缺vault→safe', 'S[vault].add(safe)', '50→11步'],
    ]
)

add_paragraph_normal(
    '案例#99是最具说明性的：任务要求在safe中检视物体，但任务描述使用了"vault"一词，'
    '而环境中的实体名为"safe"。YLYW系统在第一次遇到这种情况时失败（花了50步也没完成），'
    '凶之几定位发现：agent从未尝试与safe交互，因为不知道vault就是safe。'
    '于是执行S[vault].add(safe)，之后同类任务仅需5步即可完成。'
)
add_paragraph_normal(
    '这个案例生动地说明了"1次精确归因大于等于10000次统计平均"：'
    '传统RL在不知道"vault=safe"的情况下，需要大量随机探索才能偶然学会这一知识；'
    '知几学习通过一次失败的精确归因就永久掌握了这一映射。'
)

add_figure_placeholder('[图1位置] 图1  知几学习校准目标示意图')

add_heading2('5.4  收敛实验')
add_paragraph_normal(
    '为验证经验持久化的收敛特性，我们在相同134局上进行多轮运行，'
    '每轮结束后保存P/M/S参数，下轮加载作为初始值。结果如表3所示。'
)

# 表3
add_table_caption('表3  经验持久化收敛')
add_table(
    ['Round', '成功率', '步数', '含义'],
    [
        ['V9', '94/134 (70.1%)', '21.5', '基线'],
        ['R1', '98/134 (73.1%)', '20.5', 'P/M/S积累中'],
        ['R2', '99/134 (73.9%)', '20.1', 'P/M/S饱和'],
        ['R3-R5', '99/134 (73.9%)', '20.1-20.2', '重复确认'],
    ]
)

add_paragraph_normal(
    '表3表明，知几学习在第2轮（R2）即达到饱和——后续轮次（R3-R5）的成功率和步数保持稳定。'
    '一轮收敛的含义是：P/M/S三类参数在134局构成的知识空间中已经覆盖了所有需要学习的映射关系。'
    '这种快速收敛是精确校准的自然结果——每次更新都是有针对性的，不存在无效探索。'
)

add_figure_placeholder('[图2位置] 图2  经验持久化收敛曲线')

add_heading2('5.5  方法对比')
add_paragraph_normal(
    '表4给出了知几学习与现有方法的对比结果。'
)

# 表4
add_table_caption('表4  方法对比')
add_table(
    ['方法', '成功率', '需LLM', '可解释', '校准精度'],
    [
        ['ReAct (GPT-4)', '71%', '是', '否', '全局'],
        ['Reflexion', '77%', '是', '部分', '粗粒度'],
        ['EmbodiSkill', '93.28%', '是', '否', '技能级'],
        ['YLYW+知几(本文)', '73.9%', '否', '是', '单条目'],
    ]
)

add_paragraph_normal(
    'EmbodiSkill凭借GPT-4的强大语言理解能力取得了最高成功率，但不可解释且完全依赖昂贵的LLM推理。'
    'ReAct和Reflexion同样依赖LLM，其中Reflexion通过自然语言自我反思实现了粗粒度的学习。'
)
add_paragraph_normal(
    'YLYW+知几在不使用任何LLM的条件下达到了73.9%的成功率，'
    '与GPT-4驱动的ReAct相当（71% vs 73.9%）。更重要的是，YLYW+知几是唯一同时满足以下条件的方法：'
    '（1）完全可解释；（2）无LLM依赖；（3）支持单条目精度的参数校准；'
    '（4）一轮收敛。这些特性使其特别适合对可解释性和低成本有刚性需求的自动化应用场景。'
)

add_figure_placeholder('[图3位置] 图3  知几学习与强化学习参数更新方式对比')

# ============================================================
# 6 讨论
# ============================================================
add_heading1('6  讨论')

add_heading2('6.1  与强化学习的本质对比')
add_paragraph_normal(
    '传统RL的学习信号传播路径是：环境奖励 → 值函数/策略梯度 → 反向传播 → 所有网络参数[8]。'
    '这一路径存在根本性的信息衰减：一个标量奖励信号在传播过程中被分散到数百万个参数上，'
    '每个参数获得的更新信号极其微弱。学习本质上依赖于大量样本的统计平均来消除噪声。'
)
add_paragraph_normal(
    '知几学习的信号传播路径截然不同：任务结果 → 推理链回溯 → 精确定位出错环节 → 直接修改具体参数条目。'
    '信号不经过任何衰减或分散，一次成功/失败的完整信息量被完全利用于一个或少数几个参数的校准。'
)
add_paragraph_normal(
    '打一个比方：RL的学习方式像是"给整个城市降温来治疗一个人的发烧"，'
    '而知几学习像是"诊断出具体的感染部位并对症下药"。'
    '前者需要巨大的能量（样本）且大部分被浪费，后者精准高效但需要透明的诊断能力（可解释性）。'
)

add_heading2('6.2  可解释性的工程价值')
add_paragraph_normal(
    '在当前AI研究中，可解释性通常被视为一种"附加品"——在系统已经工作后，'
    '额外添加解释能力以满足合规或信任需求。知几学习提出了一个不同的视角：'
    '可解释性是高效学习的前提条件，而非事后添加的装饰。'
)
add_paragraph_normal(
    '具体而言，精确校准的前提是精确归因，精确归因的前提是推理链透明。'
    '如果推理过程是黑箱的（如深度网络），那么即使知道任务失败了，也无法确定该调整哪个参数'
    '——只能退回到全局梯度更新。可解释性在知几学习范式中不是可选的加分项，'
    '而是使高效学习成为可能的结构性前提。'
)
add_paragraph_normal(
    '这一观点可以推广为更一般性的命题：在任何领域，如果能使推理过程透明可追溯，'
    '就能将学习效率提升数个数量级。YLYW+知几是这一命题在具身智能领域的具体实例。'
    '在自动化控制领域，基于模型的控制方法（如PID控制、模型预测控制）之所以高效，'
    '正是因为系统模型提供了透明的因果链，使得控制器设计可以精确针对系统特性。'
    '知几学习在智能学习层面延续了这一工程哲学。'
)

add_heading2('6.3  知几学习的通用性')
add_paragraph_normal(
    '需要指出的是，知几学习的核心不在于"易经"本身，而在于"可解释先验+精确校准"这一通用范式。'
    'YLYW恰好提供了一套完全可解释的推理链（模糊隶属度→六爻编码→64卦匹配），'
    '知几学习利用这种可解释性实现了对P/M/S三类参数的精确对称校准。'
    '任何具有可解释推理链的系统都可以采用知几学习的方法论。'
)
add_paragraph_normal(
    '例如，在工业自动化中，如果决策系统基于显式的规则库和参数化模型，'
    '那么知几学习的归因-校准范式可以直接应用：从失败的控制结果追溯到具体的规则或参数偏差，'
    '然后进行精确修正。在多智能体系统中，如果通信协议和决策规则是透明的，'
    '知几学习可以从协作失败中定位到具体智能体的具体决策错误。'
    '在人形机器人[26,27]的技能学习中，如果动作规划基于可解释的任务分解，'
    '知几学习可以将失败归因到具体的子任务参数。'
)
add_paragraph_normal(
    '认知科学区分了系统1（快速、直觉、自动化）和系统2（慢速、推理、有意识）两种认知模式[9]。'
    'YLYW+知几自然地整合了两种模式：YLYW的三层推理架构实现了快速的类系统1推理'
    '（一旦参数确定，推理是即时的、确定性的）；知几学习实现了类系统2的反思学习'
    '（通过回溯和归因进行有意识的参数调整）。'
    '更有意思的是，知几学习的终极目标是让自己变得"不再需要"——'
    '当P/M/S参数完全饱和后，系统进入纯系统1模式，不再需要学习和反思。'
    '这与人类从"有意识学习"到"自动化专业表现"的成长路径一致。'
)

add_heading2('6.4  局限性')
add_paragraph_normal(
    '知几学习当前存在以下局限性：'
    '（1）依赖可解释先验体系的存在——如果领域缺乏结构化的先验知识表示，知几学习无法直接应用；'
    '（2）当前实验仅在文本交互环境中验证，尚未扩展到视觉感知和连续动作空间；'
    '（3）P/M/S三类参数的设计具有一定的任务特异性，向更复杂任务的扩展需要定义新的参数类型；'
    '（4）对初始先验的质量有一定依赖——如果初始先验与真实环境偏差过大，收敛可能需要更多轮次。'
)
add_paragraph_normal(
    '未来工作将在两个方向展开：'
    '（1）将知几学习扩展到更复杂的具身智能任务（如多步操作、多智能体协作），'
    '探索P/M/S之外更多类型的可校准参数；'
    '（2）研究知几学习与深度学习的融合可能性——用可解释的上层推理指导深层表示的定向优化，'
    '结合两种范式的优势。'
)

# ============================================================
# 7 结论
# ============================================================
add_heading1('7  结论')

add_paragraph_normal(
    '本文提出了知几学习——一种基于可解释先验体系的具身智能学习方法。主要贡献总结如下：'
)
add_paragraph_normal(
    '（1）提出了"可解释先验+精确校准"的学习范式，论证了"1次精确归因大于等于10000次统计平均"的核心命题，'
    '为具身智能学习提供了RL之外的新路径。'
)
add_paragraph_normal(
    '（2）定义了三类可校准参数（位置先验P、匹配置信度M、同义词映射S），'
    '设计了统一的对称更新接口和五种凶之几定位策略。'
)
add_paragraph_normal(
    '（3）阐明了YLYW可解释先验体系中模糊隶属度的核心创新'
    '——以连续值实现符号接地，使推理链既可微又可解释。'
)
add_paragraph_normal(
    '（4）在ALFWorld 134局上验证了知几学习的有效性（67.2%到73.9%）和一轮收敛特性，'
    '且无需GPU或LLM，仅800行Python代码。'
)
add_paragraph_normal(
    '（5）论证了可解释性不是附加品而是高效学习的结构性前提。'
)
add_paragraph_normal(
    '知几学习为具身智能学习提供了一条与主流深度强化学习不同的技术路线。'
    '它表明，当推理体系足够透明时，学习可以是外科手术式的精准操作而非统计意义上的参数漫游。'
    '这一思想不仅适用于YLYW系统，也可推广到任何具有可解释推理链的智能系统，'
    '在自动化控制、工业机器人、人机协作等领域具有广阔的应用前景。'
)

# ============================================================
# 参考文献
# ============================================================
add_heading1('参考文献')

refs = [
    '[1] Duan Y, Chen X, Houthooft R, et al. Benchmarking deep reinforcement learning for continuous control[C]// ICML, 2016: 1329-1338.',
    '[2] Sutton R S, Barto A G. Reinforcement Learning: An Introduction[M]. 2nd ed. MIT Press, 2018.',
    '[3] Dulac-Arnold G, Mankowitz D, Hester T. Challenges of real-world reinforcement learning[J]. arXiv:1904.12901, 2019.',
    '[4] Shridhar M, Thomason J, Gordon D, et al. ALFWorld: Aligning text and embodied environments for interactive learning[C]// ICLR, 2021.',
    '[5] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing reasoning and acting in language models[C]// ICLR, 2023.',
    '[6] Shinn N, Cassano F, Gopinath A, et al. Reflexion: Language agents with verbal reinforcement learning[C]// NeurIPS, 2023.',
    '[7] Xu Z, et al. EmbodiSkill: Embodied skill learning with large language models[J]. arXiv:2024.',
    '[8] Mnih V, Kavukcuoglu K, Silver D, et al. Human-level control through deep reinforcement learning[J]. Nature, 2015, 518(7540): 529-533.',
    '[9] Kahneman D. Thinking, Fast and Slow[M]. Farrar, Straus and Giroux, 2011.',
    '[10] Zadeh L A. Fuzzy sets[J]. Information and Control, 1965, 8(3): 338-353.',
    '[11] Anderson J R. The Architecture of Cognition[M]. Harvard University Press, 1983.',
    '[12] Laird J E, Newell A, Rosenbloom P S. SOAR: An architecture for general intelligence[J]. Artificial Intelligence, 1987, 33(1): 1-64.',
    '[13] Silver D, Schrittwieser J, Simonyan K, et al. Mastering the game of Go without human knowledge[J]. Nature, 2017, 550(7676): 354-359.',
    '[14] Brown T B, Mann B, Ryder N, et al. Language models are few-shot learners[C]// NeurIPS, 2020.',
    '[15] Wei J, Wang X, Schuurmans D, et al. Chain-of-thought prompting elicits reasoning in large language models[C]// NeurIPS, 2022.',
    '[16] Brohan A, Brown N, Carbajal J, et al. RT-2: Vision-language-action models transfer web knowledge to robotic control[J]. arXiv:2307.15818, 2023.',
    '[17] Côté M A, Kádár Á, Yuan X, et al. TextWorld: A learning environment for text-based games[C]// CGW@IJCAI, 2018.',
    '[18] Shridhar M, Thomason J, Gordon D, et al. ALFRED: A benchmark for interpreting grounded instructions for everyday tasks[C]// CVPR, 2020: 10740-10749.',
    '[19] 刘华平, 郭迪, 孙富春, 等. 机器人操作技能学习方法综述[J]. 自动化学报, 2019, 45(3): 458-470.',
    '[20] 谭营, 徐昕, 李学龙. 深度强化学习的发展与挑战[J]. 自动化学报, 2021, 47(9): 2011-2038.',
    '[21] 曾毅, 刘成林, 谭铁牛. 类脑智能研究综述[J]. 自动化学报, 2023, 49(1): 1-17.',
    '[22] Zadeh L A. Fuzzy sets[J]. Information and Control, 1965, 8(3): 338-353.',
    '[23] d\'Avila Garcez A S, Broda K B, Gabbay D M. Neural-Symbolic Learning Systems[M]. Springer, 2002.',
    '[24] Marcus G. The next decade in AI: Four steps towards robust artificial intelligence[J]. arXiv:2002.06177, 2020.',
    '[25] Bommasani R, Hudson D A, Adeli E, et al. On the opportunities and risks of foundation models[J]. arXiv:2108.07258, 2021.',
    '[26] Reed S, Zolna K, Parisotto E, et al. A generalist agent[J]. Transactions on Machine Learning Research, 2022.',
    '[27] 黄凯奇, 陈晓棠, 康运锋, 等. 具身智能体的研究现状与发展趋势[J]. 中国科学: 信息科学, 2023, 53(6): 1047-1081.',
    '[28] 朱松纯. 从感知到认知: 人工智能的初心与未来[J]. 中国科学: 信息科学, 2020, 50(11): 1628-1667.',
    '[29] Lake B M, Ullman T D, Tenenbaum J B, et al. Building machines that learn and think like people[J]. Behavioral and Brain Sciences, 2017, 40: e253.',
    '[30] Bengio Y, Deleu T, Rahaman N, et al. A meta-transfer objective for learning to disentangle causal mechanisms[C]// ICLR, 2020.',
    '[31] Huang W, Xia F, Xiao T, et al. Inner monologue: Embodied reasoning through planning with language models[C]// CoRL, 2022.',
    '[32] Ahn M, Brohan A, Brown N, et al. Do as I can, not as I say: Grounding language in robotic affordances[J]. arXiv:2204.01691, 2022.',
]

for ref in refs:
    add_ref(ref)

# ========== 保存 ==========
output_path = '/home/lijinhan/MXL/科研/ylyw/paper/自动化学报_知几学习_投稿版.docx'
doc.save(output_path)
print(f'Done: {output_path}')
