# -*- coding: utf-8 -*-
import json
from docx import Document
from copy import deepcopy
from docx.oxml.ns import qn

# Load texts from JSON (avoids encoding issues in Python source)
with open('texts_agi.json', 'r', encoding='utf-8') as f:
    T = json.load(f)

def make_para(parent_el, text, style_from=None, doc=None):
    new_p = parent_el.makeelement(qn('w:p'), {})
    if style_from and doc:
        for p in doc.paragraphs:
            if p.style.name == style_from:
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    new_p.append(deepcopy(pPr))
                break
    new_r = parent_el.makeelement(qn('w:r'), {})
    new_t = parent_el.makeelement(qn('w:t'), {})
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p

# Load original
doc = Document('中国科学_信息科学_YLYW完整版.docx')
print(f"Loaded: {len(doc.paragraphs)} paras, {sum(len(p.text) for p in doc.paragraphs)} chars")

# 1. Abstract addition
p7 = doc.paragraphs[7]
p7.add_run(T['abstract_addition'])
print("[1/4] Abstract done")

# 2. AGI positioning after contributions (para 38)
ref = doc.paragraphs[38]._element
el = make_para(ref, T['agi_positioning'], doc=doc)
ref.addnext(el)
print("[2/4] AGI positioning done")

# 3. Add AGI direction to 6.6
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('第三，知几学习与强化学习的混合'):
        ref2 = p._element
        el2 = make_para(ref2, T['sec66_agi_direction'], doc=doc)
        ref2.addnext(el2)
        print(f"[3/4] sec6.6 AGI direction added after para {i}")
        break

# 3b. Strengthen core insight
for i, p in enumerate(doc.paragraphs):
    if '本文的核心洞见可以概括为一句话' in p.text:
        p.add_run(T['core_insight_addition'])
        print(f"  Core insight strengthened at para {i}")
        break

# 4. Change heading + add outlook
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and p.text.strip().startswith('7') and '结论' in p.text:
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = "7 结论与展望"
        else:
            p.add_run("7 结论与展望")
        break

# Find 参考文献 heading, insert outlook before it
ref_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and '参考文献' in p.text:
        ref_idx = i
        break

ref_el = doc.paragraphs[ref_idx]._element

outlook_parts = [
    ('Heading 2', T['outlook_heading']),
    (None, T['outlook_intro']),
    (None, T['outlook_p1']),
    (None, T['outlook_p2']),
    (None, T['outlook_p3']),
    (None, T['outlook_closing']),
]

for style, text in reversed(outlook_parts):
    el = make_para(ref_el, text, style_from=style, doc=doc)
    ref_el.addprevious(el)

print("[4/4] Outlook section done")

# Save
output = '中国科学_信息科学_YLYW完整版_v2_AGI展望.docx'
doc.save(output)
doc2 = Document(output)
chars2 = sum(len(p.text) for p in doc2.paragraphs)
print(f"\n{'='*50}")
print(f"SAVED: {output}")
print(f"  {len(doc2.paragraphs)} paragraphs, {chars2} chars")
print(f"  (Original: 274 paragraphs, 28773 chars)")
