"""V2 Part3: 讨论+结论+参考文献（去掉[1][2][3]，重新编号）"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('中国科学_YLYW_ALFWorld_v2.docx')

def sup(p, n):
    r = p.add_run(f'[{n}]')
    r.font.superscript = True
    r.font.size = Pt(8)

def h1(t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: r.font.name = '黑体'
def h2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = '黑体'
def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p

# ====== 6 讨论 ======
h1('6  讨论')

h2('6.1  关于Admissible Commands的公平性')
para('admissible_commands列表实质上等价于完美的环境状态传感器——"take plate 2 from countertop 2"直接暴露了物体存在与位置，将感知问题转化为列表筛选。本文与ReAct、Reflexion、EmbodiSkill在完全相同的admissible条件下对比，结论有效。BUTLER使用更困难的generation模式（37%），与之对比时应注明条件差异。本文贡献应理解为：在给定合法动作空间的前提下，如何用极轻量的规则系统实现超越训练型方法、接近LLM方法的决策效率。')

h2('6.2  YLYW先验在导航决策中的跨域验证')
p = para('本文结果从跨域角度验证了YLYW"知几学习"范式的有效性。在物理域（抓取决策），YLYW先验是"物体物理属性→卦象→策略"映射，92.7%合理率；在导航域（ALFWorld），先验是"物体典型位置→探索优先级"映射，90.0%成功率（标注一致子集）。两者共享相同的设计原则——将人类常识编码为结构化先验，零样本驱动决策——且在不同领域呈现稳定的90%+能力水平。这一一致性并非巧合：当先验知识覆盖了任务空间的主要情况时，零样本性能自然收敛于该水平，对应《系辞传》"范围天地之化而不过"')
sup(p, 11)
p.add_run('的理想。')

h2('6.3  规则系统 vs LLM：两条技术路线的边界')
para('在相同admissible条件下，约800行规则代码（标注一致时90.0%）接近270亿参数LLM（93.28%）。规则系统的优势在于：（1）完全确定性，相同输入始终相同输出，不存在LLM的采样随机性；（2）零幻觉，严格从admissible列表选择，不生成无效动作；（3）部署成本极低，无需GPU/API，134个任务仅需约180秒纯CPU运算；（4）完全可解释，每个决策可追溯到具体的先验矩阵评分和状态机阶段。')

p = para('LLM的核心优势在于语义弹性：当task_desc说"salt shaker"但场景只有pepper shaker时，LLM能通过上下文推断修正目标。此外，EmbodiSkill的多轮迭代演化机制允许从失败中学习，这是静态规则系统无法复制的。因此，')
p.add_run('两条路线的边界条件是：当任务描述准确清晰时，规则系统可接近LLM水平；当描述存在歧义或噪声时，LLM的语义理解能力不可替代。').bold = True

h2('6.4  ALFWorld标注质量问题的启示')
para('本文发现的25.4%标注不一致率对ALFWorld基准有重要启示。这些不一致并非随机噪声，而是系统性的：（1）ALFRED数据集的task_desc由众包工人标注，与PDDL自动生成的ground truth之间存在语义gap；（2）同义词混用（mug/cup 11个、salt/pepper 3个）反映了自然语言的固有模糊性；（3）信息缺失（"Turn on the desk lamp"不说看什么物体，5个）反映了标注指南的不完善。这一发现对所有在ALFWorld上评估的方法都有影响：使用LLM的方法（如EmbodiSkill）之所以能绕过这些不一致，不是因为其算法更优，而是因为LLM具备语义弹性。建议ALFWorld社区对这34个任务进行标注修正或标记。')

h2('6.5  局限性与未来工作')
para('（1）NL解析精度：当前基于关键词的解析器面对同义词和模糊表达时能力有限，未来可结合轻量NLP模型。（2）50步限制：pick_two类型需完成两轮操作，50步下成功率仅47.1%，放宽至80步可显著提升。（3）admissible模式的现实局限：真实机器人不存在admissible oracle，推广到真实世界需结合视觉感知模块。（4）领域特化：任务模板针对ALFWorld 6种类型设计，推广到新任务类型需人工编写模板。（5）YLYW先验矩阵的构建：当前矩阵基于常识手工编码，未来可探索从少量示例中自动学习先验（对应K_calibration的自动化）。')

# ====== 7 结论 ======
h1('7  结论')
para('本文提出了一种基于YLYW先验知识框架的零样本具身决策方法，并在ALFWorld基准测试（134个valid_unseen任务）上进行了系统验证。核心成果包括：')

for c in [
    '将YLYW三层推理架构（L1八卦隶属度→L2六爻编码→L3卦象匹配）成功迁移至具身导航领域，设计了admissible-commands信号驱动的层次化状态机Agent，以约800行Python代码实现67.2%的整体成功率。',
    '揭示了ALFWorld数据集中25.4%任务（34/134）存在task_desc与PDDL标注不一致的问题。在标注一致的100个任务上，YLYW达到90.0%的成功率，与使用270亿参数LLM的EmbodiSkill（93.28%）仅差约3个百分点。',
    '验证了YLYW"知几学习"范式（K = K_prior ⊕ K_calibration）的跨域有效性：同一套"先验知识+信号驱动"方法论从物理域（92.7%）成功迁移至导航域（90.0%）。',
    '通过与EmbodiSkill的系统对比，明确了先验知识驱动与LLM驱动两条技术路线的优势边界：在任务描述准确时，精心设计的规则系统可接近大规模LLM的性能水平，同时在确定性、可解释性和部署成本方面具有显著优势。',
]:
    doc.add_paragraph(c, style='List Number')

p = para('本文的核心论点是：在具身智能的有限状态空间中，结构化先验知识+确定性规则系统是一条被严重低估的技术路径。《易经》以64卦"弥纶天地之化"的思想，在具身决策领域得到了工程化验证——这为"知识驱动的具身智能"提供了有力的实证支持。')
doc.add_paragraph()

# ====== 参考文献（去掉原[1][2][3]，重新编号） ======
h1('参考文献')
refs = [
    '[1]  Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models. In: ICLR, 2023.',
    '[2]  Shinn N, Cassano F, Gopinath A, et al. Reflexion: Language Agents with Verbal Reinforcement Learning. In: NeurIPS, 2023.',
    '[3]  Ju R, Wang X, Ding X, et al. EmbodiSkill: Skill-Aware Reflection for Self-Evolving Embodied Agents. arXiv:2605.10332, 2026.',
    '[4]  Shridhar M, Yuan X, Côté M A, et al. ALFWorld: Aligning Text and Embodied Environments for Interactive Learning. In: ICLR, 2021.',
    '[5]  Shridhar M, Thomason J, Gordon D, et al. ALFRED: A Benchmark for Interpreting Grounded Instructions for Everyday Tasks. In: CVPR, 2020.',
    '[6]  Côté M A, Kádár Á, Yuan X, et al. TextWorld: A Learning Environment for Text-Based Games. In: CGW@IJCAI, 2018.',
    '[7]  Ahn M, Brohan A, Brown N, et al. Do As I Can, Not As I Say: Grounding Language in Robotic Affordances. arXiv:2204.01691, 2022.',
    '[8]  Driess D, Xia F, Sajjadi M S M, et al. PaLM-E: An Embodied Multimodal Language Model. In: ICML, 2023.',
    '[9]  Brohan A, Brown N, Carbajal J, et al. RT-2: Vision-Language-Action Models Transfer Web Knowledge to Robotic Control. arXiv:2307.15818, 2023.',
    '[10] 《周易》. 先秦典籍.',
    '[11] 《系辞传》. 先秦典籍.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(ref).font.size = Pt(9)

doc.save('中国科学_YLYW_ALFWorld_v2.docx')
print('V2 Part3完成: 6讨论+7结论+参考文献（共11条，无团队自引）')
print('论文全文生成完毕！')
