# -*- coding: utf-8 -*-
"""生成 YLYW ALFWorld 技术论文 V6 Word文档"""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from paper_content import *

# 创建文档
doc = Document()

# ============ 设置默认样式 ============
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ============ 辅助函数 ============
def add_heading_styled(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading

def add_para(text, bold=False, first_line_indent=True, font_size=None):
    para = doc.add_paragraph()
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    if font_size:
        run.font.size = font_size
    return para

def add_bullet(text):
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Cm(1.0)
    for run in para.runs:
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return para

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_with_data(table_info):
    caption = table_info.get('caption')
    headers = table_info['headers']
    rows = table_info['rows']
    
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10)
        run.bold = True
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')
    
    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_paragraph()  # 表后空行
    return table

def render_section(items):
    """渲染一个section的内容列表"""
    for item in items:
        item_type = item[0]
        content = item[1]
        if item_type == 'p':
            add_para(content)
        elif item_type == 'p_bold':
            add_para(content, bold=True)
        elif item_type == 'b':
            add_bullet(content)
        elif item_type == 'h2':
            add_heading_styled(content, level=2)
        elif item_type == 'h3':
            add_heading_styled(content, level=3)
        elif item_type == 'table':
            add_table_with_data(content)
        elif item_type == 'code':
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(content)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9)


# ============ 论文正文开始 ============

# 标题
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run(TITLE)
title_run.font.size = Pt(16)
title_run.bold = True
title_run.font.name = '黑体'
title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 作者
author_para = doc.add_paragraph()
author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author_para.add_run(AUTHOR)
author_run.font.size = Pt(12)
author_run.font.name = '宋体'
author_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# 摘要
add_heading_styled('摘要', level=1)
add_para(ABSTRACT)

# 关键词
kw_para = doc.add_paragraph()
kw_para.paragraph_format.first_line_indent = Cm(0.74)
kw_run = kw_para.add_run('关键词：')
kw_run.bold = True
kw_run.font.name = '宋体'
kw_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
kw_run2 = kw_para.add_run(KEYWORDS)
kw_run2.font.name = '宋体'
kw_run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# 1. 引言
add_heading_styled('1. 引言', level=1)
render_section(SEC1_INTRO)
doc.add_paragraph()

# 2. 相关工作
add_heading_styled('2. 相关工作', level=1)

add_heading_styled('2.1 ALFWorld基准测试', level=2)
render_section(SEC2_1)

add_heading_styled('2.2 基于大语言模型的具身智能方法', level=2)
render_section(SEC2_2)

add_heading_styled('2.3 规则和训练型方法', level=2)
render_section(SEC2_3)

add_heading_styled('2.4 YLYW易理模糊模型', level=2)
render_section(SEC2_4)
doc.add_paragraph()

# 3. 问题定义
add_heading_styled('3. 问题定义', level=1)

add_heading_styled('3.1 ALFWorld环境形式化', level=2)
render_section(SEC3_1)

add_heading_styled('3.2 评估设置', level=2)
render_section(SEC3_2)
doc.add_paragraph()

# 4. 方法
add_heading_styled('4. 方法', level=1)

add_heading_styled('4.1 系统架构概览', level=2)
render_section(SEC4_1)

add_heading_styled('4.2 环境适配：方案B (Per-Game Env)', level=2)
render_section(SEC4_2)

add_heading_styled('4.3 目标提取模块', level=2)
render_section(SEC4_3)

add_heading_styled('4.4 层次化状态机', level=2)
render_section(SEC4_4)

add_heading_styled('4.5 Admissible-Commands信号驱动', level=2)
render_section(SEC4_5)

add_heading_styled('4.6 YLYW常识先验', level=2)
render_section(SEC4_6)

add_heading_styled('4.7 V6增强能力', level=2)
render_section(SEC4_7)

add_heading_styled('4.8 算法伪代码', level=2)
add_para('YLYW Agent的完整决策流程如算法1所示：')
algo_para = doc.add_paragraph()
algo_run = algo_para.add_run(SEC4_8_ALGO)
algo_run.font.name = 'Courier New'
algo_run.font.size = Pt(9)
doc.add_paragraph()

# 5. 实验
add_heading_styled('5. 实验', level=1)

add_heading_styled('5.1 实验设置', level=2)
render_section(SEC5_1)

add_heading_styled('5.2 主实验结果', level=2)
render_section(SEC5_2)

add_heading_styled('5.3 消融实验', level=2)
render_section(SEC5_3)

add_heading_styled('5.4 步数效率分析', level=2)
render_section(SEC5_4)

add_heading_styled('5.5 失败案例分析', level=2)
render_section(SEC5_5)

add_heading_styled('5.6 与文献方法对比', level=2)
render_section(SEC5_6)
doc.add_paragraph()

# 6. 讨论
add_heading_styled('6. 讨论', level=1)

add_heading_styled('6.1 为什么YLYW能超越LLM方法？', level=2)
render_section(SEC6_1)

add_heading_styled('6.2 YLYW先验的作用分析', level=2)
render_section(SEC6_2)

add_heading_styled('6.3 局限性', level=2)
render_section(SEC6_3)

add_heading_styled('6.4 推广性讨论', level=2)
render_section(SEC6_4)
doc.add_paragraph()

# 7. 结论
add_heading_styled('7. 结论', level=1)
render_section(SEC7)
doc.add_paragraph()

# 参考文献
add_heading_styled('参考文献', level=1)
for ref in REFERENCES:
    ref_para = doc.add_paragraph()
    ref_para.paragraph_format.first_line_indent = Cm(0)
    ref_para.paragraph_format.left_indent = Cm(0.5)
    ref_run = ref_para.add_run(ref)
    ref_run.font.name = 'Times New Roman'
    ref_run.font.size = Pt(10)
    ref_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# 附录A
add_heading_styled('附录A：完整实验数据', level=1)
render_section(APPENDIX_A)

# 附录B
add_heading_styled('附录B：YLYW先验矩阵的完整定义', level=1)
render_section(APPENDIX_B)

# 附录C
add_heading_styled('附录C：代码结构说明', level=1)
add_para('YLYW Agent V6的代码结构如下：')
code_para = doc.add_paragraph()
code_run = code_para.add_run(APPENDIX_C_CODE)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph()

# 保存文档
output_path = '/home/lijinhan/MXL/科研/ylyw/alfworld_exp/YLYW_ALFWorld_技术论文_V6.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'论文已保存到: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
