#!/usr/bin/env python3
"""将 EXPERIMENT_REPORT.md 转换为 docx 格式"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
import os

REPORT_MD = os.path.join(os.path.dirname(__file__), "EXPERIMENT_REPORT.md")
OUTPUT = os.path.join(os.path.dirname(__file__), "YLYW_ALFWorld实验报告.docx")


def parse_md_sections(md_text: str) -> list:
    """简易MD解析为段落结构"""
    lines = md_text.strip().split('\n')
    sections = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith('# ') and not line.startswith('## '):
            sections.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            sections.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            sections.append(('h3', line[4:].strip()))
        elif line.startswith('---'):
            sections.append(('hr', ''))
        elif line.startswith('|') and '|' in line[1:]:
            # 表格
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            sections.append(('table', table_lines))
            i -= 1  # 修正
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            sections.append(('code', '\n'.join(code_lines)))
        elif line.startswith('- **'):
            sections.append(('bullet_bold', line[2:].strip()))
        elif line.startswith('- '):
            sections.append(('bullet', line[2:].strip()))
        elif line.strip():
            sections.append(('para', line.strip()))
        else:
            sections.append(('blank', ''))
        i += 1
    return sections


def build_docx(sections: list, output_path: str):
    doc = Document()

    # 样式设置
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    for sec_type, content in sections:
        if sec_type == 'h1':
            p = doc.add_heading(content, level=1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
        elif sec_type == 'h2':
            p = doc.add_heading(content, level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
        elif sec_type == 'h3':
            doc.add_heading(content, level=3)
        elif sec_type == 'hr':
            doc.add_paragraph('─' * 60)
        elif sec_type == 'table':
            _add_table(doc, content)
        elif sec_type == 'code':
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.left_indent = Cm(1)
        elif sec_type == 'bullet_bold':
            p = doc.add_paragraph(style='List Bullet')
            m = re.match(r'\*\*(.+?)\*\*(.*)', content)
            if m:
                run_b = p.add_run(m.group(1))
                run_b.bold = True
                p.add_run(m.group(2))
            else:
                p.add_run(content)
        elif sec_type == 'bullet':
            doc.add_paragraph(content, style='List Bullet')
        elif sec_type == 'para':
            # 处理内联加粗
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.+?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        elif sec_type == 'blank':
            doc.add_paragraph('')

    doc.save(output_path)
    print(f"✅ DOCX saved: {output_path}")


def _add_table(doc, table_lines: list):
    """解析Markdown表格并添加到docx"""
    if len(table_lines) < 2:
        return

    # 解析表头
    headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
    # 跳过分隔行
    data_rows = []
    for tl in table_lines[2:]:
        cells = [c.strip() for c in tl.split('|')[1:-1]]
        data_rows.append(cells)

    ncols = len(headers)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 数据行
    for i, row_data in enumerate(data_rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)


def main():
    with open(REPORT_MD, 'r') as f:
        md_text = f.read()

    sections = parse_md_sections(md_text)
    build_docx(sections, OUTPUT)


if __name__ == '__main__':
    main()
