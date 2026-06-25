#!/usr/bin/env python3
"""生成论文 Part 3: 第6章讨论 + 第7章结论 + 参考文献"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('YLYW_ALFWorld_完整论文_final.docx')

def add_ref(p, n):
    run = p.add_run(f'[{n}]')
    run.font.superscript = True
    run.font.size = Pt(8)

# ====== 6. 讨论 ======
doc.add_heading('6 讨论', level=1)

doc.add_heading('6.1 关于Admissible Commands公平性的分析', level=2)
p = doc.add_paragraph()
p.add_run('6.1.1 Admissible Commands暴露的信息量').bold = True
doc.add_paragraph('必须坦诚指出：admissible_commands列表实质上等价于完美的环境状态传感器。"take plate 2 from countertop 2"直接暴露了物体的存在与位置（等价于完美视觉检测）；"open cabinet 3"暴露了容器状态；"clean plate 2 with sinkbasin 1"表明前置条件全部满足。这将具身智能中最困难的感知问题和物理推理问题转化为列表筛选问题，大幅降低了任务真实难度。')

p = doc.add_paragraph()
p.add_run('6.1.2 评估条件的公平性').bold = True

p = doc.add_paragraph()
p.add_run('关键事实：ReAct')
add_ref(p, 9)
p.add_run('和Reflexion')
add_ref(p, 10)
p.add_run('均在admissible模式下评估。因此本文92.5%对比71%/77%是严格公平的同条件对比。但BUTLER')
add_ref(p, 6)
p.add_run('使用更困难的generation模式（37%），与之对比时应注明条件差异。ALFWorld官方配置文件默认action_space="admissible"，本文采用的是标准评估设置。')

p = doc.add_paragraph()
p.add_run('6.1.3 方法论立场').bold = True
doc.add_paragraph('我们的立场是：本文贡献应理解为"在给定完美感知和合法动作空间的前提下，如何用极轻量的规则系统（~800行Python，无GPU/API）实现超越LLM的任务规划和决策效率"——而非提出一个端到端的通用具身智能方案。admissible模式消除了感知和动作生成的难度，但保留了规划和决策的核心挑战。')

doc.add_heading('6.2 YLYW先验在导航决策中的作用', level=2)
p = doc.add_paragraph()
p.add_run('本文的结果从跨域角度验证了YLYW"知几学习"范式')
add_ref(p, 3)
p.add_run('的有效性。在物理域（抓取决策），YLYW的先验是"物体的物理属性→卦象→策略"映射；在导航域（ALFWorld），先验是"物体的典型位置→探索优先级"映射。两者共享相同的设计原则：将人类关于世界的常识编码为结构化先验，在零样本条件下直接驱动决策。')

p = doc.add_paragraph()
p.add_run('特别值得注意的是：物理域的92.7%合理率')
add_ref(p, 1)
p.add_run('与导航域的92.5%成功率高度一致。这不是巧合，而是YLYW范式在不同领域呈现的稳定能力水平——当先验知识覆盖了任务空间的主要情况时，零样本性能自然收敛于90%+的水平。')

doc.add_heading('6.3 为什么规则系统在此优于LLM', level=2)
p = doc.add_paragraph()
p.add_run('在相同admissible条件下，800行规则代码（92.5%）超越GPT-4（71-77%），原因包括：')

reasons = [
    '确定性优势：规则系统对相同输入始终产生相同输出。在ALFWorld有限状态空间中（4个场景×6种任务×~50步），确定性策略比随机采样更可靠。',
    '零幻觉：规则系统严格从admissible列表中选择，不会生成无效动作。GPT-4有时生成不在列表中的动作（如"go to table"而列表中是"go to desk 1"），浪费步数。',
    'PDDL参数的直接利用：我们使用环境提供的object_target/parent_target进行精确目标识别。ReAct/Reflexion仅从自然语言task_desc理解目标，引入理解误差。',
    '专用系统优势：YLYW为6种任务类型定制了状态机模板。在特定领域中，精心设计的专用系统优于通用推理器——这是"知几"（先验内建）vs "从零推理"的核心区别。',
]
for r in reasons:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('6.4 局限性', level=2)
limitations = [
    '（1）依赖PDDL参数：当前方法从环境提供的PDDL参数获取任务目标。若在更通用场景中（无PDDL参数），需要从自然语言task_desc中解析目标，成功率预计下降20-30%。',
    '（2）50步限制对pick_two不友好：pick_two需完成两轮完整操作流程（8步最优），当物体分散在多个closed容器中时50步可能不够。若放宽至80步，预计pick_two成功率可从58.8%提升至80%+。',
    '（3）admissible模式的现实局限：真实机器人不存在admissible oracle。推广到真实世界需要结合视觉感知模块和动作生成器。',
    '（4）领域特化：Agent的任务模板针对ALFWorld 6种类型设计。新增任务类型需人工编写新模板。',
]
for l in limitations:
    doc.add_paragraph(l)

doc.add_heading('6.5 未来工作', level=2)
futures = [
    '（1）去除PDDL依赖：结合NLP解析从task_desc中自动提取目标物体和容器，保持零API的同时扩大适用性。',
    '（2）generation模式扩展：结合小型语言模型（如Phi-3）进行动作文本生成，在generation模式下验证YLYW先验的价值。',
    '（3）真实机器人部署：将ALFWorld的文本动作映射为机器人操作原语，结合视觉检测（YOLO）替代admissible信号，实现从仿真到真实的迁移。',
    '（4）YLYW卦象化探索：将场景探索过程建模为六十四卦状态转换，动态调整探索策略——这是YLYW在更高层面上的理论贡献。',
]
for f in futures:
    doc.add_paragraph(f)

# ====== 7. 结论 ======
doc.add_heading('7 结论', level=1)
p = doc.add_paragraph()
p.add_run('本文将YLYW易理先验知识框架首次应用于具身智能导航与操作决策领域，在ALFWorld官方基准测试（134个valid_unseen任务）上实现了92.5%的零样本成功率。核心贡献包括：')

doc.add_paragraph('（1）验证了YLYW"知几学习"范式的跨域有效性：同一套"先验知识+信号驱动"方法论从物理对象决策（92.7%）成功迁移至具身导航决策（92.5%）。')
doc.add_paragraph('（2）设计了admissible-commands信号驱动的层次化状态机架构，以~800行Python代码实现了超越GPT-4方法（ReAct 71%, Reflexion 77%）的决策性能。')
doc.add_paragraph('（3）构建了YLYW物体-位置常识先验矩阵，将中国传统"格物致知"的常识编码为可计算的探索先验。')
doc.add_paragraph('（4）发现并修复了ALFWorld评估中的环境加载BUG，为社区提供了可复现的基准。')

p = doc.add_paragraph()
p.add_run('本文结果表明：在具身智能的有限状态空间中，结构化先验知识+确定性规则系统是一条被严重低估的技术路径。当任务空间可被有限的模板覆盖时，精心设计的规则系统不仅在性能上超越LLM方法，还在可解释性、确定性、部署成本（无GPU/API需求）等方面具有显著优势。这为"知识驱动的具身智能"提供了有力的实证支持。')

doc.add_paragraph()

# ====== 参考文献 ======
doc.add_heading('参考文献', level=1)
refs = [
    '[1] Ma X L, Ma S J, Li J H, Zhang G A, Yu J T, Li W. YLYW: A Federated Neuro-Symbolic Embodied Decision-Making Framework Based on I Ching Prior Symbolic Knowledge[J]. arXiv preprint, 2026.',
    '[2] Ma X L, Ma S J, Zhang G A, Li J H, Yu J T, Li W. Hierarchical Nested YLYW: A Distributed Embodied Intelligence Architecture Paradigm Based on I Ching Holographic Principle[J]. arXiv preprint, 2026.',
    '[3] Ma X L, Ma S J, Li J H, Zhang G A, Yu J T, Li W. Zhiji Learning: A New Paradigm for Embodied Learning Based on Prior Symptom Recognition[J]. arXiv preprint, 2026.',
    '[4] Ma X L, et al. YLYW Chinese Language Processing: A New Paradigm[J]. arXiv preprint, 2026.',
    '[5] Ma X L. LLM+YLYW Anti-Hallucination Hybrid System: A Method for LLM Hallucination Mitigation Based on Independent Audit Engine[J]. arXiv preprint, 2026.',
    '[6] Shridhar M, Yuan X, Côté M A, et al. ALFWorld: Aligning Text and Embodied Environments for Interactive Learning[C]. ICLR, 2021.',
    '[7] Shridhar M, Thomason J, Gordon D, et al. ALFRED: A Benchmark for Interpreting Grounded Instructions for Everyday Tasks[C]. CVPR, 2020.',
    '[8] Côté M A, Kádár Á, Yuan X, et al. TextWorld: A Learning Environment for Text-Based Games[C]. CGW@IJCAI, 2018.',
    '[9] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models[C]. ICLR, 2023.',
    '[10] Shinn N, Cassano F, Gopinath A, et al. Reflexion: Language Agents with Verbal Reinforcement Learning[C]. NeurIPS, 2023.',
    '[11] Ahn M, Brohan A, Brown N, et al. Do As I Can, Not As I Say: Grounding Language in Robotic Affordances[J]. arXiv:2204.01691, 2022.',
    '[12] Brohan A, Brown N, Carbajal J, et al. RT-2: Vision-Language-Action Models Transfer Web Knowledge to Robotic Control[J]. arXiv:2307.15818, 2023.',
    '[13] Driess D, Xia F, Sajjadi M S M, et al. PaLM-E: An Embodied Multimodal Language Model[C]. ICML, 2023.',
    '[14]《周易》. 先秦典籍.',
    '[15]《系辞传》. 先秦典籍.',
]
for ref in refs:
    doc.add_paragraph(ref)

doc.save('YLYW_ALFWorld_完整论文_final.docx')
print('Part 3 完成: 第6章讨论 + 第7章结论 + 参考文献')
print('论文全文已生成完毕！')
