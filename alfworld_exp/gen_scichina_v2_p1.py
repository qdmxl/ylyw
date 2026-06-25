"""中国科学:信息科学 V2 - 理论自包含版 Part1"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

def sup(p, n):
    r = p.add_run(f'[{n}]')
    r.font.superscript = True
    r.font.size = Pt(8)

def h1(t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: r.font.name = '黑体'
def h2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = '黑体'
def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p

# ====== 标题 ======
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.first_line_indent = Cm(0)
run = t.add_run('基于YLYW先验知识的零样本具身决策方法\n及其在ALFWorld基准上的验证')
run.bold = True
run.font.size = Pt(16)

a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
a.paragraph_format.first_line_indent = Cm(0)
a.add_run('马兴录').bold = True
a.add_run('1*  ')
a.add_run('张国安').bold = True
a.add_run('1  ')
a.add_run('李金函').bold = True
a.add_run('1  ')
a.add_run('于敬涛').bold = True
a.add_run('1  ')
a.add_run('李望').bold = True
a.add_run('1  ')
a.add_run('马圣洁').bold = True
a.add_run('1')

a2 = doc.add_paragraph()
a2.alignment = WD_ALIGN_PARAGRAPH.CENTER
a2.paragraph_format.first_line_indent = Cm(0)
a2.add_run('1. 青岛科技大学 信息科学技术学院, 青岛 266061\n* 通讯作者. E-mail: maxinglu@qust.edu.cn')
doc.add_paragraph()

# ====== 中文摘要 ======
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.add_run('摘要  ').bold = True
p.add_run('当前具身智能决策方法主要依赖大语言模型（LLM），存在部署成本高、决策不可解释、输出随机等问题。本文提出一种基于YLYW（易理研物）先验知识框架的零样本具身决策方法。YLYW以《易经》六十四卦为结构化先验模板，通过L1八卦模糊隶属度→L2六爻编码→L3卦象匹配的三层推理链，实现"卦定策略类型、爻定执行参数"的分层决策。本文将这一框架迁移至具身导航与操作领域，设计了admissible-commands信号驱动的层次化状态机Agent，在ALFWorld官方TextWorld仿真器的134个valid_unseen任务上验证。该Agent仅依赖约800行Python代码和YLYW常识先验矩阵，不使用任何LLM或API调用。实验表明，该方法达到67.2%的整体成功率。更重要的是，本文揭示了ALFWorld数据集中25.4%的任务存在自然语言标注与PDDL ground truth不一致的问题。在标注一致的100个任务上，YLYW达到90.0%的成功率，接近使用270亿参数LLM的EmbodiSkill方法（93.28%），仅差约3个百分点。本文结果验证了YLYW"知几学习"范式——先验知识⊕少量校准即可实现零样本决策——在具身导航领域的跨域有效性。')

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.add_run('关键词  ').bold = True
p.add_run('YLYW; 易理先验知识; 具身智能; ALFWorld; 零样本决策; 层次化状态机; 标注一致性')
doc.add_paragraph()

# ====== 英文摘要 ======
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.add_run('Abstract  ').bold = True
p.add_run('Current embodied AI decision-making methods rely heavily on large language models (LLMs), suffering from high deployment costs, unexplainable decisions, and stochastic outputs. This paper proposes a zero-shot embodied decision-making method based on the YLYW (Yi-Li-Yan-Wu) prior knowledge framework, which uses the 64 hexagrams of I Ching as structured prior templates through a three-layer reasoning chain: L1 trigram fuzzy membership → L2 six-yao encoding → L3 hexagram matching. We design an admissible-commands-driven hierarchical state machine agent for ALFWorld\'s 134 valid_unseen tasks using only ~800 lines of Python without any LLM. The agent achieves 67.2% overall success. We reveal that 25.4% of ALFWorld tasks contain annotation inconsistencies between task descriptions and PDDL ground truth. On the 100 consistently-annotated tasks, YLYW achieves 90.0%, approaching EmbodiSkill\'s 93.28% (requiring a 27B-parameter LLM) by only ~3 percentage points.')

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.add_run('Keywords  ').bold = True
p.add_run('YLYW; I Ching prior knowledge; embodied AI; ALFWorld; zero-shot; hierarchical state machine; annotation consistency')
doc.add_paragraph()

# ====== 1 引言 ======
h1('1  引言')
p = para('具身智能（Embodied AI）要求智能体在物理或仿真环境中感知、推理并执行动作序列以完成任务，面临感知-动作闭环、长程规划和组合爆炸三重挑战。当前主流方法依赖大语言模型（LLM）：ReAct')
sup(p, 1)
p.add_run('利用GPT-4进行推理-行动循环（成功率71%），Reflexion')
sup(p, 2)
p.add_run('增加多轮反思（77%），EmbodiSkill')
sup(p, 3)
p.add_run('通过技能自演化使270亿参数LLM达93.28%。然而，这些方法均依赖大规模LLM推理，部署成本高、决策不可解释、输出具有随机性。')

p = para('一个自然的问题是：是否存在一条不依赖LLM、仅凭结构化先验知识即可实现高效具身决策的路径？')

p = para('本文基于YLYW（易理研物）先验知识框架回答这一问题。YLYW的核心思想源自《易经》的"象-数-理"认知范式')
sup(p, 10)
p.add_run('：以八卦为模糊基元、六十四卦为先验模板、爻位关系为参数修正算子，构建联邦式三层推理架构（L1八卦隶属度→L2六爻编码→L3卦象匹配）。该框架已在物理域300物体零样本抓取决策中达到92.7%合理率和0%严重错误率。本文是YLYW从"物理对象决策"向"具身导航与操作决策"的首次跨域扩展验证。')

p = para('本文在ALFWorld')
sup(p, 4)
p.add_run('官方TextWorld')
sup(p, 6)
p.add_run('仿真器上评估YLYW方法。ALFWorld是将ALFRED视觉导航基准')
sup(p, 5)
p.add_run('与TextWorld文本引擎结合的具身智能评估平台，其valid_unseen测试集包含134个任务，覆盖6种家务任务类型。本文的主要贡献包括：')

for c in [
    '提出YLYW三层先验推理架构（L1八卦隶属度→L2六爻编码→L3卦象匹配）在具身导航领域的迁移方法，设计了admissible-commands信号驱动的层次化状态机Agent。',
    '构建YLYW物体-位置常识先验矩阵和NL任务描述解析器，以约800行Python代码实现不依赖LLM的零样本决策（67.2%成功率）。',
    '揭示ALFWorld数据集25.4%任务的task_desc标注与PDDL ground truth不一致，在标注一致子集上YLYW达到90.0%，接近270亿参数LLM方法（93.28%）。',
    '验证了YLYW"知几学习"范式（K = K_prior ⊕ K_calibration）在具身导航领域的跨域有效性。',
    '发现并修复ALFWorld官方TextWorld环境的游戏加载BUG，为社区提供可复现的评估基础。',
]:
    doc.add_paragraph(c, style='List Number')
doc.add_paragraph()

# ====== 2 YLYW理论框架 ======
h1('2  YLYW先验知识框架')

h2('2.1  设计动机：先验知识驱动 vs 数据驱动')
p = para('当前具身智能决策的主流范式是数据驱动：端到端VLA模型（如RT-2')
sup(p, 9)
p.add_run('）需要数十万条机器人轨迹训练，LLM-based方法（如PaLM-E')
sup(p, 8)
p.add_run('）依赖千亿参数预训练。这些方法在零样本场景下性能受限，且决策过程不可解释。')

p = para('YLYW提出了另一条路径：将人类关于物理世界的先验知识——特别是《易经》六十四卦中编码的"观物取象"认知范式')
sup(p, 10)
p.add_run('——形式化为可计算的推理引擎。《系辞传》云："易与天地准，故能弥纶天地之道"')
sup(p, 11)
p.add_run('，表明易理试图以有限的符号系统（64卦×6爻=384种基本情境）"弥纶"（覆盖）万物变化之道。YLYW将这一哲学思想工程化：以八卦为连续模糊基元解决符号接地问题，以六十四卦为结构化先验模板提供强归纳偏置，以爻位关系为参数修正算子实现精细调节。')

h2('2.2  三层推理架构：L1八卦→L2六爻→L3六十四卦')
para('YLYW采用联邦式神经符号架构，核心设计原则为：先验知识独立（不"溶解"进神经网络权重）、连续与离散贯通（模糊隶属度而非二值判断）、推理链完全可追溯。三层架构如下：')

p = para('')
p.add_run('L1 八卦基元层：连续模糊隶属度。').bold = True
p.add_run('传统符号系统将连续物理量二值化（"重或轻"、"大或小"），丢失了物理世界的模糊性。YLYW的解决方案是：一个对象不是"属于或不属于"某卦，而是以不同程度的隶属度μ∈[0,1]同时关联多个卦象。八卦（乾坤震巽坎离艮兑）对应八种基本物理/语义原型，给定特征向量f和卦象原型p，通过高斯核函数计算隶属度。在物理域中，八卦对应8种物理特征原型（乾"健"→高稳定性、坤"顺"→低力需求、坎"陷"→含水/凹陷等）。在本文的导航域中，八卦映射为位置类型的语义原型（坤→平面/承载如countertop、坎→水域如sinkbasin、离→热源如microwave、艮→封闭如safe）。')

p = para('')
p.add_run('L2 六爻编码层：从连续量到符号向量。').bold = True
p.add_run('将8维隶属度通过加权公式聚合为6维爻值向量y∈[0,1]⁶。每爻≥0.5为阳爻（—），<0.5为阴爻（--）。六爻从初爻到上爻分别对应不同维度的语义。在物理域中：初爻=基础稳定性、二爻=可达性、三爻=力需求、四爻=脆弱性、五爻=优先级、上爻=环境约束。在本文导航域中：初爻=物体匹配度、二爻=位置相关度、三爻=操作可行性、四爻=先验置信度、五爻=阶段匹配度、上爻=探索新鲜度。六爻编码的关键设计在于每个爻值携带明确的语义锚定——这不是任意的，而是对应《周易》中"初爻代表事物根基""五爻代表事物鼎盛"的传统定位。')

p = para('')
p.add_run('L3 六十四卦规则层：结构化先验模板匹配。').bold = True
p.add_run('给定6维爻值向量y，通过余弦相似度在64个卦象的理想爻模板中搜索最佳匹配，确定策略类型。每个卦象关联一个预定义策略，包含策略类型、参数预设和注意事项。在物理域中，卦象对应抓取策略（如乾卦→标准抓取、坎卦→特殊处理湿滑物体）。在导航域中，卦象对应探索策略（如渐卦→逐步探索、师卦→目标导向搜索、蹇卦→困难转向）。卦象-策略映射的独特之处在于：理想爻模板不是简单的二值向量，而是考虑了爻位在工程语境中的相对重要性。')

p = para('该架构实现了"卦定策略类型、爻定执行参数"的分层决策体系。在300物体零样本抓取基线测试中，系统达到92.7%合理率和0%严重错误率。三维消融实验验证了三个设计选择的独立贡献：易理规则（+33.6%）、三层架构（+12.7%）、连续模糊隶属度（+23.0%）。')

h2('2.3  知几学习范式')
p = para('YLYW进一步提出"知几学习"范式，与主流的强化学习（RL）形成本质区分。其哲学基础源于《系辞下》："知几其神乎！几者，动之微，吉之先见者也。君子见几而作，不俟终日"')
sup(p, 11)
p.add_run('。核心主张：具身智能体的学习不必从白纸开始，关于世界变化规律的基本知识应作为先验内建于系统中。学习的真正功能不是"从无到有发现规律"，而是在先验框架上校准参数：')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('K = K')
run.italic = True
r2 = p.add_run('prior')
r2.font.subscript = True
r2.font.italic = True
p.add_run(' ⊕ K')
r3 = p.add_run('calibration')
r3.font.subscript = True
r3.font.italic = True

para('其中K_prior是先验知识（YLYW的64卦规则库、常识先验矩阵），K_calibration是少量经验校准。本文在ALFWorld上的应用正是知几学习的典型案例：Agent携带物体-位置先验知识和任务类型模板"出生"，在运行时通过admissible_commands信号"感知征兆"（见几），然后即时决策（而作），无需任何训练数据或试错迭代。')

doc.add_paragraph()

# ====== 3 相关工作 ======
h1('3  相关工作')

h2('3.1  ALFWorld与具身智能基准')
p = para('ALFWorld')
sup(p, 4)
p.add_run('是将ALFRED')
sup(p, 5)
p.add_run('与TextWorld')
sup(p, 6)
p.add_run('结合的具身智能评估平台。valid_unseen测试集包含134个可解任务，覆盖6种家务任务类型（查看物体、拿放、清洗后放、加热后放、冷却后放、拿两个物体放置），是评估零样本具身决策的标准基准。')

h2('3.2  LLM驱动的具身决策方法')
p = para('ReAct')
sup(p, 1)
p.add_run('利用GPT-4交替进行推理和行动，成功率71%。Reflexion')
sup(p, 2)
p.add_run('增加失败后反思机制，允许多次重试，成功率77%。EmbodiSkill')
sup(p, 3)
p.add_run('提出技能感知反思（Skill-Aware Reflection），区分技能缺陷与执行失误实现定向修正，使Qwen3.5-27B达93.28%。SayCan')
sup(p, 7)
p.add_run('和PaLM-E')
sup(p, 8)
p.add_run('将LLM与机器人操作原语结合。训练型方法BUTLER')
sup(p, 4)
p.add_run('使用DAgger训练在generation模式下达37%。这些方法要么依赖大规模LLM推理，要么需要大量训练数据。')

h2('3.3  神经符号与知识驱动方法')
p = para('神经符号AI（NeSy）试图结合神经网络的学习能力与符号系统的推理可解释性。然而，已有NeSy系统多工作在抽象符号世界（逻辑推理、知识图谱），缺乏与物理世界的直接交互。YLYW的独特定位在于：（1）以《易经》六十四卦这一紧凑完备的符号系统（2⁶=64种情境）替代通用谓词逻辑；（2）通过连续模糊隶属度而非离散真值实现符号接地；（3）直接面向物理世界的具身决策而非抽象推理。')
doc.add_paragraph()

doc.save('中国科学_YLYW_ALFWorld_v2.docx')
print('V2 Part1完成: 标题+摘要+1引言+2YLYW理论+3相关工作')
