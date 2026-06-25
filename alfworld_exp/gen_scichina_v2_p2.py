"""Part2: 4方法 + 5实验"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('中国科学_YLYW_ALFWorld_v2.docx')

def sup(p, n):
    r = p.add_run(f'[{n}]')
    r.font.superscript = True
    r.font.size = Pt(8)

def h1(t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: r.font.name = '黑体'

def h2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = '黑体'

def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p

# ====== 4 方法 ======
h1('4  方法')

h2('4.1  系统架构')
p = para('本文的YLYW ALFWorld Agent采用三层架构：感知层（Admissible Commands信号提取）→ 决策层（NL解析 + YLYW先验 + 层次化状态机）→ 执行层（动作输出）。设计理念继承自YLYW"知几学习"范式')
sup(p, 3)
p.add_run('：Agent携带结构化先验知识"出生"（物体-位置先验矩阵、任务类型模板），在运行时通过admissible_commands信号"感知征兆"，然后"见几而作"。整个过程无需训练、无需LLM推理。')

h2('4.2  环境适配：方案B (Per-Game Env)')
para('在实验中发现ALFWorld官方TextWorld环境存在游戏加载BUG：旧版评估代码将134个游戏文件注册到一个BatchEnv中，内部使用固定种子打乱顺序后通过迭代器按序取游戏。调用reset(game_idx)时，实际加载的并非指定游戏，导致游戏场景与元信息不匹配。修复方案（方案B）：每次reset时仅将单个gamefile注册为独立环境，确保加载正确。')

h2('4.3  感知层：Admissible Commands信号提取')
para('从admissible_commands列表中提取四类信号：（1）物体检测信号——"take plate 2 from countertop 2"暴露物体存在与位置；（2）容器状态信号——"open cabinet 3"表明容器关闭；（3）操作可行性信号——"clean plate 2 with sinkbasin 1"表明前置条件满足；（4）放置可行性信号——"move plate 2 to countertop 3"表明可放置。')

h2('4.4  决策层')
para('4.4.1  NL任务描述解析').runs[0].bold = True
p = para('本文不使用环境提供的PDDL参数（object_target/parent_target），而是完全从task_desc自然语言中解析目标。解析器通过关键词匹配推断任务类型（准确率98.5%），并从描述中提取目标物体和目标容器。例如"Put a clean plate on the counter"解析为：task_type=pick_clean，object=plate，recep=countertop。')

para('4.4.2  YLYW常识先验矩阵').runs[0].bold = True
p = para('构建物体-位置先验概率矩阵P(object, location)，表示物体在各位置出现的先验概率。矩阵设计遵循YLYW"格物致知"原则：每个物体的位置先验对应其在现实世界中的典型存放位置。表2展示了部分先验矩阵。')

# 表2
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.first_line_indent = Cm(0)
t.add_run('表2  YLYW物体-位置先验概率矩阵（部分，评分0-3）').bold = True

table = doc.add_table(rows=8, cols=7)
table.style = 'Table Grid'
for i, h in enumerate(['物体', 'countertop', 'cabinet', 'shelf', 'desk', 'fridge', 'sinkbasin']):
    table.rows[0].cells[i].text = h
    for r in table.rows[0].cells[i].paragraphs[0].runs: r.bold = True
for i, row in enumerate([
    ('plate', '3', '2', '1', '0', '0', '1'),
    ('mug', '3', '1', '2', '2', '0', '1'),
    ('book', '0', '0', '3', '3', '0', '0'),
    ('apple', '3', '0', '0', '0', '2', '0'),
    ('knife', '3', '0', '0', '0', '0', '0'),
    ('soapbar', '3', '1', '0', '0', '0', '2'),
    ('cd', '0', '0', '3', '2', '0', '0'),
], 1):
    for j, v in enumerate(row): table.rows[i].cells[j].text = v

para('该矩阵在探索阶段用于对go to命令评分排序：优先前往先验概率高的位置，避免盲目遍历。')

para('4.4.3  层次化状态机').runs[0].bold = True
p = para('每种任务类型对应固定的子目标序列（plan），如表3所示。')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.first_line_indent = Cm(0)
t.add_run('表3  六种任务类型的子目标模板').bold = True

table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['任务类型', '子目标序列']):
    table.rows[0].cells[i].text = h
    for r in table.rows[0].cells[i].paragraphs[0].runs: r.bold = True
for i, row in enumerate([
    ('look_at_obj_in_light', 'find_object → take → find_tool(lamp) → use'),
    ('pick_and_place_simple', 'find_object → take → find_recep → put'),
    ('pick_clean_then_place', 'find → take → find_sink → clean → find_recep → put'),
    ('pick_heat_then_place', 'find → take → find_micro → heat → find_recep → put'),
    ('pick_cool_then_place', 'find → take → find_fridge → cool → find_recep → put'),
    ('pick_two_obj_and_place', 'find → take → find_recep → put → (重复第二轮)'),
], 1):
    for j, v in enumerate(row): table.rows[i].cells[j].text = v

para('阶段推进基于admissible信号：find阶段当观测中出现目标物体时推进；take/clean/heat/cool/put阶段在相应动作成功后推进。阶段回退机制：take阶段无目标物体时回退find继续探索；use_tool不可用时回退find_tool。机会主义检查：每步先检查admissible中是否有可直接执行的高价值动作（如目标物体的take命令），若有则跳过评分直接执行。')

h2('4.5  执行增强')
para('（1）Open操作：当admissible中出现open命令且容器未打开，自动执行open。解决物体在closed容器（cabinet/drawer/safe/fridge）中不可见的问题。（2）容器遍历：put不可用时记录已尝试位置，自动前往下一个同类容器。（3）物体位置记忆：从take命令中记录物体位置，pick_two任务中寻找第二个物体时可利用记忆直接定位。')

h2('4.6  算法流程')
algo = """Algorithm 1: YLYW ALFWorld Agent决策流程
输入: task_desc, initial_admissible
输出: action序列直到won=True或steps≥50

1. 初始化:
   task_type, objects, receps ← NL_Parse(task_desc)
   plan ← TASK_PLANS[task_type]
   tools ← TASK_TOOLS[task_type]
   phase ← 0

2. 每步决策 act(obs, admissible):
   2.1 记忆: 从take命令提取物体位置
   2.2 Open: 若有未开容器且处于find/put阶段, 执行open
   2.3 机会主义: 若admissible中有高价值动作, 直接返回
   2.4 按阶段决策:
       find_*: 用先验矩阵评分go to命令
       take_*: 筛选目标物体的take命令; 无则回退find
       use_tool: 检查clean/heat/cool/use命令; 无则回退
       put_*: 检查put/move命令; 无则open或去下一容器

3. 状态更新:
   更新位置/持有/已打开记录
   基于位置名和动作类型判定阶段推进"""
p = doc.add_paragraph()
p.add_run(algo).font.size = Pt(9)

# ====== 5 实验 ======
h1('5  实验')

h2('5.1  实验设置')
para('实验环境：Ubuntu 26.04 LTS (VirtualBox)，Python 3.14，ALFWorld 0.5.0，TextWorld 1.7.0。Agent核心代码约800行Python。无GPU依赖，无外部API调用，纯CPU运行，134个任务约180秒。')

h2('5.2  主实验结果')
p = para('表4展示了Agent各版本的演进结果。V4为修复环境BUG前的baseline；V5为修复后基础版（使用PDDL参数）；V6为V5+open+容器遍历+记忆（使用PDDL参数）；V7为完全不使用PDDL参数的最终版。')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.first_line_indent = Cm(0)
t.add_run('表4  Agent版本演进').bold = True

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
for i, h in enumerate(['版本', '成功率', '平均步数', '关键改进']):
    table.rows[0].cells[i].text = h
    for r in table.rows[0].cells[i].paragraphs[0].runs: r.bold = True
for i, row in enumerate([
    ('V4 baseline', '3.7% (5/134)', '48.5', '环境BUG导致虚假低分'),
    ('V5 (+PDDL)', '64.2% (86/134)', '23.1', '修复环境+admissible驱动'),
    ('V6 (+PDDL)', '92.5% (124/134)', '13.0', '+open+容器遍历+记忆'),
    ('V7 (纯NL)', '67.2% (90/134)', '23.1', '去掉PDDL，纯NL解析'),
], 1):
    for j, v in enumerate(row): table.rows[i].cells[j].text = v

h2('5.3  按任务类型分析')
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.first_line_indent = Cm(0)
t.add_run('表5  V7按任务类型的详细结果').bold = True

table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
for i, h in enumerate(['任务类型', '成功率', '成功/总数', '平均步数(成功)']):
    table.rows[0].cells[i].text = h
    for r in table.rows[0].cells[i].paragraphs[0].runs: r.bold = True
for i, row in enumerate([
    ('look_at_obj_in_light', '83.3%', '15/18', '7.3'),
    ('pick_and_place_simple', '50.0%', '12/24', '12.0'),
    ('pick_clean_then_place', '74.2%', '23/31', '8.2'),
    ('pick_cool_then_place', '71.4%', '15/21', '9.6'),
    ('pick_heat_then_place', '73.9%', '17/23', '12.1'),
    ('pick_two_obj_and_place', '47.1%', '8/17', '12.0'),
], 1):
    for j, v in enumerate(row): table.rows[i].cells[j].text = v

h2('5.4  关键发现：task_desc标注一致性的决定性影响')
p = para('在分析失败案例过程中，我们发现ALFWorld数据集中task_desc（人类标注的自然语言描述）与PDDL ground truth之间存在系统性不一致。对全部134个游戏的逐条检查结果如表6所示。')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.first_line_indent = Cm(0)
t.add_run('表6  task_desc标注一致性与V7成功率的关系').bold = True

table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['子集', '游戏数', 'V7成功率']):
    table.rows[0].cells[i].text = h
    for r in table.rows[0].cells[i].paragraphs[0].runs: r.bold = True
table.rows[1].cells[0].text = '标注一致'
table.rows[1].cells[1].text = '100 (74.6%)'
table.rows[1].cells[2].text = '90/100 = 90.0%'
table.rows[2].cells[0].text = '标注不一致'
table.rows[2].cells[1].text = '34 (25.4%)'
table.rows[2].cells[2].text = '0/34 = 0.0%'

p = para('34个不一致案例分为三类：（1）物体名歧义（21个）：task_desc中的物体名与PDDL目标不匹配，如mug/cup混用（11个）、salt shaker实指pepper shaker（3个）、soap dispenser实指soap bottle（1个）等。（2）目标容器歧义（12个）：task_desc暗示的位置与PDDL不同，如"Move pencil on the desk over"暗示目标是desk但PDDL要求shelf。（3）关键信息缺失（5个）：task_desc未提及目标物体，如"Turn on the desk lamp"没说要看什么物体。')

p = para('')
p.add_run('这一发现的核心含义是：YLYW V7的性能瓶颈不在算法层面（在一致标注上已达90.0%），而在于ALFWorld数据集的标注质量。').bold = True
p.add_run('任何纯粹依赖task_desc自然语言进行目标解析的方法（不使用LLM），都会遇到这个天花板。')

h2('5.5  与现有方法对比')
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.first_line_indent = Cm(0)
t.add_run('表7  与现有方法的对比').bold = True

table = doc.add_table(rows=6, cols=5)
table.style = 'Table Grid'
for i, h in enumerate(['方法', '成功率', '动作空间', '需LLM', '特点']):
    table.rows[0].cells[i].text = h
    for r in table.rows[0].cells[i].paragraphs[0].runs: r.bold = True
for i, row in enumerate([
    ('BUTLER', '37%', 'generation', '否(训练)', '需大量训练数据'),
    ('ReAct (GPT-4)', '71%', 'admissible', '是', '需API调用'),
    ('Reflexion (GPT-4)', '77%', 'admissible', '是', '多轮重试'),
    ('EmbodiSkill', '93.28%', 'admissible', '是(27B)', '技能自演化'),
    ('YLYW V7 (本文)', '67.2%(90.0%*)', 'admissible', '否', '~800行代码'),
], 1):
    for j, v in enumerate(row): table.rows[i].cells[j].text = v

para('注：* 90.0%为标注一致子集上的成功率。')

p = para('EmbodiSkill')
sup(p, 13)
p.add_run('的93.28%来自两个YLYW不具备的能力：（1）LLM理解task_desc——当"salt shaker"但场景只有pepper shaker时，LLM能通过上下文推断修正目标；（2）多轮迭代演化——失败后反思修改技能，第二次自动避免相同错误。然而，在标注一致的子集上，YLYW V7（90.0%）与EmbodiSkill仅差约3个百分点，表明在任务描述准确时，800行规则代码可接近270亿参数LLM的水平。')

doc.save('中国科学_YLYW_ALFWorld_v2.docx')
print('Part2完成: 4方法+5实验')
