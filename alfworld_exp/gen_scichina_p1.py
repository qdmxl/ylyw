#!/usr/bin/env python3
"""中国科学:信息科学 论文 Part1: 标题+摘要+引言+相关工作+问题定义"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

def sup(p, n):
    r = p.add_run(f'[{n}]')
    r.font.superscript = True
    r.font.size = Pt(8)

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.name = '黑体'

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.name = '黑体'

def para(text, indent=True):
    p = doc.add_paragraph()
    if not indent: p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(text)
    return p

# ====== 标题 ======
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.first_line_indent = Cm(0)
run = t.add_run('基于YLYW先验知识的零样本具身决策方法\n及其在ALFWorld基准上的验证')
run.bold = True
run.font.size = Pt(16)

# 作者
a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
a.paragraph_format.first_line_indent = Cm(0)
a.add_run('马兴录').bold = True
a.add_run('1*  ')
a.add_run('张国安').bold = True
a.add_run('1  ')
a.add_run('李金函').bold = True
a.add_run('1  ')
a.add_run('于敬涛').bold = True
a.add_run('1  ')
a.add_run('李望').bold = True
a.add_run('1  ')
a.add_run('马圣洁').bold = True
a.add_run('1')

a2 = doc.add_paragraph()
a2.alignment = WD_ALIGN_PARAGRAPH.CENTER
a2.paragraph_format.first_line_indent = Cm(0)
a2.add_run('1. 青岛科技大学 信息科学技术学院, 青岛 266061\n* 通讯作者. E-mail: maxinglu@qust.edu.cn')

doc.add_paragraph()

# ====== 中文摘要 ======
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.add_run('摘要  ').bold = True
p.add_run('YLYW（易理研物）是一种基于《易经》先验符号知识的联邦式神经符号决策框架，已在物理域零样本抓取决策（92.7%合理率）、多智能体协调、知几学习等领域得到验证。本文将YLYW首次扩展至具身智能导航与操作领域，在ALFWorld官方TextWorld仿真器的valid_unseen测试集（134个任务）上进行验证。本文设计了一种admissible-commands信号驱动的层次化状态机决策Agent，仅依赖任务描述的自然语言解析和YLYW常识先验矩阵进行决策，不使用任何大语言模型（LLM）或API调用。实验结果表明，该方法在134个任务上达到67.2%的整体成功率。更重要的是，本文揭示了ALFWorld数据集中25.4%的任务（34/134）存在task_desc自然语言标注与PDDL ground truth不一致的问题。在标注一致的100个任务上，YLYW达到90.0%的成功率，接近使用270亿参数LLM的EmbodiSkill方法（93.28%），两者仅差约3个百分点——而YLYW仅需约800行Python代码，无需GPU。本文结果从跨域角度验证了YLYW"知几学习"范式的有效性，同时为ALFWorld基准的标注质量改进提供了实证依据。')

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.add_run('关键词  ').bold = True
p.add_run('YLYW; 易理先验知识; 具身智能; ALFWorld; 零样本决策; admissible commands; 标注一致性')

doc.add_paragraph()

# ====== 英文摘要 ======
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.add_run('Abstract  ').bold = True
p.add_run('YLYW (Yi-Li-Yan-Wu) is a federated neuro-symbolic decision-making framework based on I Ching prior symbolic knowledge, previously validated in zero-shot grasping (92.7% rationality). This paper extends YLYW to embodied navigation and manipulation, evaluating it on ALFWorld\'s valid_unseen benchmark (134 tasks). We design an admissible-commands-driven hierarchical state machine agent using only natural language parsing and YLYW commonsense priors, without any LLM or API calls. The agent achieves 67.2% overall success. Crucially, we discover that 25.4% of ALFWorld tasks (34/134) contain inconsistencies between task_desc annotations and PDDL ground truth. On the 100 consistently-annotated tasks, YLYW achieves 90.0%, approaching EmbodiSkill\'s 93.28% (which requires a 27B-parameter LLM), with only ~3pp gap—using merely ~800 lines of Python without GPU.')

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.add_run('Keywords  ').bold = True
p.add_run('YLYW; I Ching prior knowledge; embodied AI; ALFWorld; zero-shot decision-making; annotation consistency')

doc.add_paragraph()

# ====== 1 引言 ======
h1('1  引言')
p = para('具身智能（Embodied AI）要求智能体在物理或仿真环境中感知、推理并执行动作序列以完成任务。与纯语言或视觉任务不同，具身智能决策面临感知-动作闭环、长程规划和组合爆炸三重挑战。当前主流方法依赖大语言模型（LLM）进行决策：ReAct')
sup(p, 9)
p.add_run('利用GPT-4进行推理-行动循环（71%成功率），Reflexion')
sup(p, 10)
p.add_run('在此基础上增加多轮反思机制（77%），EmbodiSkill')
sup(p, 13)
p.add_run('通过技能自演化使270亿参数的Qwen3.5达到93.28%。然而，这些方法均依赖大规模LLM推理，部署成本高、决策过程不可解释、输出具有随机性。')

p = para('是否存在一条不依赖LLM、仅凭结构化先验知识即可实现高效决策的路径？YLYW（易理研物）系列研究探索了这一方向。本团队此前已提出YLYW核心架构')
sup(p, 1)
p.add_run('（300物体零样本抓取92.7%合理率）、层次化嵌套架构')
sup(p, 2)
p.add_run('、知几学习范式')
sup(p, 3)
p.add_run('（K = K_prior ⊕ K_calibration）、汉语言处理范式')
sup(p, 4)
p.add_run('及反幻觉系统')
sup(p, 5)
p.add_run('。本文是该系列的第六项工作，首次将YLYW从物理对象决策扩展至具身导航与操作决策。')

p = para('本文在ALFWorld')
sup(p, 6)
p.add_run('官方TextWorld')
sup(p, 8)
p.add_run('仿真器上评估YLYW方法。ALFWorld的valid_unseen测试集包含134个任务，覆盖6种家务任务类型和4种家庭场景。本文的主要贡献包括：')

contribs = [
    '提出admissible-commands信号驱动的层次化状态机决策架构，以约800行Python代码实现不依赖LLM的零样本具身决策Agent。',
    '设计YLYW物体-位置常识先验矩阵和NL任务描述解析器，在不使用PDDL参数的条件下实现67.2%的整体成功率。',
    '揭示ALFWorld数据集中25.4%任务存在task_desc与PDDL标注不一致的问题：在标注一致的子集上YLYW达到90.0%，接近LLM方法（93.28%）。',
    '发现并修复了ALFWorld官方TextWorld环境的游戏加载BUG（方案B: per-game env），为社区提供了可复现的评估基础。',
    '通过与EmbodiSkill的对比分析，论证了先验知识驱动 vs LLM驱动两条技术路线的优劣边界。',
]
for c in contribs:
    doc.add_paragraph(c, style='List Number')

doc.add_paragraph()

# ====== 2 相关工作 ======
h1('2  相关工作')

h2('2.1  ALFWorld与具身智能基准')
p = para('ALFWorld')
sup(p, 6)
p.add_run('是将ALFRED视觉导航基准')
sup(p, 7)
p.add_run('与TextWorld文本引擎')
sup(p, 8)
p.add_run('结合的具身智能评估平台，将家庭场景中的导航和操作任务抽象为文本交互游戏。智能体接收文本观测，执行文本动作，目标是完成如"把干净的盘子放在台面上"等家务任务。valid_unseen测试集包含134个可解任务，是评估零样本具身决策能力的标准基准。')

h2('2.2  LLM驱动的具身决策方法')
p = para('ReAct')
sup(p, 9)
p.add_run('利用GPT-4进行"推理+行动"交替循环，成功率71%。Reflexion')
sup(p, 10)
p.add_run('增加失败后的反思机制，允许多次重试，成功率77%。EmbodiSkill')
sup(p, 13)
p.add_run('提出技能感知反思（Skill-Aware Reflection），通过区分技能缺陷与执行失误实现定向技能修正，使Qwen3.5-27B达到93.28%。SayCan')
sup(p, 11)
p.add_run('和PaLM-E')
sup(p, 12)
p.add_run('等方法则将LLM与机器人操作原语结合。这些方法的共同特征是依赖大规模LLM推理，单次推理消耗数千token。')

h2('2.3  YLYW先验知识框架')
p = para('YLYW采用联邦式神经符号架构')
sup(p, 1)
p.add_run('，由三层组成：L1八卦隶属度层（物理特征→8维模糊向量）、L2六爻编码层（聚合为6维爻向量）、L3六十四卦规则层（余弦匹配→策略输出）。知几学习范式')
sup(p, 3)
p.add_run('主张具身智能体应携带先验知识"出生"，学习的功能是校准而非从零发现规律（K = K_prior ⊕ K_calibration）。本文将这一思想迁移至ALFWorld：Agent携带物体-位置先验矩阵和任务类型模板出生，在未见过的场景中零样本决策。')

doc.add_paragraph()

# ====== 3 问题定义 ======
h1('3  问题定义与评估设置')

h2('3.1  ALFWorld环境')
para('ALFWorld环境的核心组件包括：（1）观测空间——自然语言文本描述当前状态；（2）动作空间——环境每步返回admissible_commands列表，包含当前合法的所有动作；（3）PDDL后端——环境状态由规划域定义语言描述；（4）胜利条件——动作序列满足目标状态时返回won=True。')

h2('3.2  任务类型')
p = para('valid_unseen测试集包含6种任务类型，如表1所示。')

# 表1
t1_title = doc.add_paragraph()
t1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
t1_title.paragraph_format.first_line_indent = Cm(0)
t1_title.add_run('表1  ALFWorld valid_unseen任务类型分布').bold = True

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['任务类型', '数量', '描述']):
    table.rows[0].cells[i].text = h
    for r in table.rows[0].cells[i].paragraphs[0].runs: r.bold = True
for i, row in enumerate([
    ('look_at_obj_in_light', '18', '在灯光下查看物体'),
    ('pick_and_place_simple', '24', '拿起物体放到指定位置'),
    ('pick_clean_then_place', '31', '拿起→清洗→放置'),
    ('pick_heat_then_place', '23', '拿起→加热→放置'),
    ('pick_cool_then_place', '21', '拿起→冷却→放置'),
    ('pick_two_obj_and_place', '17', '拿两个物体分别放置'),
], 1):
    for j, v in enumerate(row): table.rows[i].cells[j].text = v

h2('3.3  评估设置')
para('本文采用admissible模式（从环境提供的合法动作列表中选择），与ReAct、Reflexion、EmbodiSkill相同。最大步数限制50步/游戏。测试集覆盖4个FloorPlan场景。')

h2('3.4  动作空间模式说明')
p = para('ALFWorld支持两种动作空间：admissible模式（从列表选择）和generation模式（逐词生成）。BUTLER')
sup(p, 6)
p.add_run('使用generation模式（37%），ReAct/Reflexion/EmbodiSkill和本文均使用admissible模式。关于admissible模式信息量的讨论见第6.1节。')

doc.save('中国科学_信息科学_YLYW_ALFWorld.docx')
print('Part1完成: 标题+摘要+1引言+2相关工作+3问题定义')
