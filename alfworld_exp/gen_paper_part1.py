#!/usr/bin/env python3
"""生成YLYW ALFWorld完整论文 - Part 1: 文档结构和前半部分"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 默认样式
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

def add_ref(paragraph, ref_num):
    """添加上标引用"""
    run = paragraph.add_run(f'[{ref_num}]')
    run.font.superscript = True
    run.font.size = Pt(8)

def add_text_with_ref(paragraph, text, ref_num):
    """添加带引用的文本"""
    paragraph.add_run(text)
    add_ref(paragraph, ref_num)

# ====== 标题 ======
title = doc.add_heading('', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('YLYW在ALFWorld基准测试中的应用：\n一种基于易理先验知识的零样本具身决策方法')
run.font.size = Pt(18)

# 作者信息
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('马兴录，张国安，李金函，于敬涛，李望，马圣洁*').bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('青岛科技大学 信息科学技术学院，山东 青岛 266061\n* 通讯作者: 马圣洁')

doc.add_paragraph()

# ====== 摘要 ======
doc.add_heading('摘要', level=1)
p = doc.add_paragraph()
p.add_run('YLYW（易理研物）是本团队提出的一种基于《易经》六十四卦先验符号知识的联邦式神经符号决策框架，已在物理域零样本抓取决策（92.7%合理率）')
add_ref(p, 1)
p.add_run('、层次化嵌套多智能体协调')
add_ref(p, 2)
p.add_run('、知几学习新范式')
add_ref(p, 3)
p.add_run('等多个领域得到验证。本文是YLYW系列的第六项工作，将该框架首次扩展至具身智能导航与操作领域。')

p = doc.add_paragraph()
p.add_run('本文在ALFWorld官方TextWorld仿真器的valid_unseen测试集（134个任务）上验证YLYW方法。通过将YLYW的核心思想——"先验知识驱动的信号辨识与层次化状态机"——迁移至文本型具身导航场景，设计了一种admissible-commands信号驱动的零样本决策Agent。该Agent仅依赖任务模板、PDDL参数和YLYW常识先验矩阵进行决策，')
p.add_run('不使用任何大语言模型(LLM)或API调用').bold = True
p.add_run('。')

p = doc.add_paragraph()
p.add_run('实验结果表明，YLYW Agent在134个任务上达到')
run = p.add_run('92.5%')
run.bold = True
p.add_run('的成功率（124/134），平均仅需13.0步完成任务。在相同的admissible动作空间评估条件下，显著超越了基于GPT-4的ReAct方法（71%）')
add_ref(p, 9)
p.add_run('和Reflexion方法（77%）')
add_ref(p, 10)
p.add_run('，以及训练型BUTLER基线（37%）')
add_ref(p, 6)
p.add_run('。其中look_at_obj_in_light、pick_and_place_simple和pick_heat_then_place三种任务类型达到100%成功率。本文结果从跨域角度验证了YLYW"知几学习"范式的有效性：当系统内建了关于世界变化规律的结构化先验知识时，在零样本条件下即可实现高效的具身决策。')

p = doc.add_paragraph()
p.add_run('关键词：').bold = True
p.add_run('YLYW；易理先验知识；具身智能；ALFWorld；零样本学习；状态机决策；admissible commands')

doc.add_paragraph()

# ====== 1. 引言 ======
doc.add_heading('1 引言', level=1)

doc.add_heading('1.1 具身智能导航决策的挑战', level=2)
p = doc.add_paragraph()
p.add_run('具身智能（Embodied AI）要求智能体在物理或仿真环境中感知、推理并执行动作序列以完成任务。与纯语言或视觉任务不同，具身智能决策面临三重挑战：（1）感知-动作闭环——智能体的每步决策依赖于前一步执行后的环境反馈；（2）长程规划——完成一个任务通常需要6-50步有序操作；（3）组合爆炸——环境中有数十个位置和物体，可能的动作序列呈指数增长。')

p = doc.add_paragraph()
p.add_run('当前主流方法主要依赖大语言模型(LLM)进行具身决策。PaLM-E')
add_ref(p, 13)
p.add_run('、SayCan')
add_ref(p, 11)
p.add_run('等系统利用LLM的语义理解能力进行任务分解和行动规划。然而，这些方法存在根本性问题：（1）依赖大规模API调用，单次推理成本高；（2）LLM的"幻觉"问题导致生成无效动作')
add_ref(p, 5)
p.add_run('；（3）决策过程是黑箱，不可解释；（4）输出具有随机性，相同输入可能产生不同结果。')

doc.add_heading('1.2 ALFWorld基准测试的意义', level=2)
p = doc.add_paragraph()
p.add_run('ALFWorld')
add_ref(p, 6)
p.add_run('是将ALFRED视觉导航基准')
add_ref(p, 7)
p.add_run('与TextWorld文本引擎')
add_ref(p, 8)
p.add_run('相结合的具身智能评估平台。它将家庭场景中的导航和操作任务抽象为文本交互游戏：智能体接收文本观测（如"You arrive at countertop 1. On the countertop 1, you see a plate 2, and a mug 1."），执行文本动作（如"take plate 2 from countertop 1"），目标是完成如"把干净的盘子放在台面上"等家务任务。')

p = doc.add_paragraph()
p.add_run('ALFWorld的valid_unseen测试集包含134个可解任务，覆盖6种任务类型和4种家庭场景（厨房、卧室、客厅、浴室），是评估具身智能决策能力的标准基准。该测试集的"unseen"属性保证了评估的零样本性质——智能体在测试时面对的是训练时从未见过的场景布局和物体配置。')

doc.add_heading('1.3 现有方法的局限', level=2)
p = doc.add_paragraph()
p.add_run('在ALFWorld上的代表性方法包括：')

p = doc.add_paragraph()
p.add_run('（1）ReAct')
add_ref(p, 9)
p.add_run('：利用GPT-4进行"推理+行动"交替循环。LLM先进行思维推理（Thought），再选择动作（Action），观察结果（Observation）后继续推理。成功率71%。主要局限：需要大量API调用，单个任务消耗数千token；LLM有时"幻觉"出不存在的动作。')

p = doc.add_paragraph()
p.add_run('（2）Reflexion')
add_ref(p, 10)
p.add_run('：在ReAct基础上增加"反思"机制——任务失败后，LLM反思失败原因并在下次尝试中改进。允许多次重试，成功率77%。主要局限：多次重试的评估模式在真实部署中不可接受；本质上是用更多计算换取更高成功率。')

p = doc.add_paragraph()
p.add_run('（3）BUTLER')
add_ref(p, 6)
p.add_run('：使用DAgger算法训练的seq2seq模型。在更困难的generation模式（逐词生成动作文本，不使用admissible列表）下评估，成功率37%。')

p = doc.add_paragraph()
p.add_run('这些方法的共同特征是：要么依赖昂贵的LLM推理（ReAct/Reflexion），要么需要大量训练数据（BUTLER）。是否存在一条不依赖LLM和训练数据、仅凭结构化先验知识即可实现高效决策的路径？')

doc.add_heading('1.4 YLYW系列研究的演进与本文定位', level=2)
p = doc.add_paragraph()
p.add_run('YLYW（易理研物）系列研究探索了一条知识驱动的具身智能新路径。本团队此前的工作已建立了完整的理论和实验基础：')

p = doc.add_paragraph()
p.add_run('• YLYW核心架构')
add_ref(p, 1)
p.add_run('：提出以八卦为模糊基元、六十四卦为先验模板、爻位关系为参数修正算子的三层联邦式架构。300物体零样本抓取决策中达到92.7%合理率，论证了"先验知识是零样本能力的唯一来源"。')

p = doc.add_paragraph()
p.add_run('• 层次化嵌套架构')
add_ref(p, 2)
p.add_run('：将YLYW扩展至多智能体分布式系统，提出"卦象意图"通讯协议和双模型并行博弈（主模型+安全八卦，后者拥有一票否决权）。')

p = doc.add_paragraph()
p.add_run('• 知几学习范式')
add_ref(p, 3)
p.add_run('：正式提出与强化学习本质不同的学习范式——"知几"（先验征兆辨识）vs "RL"（事后奖惩驱动）。核心公式K = K_prior ⊕ K_calibration表明：学习的真正功能不是"从无到有发现规律"，而是在先验框架上进行适配和校准。')

p = doc.add_paragraph()
p.add_run('• 反幻觉系统')
add_ref(p, 5)
p.add_run('：将YLYW迁移为LLM的独立审查引擎，三层审查（规则一致性/物理合规性/价值对齐性）达到83.3%幻觉检出率。')

p = doc.add_paragraph()
p.add_run('本文是该系列的第六项工作，首次将YLYW从"物理对象决策"（如抓取策略）扩展至"具身导航与操作决策"（如在家庭场景中寻找、拾取、操作、放置物体）。ALFWorld提供了一个理想的跨域验证平台：任务结构化程度高、状态空间有限、评估标准明确。')

doc.add_heading('1.5 本文贡献', level=2)
p = doc.add_paragraph()
p.add_run('本文的主要贡献包括：')
contributions = [
    '首次将YLYW先验知识框架应用于具身导航与操作决策领域，实现92.5%成功率，验证了YLYW的跨域泛化能力。',
    '提出admissible-commands信号驱动的层次化状态机决策架构，将环境反馈的合法动作列表作为实时状态信号，驱动阶段推进和回退。',
    '设计YLYW物体-位置常识先验矩阵，编码了30+物体类型在各类位置的先验出现概率，引导高效探索。',
    '发现并修复了ALFWorld官方TextWorld环境的游戏加载BUG（方案B: per-game env），为社区提供了可复现的正确评估基础。',
    '在相同评估条件下，不使用任何LLM/API，显著超越GPT-4驱动的ReAct(71%)和Reflexion(77%)，论证了确定性规则系统在有限状态空间中的优越性。',
]
for c in contributions:
    doc.add_paragraph(c, style='List Number')

doc.add_paragraph()

# ====== 2. YLYW理论基础 ======
doc.add_heading('2 YLYW理论基础与先验工作', level=1)

doc.add_heading('2.1 YLYW核心架构回顾', level=2)
p = doc.add_paragraph()
p.add_run('YLYW采用联邦式神经符号架构，由三层组成')
add_ref(p, 1)
p.add_run('：')

p = doc.add_paragraph()
p.add_run('L1 八卦隶属度层').bold = True
p.add_run('：以八卦（乾坤震巽坎离艮兑）为连续模糊基元。物理特征通过高斯核函数映射为对各卦的隶属度μ∈[0,1]，解决了传统符号系统将连续物理量二值化的"符号接地"困境。一个物体不是"属于或不属于"某卦，而是以不同程度同时关联多个卦象。')

p = doc.add_paragraph()
p.add_run('L2 六爻编码层').bold = True
p.add_run('：将8维隶属度经加权公式聚合为6维爻值向量y∈[0,1]⁶。每爻≥0.5为阳(—)、<0.5为阴(--)。六爻从初爻到上爻分别对应不同物理/语义维度（如初爻=基础稳定性、五爻=优先级）。')

p = doc.add_paragraph()
p.add_run('L3 卦象匹配与爻位关系层').bold = True
p.add_run('：6维爻向量通过余弦相似度在64卦模板中匹配最佳卦象以确定策略类型。同时，"乘承比应当位得中"五种爻位关系被形式化为可计算算子，用于修正执行参数。实现了"卦定策略类型、爻定执行参数"的分层决策体系。')

p = doc.add_paragraph()
p.add_run('在300物体零样本基线测试中，该架构达到92.7%策略合理率和0%严重错误率，纯先验推理仅需1.7ms/物体。三维消融实验验证了易理规则(+33.6%)、三层架构(+12.7%)和连续模糊隶属度(+23.0%)的独立贡献')
add_ref(p, 1)
p.add_run('。')

doc.add_heading('2.2 知几学习范式', level=2)
p = doc.add_paragraph()
p.add_run('知几学习')
add_ref(p, 3)
p.add_run('是本团队提出的一种与强化学习本质不同的学习范式。其哲学基础源于《易经·系辞下》："知几其神乎！几者，动之微，吉之先见者也。君子见几而作，不俟终日。"')
add_ref(p, 14)

p = doc.add_paragraph()
p.add_run('核心主张是：具身智能体的学习不必从一张白纸开始。关于物理世界变化规律的基本知识（如"倾斜意味着可能倾倒"、"盘子通常在厨房台面上"）应作为先验内建于系统中。学习的真正功能不是"从无到有发现规律"，而是在先验框架上校准参数。形式化为：')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('K = K').italic = True
run = p.add_run('prior')
run.font.subscript = True
run.font.italic = True
p.add_run(' ⊕ K')
run = p.add_run('calibration')
run.font.subscript = True
run.font.italic = True

p = doc.add_paragraph()
p.add_run('其中K_prior是先验知识（YLYW的64卦规则库），K_calibration是少量经验校准。本文在ALFWorld上的应用正是知几学习的典型案例：Agent携带物体-位置先验知识和任务类型模板"出生"，无需任何训练即可在未见过的场景中高效决策。')

doc.add_heading('2.3 层次化嵌套架构的启示', level=2)
p = doc.add_paragraph()
p.add_run('层次化嵌套YLYW')
add_ref(p, 2)
p.add_run('提出了"卦象意图"通讯协议：层级间以抽象卦象（而非具体指令）作为通讯载体，接收方根据自身约束自主解读和执行。本文的ALFWorld Agent借鉴了这一设计：任务类型（如pick_clean_then_place）作为"宏观卦象意图"确定整体计划模板，而具体的物体-位置匹配则由Agent在微观层自主决策。')

doc.add_heading('2.4 从反幻觉系统继承的设计', level=2)
p = doc.add_paragraph()
p.add_run('反幻觉系统')
add_ref(p, 5)
p.add_run('中YLYW作为LLM的独立审查引擎的设计思想，在本文中被反向应用：本文不使用LLM生成动作，而是用YLYW规则系统直接替代LLM进行决策。这种"规则替代LLM"的设计在ALFWorld的有限状态空间中被证明更加可靠——消除了LLM的幻觉风险，同时保留了确定性和可解释性。')

doc.add_paragraph()

# ====== 3. 问题定义 ======
doc.add_heading('3 ALFWorld环境与问题定义', level=1)

doc.add_heading('3.1 ALFWorld环境描述', level=2)
p = doc.add_paragraph()
p.add_run('ALFWorld')
add_ref(p, 6)
p.add_run('将ALFRED视觉导航任务')
add_ref(p, 7)
p.add_run('抽象为TextWorld')
add_ref(p, 8)
p.add_run('文本交互游戏。环境核心组件包括：')

items = [
    '观测空间：自然语言文本描述当前状态（如"You are in the middle of a room. Looking quickly around you, you see a cabinet 6, a countertop 1..."）。',
    '动作空间：环境每步返回admissible_commands列表，包含当前合法的所有动作（如"go to countertop 1", "take plate 2 from countertop 2", "clean plate 2 with sinkbasin 1"等）。',
    'PDDL后端：环境状态由PDDL（Planning Domain Definition Language）描述，任务目标是满足PDDL goal condition。',
    '胜利条件：当Agent执行的动作序列使环境到达目标状态时，环境返回won=True。',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 任务类型', level=2)
p = doc.add_paragraph()
p.add_run('valid_unseen测试集包含6种任务类型，如表1所示。')

# 表1
table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
headers = ['任务类型', '数量', '描述', '典型步数']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for run in table.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
data = [
    ('look_at_obj_in_light', '18', '拿起物体，在灯光下查看', '4-5'),
    ('pick_and_place_simple', '24', '拿起物体，放到指定容器', '4-6'),
    ('pick_clean_then_place', '31', '拿起→清洗→放置', '6-7'),
    ('pick_heat_then_place', '23', '拿起→加热→放置', '6-7'),
    ('pick_cool_then_place', '21', '拿起→冷却→放置', '6-7'),
    ('pick_two_obj_and_place', '17', '拿两个物体分别放置', '8-11'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val
p = doc.add_paragraph()
p.add_run('表1 ALFWorld valid_unseen测试集任务类型分布').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('3.3 评估设置', level=2)
p = doc.add_paragraph()
p.add_run('本文采用以下评估设置：测试集为valid_unseen（134个solvable games）；动作空间为admissible模式（从环境提供的合法列表中选择）；最大步数限制50步/游戏；场景覆盖4个FloorPlan（厨房FloorPlan10: 77个，卧室FloorPlan308: 27个，浴室FloorPlan424: 19个，客厅FloorPlan219: 11个）。')

doc.add_heading('3.4 动作空间模式说明', level=2)
p = doc.add_paragraph()
p.add_run('ALFWorld支持两种动作空间模式：（a）admissible模式——Agent从环境每步提供的合法动作列表中选择；（b）generation模式——Agent需逐词生成完整动作文本字符串。当前最具影响力的基线方法ReAct')
add_ref(p, 9)
p.add_run('和Reflexion')
add_ref(p, 10)
p.add_run('均使用admissible模式评估。BUTLER')
add_ref(p, 6)
p.add_run('使用更困难的generation模式。本文采用admissible模式，确保与ReAct/Reflexion在相同条件下对比。关于admissible模式信息量的详细讨论见第6.1节。')

# 保存中间结果
doc.save('YLYW_ALFWorld_完整论文_final.docx')
print('Part 1 完成: 标题+摘要+1引言+2理论基础+3问题定义')
