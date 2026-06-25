#!/usr/bin/env python3
"""修复论文格式：首行缩进 + 生成并插入图片"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import os

plt.rcParams['font.sans-serif'] = ['WenQuanYi Micro Hei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

output_dir = '/home/lijinhan/MXL/科研/ylyw/alfworld_exp/paper_figures'
os.makedirs(output_dir, exist_ok=True)

# ====== 生成图片 ======

# 图1: 系统架构图
fig, ax = plt.subplots(1, 1, figsize=(10, 6))
ax.set_xlim(0, 10)
ax.set_ylim(0, 7)
ax.axis('off')
ax.set_title('图1  YLYW ALFWorld Agent V6 系统架构', fontsize=14, fontweight='bold', pad=20)

# 三层方框
boxes = [
    (0.5, 5.5, 9, 1.2, '感知层: Admissible Commands 信号提取', '#E3F2FD'),
    (0.5, 3.3, 9, 1.8, '决策层: YLYW先验知识 + 层次化状态机', '#FFF3E0'),
    (0.5, 1.5, 9, 1.2, '执行层: 动作选择与输出', '#E8F5E9'),
]
for x, y, w, h, label, color in boxes:
    rect = plt.Rectangle((x, y), w, h, facecolor=color, edgecolor='black', linewidth=1.5)
    ax.add_patch(rect)
    ax.text(x + w/2, y + h/2, label, ha='center', va='center', fontsize=11, fontweight='bold')

# 内部模块
modules_mid = [
    (1.2, 3.6, '目标提取\n(PDDL参数)'),
    (3.5, 3.6, 'YLYW先验\n(物体-位置矩阵)'),
    (6.0, 3.6, '状态机\n(子目标模板)'),
    (8.2, 3.6, 'V6增强\n(open/记忆)'),
]
for x, y, label in modules_mid:
    rect = plt.Rectangle((x, y), 1.8, 1.2, facecolor='white', edgecolor='#666', linewidth=1, linestyle='--')
    ax.add_patch(rect)
    ax.text(x + 0.9, y + 0.6, label, ha='center', va='center', fontsize=8)

# 箭头
ax.annotate('', xy=(5, 5.5), xytext=(5, 5.1), arrowprops=dict(arrowstyle='->', lw=2))
ax.annotate('', xy=(5, 3.3), xytext=(5, 2.7), arrowprops=dict(arrowstyle='->', lw=2))

# 顶部环境
ax.text(5, 6.7, 'ALFWorld TextWorld 环境', ha='center', va='center', fontsize=12, 
        bbox=dict(boxstyle='round,pad=0.3', facecolor='#BBDEFB', edgecolor='black'))
ax.annotate('', xy=(5, 6.4), xytext=(5, 5.8), arrowprops=dict(arrowstyle='->', lw=2, color='blue'))
ax.text(3.5, 6.1, 'obs + admissible', ha='center', fontsize=9, color='blue')

# 底部输出
ax.text(5, 0.8, 'action → 环境', ha='center', va='center', fontsize=11,
        bbox=dict(boxstyle='round,pad=0.3', facecolor='#C8E6C9', edgecolor='black'))
ax.annotate('', xy=(5, 1.5), xytext=(5, 1.1), arrowprops=dict(arrowstyle='->', lw=2, color='green'))

plt.tight_layout()
plt.savefig(f'{output_dir}/fig1_architecture.png', dpi=150, bbox_inches='tight')
plt.close()
print('图1 系统架构 ✓')

# 图2: 成功率对比柱状图
fig, ax = plt.subplots(1, 1, figsize=(8, 5))
methods = ['BUTLER\n(generation)', 'ReAct\n(GPT-4)', 'Reflexion\n(GPT-4)', 'YLYW V5', 'YLYW V6\n(本文)']
rates = [37, 71, 77, 64.2, 92.5]
colors = ['#90A4AE', '#FFA726', '#FF7043', '#66BB6A', '#E53935']
bars = ax.bar(methods, rates, color=colors, edgecolor='black', linewidth=0.8)
ax.set_ylabel('成功率 (%)', fontsize=12)
ax.set_title('图2  ALFWorld valid_unseen 成功率对比', fontsize=13, fontweight='bold')
ax.set_ylim(0, 105)
ax.axhline(y=92.5, color='red', linestyle='--', alpha=0.5)
for bar, rate in zip(bars, rates):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1.5, f'{rate}%', 
            ha='center', va='bottom', fontsize=11, fontweight='bold')
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
plt.tight_layout()
plt.savefig(f'{output_dir}/fig2_comparison.png', dpi=150, bbox_inches='tight')
plt.close()
print('图2 方法对比 ✓')

# 图3: V6按任务类型成功率
fig, ax = plt.subplots(1, 1, figsize=(9, 5))
types = ['look_at\n_in_light', 'pick_and\n_place', 'pick_heat\n_place', 'pick_clean\n_place', 'pick_cool\n_place', 'pick_two\n_obj']
v5_rates = [100, 50, 87, 58.1, 71.4, 17.6]
v6_rates = [100, 100, 100, 96.8, 90.5, 58.8]
x = np.arange(len(types))
w = 0.35
bars1 = ax.bar(x - w/2, v5_rates, w, label='V5', color='#90CAF9', edgecolor='black', linewidth=0.5)
bars2 = ax.bar(x + w/2, v6_rates, w, label='V6', color='#EF5350', edgecolor='black', linewidth=0.5)
ax.set_ylabel('成功率 (%)', fontsize=12)
ax.set_title('图3  V5 vs V6 按任务类型成功率', fontsize=13, fontweight='bold')
ax.set_xticks(x)
ax.set_xticklabels(types, fontsize=9)
ax.set_ylim(0, 115)
ax.legend(fontsize=11)
for bar in bars2:
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1, f'{bar.get_height():.0f}%', 
            ha='center', va='bottom', fontsize=9)
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
plt.tight_layout()
plt.savefig(f'{output_dir}/fig3_by_type.png', dpi=150, bbox_inches='tight')
plt.close()
print('图3 按类型对比 ✓')

# 图4: 状态机流程图（pick_clean示例）
fig, ax = plt.subplots(1, 1, figsize=(12, 3.5))
ax.set_xlim(0, 12)
ax.set_ylim(0, 3.5)
ax.axis('off')
ax.set_title('图4  pick_clean_then_place_in_recep 状态机示例', fontsize=12, fontweight='bold', pad=10)

states = [
    (0.3, 1.5, 'find\nobject'),
    (2.2, 1.5, 'take\nobject'),
    (4.1, 1.5, 'find\ntool'),
    (6.0, 1.5, 'use\ntool'),
    (7.9, 1.5, 'find\nrecep'),
    (9.8, 1.5, 'put\nobject'),
]
for i, (x, y, label) in enumerate(states):
    circle = plt.Circle((x+0.7, y+0.5), 0.55, facecolor='#FFF9C4' if i < 6 else '#C8E6C9', 
                        edgecolor='black', linewidth=1.5)
    ax.add_patch(circle)
    ax.text(x+0.7, y+0.5, label, ha='center', va='center', fontsize=8, fontweight='bold')
    if i < len(states) - 1:
        ax.annotate('', xy=(states[i+1][0]+0.15, y+0.5), xytext=(x+1.25, y+0.5),
                   arrowprops=dict(arrowstyle='->', lw=1.5, color='#333'))

# 回退箭头
ax.annotate('', xy=(0.5, 0.8), xytext=(2.5, 0.8),
           arrowprops=dict(arrowstyle='->', lw=1, color='red', linestyle='dashed'))
ax.text(1.5, 0.5, '回退', ha='center', fontsize=8, color='red')

# 成功标记
ax.text(11.2, 2.0, '✅ WON', ha='center', va='center', fontsize=14, color='green', fontweight='bold')
ax.annotate('', xy=(11.0, 1.8), xytext=(10.6, 2.0),
           arrowprops=dict(arrowstyle='->', lw=1.5, color='green'))

# 底部信号说明
ax.text(0.3, 0.1, '推进信号: obs中出现plate → take命令中有plate → 到达sinkbasin → clean出现 → 到达countertop → move出现', 
        fontsize=8, color='#555')

plt.tight_layout()
plt.savefig(f'{output_dir}/fig4_state_machine.png', dpi=150, bbox_inches='tight')
plt.close()
print('图4 状态机流程 ✓')

# 图5: 步数分布箱线图
fig, ax = plt.subplots(1, 1, figsize=(8, 4.5))
import json
with open('ylyw_agent_v6_results.json') as f:
    data = json.load(f)
wins = [r for r in data['results'] if r['won']]
type_steps = {}
for r in wins:
    t = r['task_type_real'].replace('pick_', 'p_').replace('_then_place_in_recep', '').replace('_simple','').replace('_and_place','&place').replace('look_at_obj_in_light','look_light')
    type_steps.setdefault(t, []).append(r['steps'])
labels = sorted(type_steps.keys())
box_data = [type_steps[l] for l in labels]
bp = ax.boxplot(box_data, labels=labels, patch_artist=True)
colors_box = ['#BBDEFB', '#C8E6C9', '#FFF9C4', '#FFCCBC', '#E1BEE7', '#B2DFDB']
for patch, color in zip(bp['boxes'], colors_box):
    patch.set_facecolor(color)
ax.set_ylabel('完成步数', fontsize=11)
ax.set_title('图5  各任务类型成功案例的步数分布', fontsize=12, fontweight='bold')
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
plt.tight_layout()
plt.savefig(f'{output_dir}/fig5_steps_box.png', dpi=150, bbox_inches='tight')
plt.close()
print('图5 步数分布 ✓')

print(f'\n所有图片已生成到 {output_dir}/')

# ====== 修复论文格式 ======
doc = Document('YLYW_ALFWorld_完整论文_final.docx')

# 1. 设置正文段落首行缩进2字符
for p in doc.paragraphs:
    if p.style.name == 'Normal' and p.text.strip():
        p.paragraph_format.first_line_indent = Cm(0.74)  # 约2个中文字符
    # List样式也缩进
    if 'List' in p.style.name and p.text.strip():
        p.paragraph_format.first_line_indent = Cm(0.74)

# 2. 在合适的位置插入图片
# 找到各章节的关键段落位置
def find_paragraph_containing(doc, text):
    for i, p in enumerate(doc.paragraphs):
        if text in p.text:
            return i
    return -1

# 图1: 在4.1系统架构总览后面插入
idx = find_paragraph_containing(doc, '整体决策流程为')
if idx > 0:
    # 在该段落后插入图片
    p = doc.paragraphs[idx]
    # 添加新段落放图片
    run = p.add_run('\n')
    # 我们需要在段落之后添加，用add_picture方式
    pass

# 用更简单的方法：在文档末尾不行，用paragraph的insert方法也不方便
# 改用：在关键段落的run中添加图片
# python-docx的add_picture只能在文档末尾或InlineShape方式

# 采用：在特定段落后面新建段落插入图片
from docx.shared import Inches

# 重新组织：找到关键位置，在其后添加图片段落
# 方法：遍历找位置，记录，然后从后往前插入（避免索引偏移）

insert_points = []

# 图1位置：4.1 系统架构总览 的第二段之后
idx = find_paragraph_containing(doc, '在运行时通过admissible_commands信号')
if idx > 0:
    insert_points.append((idx, f'{output_dir}/fig1_architecture.png', '图1  YLYW ALFWorld Agent V6 系统架构'))

# 图4位置：4.5 决策层 子目标模板后面
idx = find_paragraph_containing(doc, '表3 六种任务类型的子目标模板')
if idx > 0:
    insert_points.append((idx, f'{output_dir}/fig4_state_machine.png', '图4  pick_clean_then_place_in_recep 状态机示例'))

# 图2位置：5.5 与文献对比 表8后面
idx = find_paragraph_containing(doc, '表8 与文献方法的对比')
if idx > 0:
    insert_points.append((idx, f'{output_dir}/fig2_comparison.png', '图2  ALFWorld valid_unseen 各方法成功率对比'))

# 图3位置：5.2 主实验结果 表5后面  
idx = find_paragraph_containing(doc, '表5 V6按任务类型的详细结果')
if idx > 0:
    insert_points.append((idx, f'{output_dir}/fig3_by_type.png', '图3  V5 vs V6 按任务类型成功率对比'))

# 图5位置：5.3 消融分析前面
idx = find_paragraph_containing(doc, '表6 V6按场景的成功率')
if idx > 0:
    insert_points.append((idx, f'{output_dir}/fig5_steps_box.png', '图5  各任务类型成功案例的步数分布'))

# 从后往前插入图片（避免索引变化）
insert_points.sort(key=lambda x: x[0], reverse=True)
for idx, img_path, caption in insert_points:
    # 在idx段落之后插入
    target_para = doc.paragraphs[idx]
    # 在target_para后面添加新段落
    new_para = doc.add_paragraph()  # 先加到末尾
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    # 移动到正确位置：python-docx不支持直接移动，所以用element操作
    target_el = target_para._element
    new_el = new_para._element
    target_el.addnext(new_el)

doc.save('YLYW_ALFWorld_完整论文_final.docx')
print('\n✅ 论文格式修复完成：首行缩进 + 5张图片插入')
