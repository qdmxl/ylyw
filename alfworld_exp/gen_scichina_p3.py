"""Part3: 6讨论 + 7结论 + 参考文献"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('中国科学_信息科学_YLYW_ALFWorld.docx')

def sup(p, n):
    r = p.add_run(f'[{n}]')
    r.font.superscript = True
    r.font.size = Pt(8)

def h1(t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: r.font.name = '黑体'

def h2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = '黑体'

def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p

# ====== 6 讨论 ======
h1('6  讨论')

h2('6.1  关于Admissible Commands的公平性')
para('必须指出：admissible_commands列表实质上等价于完美的环境状态传感器——"take plate 2 from countertop 2"直接暴露了物体存在与位置，将感知问题转化为列表筛选。本文与ReAct、Reflexion、EmbodiSkill在完全相同的admissible条件下对比，结论有效。但BUTLER使用更困难的generation模式（37%），与之对比时应注明条件差异。')

h2('6.2  YLYW先验在导航决策中的价值')
p = para('本文结果从跨域角度验证了YLYW"知几学习"范式')
sup(p, 3)
p.add_run('的有效性。在物理域（抓取决策），YLYW先验是"物体物理属性→卦象→策略"映射；在导航域，先验是"物体典型位置→探索优先级"映射。两者共享相同的设计原则：将人类常识编码为结构化先验，零样本驱动决策。物理域92.7%合理率')
sup(p, 1)
p.add_run('与导航域90.0%成功率（标注一致子集）高度一致，表明YLYW范式在不同领域呈现稳定的90%+能力水平。')

h2('6.3  规则系统 vs LLM：两条技术路线的边界')
para('在相同admissible条件下，800行规则代码（标注一致时90.0%）接近270亿参数LLM（93.28%）。规则系统的优势在于：（1）完全确定性，相同输入始终相同输出；（2）零幻觉，严格从列表选择，不生成无效动作；（3）部署成本极低，无需GPU/API；（4）完全可解释，每个决策可追溯到具体规则。')

p = para('LLM的优势在于语义弹性：当task_desc说"salt shaker"但场景只有pepper shaker时，LLM能通过上下文推断修正。这种能力是规则系统无法复制的。因此，两条路线的边界在于：')
p.add_run('当任务描述准确时，规则系统可接近LLM水平；当描述存在歧义时，LLM的语义理解能力不可替代。').bold = True

h2('6.4  ALFWorld标注质量问题的启示')
para('本文发现的25.4%标注不一致率对ALFWorld基准有重要启示。这些不一致并非随机噪声，而是系统性的：（1）ALFRED数据集的task_desc由众包工人标注，与PDDL自动生成的ground truth存在语义gap；（2）同义词混用（mug/cup、salt/pepper）反映了自然语言的固有模糊性；（3）信息缺失（"Turn on the desk lamp"不说看什么）反映了标注指南的不完善。建议ALFWorld社区对这34个任务的标注进行修正或标记，以提供更公平的评估基准。')

h2('6.5  局限性与未来工作')
para('（1）NL解析精度：当前基于关键词的解析器在面对同义词和模糊表达时能力有限。未来可结合轻量NLP模型提升解析精度，同时保持不依赖LLM的优势。（2）50步限制：pick_two类型需完成两轮完整操作，50步限制下成功率仅47.1%。放宽至80步预计可显著提升。（3）admissible模式的现实局限：真实机器人不存在admissible oracle，推广到真实世界需结合视觉感知模块。（4）领域特化：任务模板针对ALFWorld 6种类型设计，新任务类型需人工编写新模板。')

# ====== 7 结论 ======
h1('7  结论')
p = para('本文将YLYW先验知识框架首次应用于具身智能导航与操作决策领域，在ALFWorld基准测试（134个valid_unseen任务）上取得了以下成果：')

contribs = [
    '设计了admissible-commands信号驱动的层次化状态机Agent，以约800行Python代码、无LLM/API依赖，实现67.2%的整体成功率。',
    '揭示了ALFWorld数据集中25.4%任务（34/134）存在task_desc与PDDL标注不一致的问题。在标注一致的100个任务上，YLYW达到90.0%的成功率，与使用270亿参数LLM的EmbodiSkill（93.28%）仅差约3个百分点。',
    '验证了YLYW"知几学习"范式的跨域有效性：同一套"先验知识+信号驱动"方法论从物理对象决策（92.7%）成功迁移至具身导航决策（90.0%）。',
    '发现并修复了ALFWorld评估中的环境加载BUG，为社区提供了可复现的基准。',
    '通过与EmbodiSkill的系统对比，明确了先验知识驱动 vs LLM驱动两条技术路线的各自优势和边界条件。',
]
for c in contribs:
    doc.add_paragraph(c, style='List Number')

p = para('本文的核心论点是：在具身智能的有限状态空间中，当任务描述准确时，精心设计的规则系统可以接近大规模LLM的性能水平，同时在确定性、可解释性、部署成本等方面具有显著优势。这为"知识驱动的具身智能"提供了有力的实证支持。')

doc.add_paragraph()

# ====== 参考文献 ======
h1('参考文献')
refs = [
    '[1]  Ma X L, Ma S J, Li J H, et al. YLYW: A Federated Neuro-Symbolic Embodied Decision-Making Framework Based on I Ching Prior Symbolic Knowledge. arXiv preprint, 2026.',
    '[2]  Ma X L, Ma S J, Zhang G A, et al. Hierarchical Nested YLYW: A Distributed Embodied Intelligence Architecture Paradigm Based on I Ching Holographic Principle. arXiv preprint, 2026.',
    '[3]  Ma X L, Ma S J, Li J H, et al. Zhiji Learning: A New Paradigm for Embodied Learning Based on Prior Symptom Recognition. arXiv preprint, 2026.',
    '[4]  Ma X L, et al. YLYW Chinese Language Processing: A New Paradigm. arXiv preprint, 2026.',
    '[5]  Ma X L. LLM+YLYW Anti-Hallucination Hybrid System. arXiv preprint, 2026.',
    '[6]  Shridhar M, Yuan X, Côté M A, et al. ALFWorld: Aligning Text and Embodied Environments for Interactive Learning. In: ICLR, 2021.',
    '[7]  Shridhar M, Thomason J, Gordon D, et al. ALFRED: A Benchmark for Interpreting Grounded Instructions for Everyday Tasks. In: CVPR, 2020.',
    '[8]  Côté M A, Kádár Á, Yuan X, et al. TextWorld: A Learning Environment for Text-Based Games. In: CGW@IJCAI, 2018.',
    '[9]  Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models. In: ICLR, 2023.',
    '[10] Shinn N, Cassano F, Gopinath A, et al. Reflexion: Language Agents with Verbal Reinforcement Learning. In: NeurIPS, 2023.',
    '[11] Ahn M, Brohan A, Brown N, et al. Do As I Can, Not As I Say: Grounding Language in Robotic Affordances. arXiv:2204.01691, 2022.',
    '[12] Driess D, Xia F, Sajjadi M S M, et al. PaLM-E: An Embodied Multimodal Language Model. In: ICML, 2023.',
    '[13] Ju R, Wang X, Ding X, et al. EmbodiSkill: Skill-Aware Reflection for Self-Evolving Embodied Agents. arXiv:2605.10332, 2026.',
    '[14] 《周易》. 先秦典籍.',
    '[15] 《系辞传》. 先秦典籍.',
    '[16] Brohan A, Brown N, Carbajal J, et al. RT-2: Vision-Language-Action Models Transfer Web Knowledge to Robotic Control. arXiv:2307.15818, 2023.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(ref).font.size = Pt(9)

doc.save('中国科学_信息科学_YLYW_ALFWorld.docx')
print('Part3完成: 6讨论+7结论+参考文献')
print('论文全文生成完毕！')
