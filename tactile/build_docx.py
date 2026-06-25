#!/usr/bin/env python3
"""Build tactile experiment plan docx"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re, os

doc = Document()
outdir = os.path.dirname(os.path.abspath(__file__))

for sec in doc.sections:
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.first_line_indent = Cm(0.74)

def P(text, bold=False, size=None, align=None, indent=True, sa=None):
    p = doc.add_paragraph()
    if not indent: p.paragraph_format.first_line_indent = Cm(0)
    if sa is not None: p.paragraph_format.space_after = Pt(sa)
    if align: p.alignment = align
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        else:
            r = p.add_run(part)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size or 11)
    return p

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.first_line_indent = Cm(0)
    for r in h.runs: r.font.name = 'Times New Roman'
    return h

def table(data, widths=None):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            c = t.rows[i].cells[j]
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(str(cell_text))
            r.font.name = 'Times New Roman'; r.font.size = Pt(9)
            if i == 0: r.bold = True
    doc.add_paragraph()

# ===== TITLE =====
P('YLYW触觉：基于易经时序推理的触觉传感器压力场估计 — 实验方案',
  bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, sa=8)
P('版本 v0.1  |  2026-06-05', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, sa=12)

# ===== 1 =====
H('1 问题定义', 1)

H('1.1 传感器原理', 2)
P('触觉传感器由压致变色薄膜构成：薄膜在压力作用下发生形变，其反射/透射光谱发生可观测变化。通过工业相机采集薄膜表面的视频图像，从光谱变化反演压力场的时空分布。')

H('1.2 核心挑战', 2)
P('欠定反问题：视频像素值(RGB/多光谱) → 压力值，信息压缩比极高。时序依赖性：压力场在时间上连续演化，单帧信息不足，需要时序建模。空间耦合：薄膜的力学响应存在空间耦合，相邻区域相互影响。')

H('1.3 YLYW的独特优势', 2)
P('与传统的图像→压力回归（CNN、物理模型反演）不同，YLYW提供了一种符号先验驱动的时序推理路线。六爻编码天然适合时空表示——6个爻位对应6个空间/物理维度，每帧一个卦象。变卦分析直接对应压力演化——爻值变化即压力变化，变爻位置即压力变化位置。爻位关系建模空间耦合——乘承比应五种算子天然描述相邻/远距区域的力学交互。物极必反检测边界状态——爻值逼近0/1即压力达到薄膜响应极限。')

# ===== 2 =====
H('2 系统架构', 1)

P('系统分为四个层次：L1逐帧六爻编码（视频帧→6维触觉特征向量）→ L2六十四卦时序匹配（本卦H(t) + 变卦分析）→ L3爻位关系空间耦合分析（乘承比应当位得中 → 力修正系数）→ L3+时序演化分析（物极必反、否极泰来、变卦序列→动作识别）→ 输出层（全局压力、压力梯度场、接触状态、力控参数）。')

# ===== 3 =====
H('3 六爻编码设计（触觉域）', 1)

H('3.1 逐帧特征提取', 2)
P('每帧图像预处理后（去噪、光照归一化），计算以下6维触觉特征。')

P('表1. 触觉域六爻编码', bold=True, indent=False)
table([
    ['爻位', '触觉语义', '计算方法', '[0,1]映射'],
    ['初爻', '全局压力', '全幅光谱偏移量 Δλ/λ₀', 'Sigmoid归一化'],
    ['二爻', '中心压力', 'ROI中心区的光谱偏移', '尺寸归一化'],
    ['三爻', '压力梯度', 'Sobel梯度幅值的空间均值', '对数归一化'],
    ['四爻', '压力集中度', '高压力区(前20%)的空间CV', '1−exp(−CV)'],
    ['五爻', '接触稳定性', '相邻帧光谱变化的时间标准差', '1−tanh(σ)'],
    ['上爻', '边界压力', '图像四周20%边缘区的光谱偏移', '同上'],
])

H('3.2 卦象模板语义', 2)
P('利用物理域YLYW的64卦模板，通过余弦相似度匹配y(t)到最近卦象。触觉域的模板语义重新定义为：乾宫卦（上卦乾，刚健）→ 高压力、稳定接触；坤宫卦（上卦坤，柔顺）→ 低压力、轻微接触；震宫卦（上卦震，动）→ 压力快速变化；巽宫卦（上卦巽，入）→ 压力逐渐渗透；坎宫卦（上卦坎，陷）→ 局部压力集中；离宫卦（上卦离，丽）→ 多区域同时受压；艮宫卦（上卦艮，止）→ 压力到达稳态；兑宫卦（上卦兑，说）→ 压力释放/反弹。')

H('3.3 爻位关系 — 压力场空间耦合', 2)
P('对y(t)执行乘承比应当位得中分析：乘（阴乘阳）→ 低压区压制高压区，力学约束验证；承（阴承阳）→ 压力从前区传导后区，力传导路径追踪；比（邻爻同异）→ 相邻区域的压力一致性；应（初↔四、二↔五、三↔上）→ 远距力耦合；当位/得中 → 全局压力分布合理性评估。输出力修正系数 m(t) ∈ [0.75, 1.05]。')

H('3.4 时序演化分析', 2)
P('物极必反检测：爻值(t) > 0.95 且下降趋势 → 压力饱和预警。否极泰来检测：连续K帧爻值 < 0.1 → 即将接触/脱开。变卦序列分析：H(t-2)→H(t-1)→H(t) 三连卦 → 识别动作类型（接近、接触、抓紧、滑动、释放、保持）。')

# ===== 4 =====
H('4 实验设计', 1)

H('4.1 阶段一：标定数据采集', 2)
P('设备：触觉传感器薄膜 + 工业相机（≥30fps）+ 标准力传感器（精度0.01N）+ 可控压头（步进电机，力控精度0.1N）。')

P('表2. 数据采集协议', bold=True, indent=False)
table([
    ['场景', '描述', '参数', '帧数'],
    ['A: 单点静态', '0→0.5→1→2→5→10→5→2→1→0.5→0N', '每级3秒, 30fps', '~990'],
    ['B: 单点动态', '0→10N匀速(5s)→保持(2s)→10→0N(5s)', '重复10次', '~3,600'],
    ['C: 多点顺序', '3压头(间距15mm), 依次→同时→依次释放', '重复10次', '~3,000'],
    ['D: 多点复杂', '3压头随机力值(0~8N, 0.5~3s)', '重复20次', '~18,000'],
    ['总计', '', '', '~25,000'],
])

H('4.2 阶段二：压力场标定与验证', 2)
P('实验1 — 全局压力回归：输入单帧6爻向量y(t)，输出全局压力标量P̄(t)，对比力传感器读数。基线方法包括线性回归和多项式回归，YLYW方法为P̄ = f(卦象类型, 爻位修正系数)。指标：MAE、RMSE、R²。')
P('实验2 — 压力变化检测：输入帧对y(t-1)和y(t)，得到变爻向量d(t)，输出压力变化量ΔP(t)及变化区域。对比力传感器差值ΔF(t)。指标：变化检测准确率、ΔP回归MAE。')
P('实验3 — 接触状态识别：输入时间窗口[y(t-K), ..., y(t)]的变卦序列，输出接触状态∈{接近, 接触, 抓紧, 滑动, 释放, 保持}。对比人工标注或力传感器阈值规则。指标：状态分类准确率、状态转换时延。')

H('4.3 阶段三：压力场空间分辨率', 2)
P('方法A — 分区域6爻编码：将图像划分为M×N网格，每格独立提取6爻特征，全局压力场P(x,y) = {P_{ij}}。爻位关系提供网格间的空间一致性约束（乘承保证垂直连续、比保证水平连续、应保证远距对称）。')
P('方法B — 时序超分辨率：利用多帧信息提高空间分辨率。帧间亚像素位移→光谱变化→6爻变化→压力变化，从上采样网格的爻值变化细微空间模式中恢复高分辨率压力场。')

H('4.4 阶段四：灵巧手力控闭环', 2)
P('控制回路：触觉传感器→YLYW(每帧)→力修正系数m(t)+接触状态S(t)→灵巧手控制器（目标力×m(t)→关节力矩）+状态机（S(t)变化→动作阶段切换）。评测任务包括：抓取易碎物体（鸡蛋，m(t)自动降低防碎）、抓取不规则物体（石头，m(t)自适应局部压力分布）、滑动检测与补偿（物极必反检测→增加握力）。')

# ===== 5 =====
H('5 评估指标', 1)

table([
    ['指标', '公式/描述', '目标值'],
    ['压力MAE', 'Σ|P_pred − P_true| / N', '< 0.1N'],
    ['压力场PSNR', '20·log₁₀(MAX/√MSE)', '> 30dB'],
    ['接触状态准确率', '正确帧数/总帧数', '> 90%'],
    ['状态转换时延', 't_detected − t_true', '< 100ms'],
    ['变爻→ΔP相关系数', 'Pearson r(d(t), ΔF(t))', '> 0.85'],
    ['力控成功率', '成功抓取次数/总次数', '> 95%'],
])

# ===== 6 =====
H('6 实施计划', 1)

table([
    ['阶段', '任务', '预估时间', '产出'],
    ['一', '标定数据采集', '1周', '25K帧 + 力传感器真值'],
    ['二.1', '全局压力回归实验', '3天', '压力估计精度报告'],
    ['二.2', '压力变化检测', '2天', '变爻→ΔP相关分析'],
    ['二.3', '接触状态识别', '3天', '6状态分类器'],
    ['三', '空间压力场恢复', '1周', '压力场可视化+精度'],
    ['四', '灵巧手闭环控制', '1周', '抓取成功率报告'],
    ['—', '论文撰写', '1周', '触觉YLYW论文'],
])

# ===== 7 =====
H('7 预期创新点', 1)

P('（1）首次将易经符号推理用于触觉感知：开辟"符号先验+时序物理信号"新范式。')
P('（2）爻位关系建模空间力耦合：乘承比应算子为压力场空间一致性提供可解释约束。')
P('（3）变卦序列分析→接触状态机：64卦的时序转移模式天然对应物理接触的阶段转换。')
P('（4）物极必反检测滑动：爻值逼近边界提供物理极限的早期预警。')
P('（5）零样本/小样本通用性：不同于需要大量标定数据的CNN方法，YLYW路线在先验知识驱动下仅需少量标定即可工作。')

# Save
out = os.path.join(outdir, '实验方案_触觉YLYW.docx')
doc.save(out)
print(f'Saved: {out}')
