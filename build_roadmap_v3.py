#!/usr/bin/env python3
"""生成 YLYW 技术路线图 — 嵌入 docx"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from matplotlib.font_manager import FontProperties
import numpy as np

# 中文字体
zh_font = FontProperties(fname='/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc', size=10)
zh_font_sm = FontProperties(fname='/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc', size=8)
zh_font_lg = FontProperties(fname='/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc', size=12)
zh_font_title = FontProperties(fname='/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc', size=18, weight='bold')

# 画布设置
fig, ax = plt.subplots(1, 1, figsize=(20, 14))
ax.set_xlim(0, 20)
ax.set_ylim(0, 14)
ax.axis('off')
ax.set_facecolor('#F8F9FA')

# 配色
COLOR_VISION = '#27AE60'      # 远景-绿
COLOR_DONE = '#1A5276'        # 已完成-深蓝
COLOR_PROGRESS = '#E74C3C'    # 进行中-红
COLOR_PLAN = '#7F8C8D'        # 规划-灰
COLOR_BRANCH = '#8E44AD'      # 分支-紫
COLOR_MAIN = '#2C3E50'        # 主干-深灰
COLOR_BG_BOX = '#FFFFFF'
COLOR_ARROW = '#95A5A6'

def draw_box(ax, x, y, w, h, text, color, fontsize=8, bold=False, text_color='white', border_color=None, alpha=1.0):
    """画圆角矩形框"""
    box = FancyBboxPatch((x, y), w, h,
                          boxstyle="round,pad=0.08",
                          facecolor=color, edgecolor=border_color or color,
                          linewidth=1.5, alpha=alpha, zorder=3)
    ax.add_patch(box)
    fp = zh_font if fontsize <= 9 else zh_font_lg
    weight = 'bold' if bold else 'normal'
    ax.text(x + w/2, y + h/2, text, ha='center', va='center',
            fontproperties=fp, fontsize=fontsize, color=text_color, fontweight=weight, zorder=4)

def draw_branch_box(ax, x, y, w, h, title, subtitle, color, fontsize=9):
    """画分支框（标题+副标题）"""
    box = FancyBboxPatch((x, y), w, h,
                          boxstyle="round,pad=0.08",
                          facecolor=color, edgecolor=color,
                          linewidth=1.5, alpha=0.12, zorder=3)
    ax.add_patch(box)
    box2 = FancyBboxPatch((x, y), w, h,
                           boxstyle="round,pad=0.08",
                           facecolor='none', edgecolor=color,
                           linewidth=1.5, alpha=0.6, zorder=3)
    ax.add_patch(box2)
    ax.text(x + w/2, y + h - 0.25, title, ha='center', va='top',
            fontproperties=zh_font, fontsize=fontsize, color=color, fontweight='bold', zorder=4)
    ax.text(x + w/2, y + 0.25, subtitle, ha='center', va='bottom',
            fontproperties=zh_font_sm, fontsize=7.5, color='#555', zorder=4)

def draw_arrow(ax, x1, y1, x2, y2, color=COLOR_ARROW, lw=1.5, style='simple'):
    """画箭头"""
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=lw,
                               connectionstyle='arc3,rad=0'))

# ============================================================
# TITLE
# ============================================================
ax.text(10, 13.5, 'YLYW 技术路线总图', ha='center', va='center',
        fontproperties=zh_font_title, fontsize=18, color='#1A1A2E')
ax.text(10, 13.1, '2026.06.07 更新  |  已完成  进行中  规划中',
        ha='center', va='center', fontproperties=zh_font_sm, fontsize=9, color='#888')

# ============================================================
# 远景愿景 — TOP BANNER
# ============================================================
draw_box(ax, 0.5, 11.8, 19, 0.9,
         '[远景] 层次化嵌套 YLYW → 道器合一的通用人工智能  |  三步走：智能元胞 → 智能组织 → 智能生态',
         COLOR_VISION, fontsize=10, bold=True, alpha=0.9)

# ============================================================
# 主干：阶段一 → 阶段二 → 阶段三 → 阶段四
# ============================================================
main_y = 9.8
box_h = 0.9
box_w = 2.6
gap = 0.5
start_x = 1.0

stages = [
    ('阶段一\n符号先验推理', '[完成] 92.7% 合理率', COLOR_DONE),
    ('阶段二\n双八卦安全架构', '[完成] 0% 严重错误', COLOR_DONE),
    ('阶段三\n通用行动扩展', '[待] 抓取→推物→倒水', COLOR_PROGRESS),
    ('阶段四\n分层泛化验证', '[待] 物理/类比/组合', COLOR_PROGRESS),
]

for i, (title, sub, color) in enumerate(stages):
    x = start_x + i * (box_w + gap)
    draw_box(ax, x, main_y, box_w, box_h, title, color,
             fontsize=9, bold=True, border_color=color)

    # subtitle below
    status_color = '#27AE60' if '[完成]' in sub else '#E67E22'
    ax.text(x + box_w/2, main_y - 0.2, sub, ha='center', va='top',
            fontproperties=zh_font_sm, fontsize=7, color=status_color)

    # arrow between stages
    if i < len(stages) - 1:
        draw_arrow(ax, x + box_w + 0.05, main_y + box_h/2,
                   x + box_w + gap - 0.05, main_y + box_h/2, COLOR_ARROW, 1.8)

# ============================================================
# 分支 - 视觉分类（上部，已完成）
# ============================================================
vision_y = 11.0
draw_branch_box(ax, 7.0, vision_y, 6.0, 0.65,
                '视觉分类分支  [完成]',
                'STL-10 37.0% Top-1 (3×随机基线) | 论文 v2.0 | 爻位关系验证',
                '#27AE60')
draw_arrow(ax, 8.0, vision_y, 3.6, main_y + box_h, '#27AE60', 1.2)
draw_arrow(ax, 8.0, vision_y, 11.5, 11.8, '#27AE60', 1.0)  # connect to vision banner

# ============================================================
# 分支行 — 下方
# ============================================================
branch_data = [
    # (x, y, w, title, subtitle, color, connect_x)
    (0.5, 8.3, 4.3, '触觉感知  🆕 方案完成',
     '压致变色薄膜→压力场\n爻位关系→空间力耦合',
     '#E67E22', 2.0),
    (5.5, 8.3, 4.3, '芯片化  🆕 远期规划',
     '纯逻辑门 <50mW <100ns\nFPGA→ASIC 流片路线',
     '#3498DB', 5.5),
    (10.5, 8.3, 4.3, '量子计算  🆕 远景规划',
     '64卦=6-qubit基矢\n量子爻位关系→AGI',
     '#9B59B6', 8.0),
    (15.5, 8.3, 4.3, '自适应学习  🆕 方案完成',
     '443参数可解释修正\n1次失败→精确定位→修正',
     '#1ABC9C', 11.0),
]

for x, y, w, title, sub, color, cx in branch_data:
    draw_branch_box(ax, x, y, w, 0.85, title, sub, color)
    # connect to main
    draw_arrow(ax, x + w/2, y + 0.85, cx, main_y, color, 1.0)

# ============================================================
# 图例
# ============================================================
legend_y = 7.5
for i, (label, color) in enumerate([
    ('[完成] 已完成', COLOR_DONE), ('[待] 进行中 / 方案完成', '#E67E22'),
    ('[新] 远期/远景规划', '#3498DB'), ('[远景] 终极愿景', COLOR_VISION)
]):
    ax.add_patch(plt.Rectangle((1.0 + i*5, legend_y - 0.1), 0.3, 0.18,
                                facecolor=color, edgecolor=color, linewidth=1))
    ax.text(1.4 + i*5, legend_y, label, fontproperties=zh_font_sm, fontsize=8, color='#555', va='center')

# ============================================================
# 实物验证提示
# ============================================================
draw_box(ax, 0.5, 6.5, 19, 0.55,
         '[近期] 灵犀 X2 + OmniHand 2025  抓取实物验证（零样本物理常识）— 路径图的智能元胞基石',
         '#E74C3C', fontsize=9, alpha=0.85)

# ============================================================
# 底部：三步走路径
# ============================================================
path_y = 5.5
steps = [
    ('智能的"元胞"', '单个 YLYW 模型\n常识+本能+底线\n有限但可靠', COLOR_DONE),
    ('智能的"组织"', '层次化嵌套\n多元胞 → 复杂智能体\n每层同样易理法则', '#E67E22'),
    ('智能的"生态"', '分布式智能系统\n共通"天道"语言\n自组织/自和谐', '#3498DB'),
]

for i, (title, sub, color) in enumerate(steps):
    bx = 1.5 + i * 6.2
    draw_box(ax, bx, path_y, 2.8, 0.9, title, color, fontsize=9, bold=True, alpha=0.85)
    ax.text(bx + 1.4, path_y - 0.15, sub, ha='center', va='top',
            fontproperties=zh_font_sm, fontsize=7.5, color='#555')
    if i < 2:
        draw_arrow(ax, bx + 2.8 + 0.1, path_y + 0.45,
                   bx + 6.2 - 0.1, path_y + 0.45, COLOR_ARROW, 2)

# ============================================================
# 底部：核心主张
# ============================================================
ax.text(10, 4.0,
        '通用人工智能的基石，或许不在于无限的算力与数据，\n'
        '而在于一套内建的、自洽的、关于世界如何变化与运作的根本法则。',
        ha='center', va='center', fontproperties=zh_font, fontsize=9.5, color='#1A5276',
        style='italic')

# ============================================================
# Footer
# ============================================================
ax.text(10, 0.3, '青岛科技大学 信息科学技术学院  |  YLYW Project  |  2026',
        ha='center', fontproperties=zh_font_sm, fontsize=7.5, color='#AAA')

plt.tight_layout(pad=0.5)
plt.savefig('/home/lijinhan/MXL/科研/ylyw/技术路线图.png', dpi=200,
            bbox_inches='tight', facecolor='#F8F9FA')
plt.close()

# ============================================================
# 嵌入到 docx
# ============================================================
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)

md_path = '/home/lijinhan/MXL/科研/ylyw/技术路线_实际进展.md'
with open(md_path, 'r') as f:
    lines = f.readlines()

def add_heading_text(text, level):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text.lstrip('#').strip())
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = '黑体'
    elif level == 2:
        run = p.add_run(text.lstrip('#').strip())
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = '黑体'
        run.font.color.rgb = None
    elif level == 3:
        run = p.add_run(text.lstrip('#').strip())
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = '黑体'

import re

i = 0
in_code = False
in_table = False
table_rows = []

while i < len(lines):
    line = lines[i]

    if line.strip().startswith('```'):
        if in_code:
            in_code = False
        else:
            in_code = True
        i += 1
        continue

    if in_code:
        i += 1
        continue

    # 检测总体路线图区域——替换为图片
    if line.strip().startswith('## 总体路线图'):
        add_heading_text(line, 2)
        # 找到代码块起始和结束
        j = i + 1
        code_start = -1
        code_end = -1
        while j < len(lines):
            if lines[j].strip().startswith('```') and code_start == -1:
                code_start = j
            elif lines[j].strip().startswith('```') and code_start != -1:
                code_end = j
                break
            j += 1

        if code_start >= 0 and code_end >= 0:
            # 插入图片
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture('/home/lijinhan/MXL/科研/ylyw/技术路线图.png',
                           width=Inches(6.0))
            i = code_end + 1
            continue

    # 表格
    if line.strip().startswith('|') and line.strip().endswith('|'):
        if '---' not in line:
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            table_rows.append(cells)
        i += 1
        continue
    elif table_rows:
        if len(table_rows) >= 1:
            table = doc.add_table(rows=1, cols=len(table_rows[0]))
            table.style = 'Light Grid Accent 1'
            for j, cell_text in enumerate(table_rows[0]):
                table.rows[0].cells[j].text = cell_text
                for p in table.rows[0].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(8.5)
            for row in table_rows[1:]:
                row_cells = table.add_row().cells
                for j, cell_text in enumerate(row):
                    if j < len(row_cells):
                        row_cells[j].text = cell_text
                        for p in row_cells[j].paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(8.5)
            doc.add_paragraph()
        table_rows = []
        continue

    # 标题
    if line.startswith('# ') and not line.startswith('## '):
        add_heading_text(line, 1)
    elif line.startswith('## '):
        add_heading_text(line, 2)
    elif line.startswith('### '):
        add_heading_text(line, 3)
    elif line.startswith('>'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(line.lstrip('> ').rstrip('\n'))
        run.italic = True
        run.font.size = Pt(9.5)
    elif line.strip() == '---':
        p = doc.add_paragraph()
        run = p.add_run('─' * 60)
        run.font.size = Pt(6)
    elif line.strip() == '':
        pass
    elif line.strip():
        p = doc.add_paragraph()
        text = line.rstrip('\n')
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        if re.match(r'^\s*[\-\d\.]\s', text):
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(10.5)

    i += 1

# 处理末尾表格
if table_rows:
    table = doc.add_table(rows=1, cols=len(table_rows[0]))
    table.style = 'Light Grid Accent 1'
    for j, cell_text in enumerate(table_rows[0]):
        table.rows[0].cells[j].text = cell_text
        for p in table.rows[0].cells[j].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
    for row in table_rows[1:]:
        row_cells = table.add_row().cells
        for j, cell_text in enumerate(row):
            if j < len(row_cells):
                row_cells[j].text = cell_text
                for p in row_cells[j].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8.5)

output = '/home/lijinhan/MXL/科研/ylyw/YLYW技术路线_v3.0.docx'
doc.save(output)
print(f'✅ Saved: {output} (with roadmap diagram embedded)')
