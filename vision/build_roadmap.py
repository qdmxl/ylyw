#!/usr/bin/env python3
"""Build YLYW technical roadmap docx from markdown"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re, os

doc = Document()
outdir = os.path.dirname(os.path.abspath(__file__))

for sec in doc.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
style.paragraph_format.line_spacing = 1.1
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.first_line_indent = Cm(0.5)

def P(text, bold=False, size=None, align=None, indent=True, font='Times New Roman'):
    p = doc.add_paragraph()
    if not indent: p.paragraph_format.first_line_indent = Cm(0)
    if align: p.alignment = align
    if isinstance(text, str):
        r = p.add_run(text); r.font.name = font; r.font.size = Pt(size or 10)
        if bold: r.bold = True
    else:
        for seg in text:
            r = p.add_run(seg[0])
            r.font.name = font; r.font.size = Pt(size or 10)
            if len(seg) > 1 and seg[1]: r.bold = True
    return p

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.first_line_indent = Cm(0)
    for r in h.runs: r.font.name = 'Times New Roman'
    return h

def Code(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(8)
    return p

def T(data):
    if not data or len(data) < 2: return
    ncols = len(data[0])
    # Filter out rows with wrong column count
    data = [r for r in data if len(r) == ncols]
    if len(data) < 2: return
    t = doc.add_table(rows=len(data), cols=ncols)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(data):
        for j, ct in enumerate(row):
            c = t.rows[i].cells[j]
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(str(ct))
            r.font.name = 'Times New Roman'; r.font.size = Pt(7.5)
            if i == 0: r.bold = True
    doc.add_paragraph()

# ===== READ & PARSE MD =====
path = os.path.join(os.path.dirname(outdir), '技术路线_实际进展.md')
with open(path) as f:
    lines = f.readlines()

i = 0
in_table = False
table_rows = []
in_code = False
code_lines = []

while i < len(lines):
    line = lines[i].rstrip()

    # Code block
    if line.startswith('```'):
        if in_code:
            Code('\n'.join(code_lines))
            code_lines = []
            in_code = False
        else:
            if in_table:
                T(table_rows)
                table_rows = []
                in_table = False
            in_code = True
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # Empty
    if not line.strip():
        if in_table:
            T(table_rows)
            table_rows = []
            in_table = False
        i += 1
        continue

    # Separator
    if line.strip() == '---':
        i += 1
        continue

    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        H(line[2:], 1)
        i += 1; continue
    if line.startswith('## '):
        H(line[3:], 2)
        i += 1; continue
    if line.startswith('### '):
        H(line[4:], 3)
        i += 1; continue

    # Table (must have at least 3 pipe chars and look like a data table, not ASCII art)
    pipe_count = line.count('|')
    if pipe_count >= 3 and not any(c in line for c in '┌┐└┘├┤┬┴┼─│'):
        if '---' in line:
            i += 1; continue
        # Skip lines that are box-drawing or have too many non-table chars
        if re.match(r'^\s*\|', line):
            in_table = True
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue

    # Regular text
    if in_table:
        T(table_rows)
        table_rows = []
        in_table = False

    # Handle bold **text**
    parts = re.split(r'(\*\*.*?\*\*)', line)
    if len(parts) > 1:
        segs = []
        for p in parts:
            if p.startswith('**') and p.endswith('**'):
                segs.append((p[2:-2], True))
            else:
                segs.append((p, False))
        P(segs)
    else:
        P(line)
    i += 1

# Trailing
if in_table and table_rows:
    T(table_rows)
if in_code and code_lines:
    Code('\n'.join(code_lines))

out = os.path.join(outdir, '..', 'YLYW技术路线.docx')
doc.save(out)
print(f'Saved: {os.path.abspath(out)}')
