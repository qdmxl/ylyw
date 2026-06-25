#!/usr/bin/env python3
"""Convert paper_vision.md to formatted .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(4)

# Read markdown
with open('paper_vision.md', 'r') as f:
    text = f.read()

lines = text.split('\n')

def add_heading_cn(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

def add_para(doc, text, bold=False, italic=False, size=None, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if align: p.alignment = align
    return p

def add_table_from_text(doc, text, header=True):
    """Parse a markdown-style table and add to doc"""
    rows = [row.strip() for row in text.strip().split('\n') if '|' in row and '---' not in row]
    if not rows: return
    
    data = []
    for row in rows:
        cells = [c.strip() for c in row.split('|')[1:-1]]
        data.append(cells)
    
    if not data: return
    
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)
                    if i == 0 and header:
                        run.bold = True
    
    doc.add_paragraph()  # spacing

# Process markdown
i = 0
in_table = False
table_lines = []

while i < len(lines):
    line = lines[i]
    
    # Skip empty
    if not line.strip():
        if in_table:
            add_table_from_text(doc, '\n'.join(table_lines))
            table_lines = []
            in_table = False
        i += 1
        continue
    
    # Title
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:]
        add_para(doc, title, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue
    
    # Section heading
    if line.startswith('## '):
        if in_table:
            add_table_from_text(doc, '\n'.join(table_lines))
            table_lines = []
            in_table = False
        add_heading_cn(doc, line[3:], level=1)
        i += 1
        continue
    
    # Subsection
    if line.startswith('### '):
        if in_table:
            add_table_from_text(doc, '\n'.join(table_lines))
            table_lines = []
            in_table = False
        add_heading_cn(doc, line[4:], level=2)
        i += 1
        continue
    
    # Table
    if '|' in line:
        in_table = True
        table_lines.append(line)
        i += 1
        continue
    
    # Bold marker
    if '**' in line:
        if in_table:
            add_table_from_text(doc, '\n'.join(table_lines))
            table_lines = []
            in_table = False
        parts = re.split(r'(\*\*.*?\*\*)', line)
        p = doc.add_paragraph()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        i += 1
        continue
    
    # Code blocks
    if line.startswith('```'):
        if in_table:
            add_table_from_text(doc, '\n'.join(table_lines))
            table_lines = []
            in_table = False
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        i += 1
        continue
    
    # Regular paragraph
    if in_table:
        add_table_from_text(doc, '\n'.join(table_lines))
        table_lines = []
        in_table = False
    
    p = doc.add_paragraph()
    # Handle inline formatting
    text = line
    # Bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    i += 1

# Handle trailing table
if in_table and table_lines:
    add_table_from_text(doc, '\n'.join(table_lines))

# Save
outpath = 'YLYW视觉论文_v1.0.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
