#!/usr/bin/env python3
"""Build YLYW research memo from all session records"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
outdir = os.path.dirname(os.path.abspath(__file__))

for sec in doc.sections:
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.first_line_indent = Cm(0.74)

def P(text, bold=False, size=None, align=None, indent=True):
    p = doc.add_paragraph()
    if not indent: p.paragraph_format.first_line_indent = Cm(0)
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size or 10.5)
    if bold: r.bold = True
    return p

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.first_line_indent = Cm(0)
    for r in h.runs: r.font.name = 'Times New Roman'
    return h

def T(data):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(data):
        for j, ct in enumerate(row):
            c = t.rows[i].cells[j]; c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(str(ct))
            r.font.name = 'Times New Roman'; r.font.size = Pt(8)
            if i == 0: r.bold = True
    doc.add_paragraph()

def section(title):
    H(title, 1)

def subsection(title):
    H(title, 2)

def entry(date, time, content):
    P(f'[{date} {time}] {content}', indent=False)

def files(files_list):
    for f in files_list:
        P(f'    {f}', indent=False, size=9)

# ===== TITLE PAGE =====
P('', indent=False)
P('YLYW（易理研物）研究过程备忘录', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
P('Research Process Memo', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
P('', indent=False)
P('课题负责人：马老师', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
P('青岛科技大学 信息科学技术学院', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
P('软件工程专业', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
P('', indent=False)
P('记录期间：2026年5月21日 — 2026年6月5日', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
P('生成日期：2026年6月5日', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ===== OVERVIEW =====
section('一、项目概述')

P('YLYW（易理研物）是一个将《易经》六十四卦符号系统形式化为可计算先验知识引擎的研究项目。核心思想是将八卦原型匹配、六爻编码、六十四卦策略映射和爻位关系运算（乘承比应当位得中）构建为联邦式神经符号架构，应用于具身智能决策、运动控制、图像分类和触觉感知等多个领域。')

P('项目位于：/home/lijinhan/MXL/科研/ylyw/')

# ===== PHASE 1 =====
section('二、阶段一：符号先验推理系统（物理域）')
P('时间：2026年5月21日 — 6月1日 | 状态：✅ 已完成')

subsection('2.1 核心成果')
P('实现了将《易经》六十四卦符号系统形式化为可计算的先验决策引擎，完成零样本具身智能推理。')

T([
    ['里程碑', '成果', '关键数据'],
    ['20卦初始系统', '直觉模板→卦象匹配', '48.0%合理率, 2.0%错误率'],
    ['A档22卦补全(42卦)', '策略类型18→39种, 质心优化', '91.7%合理率, 0%错误率'],
    ['B档22卦补全(64卦)', '策略近邻模板, 全覆盖', '92.7%合理率, 0%错误率'],
    ['爻位关系运算', '乘承比应当位得中形式化', '力修正系数均值0.94'],
    ['消融实验', '三维消融验证', '易理+33.6%, 架构+12.7%, 模糊+23.0%'],
    ['对比实验', 'vs MLP, vs 随机', '92.7% vs 28.7% vs 37.5%'],
])

subsection('2.2 关键决策')
P('（1）英文缩写保持YLYW不变（曾讨论YLFM，暂缓修改）')
P('（2）爻位关系权重：当位0.40/得中0.20/乘承0.15/比0.10/应0.15')
P('（3）B档卦用策略近邻模板，避免干扰A档匹配')
P('（4）小样本微调（64卦）因scipy优化耗时长而延后')

subsection('2.3 阶段二：双八卦安全架构')
P('时间：2026年6月2-3日 | 状态：✅ 已完成')
P('原计划使用逆动力学硬性映射实现物理约束层。实际提出了更优雅的双八卦安全架构：策略八卦（决定"抓什么怎么抓"）+ 安全八卦（验证"能否安全执行"），两个64卦系统并行运行，通过仲裁层合并输出。安全八卦将6条物理安全检查公式（摩擦力/破坏阈值/摩擦锥/关节力矩/力封闭/穿透风险）符号化为六爻向量，L3层映射为5级安全等级（SAFE/CAUTION/WARNING/DANGER/CRITICAL），在CRITICAL级别触发变卦重新选策。304物体基线测试：SAFE 38.8%、WARNING ~52%、CRITICAL ~0.3%。')
P('关键创新：（1）6条物理公式仅用于L2层阴阳二值化，L3后完全符号空间推理，零数值计算；（2）双64卦系统独立运行、共享架构、语义互补；（3）CRITICAL触发变卦——安全约束反哺策略选择。')

subsection('2.4 产出文件')
files([
    'paper/YLYW技术论文_v0.1~v0.6.docx（多版本迭代）',
    'paper/YLYW_Technical_Paper_v0.5_EN.docx（英文版）',
    'experiment_phase1/ylyw_core/（L1八卦+L2六爻+L3六十四卦+L3+爻位关系）',
    'scripts/fewshot_finetune.py（小样本微调）',
    '技术路线_实际进展.md',
])

# ===== MOTION CONTROL =====
section('三、运动控制仿真（物理域扩展）')
P('时间：2026年6月3日 | 状态：✅ 初步完成')

subsection('3.1 核心成果')
P('将YLYW推理架构应用于人形机器人步态生成：L1八卦→运动基元映射（乾驱动力/坤柔顺性/震速度等），L2六爻→6维状态编码（姿态/质心/力分布/ZMP裕度/扰动/地形），L3六十四卦→14种步态原型模板库。')

subsection('3.2 环境搭建过程')
P('PyBullet仿真环境搭建历经多次调试：约束式机器人塌陷→改用运动学动画→人体URDF模型错位→改为基础形状搭建→MuJoCo升级。最终实现站立→慢走→行走→小跑→奔跑→恢复→爬坡→平路→站立的10种步态演示。')

subsection('3.3 产出文件')
files([
    'motion_control/launch_sim.py（PyBullet仿真）',
    'motion_control/launch_mujoco.py（MuJoCo仿真）',
    'motion_control/hexagram_gait_rules.py（卦象步态规则）',
    'motion_control/ylyw_locomotion.py（步态推理引擎）',
    'motion_control/lingxi/（灵犀X2适配）',
])

# ===== PAPER =====
section('四、学术论文')
P('时间：2026年6月1日 — 6月3日 | 状态：✅ 多版本完成')

subsection('4.1 论文迭代')
P('技术论文从v0.1迭代至v0.6，内容逐步完善：16张表+4张图+附录A（64卦策略全表）+30条参考文献。§2国内外研究现状从3个子节扩展为5大子节，新增具身智能决策算法、NeSy系统、易经跨学科研究、国际同行对比等内容。')

subsection('4.2 IEEE格式')
P('生成IEEE标准格式版本：两栏布局、10pt Times New Roman、首字下沉、图表规范、IEEE版权行。')

files([
    'paper/YLYW技术论文_v0.6.docx（最新中文版）',
    'paper/YLYW_Technical_Paper_v0.5_EN.docx（英文版）',
    'paper/build_final_docx.py（中英文构建脚本）',
    'paper/技术论文_v0.6.md（Markdown源文件）',
    'paper/国内外研究现状_v0.6.md',
])

# ===== VISION =====
section('五、视觉分类分支（跨域泛化验证）')
P('时间：2026年6月5日 | 状态：✅ 论文完成')

subsection('5.1 研究动机')
P('核心问题：易经符号先验是否具备跨域泛化能力？将物理域的八卦原型（刚性/柔性、重/轻等）替换为视觉域的八卦原型（结构几何、平滑均匀、高对比方向、细纹理、曲线流动、亮度辐射、块状厚重、反射高光），验证零样本图像分类的可行性。这是对易理模型泛化能力的关键验证，也是独立可发表的研究方向。')

subsection('5.2 特征工程迭代')
P('经历了三个版本的迭代优化：')

T([
    ['版本', '特征维度', '方法', 'Brodatz Top-1', '说明'],
    ['v3原始', '6D', '纹理均匀度/边缘清晰度/局部对比度/形状规整度/显著性/背景复杂度', '12.2%', '共享特征，信息瓶颈'],
    ['v4专用检测器', '8D', '每卦一个算子+softmax', '22.7%', '检测器输出未校准'],
    ['8D简单算子', '8D', '一卦一算子+Z-score归一化', '14.0%', '简洁但纹理判别不足'],
    ['Rich 52D', '52D', 'GLCM(18)+Gabor(24)+LBP(10)', '30.2%', '特征最丰富但维度过高'],
    ['合成自然物体', '8D', '8D简单算子+5-fold CV', '89.6%', '架构验证通过'],
    ['STL-10(真实)', '8D', '8D简单算子+Z-score', '37.0%', '标准基准，3×随机基线'],
])

subsection('5.3 STL-10标准数据集结果')
P('在STL-10标准基准（10,400张真实照片，8类自然物体）上：零样本Top-1准确率37.0%（随机基线12.5%，3倍提升），Top-3准确率75.0%。各类别表现：震(horse)57.9%、离(car)53.0%、乾(airplane)51.5%、坎(deer)40.0%、坤(ship)29.0%、艮(truck)25.4%、兑(cat)24.0%、巽(bird)15.5%。')

subsection('5.4 爻位关系在视觉域的分析')
P('关键发现：爻位关系（乘承比应当位得中）在静态图像分类中修正作用接近于零——力修正系数0.972±0.059，接近恒等映射1.0。82%样本的谨慎级别为normal。这验证了理论预期：爻位关系设计用于处理特征间的动态协调关系（如物理域中"力需求与脆弱性的矛盾"），而非提取静态判别信息。其真正价值在时序变化分析中——为后续时序视觉和触觉感知研究提供了明确方向。')

subsection('5.5 视觉论文')
P('生成独立学术论文YLYW视觉论文_v2.0.docx，包含摘要、引言、方法（架构/算子/爻编码/爻位关系）、实验（STL-10/Brodatz/合成）、结果分析（分类/混淆/跨数据集对比/爻位）、讨论（泛化机制/定位差异/爻位定位/局限）、结论和6篇参考文献。4张嵌入式图片：架构图、各类别准确率柱状图、混淆矩阵热力图、跨数据集对比图。')

subsection('5.6 产出文件')
files([
    'vision/paper_vision.md（论文Markdown源文件）',
    'vision/YLYW视觉论文_v2.0.docx（Word格式，含图片）',
    'vision/build_paper_v2.py（论文构建脚本）',
    'vision/simple_8d.py（8D简单算子提取器）',
    'vision/specialized_detectors_v2.py（专用检测器v2）',
    'vision/rich_features.py（52D丰富特征提取器）',
    'vision/feature_extractor_vision.py（6D特征提取器）',
    'vision/classifier.py（分类器v3）',
    'vision/classifier_v4.py（分类器v4）',
    'vision/test_stl10.py（STL-10测试）',
    'vision/test_stl10_yao.py（STL-10+爻位关系测试）',
    'vision/test_brodatz.py（Brodatz测试）',
    'vision/test_8d_simple.py（8D简单测试）',
    'vision/test_fewshot.py（小样本微调测试）',
    'vision/test_rich_features.py（丰富特征测试）',
    'vision/generate_figures.py（论文图表生成）',
])

# ===== TACTILE =====
section('六、触觉传感器方向（新方向）')
P('时间：2026年6月5日 | 状态：🆕 实验方案完成')

subsection('6.1 研究背景')
P('研制新型触觉传感器：压致变色薄膜在压力作用下光谱发生变化，通过工业相机采集视频图像，从光谱变化反演压力场的时空分布。触觉信息用于灵巧手的精细力控。核心算法挑战是从视频图像的光谱变化估算压力场——这是一个欠定的时序反问题。')

subsection('6.2 YLYW适配方案')
P('触觉传感器是YLYW的理想应用场景，因为信号是时序变化的，完美匹配爻位关系的设计目的。系统架构分为四层：L1逐帧六爻编码（6维触觉特征：全局压力/中心压力/压力梯度/压力集中度/接触稳定性/边界压力）→ L2六十四卦时序匹配（本卦+变卦分析）→ L3爻位关系空间力耦合（乘承比应建模力学交互）→ L3+时序演化分析（物极必反/否极泰来/变卦序列→动作识别）。')

subsection('6.3 爻位关系的真正用武之地')
P('与静态视觉中爻位修正系数≈1.0（几乎无用）形成鲜明对比，在触觉时序信号中：')
P('乘：低压区压制高压区→力学约束异常检测')
P('承：压力从前区传导后区→力传导路径追踪')
P('应（初↔四）：基底压力和集中度呼应→力分布合理性验证')
P('物极必反：爻值>0.95→薄膜即将饱和→减力预警')
P('否极泰来：连续低爻值→即将接触→预紧力准备')
P('每帧都产出可用的力控修正系数，形成感知→控制闭环。')

subsection('6.4 实验设计')
P('四个阶段：阶段一（标定数据采集，4种场景约25,000帧）→阶段二（压力回归+变化检测+接触状态识别，3个实验）→阶段三（分区域6爻编码实现压力场空间恢复+时序超分辨率）→阶段四（灵巧手力控闭环：抓鸡蛋/石头，滑动检测与补偿）。评估指标包括压力MAE<0.1N、压力场PSNR>30dB、状态识别>90%、力控成功率>95%。')

subsection('6.5 产出文件')
files([
    'tactile/实验方案_触觉YLYW.md（方案源文件）',
    'tactile/实验方案_触觉YLYW.docx（Word格式）',
])

# ===== NAMING =====
section('七、命名与术语')
P('时间：2026年6月5日')

P('（1）英文缩写讨论：提议YLFM（易理、模糊、模型），但最终决定保持YLYW不变，暂不修改。')
P('（2）核心术语：八卦原型（Trigram Base）、六爻编码（Yao Encoding）、六十四卦规则（Hexagram Rules）、爻位关系（Yao Relations）——包含五种算子：当位、得中、乘、承、比、应。')
P('（3）视觉域术语：8D简单算子（Simple8DExtractor）、一卦一算子设计原则、零样本最近原型分类。')

# ===== ARCHITECTURE =====
section('八、技术架构总结')

subsection('8.1 三层架构（跨域统一）')
T([
    ['层次', '物理域', '运动控制域', '视觉域', '触觉域（规划）'],
    ['L1八卦原型', '8类物体物理质心', '8卦运动基元', '8个视觉算子', '6维触觉特征'],
    ['L2六爻编码', '13维物理特征加权', '6维状态编码', '8D→6爻映射', '光谱→6爻映射'],
    ['L3卦象匹配', '余弦匹配+策略查表', '余弦匹配+步态查表', '最近原型距离', '余弦匹配+变卦'],
    ['L3+爻位关系', '力修正系数', '步态参数微调', '置信度修正(≈1.0)', '力控修正+空间耦合'],
])

subsection('8.2 爻位关系的域适应性')
P('爻位关系的有效性高度依赖信号域的时序特性：物理域（力修正系数0.94，有效）> 触觉域（预期高有效，待验证）> 运动控制域（步态参数微调，中等有效）> 视觉静态域（0.97，几乎无效）。这从实验上验证了爻位关系的设计本质：处理变化规律而非静态判别。')

# ===== TIMELINE =====
section('九、研究时间线')

T([
    ['日期', '主要进展'],
    ['5月21日', '项目启动，确定三研究方向，首次对话'],
    ['5月21日', '模块化具身智能原理实验台PPT（17页）'],
    ['6月1日', '阶段一完成：64卦全覆盖，92.7%合理率'],
    ['6月1日', '爻位关系运算首次形式化实现'],
    ['6月1日', '完整学术论文v0.4 + git版本管理'],
    ['6月2-3日', '阶段二完成：双八卦安全架构（策略八卦+安全八卦→仲裁层）'],
    ['6月3日', '运动控制仿真完成（PyBullet+MuJoCo）'],
    ['6月3日', '论文v0.6：国内外研究现状重写，30条参考文献'],
    ['6月5日', '视觉分类分支启动：架构设计+特征工程'],
    ['6月5日', 'Brodatz纹理测试：6D/52D/8D多版本对比'],
    ['6月5日', 'STL-10标准基准：37.0%零样本（3×随机基线）'],
    ['6月5日', '视觉论文v2.0完成（4张图嵌入，首行缩进，上标引用）'],
    ['6月5日', '爻位关系跨域分析：静态视觉中作用微弱'],
    ['6月5日', '触觉传感器新方向确立：时序压力场估计'],
    ['6月5日', '触觉实验方案设计完成（4阶段，25K帧，6项指标）'],
])

# ===== FILES =====
section('十、项目文件树')

P('主要目录结构：', bold=True, indent=False)
files([
    'ylyw/',
    '├── paper/                              # 学术论文（v0.1~v0.6, 中英文）',
    '├── experiment_phase1/                  # 阶段一：物理域推理引擎',
    '│   ├── ylyw_core/                      # 核心推理模块（L1~L3+）',
    '│   │   ├── trigram_base.py             # L1: 八卦基元',
    '│   │   ├── yao_encoder.py              # L2: 六爻编码器',
    '│   │   ├── hexagram_rules.py           # L3: 六十四卦规则库',
    '│   │   └── yao_relations.py            # L3+: 爻位关系运算',
    '│   ├── perception/                     # 物理特征提取',
    '│   ├── scripts/                        # 实验与评估脚本',
    '│   └── adapter/                        # 灵犀X2适配层',
    '├── motion_control/                     # 运动控制仿真',
    '│   ├── launch_sim.py                   # PyBullet仿真主程序',
    '│   ├── hexagram_gait_rules.py          # 卦象步态规则',
    '│   └── lingxi/                         # 灵犀X2人形机器人适配',
    '├── vision/                             # 视觉分类分支',
    '│   ├── simple_8d.py                    # 8D简单算子提取器',
    '│   ├── rich_features.py                # 52D丰富特征提取器',
    '│   ├── specialized_detectors_v2.py     # 专用检测器v2',
    '│   ├── feature_extractor_vision.py     # 6D视觉特征提取器',
    '│   ├── classifier.py / classifier_v4.py # 分类器v3/v4',
    '│   ├── test_stl10.py / test_brodatz.py # 标准数据集测试',
    '│   ├── generate_figures.py             # 论文图表生成',
    '│   ├── YLYW视觉论文_v2.0.docx          # 视觉论文',
    '│   └── stl10/                          # STL-10数据集',
    '├── tactile/                            # 触觉传感器方向',
    '│   └── 实验方案_触觉YLYW.docx           # 实验方案',
    '├── 技术路线_实际进展.md                  # 技术路线与进展跟踪',
    '└── arxiv_submission/                   # arXiv投稿准备',
])

# ===== SIGNATURES =====
doc.add_page_break()
section('十一、备忘说明')
P('本备忘录基于2026年5月21日至6月5日期间的会话记录整理生成，涵盖YLYW项目的全部研究方向、实验数据、关键决策和文件产出。所有时间均为北京时间（GMT+8）。')
P('', indent=False)
P('记录人：AI科研助手', indent=False)
P('日期：2026年6月5日', indent=False)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'YLYW研究备忘录.docx')
doc.save(out)
print(f'Saved: {os.path.abspath(out)}')
