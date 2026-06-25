# -*- coding: utf-8 -*-
"""Generate publication-quality docx for submission"""
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from copy import deepcopy

# Load source
src = Document('LLM_YLYW\u53cd\u5e7b\u89c9_\u5b8c\u6574\u7248_\u542b\u77e5\u51e0\u5b66\u4e60.docx')

# Create new doc
doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Style setup
style_normal = doc.styles['Normal']
style_normal.font.name = 'SimSun'
style_normal.font.size = Pt(10.5)
style_normal.paragraph_format.first_line_indent = Cm(0.74)
style_normal.paragraph_format.line_spacing = 1.5

# Collect source content (skip the misplaced zhiji abstract para at index 5)
paragraphs_data = []
for i, p in enumerate(src.paragraphs):
    if i == 5:  # Skip misplaced zhiji paragraph
        continue
    paragraphs_data.append({
        'text': p.text,
        'style': p.style.name,
        'index': i
    })

# New abstract
NEW_ABSTRACT = (
    "\u5927\u8bed\u8a00\u6a21\u578b\uff08LLM\uff09\u7684\u5e7b\u89c9\u95ee\u9898\u5df2\u6210\u4e3a\u5236\u7ea6\u4eba\u5de5\u667a\u80fd"
    "\u53ef\u4fe1\u90e8\u7f72\u7684\u6838\u5fc3\u969c\u788d\u3002\u73b0\u6709\u65b9\u6848\u2014\u2014\u68c0\u7d22\u589e\u5f3a\u751f\u6210"
    "\uff08RAG\uff09\u3001\u81ea\u4e00\u81f4\u6027\u9a8c\u8bc1\uff08SelfCheckGPT\uff09\u3001\u591a\u667a\u80fd\u4f53\u4e92\u5ba1"
    "\uff08LLM-as-Judge\uff09\u7b49\u2014\u2014\u5747\u672a\u8df3\u51fa\u201c\u4ee5LLM\u5ba1\u67e5LLM\u201d\u7684\u903b\u8f91"
    "\u95ed\u73af\uff0c\u5ba1\u67e5\u8005\u4e0e\u88ab\u5ba1\u67e5\u8005\u5171\u4eab\u76f8\u540c\u7684\u8ba4\u77e5\u504f\u5dee\u3002"
    "\u672c\u6587\u63d0\u51fa\u201c\u751f\u6210-\u5ba1\u67e5\u5206\u79bb\u201d\u8303\u5f0f\uff0c\u4e3b\u5f20\u5e7b\u89c9\u5ba1\u67e5"
    "\u5fc5\u987b\u7531\u72ec\u7acb\u4e8eLLM\u7684\u5f02\u6784\u7cfb\u7edf\u5b8c\u6210\u3002\u57fa\u4e8e\u8fd9\u4e00\u8303\u5f0f\uff0c"
    "\u6784\u5efaYLYW\u72ec\u7acb\u5ba1\u67e5\u5f15\u64ce\uff0c\u91c7\u7528\u201c\u9053\u6cd5\u672f\u5668\u201d\u56db\u5c42\u4f53\u7cfb"
    "\u7ec4\u7ec7\u67b6\u6784\uff1a\u4ee5\u201c\u9053\u201d\u786e\u7acb\u4e09\u5c42\u9012\u8fdb\u5ba1\u67e5\u903b\u8f91"
    "\uff08\u4e8b\u5b9e\u2192\u7269\u7406\u2192\u4ef7\u503c\uff09\u4e0e\u201c\u8d8b\u5409\u907f\u51f6\u201d\u4ef7\u503c\u6839\u57fa\uff1b"
    "\u4ee5\u201c\u6cd5\u201d\u89c4\u5b9a\u56db\u7ea7\u5224\u5b9a\uff08\u7ea2/\u9ec4/\u84dd/\u7eff\uff09\u4e0e\u95ed\u73af\u4fee\u6b63"
    "\u65b9\u6cd5\u8bba\uff1b\u4ee5\u201c\u672f\u201d\u6784\u5efa\u53ef\u6269\u5145\u7684\u4e09\u5c42\u89c4\u5219\u77e5\u8bc6\u5e93"
    "\uff08\u4e8b\u5b9e30+\u5b9e\u4f53\u300114\u6761\u7269\u7406\u7ea6\u675f\u300115\u6761\u4ef7\u503c\u89c4\u5219\uff09\uff1b"
    "\u4ee5\u201c\u5668\u201d\u5b9e\u73b0\u96f6\u4f9d\u8d56\u7eaf\u89c4\u5219\u5f15\u64ce\uff08717\u884cPython\uff0c<1ms\uff0c"
    "\u96f6API\uff09\u3002\u66f4\u91cd\u8981\u7684\u662f\uff0c\u672c\u6587\u63d0\u51fa\u77e5\u51e0\u5b66\u4e60\uff08Zhiji Learning\uff09"
    "\u8d4b\u4e88\u5ba1\u67e5\u5f15\u64ce\u81ea\u8fdb\u5316\u80fd\u529b\uff1a\u901a\u8fc7\u201c\u5409\u4e4b\u51e0\u201d\u5f3a\u5316"
    "\u6709\u6548\u89c4\u5219\u3001\u201c\u51f6\u4e4b\u51e0\u201d\u7cbe\u786e\u5f52\u56e0\u5e76\u81ea\u52a8\u65b0\u589e\u89c4\u5219\uff0c"
    "\u5b9e\u73b0K=K_prior\u2295K_calibration\u7684\u6301\u7eed\u77e5\u8bc6\u79ef\u7d2f\u3002\u5b9e\u9a8c\u8868\u660e\uff0c"
    "\u7cfb\u7edf\u572812\u4e2a\u5178\u578b\u7528\u4f8b\u4e0a\u5b9e\u73b083.3%\u68c0\u51fa\u7387\u30017\u4e2a\u4e25\u91cd\u5e7b\u89c9"
    "\u5168\u90e8\u62e6\u622a\u3001\u96f6\u5047\u9633\u6027\uff1b\u77e5\u51e0\u5b66\u4e60\u4f7f\u7cfb\u7edf\u4ec51\u6b21\u4eba\u7c7b"
    "\u53cd\u9988\u5373\u53ef\u4ece\u6f0f\u653e\u6536\u655b\u4e3a\u6b63\u786e\u68c0\u51fa\u3002\u672c\u6587\u6838\u5fc3\u8bba\u65ad\uff1a"
    "\u5ba1\u67e5\u8005\u4e0d\u80fd\u662f\u88ab\u5ba1\u67e5\u8005\u7684\u540c\u7c7b\u2014\u2014\u8fd9\u5e94\u6210\u4e3aAI\u5b89\u5168"
    "\u7684\u57fa\u7840\u516c\u7406\u3002"
)

KEYWORDS_CN = "\u5927\u8bed\u8a00\u6a21\u578b\uff1b\u5e7b\u89c9\u68c0\u6d4b\uff1b\u72ec\u7acb\u5ba1\u67e5\uff1b\u751f\u6210-\u5ba1\u67e5\u5206\u79bb\uff1bYLYW\uff1b\u9053\u6cd5\u672f\u5668\uff1b\u77e5\u51e0\u5b66\u4e60\uff1b\u81ea\u8fdb\u5316\u5ba1\u67e5\uff1bAI\u5b89\u5168"

# Table captions
TABLE_NAMES = [
    "\u88681 \u56db\u7ea7\u5224\u5b9a\u673a\u5236\u4e0e\u5904\u7406\u52a8\u4f5c",
    "\u88682 \u7aef\u5230\u7aef\u6d4b\u8bd5\u7528\u4f8b\u8bbe\u8ba1",
    "\u88683 \u6d4b\u8bd5\u7ed3\u679c\u6c47\u603b",
    "\u88684 \u5ba1\u67e5\u5f15\u64ce\u6027\u80fd\u6307\u6807",
    "\u88685 \u77e5\u51e0\u5b66\u4e60\u81ea\u8fdb\u5316\u5b9e\u9a8c\u7ed3\u679c",
    "\u88686 \u4e0e\u73b0\u6709\u65b9\u6cd5\u7684\u7cfb\u7edf\u5bf9\u6bd4",
]

# Figure insertion points (section keywords to insert after)
FIGURES = {
    '\u201c\u9053\u6cd5\u672f\u5668\u201d\u56db\u5c42\u4f53\u7cfb': ('\u56fe1 \u57fa\u4e8e\u201c\u9053\u6cd5\u672f\u5668\u201d\u4f53\u7cfb\u7684LLM+YLYW\u53cd\u5e7b\u89c9\u7cfb\u7edf\u603b\u4f53\u67b6\u6784', '[\u56fe1\u4f4d\u7f6e]'),
    '\u6cd5\u7684\u6f14\u5316\uff1a\u77e5\u51e0\u5b66\u4e60': ('\u56fe2 \u77e5\u51e0\u5b66\u4e60\u5bf9\u79f0\u6821\u51c6\u95ed\u73af', '[\u56fe2\u4f4d\u7f6e]'),
    '\u77e5\u51e0\u5b66\u4e60\u81ea\u8fdb\u5316\u5b9e\u9a8c': ('\u56fe3 \u9759\u6001\u7cfb\u7edf\u4e0e\u77e5\u51e0\u5b66\u4e60\u7cfb\u7edf\u7684\u5bf9\u6bd4', '[\u56fe3\u4f4d\u7f6e]'),
}

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.first_line_indent = Cm(0)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    # Make references superscript
    if re.search(r'\[\d', text):
        # Clear and rebuild with superscript refs
        p.clear()
        parts = re.split(r'(\[\d[\d,\-]*\])', text)
        for part in parts:
            if re.match(r'\[\d[\d,\-]*\]', part):
                run = p.add_run(part)
                run.font.superscript = True
                run.font.size = Pt(8)
            else:
                run = p.add_run(part)
                run.font.size = Pt(10.5)
    return p

def add_table_caption(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(9)
    return p

def add_figure_placeholder(doc, caption, placeholder):
    # Placeholder
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.first_line_indent = Cm(0)
    run = p1.add_run(placeholder)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    # Caption
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Cm(0)
    run = p2.add_run(caption)
    run.bold = True
    run.font.size = Pt(9)

# ============================================================
# BUILD DOCUMENT
# ============================================================
# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('\u57fa\u4e8e\u201c\u9053\u6cd5\u672f\u5668\u201d\u4f53\u7cfb\u7684\u5927\u8bed\u8a00\u6a21\u578b\u5e7b\u89c9\u5ba1\u67e5\u6846\u67b6\uff1aYLYW\u72ec\u7acb\u5ba1\u67e5\u5f15\u64ce')
run.bold = True
run.font.size = Pt(16)

# Authors
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('\u9a6c\u5174\u5f55*\uff0c\u674e\u91d1\u51fd\uff0c\u5f20\u56fd\u5b89\uff0c\u4e8e\u656c\u6d9b\uff0c\u674e\u671b\uff0c\u9a6c\u5723\u6d01')
run.font.size = Pt(10.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('\uff08\u9752\u5c9b\u79d1\u6280\u5927\u5b66 \u81ea\u52a8\u5316\u4e0e\u7535\u5b50\u5de5\u7a0b\u5b66\u9662\uff0c\u5c71\u4e1c \u9752\u5c9b 266061\uff09')
run.font.size = Pt(9)

# Abstract
add_heading(doc, '\u6458\u8981', level=1)
add_body(doc, NEW_ABSTRACT)

# Keywords
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
run = p.add_run('\u5173\u952e\u8bcd\uff1a')
run.bold = True
run.font.size = Pt(10.5)
run = p.add_run(KEYWORDS_CN)
run.font.size = Pt(10.5)

# Process rest of source content
table_idx = 0
figure_inserted = set()
skip_until_next_heading = False
in_abstract = False

for pd in paragraphs_data:
    text = pd['text']
    style = pd['style']
    
    # Skip original abstract section (we already wrote new one)
    if 'Heading' in style and '\u6458\u8981' in text and 'Abstract' not in text:
        in_abstract = True
        continue
    if in_abstract:
        if 'Heading' in style and ('Abstract' in text or '1' in text):
            in_abstract = False
            if 'Abstract' in text:
                # Skip English abstract for now (or include it)
                add_heading(doc, 'Abstract', level=1)
                continue
        else:
            continue  # Skip old abstract content
    
    # Skip English abstract keywords line (already have new keywords)  
    if text.startswith('Keywords:') or text.startswith('\u5173\u952e\u8bcd'):
        continue
    
    if 'Heading' in style:
        level_map = {'Heading 0': 0, 'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3}
        level = 2  # default
        for k, v in level_map.items():
            if style.startswith(k) or style == k:
                level = v
                break
        if 'Heading 1' in style:
            level = 1
        elif 'Heading 2' in style:
            level = 2
        elif 'Heading 3' in style:
            level = 3
            
        h = add_heading(doc, text, level=min(level, 3))
        
        # Check if we need to insert a figure after this heading
        for key, (fig_caption, fig_placeholder) in FIGURES.items():
            if key in text and key not in figure_inserted:
                figure_inserted.add(key)
                add_figure_placeholder(doc, fig_caption, fig_placeholder)
                break
    else:
        if text.strip():
            add_body(doc, text)

# Now add tables from source at appropriate positions
# (Tables are complex to position perfectly; for now add a note)

# Copy tables from source
for ti, table in enumerate(src.tables):
    if ti < len(TABLE_NAMES):
        add_table_caption(doc, TABLE_NAMES[ti])
    
    rows = len(table.rows)
    cols = len(table.columns)
    new_table = doc.add_table(rows=rows, cols=cols)
    new_table.style = 'Table Grid'
    new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for ri in range(rows):
        for ci in range(cols):
            cell_text = table.rows[ri].cells[ci].text
            new_table.rows[ri].cells[ci].text = cell_text
    
    # Add spacing after table
    doc.add_paragraph()

# Save
output = 'LLM_YLYW\u53cd\u5e7b\u89c9_\u6295\u7a3f\u7248.docx'
doc.save(output)

# Verify
doc2 = Document(output)
total_chars = sum(len(p.text) for p in doc2.paragraphs)
print(f"Generated: {output}")
print(f"  Paragraphs: {len(doc2.paragraphs)}")
print(f"  Tables: {len(doc2.tables)}")
print(f"  Characters: {total_chars}")
import os
print(f"  Size: {os.path.getsize(output)/1024:.1f} KB")
