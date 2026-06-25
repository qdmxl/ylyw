#!/usr/bin/env python3
"""将论文Markdown转换为规范排版docx"""
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

SRC = "/home/lijinhan/MXL/科研/ylyw/anti_hallucination/paper_and_code/LLM_YLYW反幻觉混合系统_技术论文.md"
DST = "/home/lijinhan/MXL/科研/ylyw/anti_hallucination/paper_and_code/LLM_YLYW反幻觉混合系统_技术论文.docx"

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── 样式定义 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    rPr = style.element.get_or_add_rPr()
    rPr_elem = rPr
else:
    rPr_elem = style.element.rPr
rPr_elem.get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

# 标题样式
for i in range(1, 4):
    sname = f'Heading {i}'
    if sname in doc.styles:
        hstyle = doc.styles[sname]
    else:
        hstyle = doc.styles.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)
    hfont = hstyle.font
    hfont.name = '黑体'
    hfont.bold = True
    hfont.color.rgb = RGBColor(0, 0, 0)
    hstyle.paragraph_format.first_line_indent = Cm(0)
    hstyle.paragraph_format.space_before = Pt(12)
    hstyle.paragraph_format.space_after = Pt(6)
    hrPr = hstyle.element.get_or_add_rPr()
    hrPr.get_or_add_rFonts().set(qn('w:eastAsia'), '黑体')
    if i == 1:
        hfont.size = Pt(16)
        hstyle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif i == 2:
        hfont.size = Pt(14)
    else:
        hfont.size = Pt(13)

# ── 解析Markdown ──
with open(SRC, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code_block = False
code_buffer = []
in_table = False
table_rows = []
current_list_items = []
skip_until_next = False

def flush_list():
    global current_list_items
    for item in current_list_items:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        # 处理粗体
        parts = re.split(r'(\*\*.*?\*\*)', item)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
    current_list_items = []

def flush_code_block():
    if code_buffer:
        for line in code_buffer:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.74)
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        doc.add_paragraph()  # 空行
    code_buffer.clear()

def add_table_to_doc(rows):
    if not rows:
        return
    # 解析表头和对齐
    header = rows[0]
    align = rows[1] if len(rows) > 1 else None
    data_start = 2 if align and all(c in ':|- ' for c in align) else 1
    data_rows = rows[data_start:] if align and all(c in ':|- ' for c in align) else rows[1:]
    
    cols = len(header)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Table Grid'
    table.autofit = True
    
    # 表头
    for ci, cell_text in enumerate(header):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 数据行
    for ri, row_cells in enumerate(data_rows):
        for ci, cell_text in enumerate(row_cells):
            if ci < cols:
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text.strip())
                run.font.size = Pt(10)
                run.font.name = '宋体'
                p.paragraph_format.first_line_indent = Cm(0)
    
    doc.add_paragraph()  # 表后空行

while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    
    # 跳过YAML front matter / ---
    if stripped == '---' and i == 0:
        i += 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1
        continue
    
    if stripped == '---' and not in_code_block:
        i += 1
        continue
    
    # 代码块
    if stripped.startswith('```'):
        if in_code_block:
            flush_code_block()
            in_code_block = False
        else:
            flush_list()
            in_code_block = True
        i += 1
        continue
    
    if in_code_block:
        code_buffer.append(line.rstrip())
        i += 1
        continue
    
    # 表格
    if stripped.startswith('|') and stripped.endswith('|'):
        flush_list()
        if not in_table:
            in_table = True
            table_rows = []
        cells = [c.strip() for c in stripped[1:-1].split('|')]
        table_rows.append(cells)
        i += 1
        continue
    else:
        if in_table:
            add_table_to_doc(table_rows)
            table_rows = []
            in_table = False
    
    # 标题
    if stripped.startswith('# ') and not stripped.startswith('## '):
        flush_list()
        title_text = stripped[2:].strip()
        p = doc.add_paragraph()
        p.style = doc.styles['Heading 1']
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        i += 1
        continue
    
    if stripped.startswith('## ') and not stripped.startswith('### '):
        flush_list()
        title_text = stripped[3:].strip()
        p = doc.add_paragraph()
        p.style = doc.styles['Heading 2']
        run = p.add_run(title_text)
        i += 1
        continue
    
    if stripped.startswith('### ') and not stripped.startswith('#### '):
        flush_list()
        title_text = stripped[4:].strip()
        p = doc.add_paragraph()
        p.style = doc.styles['Heading 3']
        run = p.add_run(title_text)
        i += 1
        continue
    
    if stripped.startswith('#### '):
        flush_list()
        title_text = stripped[5:].strip()
        p = doc.add_paragraph()
        p.style = doc.styles['Heading 3']
        run = p.add_run(title_text)
        run.font.size = Pt(12)
        i += 1
        continue
    
    # 有序列表
    ol_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
    if ol_match:
        flush_list()
        item_text = ol_match.group(2)
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        # 处理子列表（如 - **xxx**）
        sub_match = re.match(r'^-\s+\*\*(.+?)\*\*[：:]?\s*(.*)$', item_text)
        if sub_match:
            run = p.add_run(f'{ol_match.group(1)}. ')
            run = p.add_run(sub_match.group(1))
            run.bold = True
            if sub_match.group(2):
                p.add_run(f'：{sub_match.group(2)}')
        else:
            parts = re.split(r'(\*\*.*?\*\*)', item_text)
            p.add_run(f'{ol_match.group(1)}. ')
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        i += 1
        continue
    
    # 无序列表 / 子列表
    ul_match = re.match(r'^(\s*)-\s+(.+)$', stripped)
    if ul_match:
        indent_level = len(ul_match.group(1))
        item_text = ul_match.group(2)
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74 + indent_level * 0.5)
        parts = re.split(r'(\*\*.*?\*\*)', item_text)
        p.add_run('• ')
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1
        continue
    
    # 引用块 >
    if stripped.startswith('>'):
        flush_list()
        quote_text = stripped[1:].strip()
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(quote_text)
        run.italic = True
        run.font.size = Pt(11)
        i += 1
        continue
    
    # 作者信息（**马兴录**）
    if stripped.startswith('**') and ('马兴录' in stripped or '青岛科技' in stripped):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(stripped.replace('**', ''))
        run.font.size = Pt(12)
        i += 1
        continue
    
    # 摘要/关键词粗体标签行
    if stripped.startswith('**摘要**') or stripped.startswith('**关键词**'):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1
        continue
    
    # 参考文献
    if stripped.startswith('[') and re.match(r'^\[\d+\]', stripped):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        run = p.add_run(stripped)
        run.font.size = Pt(10.5)
        i += 1
        continue
    
    # 收稿日期/基金
    if stripped.startswith('*收稿日期') or stripped.startswith('*基金项目'):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(stripped.replace('*', '').strip())
        run.italic = True
        run.font.size = Pt(10.5)
        i += 1
        continue
    
    # 空行
    if not stripped:
        flush_list()
        i += 1
        continue
    
    # 普通段落（含粗体处理）
    flush_list()
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    
    # 处理行内粗体
    parts = re.split(r'(\*\*.*?\*\*)', stripped)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            # 处理行内代码 `...`
            code_parts = re.split(r'(`.*?`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = p.add_run(cp[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    p.add_run(cp)
    i += 1

# 清理
flush_list()
flush_code_block()
if in_table:
    add_table_to_doc(table_rows)

# ── 保存 ──
doc.save(DST)
print(f"✅ 已生成: {DST}")
