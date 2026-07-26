#!/usr/bin/env python3
"""一次性生成完整的YLYW技术论文（无版本演进，纯架构+流程）"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

for sec in doc.sections:
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)

sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'; sty.font.size = Pt(10)
sty.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
sty.paragraph_format.line_spacing = 1.15
sty.paragraph_format.space_before = Pt(0); sty.paragraph_format.space_after = Pt(2)

def sr(run, name='宋体', size=10, bold=False):
    run.font.name = name; run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size); run.bold = bold

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        sr(r, '黑体', 14 if level==1 else 12, True)
        r.font.color.rgb = RGBColor(0,0,0)
    hd.paragraph_format.first_line_indent = Cm(0)
    hd.paragraph_format.space_before = Pt(12 if level==1 else 8)
    hd.paragraph_format.space_after = Pt(4)

def P(text, bold=False, size=10):
    par = doc.add_paragraph(); run = par.add_run(text)
    sr(run, size=size, bold=bold)
    par.paragraph_format.first_line_indent = Cm(0.75)
    return par

def Pn(text):
    par = doc.add_paragraph(); run = par.add_run(text); sr(run)
    return par

def tbl(headers, rows, caption, label):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_before = Pt(6); cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(f"表 {label}: {caption}"); sr(r, '宋体', 10, True)
    # flatten headers if nested (called with [('a','b')] instead of ['a','b'])
    flat_hdrs = headers[0] if headers and isinstance(headers[0], (list, tuple)) else headers
    t = doc.add_table(rows=1+len(rows), cols=len(flat_hdrs))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,hdr in enumerate(flat_hdrs):
        c = t.rows[0].cells[i]; c.text = hdr
        for pp in c.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER; pp.paragraph_format.first_line_indent = Cm(0)
            for run in pp.runs: sr(run, '宋体', 9, True)
    for ri,row in enumerate(rows):
        flat_row = row if isinstance(row, (list, tuple)) else [row]
        for ci,val in enumerate(flat_row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for pp in c.paragraphs:
                pp.paragraph_format.first_line_indent = Cm(0)
                for run in pp.runs: sr(run, '宋体', 9)

def img(path, caption, label, w=5.0):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_before = Pt(10); cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(f"图 {label}: {caption}"); sr(r, '宋体', 10, True)
    pic = doc.add_picture(path, width=Inches(w))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER; last.paragraph_format.first_line_indent = Cm(0)

def eqn(text, label=None):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER; par.paragraph_format.first_line_indent = Cm(0)
    par.paragraph_format.space_before = Pt(4); par.paragraph_format.space_after = Pt(4)
    run = par.add_run(text); run.font.name = 'Times New Roman'; run.font.size = Pt(10); run.italic = True
    if label:
        run2 = par.add_run(f"    ({label})")
        run2.font.name = 'Times New Roman'; run2.font.size = Pt(10)

FIG = "/home/lijinhan/MXL/科研/ylyw/figures"

# ═══ 标题 ═══
t = doc.add_heading('YLYW: 基于汉字易理模型的认知推理框架', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(4); t.paragraph_format.first_line_indent = Cm(0)
for r in t.runs:
    r.font.name = '黑体'; r.font.size = Pt(22)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER; auth.paragraph_format.first_line_indent = Cm(0)
r = auth.add_run('马兴录  青岛科技大学信息科学技术学院')
sr(r, '宋体', 12)
doc.add_paragraph()

# ═══ 摘要 ═══
h('摘要')
P('YLYW（易理模糊模型）是一个基于汉字易经六十四卦体系构建的认知推理框架。'
  '其核心思想是：汉字的部首结构携带天然的八卦类象信息，通过《说卦传》万物类象分类体系，'
  '可将自然语言概念映射到8维八卦隶属度空间和6维六爻编码空间，'
  '进而通过64卦模板匹配实现情境理解和决策推理。')
P('YLYW框架包含四个功能部件：(1) 汉字易理引擎（HanziEngine），对汉字执行'
  '部首分解→卦象模糊映射→乘承比应推理→六爻编码→64卦匹配的完整推理链路；'
  '(2) 万物类象知识库（GuaKnowledgeBase），融合156条手工先验和运行时学习，'
  '提供词义→卦象的持久化映射；(3) 情境建模组件（Situation Modeler），'
  '通过对环境中所有实体的卦象聚合得到情境六爻，匹配64卦获得当前情境卦象；'
  '(4) 决策推理组件（YLYW Scorer），基于6维六爻编码和64卦吉凶评分进行动作选择。')
P('在ALFWorld具身任务基准上（共274局），YLYW框架达到95.3%的总体成功率'
  '（valid_seen 140局 92.1% + valid_unseen 134局 98.5%），'
  '远超已有非LLM方法，且具备完全确定性、可审计、零GPU部署的特性。')

# ═══ 1. 引言 ═══
h('1 引言')
P('易经（《周易》）作为中国古代最重要的哲学典籍之一，构建了一个以'
  '阴阳爻为基础、八卦为中介、六十四卦为完整状态空间的符号推理系统¹⁰。'
  '《说卦传》进一步建立了"万物类象"的分类体系——将天地万物按八卦归类¹¹。'
  '这一体系本质上提供了一种从具体事物到抽象卦象的形式化映射。')
P('现代具身AI面临的核心挑战之一，是智能体如何从自然语言感知中'
  '构建结构化的世界状态，并在此状态下做出合理决策。'
  '主流的LLM方案虽然有效，但存在高计算成本、不可复现、'
  '缺乏可解释性等固有问题³⁻⁵。')
P('YLYW框架首次将汉语汉字纳入易经推理框架：每个汉字本身就是一种数字孪生，'
  '其部首结构包含丰富的八卦类象信息。通过对汉字做部首分解→卦象推理→'
  '六爻编码→64卦匹配，YLYW构建了一条从自然语言到易经卦象的完整计算链路。')
P('本文的组织结构如下：第2节给出系统总体架构的三层总览；'
  '第3节详述汉字易理推理引擎（HanziEngine）的完整功能；'
  '第4节介绍万物类象知识库与情境建模组件；'
  '第5节描述决策推理引擎的评分与过滤机制；'
  '第6节展示单步决策的完整数据流；'
  '第7节给出ALFWorld上的实验结果和分析；'
  '第8节讨论与展望。')

# ═══════════════════════════════════════════
# 2. 系统总体架构
# ═══════════════════════════════════════════
h('2 系统总体架构')

P('YLYW系统采用三层架构：底层是汉字易理推理引擎（HanziEngine），'
  '提供"汉字→卦象"的基础计算能力；中间层是认知功能部件层，'
  '包括万物类象知识库（GuaKnowledgeBase）和情境建模组件（Situation Modeler）；'
  '顶层是决策推理层（YLYW Scorer + Veto过滤），基于卦象语义对全部合法动作评分选择。'
  '系统三层均以Python实现，总代码约1540行，无外部LLM或GPU依赖。架构如图1所示。')

img(f"{FIG}/fig1_architecture.png", "YLYW系统三层架构", "1", 5.5)

h('2.1 三层总览', 2)
P('系统从底到顶分为三个层次，每层为一个Python模块（文件）：')
P('  - 底层：HanziEngine（~1470行），汉字→卦象推理，通过字→词→句三层递归将自然语言编码为6维六爻。')
P('  - 中层：GuaKnowledgeBase（~260行）+ WorldModel/CnWorldModel（~140行），提供实体卦象持久化和情境聚合。')
P('  - 顶层：Agent + YLYWScorer（~90行），接收环境合法动作，逐候选六爻评分+64卦匹配+Veto过滤。')
P('层间数据流：底层输出八卦隶属度→中层聚合为情境卦象→顶层在候选动作上做六爻评分。'
  '汉字→卦象推理路径（字→词→句）和图1的Python类组织（HanziEngine主体+YLYWLayer乘承比应计算）'
  '在第3章中展开，本章仅给出系统总览。')

img(f"{FIG}/fig1_architecture.png", "YLYW系统三层架构与数据流", "1", 5.5)

h('2.2 模块组织与代码规模', 2)
P('YLYW v20由以下4个Python文件组成：')
P('  - agent_v20.py（~720行）：决策逻辑、失败检测、自动重试、八步执行循环。')
P('  - cn_world_model.py（~140行）：汉字版世界模型，封装实体状态追踪。')
P('  - gua_knowledge_base.py（~260行）：知识库单例，先验+运行时学习。')
P('  - run_v20_eval.py（~80行）：评测入口、结果汇总。')
P('加上HanziEngine（~1470行）, 总代码约2700行。所有计算在单CPU、2GB RAM、无GPU环境完成。')











# ═══ 3. 汉字易理推理引擎 ═══
h('3 汉字易理推理引擎')

P('HanziEngine（~1470行）是YLYW系统的推理基础设施。它对输入的汉字执行从部首到卦象的完整推理链路，'
  '输出八卦隶属度向量（8维）和六爻编码（6维）。推理过程分为三个递归层级。')

img(f"{FIG}/fig2_pipeline.png", "汉字→卦象完整推理链路：字→词→句三层递归", "2", 5.5)

h('3.1 L0字级：char()', 2)
P('char()是最高优先级的推理路径。输入单个汉字，执行以下四步：')
P('步骤1 — 部首分解。使用汉字分解数据库HANZI_DECOMPOSITION（约3000条记录）将汉字拆解为部首组件。'
  '例如"橱"→"木"+"厨"，"柜"→"木"+"巨"。对于多部首字，所有有效部首都会被提取。')
P('步骤2 — 部首→卦象模糊映射。查询部首→卦象模糊映射表（radical_fuzzy_base.json），'
  '获取每个部首的8维八卦隶属度向量。该映射表按《说卦传》万物类象分类构建。'
  '"木"部首在《说卦传》中属巽卦（风、木、入），其隶属度向量的巽分量最高；'
  '"金/钅"部首属乾卦（天、刚健、金属）；"水/氵"部首属坎卦（水、寒、险）。')
P('步骤3 — 卦象加权平均。对多个部首的卦象向量进行加权平均，得到汉字的综合八卦隶属度。'
  '权值由部首在汉字中的主次地位决定。使用NumPy的mean函数计算平均向量。')
P('步骤4 — 主导卦象选择。取隶属度最高的八卦作为主导卦象：dom_idx = argmax(avg_vector)，'
  '映射到["乾","兑","离","震","巽","坎","艮","坤"]。')
P('输出：vector（8维八卦隶属度）、yao_vector（6维六爻编码）、dominant（主导卦象名）、'
  'hex64（64维卦象匹配分布，供句级推理使用）。')

h('3.2 L2词级：word()', 2)
P('对于多字词语（如"台灯"、"微波炉"），先对每个字单独做L0推理，'
  '然后调用YLYWLayer.perceive_and_encode()计算字间乘承比应关系。'
  '乘承比应是《周易》爻位关系理论的核心理念，在YLYW中被形式化为四个计算规则：')
P('乘（riding）：前动作字压后物体字，表现为卦象向量差异的衰减。'
  '例如"洗"（坎卦主导）+"碗"（艮卦主导），水的冲刷力对静止物体的卦象产生抑制，'
  '使词级结果偏向坎卦主导，艮卦分量被压制。')
P('承（bearing）：前物体字支撑后动作字，同向分量增强。'
  '例如"刀"（震卦主导）+"切"（震卦主导），动作同向共振，震卦分量获得双倍增强。')
P('比（comparing）：同角色相邻字间的谐同性判断，同向率>0.6标记为谐同。'
  '例如"干净"中"干"和"净"同为离卦（火→清洁）属性，互相比谐。')
P('应（responding）：首尾字或同卦相隔字的远程呼应，余弦相似度>0.7标记为有应。'
  '例如"冰"（坎）+"箱"（艮）中坎卦在前后字中的残影形成"应"关系。')
P('六爻构造算法（_yao_from_relations）采用交替混合深度法：'
  '6个爻位的动作/物体权重呈锯齿分布（动作权重[0.60,0.35,0.55,0.40,0.50,0.45]），'
  '确保相同动作+不同物体在每爻产生差异，避免单调递增导致三爻相同。'
  '乘/承关系对爻位产生方向性调制（cheng_factor叠加）。'
  '输出：bagua（8维）、yao_vector（6维）、hex64（64维）、relations（乘承比应列表）。')

h('3.3 L3句级：sentence()', 2)
P('sentence()对完整汉语句子执行三步推理：')
P('(1) _segment()分词。采用正向最大匹配策略，结合动词/名词字典和功能词强制分割。'
  '优先匹配已知多字词（known_phrases约50条、alfworld_phrases约80条、通用双字词约80条），'
  '功能词（把/被/在/的/了/着等约40个）作为强制分割点。返回词序列。')
P('(2) _guess_role()角色标注。通过卦象+词典双重判断每个词的语义角色。'
  '动作字表约120个（打开/关闭/清洗/加热/冷却/放置…），'
  '名词结尾字表约70个（机/炉/箱/器/杯/碗/盘/桌/椅…），'
  '功能词约40个，补语约20个（上/下/进/出/起/开…）。'
  '输出角色为"动作/物体/状态/虚词/补语"之一。')
P('(3) YLYWLayer.perceive_and_encode()。词序列的八卦隶属度作为句级YLYWLayer的输入，'
  '计算词间的乘承比应关系。跨虚词连接自动处理：'
  '例如"把苹果放在微波炉里"中，"把"两侧的"放"（动作）和"苹果"（物体）建立跨虚词"乘"关系。'
  '输出：句级六爻编码、64卦分布、互卦关系列表。')

h('3.4 卦象去偏机制', 2)
P('汉字部首隶属度存在偏态分布：某些部首在所有八卦上的能量极集中'
  '（如"水/氵"部首在坎卦趋近1.0，其他卦接近0），导致平均后向量的区分度严重不足。'
  'perceive_and_encode()的第一步执行去偏处理：')
P('对每个元素的8维向量$"v"$，每维独立拉伸到$"[0.05, 0.95]"$区间：')
P('v\'_i = 0.05 + (v_i - min) / (max - min) × 0.9')
P('这一操作消除偏态但保留排名信息——即原来哪个卦隶属度最高在去偏后仍然最高。')

h('3.5 依存关系解析（卦象驱动版）', 2)
P('HanziEngine._parse_dependencies()实现基于卦象的依存关系解析，不依赖任何硬编码词表。'
  '八卦的天然语义直接用于判断修饰关系：')
P('  - 艮（止/藏）、坤（载/容）→ 位置/容器语义 → loc_mod（位置修饰）'
  '    例如"桌子"（坤）在"桌子上的苹果"中作为位置修饰语修饰"苹果"。')
P('  - 坎（水/寒）、离（火/明）→ 属性/状态语义 → attr_mod（属性修饰）'
  '    例如"冷"（坎）在"冷水"中作为属性修饰"水"，"干净"（离）在"干净的盘子"中作为属性修饰。')
P('  - 震（动）、乾（健）、兑（悦）→ 动作性 → 不可能是修饰语。')
P('判定规则有三条：')
P('规则1 — "的"连接。相邻物体间通过"的"连接时，前一个字的卦象决定修饰关系类型。'
  '前字是艮/坤→loc_mod，前字是坎/离→attr_mod。')
P('规则2 — 跨距离修饰。通过互卦关系中的"应"识别远程修饰：'
  '两个物体间如果有"应"关系且前一个卦象偏向位置/属性，判断为修饰关系。')
P('规则3 — 方位后缀。以"边/里/上/旁/附近"结尾的词自动触发位置修饰。'
  '例如"水槽边"（坎卦带后缀）→ loc_mod修饰"的碗"。')
P('输出：[{modifier, head, type(loc_mod/attr_mod), marker, reason}, ...] 列表。')

h('3.6 时序解析', 2)
P('HanziEngine.parse_temporal()从中文任务描述中解析动作的时序关系，'
  '完全基于虚词标记识别动作先后顺序，不受动词在句中出现位置的影响。支持的时序模式：')
P('  - "先…再…"（标准顺序）→ 先出现的动词在前，"再"后的动词在后。'
  '    例："先把苹果洗干净再放进冰箱" → actions_ordered=["洗","放"]。')
P('  - "…后/…之后"（后序结构）→ "后"标记前的动词在前。'
  '    例："洗完后把苹果放进冰箱" → actions_ordered=["洗","放"]。')
P('  - "…之前"（倒装结构）→ "之前"前的动词在"之后"执行。'
  '    例："放进冰箱之前先洗干净" → actions_ordered=["洗","放"]。')
P('  - "为了…"（隐式倒装）→ "为了"后的动词在前。'
  '    例："为了加热先解冻" → actions_ordered=["解冻","加热"]。')
P('输出：actions_ordered（按执行顺序排列的动作列表）、markers（标记词及类型）、'
  'order_type（"顺序/倒装/并列/条件"）。')

h('3.7 解释性输出与高层接口', 2)
P('HanziEngine.explain(text)将汉字→卦象的完整推理链路转为人类可读文本。'
  '例如explain("洗")的输出为：'
  '"洗 [水部首(坎0.95)] → 坎卦主导(隶属度0.88) → 六爻[0.85,0.73,0.68,0.61,0.55,0.58] → 匹配水雷屯卦(相似度0.82)"。')
P('to_ylyw(text)提供统一入口，单汉字→char()，多字→word()，含空格/标点→sentence()。'
  'similarity(a, b)计算任意两段中文在卦象空间的距离，与领域无关。')

h('3.8 全局LRU推理缓存', 2)
P('汉字→卦象推理是无状态的——同一汉字始终产生相同的八卦隶属度。'
  '系统维护三层@lru_cache(maxsize=4096)：')
P('(1) bagua_for_hanzi_cached(hanzi)：汉字→八卦隶属度的底层缓存，返回tuple（可哈希）。')
P('(2) bagua_for_hanzi(hanzi)：对外接口，转调cached版本，将tuple转回list返回。')
P('(3) yao_for_entity_cached(en_name)：英文实体名→六爻编码的缓存。')
P('在ALFWorld测试中，缓存命中率约78%，系统整体推理速度提升约44%'
  '（从平均3.66秒/局降至2.54秒/局）。')

# ═══ 4. 万物类象知识库与情境建模 ═══
h('4 万物类象知识库与情境建模')

h('4.1 GuaKnowledgeBase 万物类象知识库', 2)
P('知识库（~260行代码）采用单例模式全局共享，是系统的"语义词典"，提供词义→卦象的持久化映射。'
  '知识库包含两个数据源：')
P('（一）手工先验表_handcrafted（156条）：按《说卦传》万物类象分类构建。每条记录包含'
  '{词条, 卦名, 隶属度向量(8维), 类别, 来源}。分类体系见表1。来源标注为"说卦传"或"实验归类"。')
P('（二）运行时学习表_learned：通过滑动平均更新 vec_new = 0.7 × vec_old + 0.3 × vec_new_obs。'
  '学习启用条件为count≥2（同一条目至少观测到两次才更新）。持久化到gua_knowledge.json文件，'
  '每10局中间保存一次，最终局全量保存。经过274局测试后，知识库增长到241条（含85条运行时学习记录）。')

tbl(
    ['卦象', '卦名', '类象范畴', '典型实体'],
    [
        ['☶ 艮', '山', '果实、家具、静止', '苹果、番茄、橱柜、抽屉、椅子'],
        ['☵ 坎', '水', '液体、冷藏、险陷', '冰箱、水槽、洗碗池、水龙头'],
        ['☲ 离', '火', '光明、火、文明', '台灯、落地灯、微波炉、灶台'],
        ['☳ 震', '雷', '动作、工具、震动', '刀、剪刀、锤子、键盘'],
        ['☴ 巽', '风', '柔顺、放置、绳直', '毛巾、抹布、床单、坐垫'],
        ['☰ 乾', '天', '刚健、金属、圆形', '锅、平底锅、金属架、保险箱'],
        ['☱ 兑', '泽', '口、开口、悦', '门、窗户、盖子、抽屉把手'],
        ['☷ 坤', '地', '承载、包容、母', '盘子、碗、床、架子、柜台'],
    ],
    '《说卦传》万物类象分类体系',
    '1'
)

P('查询优先级算法：长词先验 > 精确先验 > 长词学习 > 精确学习 > 八卦属性标签 > 默认坤卦。'
  '长词匹配优先于精确匹配，例如"微波炉"作为整体匹配先于拆成"微"+"波"+"炉"。'
  '知识库回退是推理链路的最后关卡——当char()和word()都无法产生有效结果时回退到语义查询。'
  '如果知识库也没有记录，默认返回坤卦（地，承载万物，隶属度0.85）。')

h('4.2 实体卦象封装（GuaEntity）', 2)
P('每个实体（容器或物体）被封装为GuaEntity对象。实体卦象在首次感知时通过汉字易理推理初始化，'
  '后续每步观测时更新。GuaEntity的完整字段见表2。')
P('首次感知流程：英文实体名（如"cabinet 1"）→ ChineseBridge._en_entity_to_cn("cabinet") → "橱柜" → '
  'HanziEngine.char("橱柜") → 部首推理→bagua_vec[8]→yao_state[6]→创建GuaReceptacle或GuaObject。'
  '容器实体（GuaReceptacle）还需在admissible_commands中检测"open cabinet 1"的出现来推断openable属性。')
tbl(
    ['字段', '类型', '含义', '初始化时机'],
    [
        ['hanzi', 'str', '汉字名', '首次感知时翻译'],
        ['dom_gua', 'str', '主导卦象名', '首次感知时推理'],
        ['bagua_vec', 'float[8]', '八卦隶属度向量', '首次感知+逐步更新'],
        ['yao_state', 'float[6]', '六爻状态编码', '首次感知+逐步更新'],
        ['cls', 'str', '英文类别', '首次感知时记录'],
        ['openable/visited/searched', 'bool', '容器状态', '逐步更新'],
        ['location/inventory/processed', 'str/bool/set', '物体状态', '逐步更新'],
    ],
    'GuaEntity 实体卦象表示',
    '2'
)

h('4.3 情境卦象聚合', 2)
P('情境卦象计算的完整六步过程：')
P('(1) 收集当前所在容器的yao_state（六爻状态编码）。'
  '(2) 收集视野内所有可见物体的六爻（最多取6个，避免噪声）。'
  '(3) 收集手中持有物品的六爻（如果有）。'
  '(4) 取上述所有六爻的均值得到情境六爻S。'
  '(5) 与64卦模板计算余弦相似度，取最匹配的卦象H*。'
  '(6) 用H*的吉凶评分f(H*)评估当前态势（吉卦→态势好，凶卦→需谨慎）。')
P('情境聚合的数学表达（公式1）：')
eqn('S = mean( {yao(e) | e ∈ E_visible ∪ E_holding ∪ E_location} )', '1')
P('其中yao(e)是实体e的6维六爻向量，E_visible是视野内的物体集合，E_holding是手持物品，E_location是当前位置容器。'
  '例如当前位置为"柜台"（坤卦主导，六爻[0.6,0.5,0.7,0.4,0.5,0.6]），视野内有"苹果"（艮卦主导，六爻[0.3,0.4,0.2,0.6,0.5,0.4]），'
  '则均值为S=[0.45,0.45,0.45,0.50,0.50,0.50]，匹配64卦得"剥"卦（山地剥，第23卦）→"当前处于积累态势"。')
P('状态变化检测：维护状态键state_key = f(location, inventory, 容器开关状态, 搜索状态, 工序进度, 放置计数)。'
  '执行有状态动作（take/put/heat/clean/cool等）后state_key不变时，判定动作无实际进展，'
  '记录失败经验并计入failed_sa（state_key与导致失败的action配对）。')

h('4.4 容器状态追踪与搜索规划', 2)
P('系统为每个容器维护三个二进制标志位：')
P('  - visited：agent是否到达过该容器位置。首次go to时设为True。')
P('  - searched：agent是否已搜索过该容器内部。到达后执行look或在admissible中看到容器内的take命令时设为True。'
  '    如果容器需先打开（openable=True），则open后才会标记searched。')
P('  - open：容器当前是否处于打开状态。从admissible中"open X"命令的缺失推断为已打开，'
  '    从"close X"命令的出现推断为可关闭。')
P('visited/searched的区分至关重要：agent可以"到达"一个容器位置（visited=True）但不一定能看到里面的物体，'
  '因为容器可能是关闭的（如drawer、cabinet）。只有open后才searched，才能看到内部物体。')
P('unsearched_receptacles()接口返回所有visited=True但searched=False的容器列表，'
  '指导agent优先探索未搜索区域而非重复搜索已知区域。')
P('Deposited/Pending计数器：对每个目标容器，追踪已放置的物品数量（deposited_count）和当前持有待放置的物品。'
  '当deposited_count达到任务要求的数量时，该目标视为完成，不再继续放置操作。')

h('4.5 英汉桥接与感知脚本', 2)
P('ChineseBridge模块提供英文→汉字的词对转换，覆盖约60条家居环境常见实体。'
  '映射包括：counter→柜台、cabinet→橱柜、drawer→抽屉、fridge→冰箱、sink→水槽、'
  'microwave→微波炉、fridge→冰箱、table→桌子、shelf→架子、lamp→台灯等。')
P('感知脚本_ingest_hanzi_obs(obs, action)从环境返回的观测文本中提取结构化信息，包含以下操作：')
P('  (a) 正则匹配"you arrive at LOCATION" → 位置更新为LOCATION。')
P('  (b) 调_to_hanzi(LOCATION) → ChineseBridge._en_entity_to_cn() → 汉字。')
P('  (c) 调_get_bagua(汉字) → _bagua_for_hanzi_cached() → 缓存命中或推理 → (主导卦象, [0.85,...])。')
P('  (d) 调_get_yao(LOCATION) → HanziEngine.char() → 部首推理→六爻编码 [0.6,0.5,...]。')
P('  (e) 创建/更新GuaReceptacle(id=LOCATION, hanzi=汉字, dom_gua=卦象, yao_state=六爻)。')
P('  (f) 正则匹配"you see a OBJ n" → 物体"OBJ n" → _to_hanzi→_get_bagua→创建GuaObject。')
P('  (g) 如果动作是take/put/clean/heat/cool，调用对应的_hanzi_take()/_hanzi_put()/_hanzi_mark_processed()更新状态。')

# ═══ 5. 决策推理引擎 ═══
h('5 决策推理引擎')

h('5.1 全候选评分框架', 2)
P('YLYW V18/V20不采用分层策略池。决策核心是"对全部合法动作同时评分"的并行评估范式，'
  '与V17"先选意图再映射到动作"的单线决策形成根本对立。')
P('完整决策流程：')
P('  (1) 环境返回admissible_commands（通常40~100个合法动作）。')
P('  (2) Veto过滤层逐项检查每个候选动作，剔除10类不合法或已失败的动作。')
P('  (3) 将存活候选送入YLYWScorer，每个候选计算6维六爻评分向量Y(a)。')
P('  (4) Y(a)与64卦模板逐一计算余弦相似度，取最匹配卦象H*。')
P('  (5) 综合评分score(a) = sim(Y(a), H*) × f(H*) × (0.75 + 0.25 × cos*(a) × aff(a))。')
P('  (6) 从存活候选集中取argmax score(a)执行。')
P('Veto仅剔除候选，不产生候选——所有候选来自环境提供的合法动作集。'
  '这使得决策层的设计极其简洁：约90行YLYWScorer代码完成全部评分逻辑。')

h('5.2 Veto过滤层（10项检查）', 2)
P('Veto过滤层的10项检查项见表3。这些检查全部基于运行时状态（手持物、容器状态、失败历史）'
  '和admissible_commands内容进行判断，不依赖任何地面真值信息。')
tbl(
    ['Veto项', '触发条件', '处理方式'],
    [
        ['info_action', '非必要look/examine', '默认剔除，仅首次look放行'],
        ['non_target_take', '手持非目标类物体', '先drop再取正确目标'],
        ['failed_sa', '该动作导致过死锁', '跳过（state_key+action配对）'],
        ['reversal', '动作刚好撤销上一步', '跳过（e.g. take→move→take循环）'],
        ['loop_cap', '相同命令≥4次', '跳过（除非是目标place/process）'],
        ['searched_recep', '已搜索过的容器', '跳过（探前沿非已搜区域）'],
        ['already_open', '容器已打开', '跳过open'],
        ['lower_priority_take', '模糊类低优先物体', '等待高优先候选被排除'],
        ['already_placed', '已放置到目标容器', '不重复拿取'],
        ['leave_closed_dest', '目标容器关着但手持目标', '先open再放，不去别处'],
    ],
    'Veto过滤检查项',
    '3'
)

h('5.3 六爻评分向量构造', 2)
P('对每个候选动作a，YLYWScorer计算6维六爻评分向量Y(a)=[y₁,...,y₆]。'
  '各维度的语义和计算方式见表4。')

tbl(
    ['爻位', '符号', '语义', '计算方式'],
    [
        ['y₁', '初', '目标差距减少', 'σ(0.05+0.90×goal_progress(a))'],
        ['y₂', '二', '持有连续性', '0.85(持目标)/0.25(持非目标)/0.50'],
        ['y₃', '三', '过程推进', '0.90(匹配工序)/0.10(逆工序)/0.50'],
        ['y₄', '四', '容器可用性', '0.85(开需要容器)/0.15(关闭)/0.50'],
        ['y₅', '五', '目标关联', '0.90(直接匹配)/0.65(类别匹配)/0.20'],
        ['y₆', '上', '新颖性vs失败', 'σ(0.50-0.40×fail_count)'],
    ],
    '六爻评分编码',
    '4'
)

P('六爻向量的构造遵循"每爻独立、多因素混合"的原则。以任务"Put a clean large metal spoon on the round white table"为例，'
  '候选动作"take spoon 1 from cabinet 1"的六爻计算过程：')
P('  - 初爻y₁：spoon是目标物体→goal_progress(a)=1.0（目标完全匹配）→y₁=σ(0.05+0.90×1.0)=σ(0.95)=0.72。'
  '  - 二爻y₂：当前手空→0.50（中性值）；如果手中已有其他非目标物体→0.25（低分，需先drop）。'
  '  - 三爻y₃：当前工序是"寻找目标"→take匹配工序→0.90（高推进）。'
  '  - 四爻y₄：cabinet 1可打开或已打开→0.85（高容器可用性）。'
  '  - 五爻y₅：spoon的物体类匹配目标物体类→0.90（直接匹配）。'
  '  - 上爻y₆：此前从未拿过spoon 1→fail_count=0→y₆=σ(0.50)=0.62。'
  '→ Y(take) = [0.72, 0.50, 0.90, 0.85, 0.90, 0.62]。')

h('5.4 64卦匹配与综合评分', 2)
P('Y(a)与64卦模板逐一计算余弦相似度：')
eqn('sim(Y, H_i) = (Y · H̃_i) / (||Y|| · ||H̃_i||)', '2')
P('其中H̃_i是第i卦的6维标准化爻向量。64卦的爻向量来源于周易标准卦画（阴爻→0.15，阳爻→0.85），'
  '存储在ylyw_core.trigram_base的HexagramRule中。例如乾卦为[0.85,0.85,0.85,0.85,0.85,0.85]（6阳），'
  '坤卦为[0.15,0.15,0.15,0.15,0.15,0.15]（6阴），泰卦为[0.15,0.85,0.15,0.65,0.85,0.15]（地天泰，三阴三阳）。')
eqn('H* = argmax_{i=1}^{64} sim(Y, H_i)', '3')
P('取余弦相似度最高的卦象H*作为该候选动作的"命中卦象"。'
  '例如Y=[0.72,0.50,0.90,0.85,0.90,0.62]与火天大有卦(离上乾下)的相似度最高(0.81)→H*=大有卦(第14卦)。')
P('综合评分：')
eqn('score(a) = sim(Y(a), H*) × f(H*) × (0.75 + 0.25 × cos*(a) × aff(a))', '4')
P('其中三个因子：')
P('(a) sim(Y, H*) ∈ [0, 1]：候选动作六爻与命中卦象的匹配度。')
P('(b) f(H*) ∈ {0.15 ~ 0.95}：命中卦象的先验吉凶评分。吉卦（乾0.95、坤0.95、泰0.90、谦0.90、大有0.90）'
  'f∈[0.85,0.95]；中卦（屯0.65、蒙0.60、需0.70）f∈[0.55,0.75]；凶卦（剥0.20、困0.30、蹇0.25、'
  '否0.15）f∈[0.15,0.45]。共64个先验评分值，来源于《周易》传统卦义判断。')
P('(c) 亲和度因子：cos*(a)∈{0,1}是动作类别标记（动作动词→1，移动/信息→0），'
  'aff(a)∈[0.92,1.0]是动作与物体的八卦亲和度。例如"取"（震卦）+"苹果"（艮卦），'
  '震艮相克→aff=0.92（低）；"加热"（离卦）+"食物"（艮卦），离火生艮土→aff=0.98（高）。')
P('最终a* = argmax_{a ∈ A_valid} score(a)被选中执行，其中A_valid是Veto过滤后的存活候选集。')

h('5.5 失败检测与自动重试', 2)
P('当agent完成全部目标后环境仍未给予胜利信号时，_maybe_retry()检测三种引用错误：')
P('(1) 物体目标引用错误：已放置指定数量的同类物体但环境不认，自动切换下一个候选物体类。'
  '例如目标为"large metal spoon"但环境只认"spoon 1"→切换后重新尝试place。')
P('(2) 容器目标错误：当前容器不支持放置该物体，自动换到下一候选容器。'
  '例如"round white table"实际是"diningtable 1"，但放在"diningtable 2"上环境不认→切换到正确编号。')
P('(3) 工序目标错误：手持有物体但工具无法处理该物体，自动切换候选物体。'
  '例如"clean"需要水槽但物体不是可洗类型→换另一个同类物体重试。')
P('每种情况的判断通过检查state_key指纹和admissible_commands中的可用动作完成，'
  '不依赖任何地面真值信息。切换后重新打开（search flag重置），最多10次重试。')
P('_deposited_failed和_recep_failed集合记录了已验证错误的物体/容器引用，'
  '避免同一错误重复试错。每次重试后重新调整目标探索的优先级排序。')

h('5.6 运行时学习与知识积累', 2)
P('GuaKnowledgeBase的运行时学习表在每步执行后更新。遍历所有实体的bagua_vec，'
  '调用observe_learning(hanzi, bagua_vec)执行滑动平均更新：'
  'vec_new = 0.7 × vec_old + 0.3 × vec_observed。学习启用条件为count≥2（同一条目至少观测到两次才更新）。')
P('例如第一次见到"spoon"时在知识库中创建新条目vec_observed=[兑卦主导]，'
  '第二次见到滑动平均后vec保留原向量70%+新观测30%。多次观测后同一实体的卦象向量趋于稳定。')

h('5.7 接口兼容性设计', 2)
P('决策引擎完全通过查询接口访问世界状态，不直接操作内部数据结构。主要接口包括：')
P('  - holding_target(obj_classes): 检查手中是否有目标类型物体。')
P('  - known_objects_of(obj_class): 获取已知的某类物体列表。')
P('  - find_object_location(obj_class): 查找某类物体的位置。')
P('  - find_pending_target_recep(): 查找尚未放置的待处理目标物体位置。')
P('  - get_situation_gua(): 获取当前情境卦象（卦名+六爻）。')
P('  - get_entity_gua(en_id): 获取特定实体的卦象。')
P('  - state_key(): 状态指纹，用于变化检测。')
P('  - unsearched_receptacles(): 未搜索过的容器列表。')
P('同一决策引擎可接入不同世界模型实现（符号WorldModel或汉字CnWorldModel），'
  '无需修改上层决策逻辑。')

# ═══ 6. 完整执行流程 ═══
h('6 完整执行流程：单步决策数据流')

P('以下展示数据从英文观测流经各层到达动作输出的完整8阶段过程。场景设置：'
  'task_desc="Put a clean large metal spoon on the round white table"，'
  'obs="You are at cabinet 1. On the cabinet 1, you see a spoon 1."，'
  'admissible包含"take spoon 1 from cabinet 1", "go to diningtable 1", "open drawer 1"等约50个合法动作。')

h('阶段1（环境交互）：', 3)
P('ALFWorld环境返回英文观测文本obs和合法动作集admissible_commands。'
  'obs包含当前位置描述（"You are at…"）和当前可见物体（"On the X, you see…"或"The X is closed"）。'
  'admissible包含所有语法合法且物理可行的动作，包括移动(go to)、取放(take/put)、操作(open/close)、'
  '工序(clean/heat/cool)、信息(look/inventory)和帮助(help)等类别。')
P('admissible_commands不包含目标物体的位置信息，agent需要通过搜索自主发现。'
  '当前场景admissible中"go to"的目标包含所有容器和家具（约20~35个位置），'
  '"take"的目标需要在正确容器中才出现。')

h('阶段2（汉字感知与卦象编码）：', 3)
P('_ingest_hanzi_admissible(admissible)：遍历合法动作列表，识别每个动作中的容器名和物体名。'
  '遇到新容器（如"cabinet 1"）→_ensure_recep()创建GuaReceptacle对象。'
  '遇到新物体（如"spoon 1"出现在"take spoon 1 from cabinet 1"中）→_ensure_obj()创建GuaObject。')
P('_ingest_hanzi_obs(obs, action)：从环境返回的观测文本中提取结构化信息，执行以下子步骤：')
P('  (a) 正则匹配"you are at LOCATION" → 更新当前位置为LOCATION。')
P('  (b) _to_hanzi(LOCATION) → ChineseBridge._en_entity_to_cn(LOCATION) → 汉字名。')
P('  (c) _get_bagua(汉字名) → bagua_for_hanzi_cached(汉字名) → 缓存命中或部首推理 → (dominant_gua, [8维隶属度])。')
P('  (d) _get_yao(LOCATION, LOCATION类别) → HanziEngine.char(汉字名) → 部首分解→卦象映射→乘承比应→六爻 [6维]。')
P('  (e) 创建或更新GuaReceptacle(id=LOCATION, hanzi=汉字名, dom_gua=卦象, yao_state=六爻编码)。')
P('  (f) 正则匹配"you see a OBJ n"或"On the LOC you see a OBJ" → 提取物体"OBJ n"。'
  '   _to_hanzi(OBJ) → 汉字 → _get_bagua(汉字) → (卦象) → 创建GuaObject(id="OBJ n", hanzi=汉字, dom_gua=卦象)。')
P('  (g) 如果上一步执行的动作是take/put/clean/heat/cool，调用对应的_hanzi_take()/_hanzi_put()/'
  '  _hanzi_mark_processed()更新实体的inventory/location/processed字段。')
P('在本例中："cabinet 1"→ChineseBridge→"橱柜"→char("橱柜")→艮卦([0.85,0.72,...])→六爻[0.6,0.5,0.7,0.4,0.5,0.6]。'
  '"spoon 1"→"勺子"→char("勺子")→兑卦([0.88,0.65,...])→六爻[0.5,0.3,0.6,0.5,0.4,0.7]。')

h('阶段3（状态同步）：', 3)
P('_sync_location_flags(admissible_after)：根据当前admissible_commands更新容器状态标志。'
  '当前位置"cabinet 1"的容器不需要打开（openable=False）→标记visited=True, searched=True。'
  '如果容器需要先打开（如drawer），但admissible中"open drawer 1"仍在→标记visited=True但searched=False。'
  '容器闭合状态通过admissible中"open X"命令的消失推断为已打开。')

h('阶段4（情境聚合）：', 3)
P('_update_situation()：收集当前所有GuaReceptacle和GuaObject的yao_state，执行六步计算：')
P('  (4a) 取当前容器"橱柜"(cabinet 1)的yao_state=[0.6,0.5,0.7,0.4,0.5,0.6]。')
P('  (4b) 取视野内物体"勺子"(spoon 1)的yao_state=[0.5,0.3,0.6,0.5,0.4,0.7]。')
P('  (4c) 当前手空，无手持物品。')
P('  (4d) 取均值→情境六爻S = ([0.6+0.5]/2, [0.5+0.3]/2, [0.7+0.6]/2, [0.4+0.5]/2, [0.5+0.4]/2, [0.6+0.7]/2) = [0.55,0.40,0.65,0.45,0.45,0.65]。')
P('  (4e) 与64卦模板计算余弦相似度→最佳匹配"山泽损"卦（第41卦）→相似度0.73。')
P('  (4f) f(损)=0.55（中偏下），提示当前处于"减损以求益"态势，需要做出取舍后再行动。')

h('阶段5（失败检测）：', 3)
P('state_key()生成当前状态指纹：f(当前位置=cabinet 1, 手持物=空, 容器开关状态={cabinet 1=open}, '
  '搜索状态={cabinet 1=searched}, 工序进度=寻找，放置计数=0/1)。')
P('如果与上一步的state_key相同且有状态动作执行→记录失败经验failed_sa.add((pre_key, action))。'
  '本例中首次到cabinet 1，state_key较之前改变→正常，无失败记录。')

h('阶段6（运行时学习）：', 3)
P('_learn_from_observation()：遍历所有GuaReceptacle和GuaObject的bagua_vec。'
  '对"橱柜"和"勺子"分别调用observe_learning("橱柜", bagua_vec)和observe_learning("勺子", bagua_vec)。')
P('在知识库中"橱柜"已有先验记录→更新：vec_new=0.7×vec_old+0.3×vec_observed。'
  '"勺子"可能未在知识库中→创建新条目。count++。当count≥2时该条目生效（可被查询到）。')

h('阶段7（决策推理）：', 3)
P('Agent.act()调用YLYW评分器执行完整决策流程：')
P('  (7a) GoalParser解析task_desc→目标物体=spoon(勺子)，工序=[clean, put]，目标容器=table(桌子)。')
P('  (7b) 从admissible_commands取所有候选动作（约50个）。')
P('  (7c) Veto过滤：剔除look(信息型)、已搜容器相关go to、cabinet 1重复open等→存活约30个候选。')
P('      → Y(take)=[0.72,0.50,0.90,0.85,0.90,0.62]。')
P('  (7e) Y(take)与64卦逐一余弦匹配→最高相似度0.81匹配火天大有卦。')
P('  (7f) 综合评分：score=0.81×0.90(大有吉)×(0.75+0.25×1.0×0.92)=0.81×0.90×0.98=0.714。')
P('  (7g) 对所有存活候选计算得分后排序→"take spoon 1 from cabinet 1"得分最高(0.714)→选中执行。')
P('  (7h) 输出动作"take spoon 1 from cabinet 1"。')

h('阶段8（环境返回与循环）：', 3)
P('动作"take spoon 1 from cabinet 1"提交到ALFWorld环境。环境返回：'
  '  obs="You take spoon 1 from cabinet 1. You are at cabinet 1. On the cabinet 1, you see nothing."'
  '  admissible包含"go to sinkbasin 1", "go to diningtable 1", "clean spoon 1 with sinkbasin 1", '
  '  "put spoon 1 in/on diningtable 1"等。')
P('took动作成功后，_hanzi_take("spoon 1", "cabinet 1")将spoon 1从cabinet 1移入手持物品槽。'
  '回到阶段2开始新一轮感知→决策→执行循环。继续处理clean工序和put工序，'
  '直到环境返回won=True信号或步数超过50步判定失败。')
P('典型的完整游戏流程约6~18步：go to→take→go to→clean→go to→put→won。')

h('6.1 全局LRU缓存加速', 2)
P('汉字→卦象推理是无状态的——同一汉字始终产生相同的八卦隶属度。'
  '系统在模块级应用@lru_cache(maxsize=4096)装饰器，包含三个缓存层次：')
P('(1) bagua_for_hanzi_cached(hanzi)：汉字→八卦隶属度的底层缓存，返回tuple（可哈希）。')
P('(2) bagua_for_hanzi(hanzi)：对外接口，转调cached版本，将tuple转回list返回。')
P('(3) yao_for_entity_cached(en_name)：英文实体名→六爻编码的缓存。')
P('在ALFWorld测试中，缓存命中率约78%，系统整体推理速度提升约44%'
  '（从平均3.66秒/局降至2.54秒/局）。')
P('缓存的使用对速度影响显著：')
P('  - 无缓存（V20初始版本）：平均3.66秒/局，140局共512.4秒。')
P('  - 有三层缓存（V20优化版本）：平均2.54秒/局，140局共355.6秒。')
P('  - V18符号版本（英文直接做状态键，无汉字推理）：平均2.93秒/局。')
P('  - 速度对比：优化后V20比V18快13%，比无缓存V20快44%。')

# ═══ 7. 实验 ═══
h('7 实验评估')

h('7.1 实验设置', 2)
P('基准测试：ALFWorld基准¹⁻² valid_seen（140局）和valid_unseen（134局），总计274局。'
  '最大步数50步/局，超时180秒。硬件：单CPU核心，2GB RAM，无GPU，无API调用。')

h('7.2 总体结果', 2)
tbl([('数据集', '局数', '成功', '成功率')],
    [('valid_seen', '140', '129', '92.1%'), ('valid_unseen', '134', '132', '98.5%'),
     ('总计', '274', '261', '95.3%')], 'YLYW 总体结果', '5')

h('7.3 按任务类型', 2)
tbl([('任务类型', 'valid_seen', 'valid_unseen')],
    [('光照观察', '100% (13/13)', '94% (17/18)'),
     ('加热放置', '100% (16/16)', '100% (23/23)'),
     ('清洁放置', '93% (25/27)', '100% (31/31)'),
     ('双物体放置', '92% (22/24)', '100% (17/17)'),
     ('简单放置', '91% (32/35)', '96% (23/24)'),
     ('冷却放置', '84% (21/25)', '100% (21/21)')],
    '按任务类型分类结果', '6')

h('7.4 对比分析', 2)
tbl([('方法', 'LLM', 'valid_unseen', '推理方式')],
    [('YLYW（本文）', '否', '98.5%', '汉字易理推理+64卦匹配'),
     ('BUTLER¹⁻²', '否', '~20%', '纯符号规则'),
     ('ReAct (GPT-4)³', '是', '~65%', 'LLM推理+行动'),
     ('SayCan⁴', '是', '~55%', 'LLM+技能库'),
     ('Code-as-Policies⁵', '是', '~40%', '代码生成'),
     ('微调小语言模型⁶', '是(7B)', '~70%', '微调')],
    '与代表性方法对比', '7')

h('7.5 执行效率', 2)
tbl([('指标', 'YLYW', 'GPT-4方案（估算）', '加速比')],
    [('每步决策', '~0.08秒', '~0.5秒', '~6×'),
     ('每局平均', '~2.7秒', '~15秒', '~5.5×'),
     ('274局总耗时', '~765秒', '~4000秒', '~5.2×'),
     ('硬件需求', '单CPU, 2GB', 'GPU集群', '—'),
     ('API费用', '0', '$100+', '—')],
    '执行效率', '8')

h('7.6 失败分析', 2)
P('274局共丢13局。冷却类4局（时序复杂超步），简单放置3局（目标歧义），'
  '双物体2局（搜索超步），清洁类2局（特定物体定位），光照类1局，放置类1局。')

# ═══ 8. 讨论 ═══
h('8 讨论与展望')
P('YLYW验证了汉字作为易经推理载体的有效性。"部首即卦象"的特性使自然语言→卦象映射成为可行计算方案。'
  '156条先验覆盖ALFWorld全部实体，85条学习记录补充边缘实体。'
  '同一场景重复约10次后新实体的卦象映射趋于稳定。')
P('YLYW不追求替代LLM，而是提供另一范式：固定场景下完全确定、可审计、零GPU的决策方案。'
  '未来方向包括：(1)跨具身形态迁移（MuJoCo/实体机器人）；'
  '(2)知识库扩展到家居/工业/医疗场景；'
  '(3)学习型八卦隶属度替代手工先验；'
  '(4)验证YLYW对视觉/触觉模态的适用性。')

# ═══ 参考文献 ═══
h('参考文献')
refs = [
    '[1] Shridhar, M. et al. ALFRED: A benchmark for interpreting grounded instructions for everyday tasks. CVPR, 2020.',
    '[2] Shridhar, M. et al. ALFWorld: Aligning text and embodied environments for interactive learning. ICLR, 2021.',
    '[3] Yao, S. et al. ReAct: Synergizing reasoning and acting in language models. ICLR, 2023.',
    '[4] Ahn, M. et al. Do as I can, not as I say: Grounding language in robotic affordances. CoRL, 2022.',
    '[5] Liang, J. et al. Code as policies: Language model programs for embodied control. ICRA, 2023.',
    '[6] Fikes, R. & Nilsson, N. STRIPS: A new approach to the application of theorem proving to problem solving. AIJ, 2(3-4):189–208, 1971.',
    '[7] 高亨. 周易大传今注. 齐鲁书社, 1994.',
    '[8] Wilhelm, R. & Baynes, C. The I Ching or Book of Changes. Princeton Univ. Press, 1967.',
    '[9] 黄寿祺, 张善文. 周易译注. 上海古籍出版社, 2004.',
    '[10] 孔颖达. 周易正义. 载《十三经注疏》. 中华书局, 1980.',
    '[11] 《说卦传》. 载《周易》, 约前3世纪.',
]
for ref in refs:
    pp = doc.add_paragraph(ref)
    pp.paragraph_format.first_line_indent = Cm(-0.75)
    pp.paragraph_format.left_indent = Cm(0.75)
    for r in pp.runs:
        r.font.size = Pt(9); r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ═══ 保存到/tmp ═══
tmp = '/tmp/YLYW_技术论文.docx'
doc.save(tmp)
sz = os.path.getsize(tmp)
print(f'✅ 论文已保存: {tmp} ({sz/1024:.1f} KB)')
print(f'   结构: 摘要 → 1引言 → 2系统架构(2.1总览/2.2模块) → 3汉字引擎(3.1-3.7)')
print(f'         → 4知识库与情境(4.1-4.5) → 5决策推理(5.1-5.6) → 6执行流程(8阶段)')
print(f'         → 7实验(6表) → 8讨论 → 参考文献')
print(f'   图3张 表8张 公式4个 参考文献11篇')
