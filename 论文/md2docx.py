#!/usr/bin/env python3
"""Markdown论文 → docx 格式转换"""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def md_to_docx(md_path, docx_path):
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    in_table = False
    table_lines = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块跳过
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            i += 1
            continue

        # 表格
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            table_lines.append(stripped)
            in_table = True
            i += 1
            continue
        else:
            if in_table and len(table_lines) >= 2:
                _add_table(doc, table_lines)
                table_lines = []
                in_table = False
                # 表格后的空行跳过
                if not stripped:
                    i += 1
                    continue
        
        if in_table and not stripped:
            i += 1
            continue

        # 标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
        elif stripped.startswith('## '):
            _add_heading(doc, stripped[3:], 2)
        elif stripped.startswith('### '):
            _add_heading(doc, stripped[4:], 3)
        elif stripped.startswith('#### '):
            _add_heading(doc, stripped[5:], 4)
        elif stripped.startswith('---') or stripped.startswith('==='):
            doc.add_paragraph('─' * 60)
        elif stripped.startswith('**摘要**') or stripped.startswith('**关键词**'):
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
        elif stripped.startswith('- ') or stripped.startswith('* '):
            # 列表项
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
            _format_inline(p)
        elif stripped and stripped[0].isdigit() and '. ' in stripped[:4]:
            # 编号列表
            p = doc.add_paragraph(stripped, style='List Number')
            _format_inline(p)
        else:
            if stripped:
                p = doc.add_paragraph(stripped)
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.paragraph_format.line_spacing = 1.5
                _format_inline(p)

        i += 1

    # 如果末尾有未处理的表格
    if in_table and len(table_lines) >= 2:
        _add_table(doc, table_lines)

    doc.save(docx_path)
    print(f"✅ {docx_path}")


def _add_heading(doc, text, level):
    sizes = {2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    if level <= 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    # 添加下划线（一级、二级标题）
    if level <= 2:
        pBdr = p.paragraph_format.element.makeelement(
            qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '1',
            qn('w:color'): '333333',
        })
        pBdr.append(bottom)
        p.paragraph_format.element.append(pBdr)


def _format_inline(paragraph):
    """处理行内的加粗、斜体"""
    text = paragraph.text
    # 加粗替换 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    if len(parts) > 1:
        paragraph.clear()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)


def _add_table(doc, lines):
    """将 Markdown 表格添加到文档"""
    # 跳过分隔行
    data_lines = [l for l in lines if not re.match(r'^[\s\|:\-]+$', l)]
    if len(data_lines) < 2:
        return
    headers = [h.strip() for h in data_lines[0].strip('|').split('|')]
    rows = []
    for l in data_lines[1:]:
        cells = [c.strip() for c in l.strip('|').split('|')]
        if cells:
            rows.append(cells)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # 数据行
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < len(headers):
                cell = table.rows[i + 1].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # 表后空行


if __name__ == '__main__':
    import sys
    md_path = sys.argv[1] if len(sys.argv) > 1 else \
        '/home/lijinhan/MXL/科研/ylyw/论文/跨本体机器人取放的YLYW爻参数自适应系统.md'
    docx_path = sys.argv[2] if len(sys.argv) > 2 else md_path.replace('.md', '.docx')
    md_to_docx(md_path, docx_path)
