#!/usr/bin/env python3
"""合并第1-3章所有md为一个完整docx，含书名页、目录、插图"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
import re, os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)

BASE = '/home/lijinhan/MXL/科研/ylyw/monograph'

# ============================================================
# 书名页
# ============================================================
for _ in range(8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('易 理 研 物'); r.bold = True; r.font.size = Pt(28); r.font.name = '黑体'

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('—— 从符号先验到通用智能的工程研究 ——'); r.font.size = Pt(16); r.font.name = '宋体'

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('马兴录  著'); r.font.size = Pt(16); r.font.name = '宋体'

doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('青岛科技大学  信息科学技术学院'); r.font.size = Pt(14); r.font.name = '宋体'

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026年'); r.font.size = Pt(14); r.font.name = '宋体'

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
def toc_chapter(num, title, size=14):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(12); pf.space_after = Pt(4)
    r = p.add_run(f'{title}'); r.bold = True; r.font.size = Pt(size); r.font.name = '黑体'

def toc_section(text, indent=0.8, size=11, bold=False):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.left_indent = Cm(indent); pf.space_after = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = '宋体'

def toc_subsection(text, indent=1.6, size=10.5):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.left_indent = Cm(indent); pf.space_after = Pt(1)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = '宋体'

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('目    录'); r.bold = True; r.font.size = Pt(20); r.font.name = '黑体'
doc.add_paragraph()

toc_chapter('', '序言：被忽视的先验——为什么《易经》值得严肃的工程研究', 12)
toc_section('0.1  这本书要回答的核心问题')
toc_section('0.2  面向什么读者')
toc_section('0.3  本书的结构导航')
toc_section('0.4  如何使用本书')

toc_chapter('', '第1章  易理研物的哲学基础', 12)
toc_section('1.1  易经作为先验知识体系')
toc_section('1.2  从符号先验到通用智能')
toc_section('1.3  核心哲学命题')

toc_chapter('', '第2章  八卦隶属度：符号化接地方法', 12)
toc_section('2.1  符号接地问题的重定义')
toc_section('2.2  八卦—物理原型映射体系')

toc_chapter('', '第3章  六十四卦联邦架构设计', 12)
toc_section('3.1  联邦式神经符号架构设计原则')
toc_section('3.2  三层计算架构的完整定义')
toc_section('3.3  六十四卦策略映射的工程转译')
toc_section('3.4  爻位关系运算：从《周易》到可计算算子')
toc_section('3.5  与主流架构的系统对比')
toc_section('3.6  输出空间的连续性：驳"64种输出"的误解')

toc_chapter('', '第4章  知己学习：从先验征兆辨识到在线自适应修正', 12)
toc_section('4.1  知己学习的哲学根基与范式定义')
toc_section('4.2  知己学习 vs 强化学习：六个维度的系统对比')
toc_section('4.3  YLYW作为知己学习的工程实现')
toc_section('4.4  知几：征兆辨识与变卦预判')
toc_section('4.5  知耻：失败诊断与定向参数修正')
toc_section('4.6  自适应学习实验')
toc_section('4.7  范式定位：知己学习不是要替代RL')

toc_chapter('', '第5章  天生安全：双八卦约束与反幻觉', 12)
toc_section('5.1  安全问题的根本困难')
toc_section('5.2  双八卦并行仲裁架构')
toc_section('5.3  零样本物理合规实验')
toc_section('5.4  反幻觉审查：从物理决策到语言审查的跨域迁移')
toc_section('5.5  安全机制的嵌入式实现')
toc_section('5.6  本章核心贡献')

toc_chapter('', '第6章  跨域实证：从感知到行动', 12)
toc_section('6.1  实验设计的整体思路')
toc_section('6.2  物理域——YCB物理力学评估')
toc_section('6.3  运动控制域——从抓取到步态的跨域迁移')
toc_section('6.4  ALFWorld具身导航域')
toc_section('6.5  补充域与消融实验')
toc_section('6.6  跨域实验总结')

toc_chapter('', '第7章  层次化嵌套：从单智能体到多智能体', 12)
toc_section('7.1  从单智能体到多智能体')
toc_section('7.2  三层递归嵌套架构')
toc_section('7.3  卦象意图通讯协议')
toc_section('7.4  双模型并行博弈')
toc_section('7.5  MVP原型与复杂度可控性')

toc_chapter('', '第8章  定位与展望', 12)
toc_section('8.1  范式划分与六维对比')
toc_section('8.2  与端到端VLA范式的深度对比')
toc_section('8.3  与世界模型路线的对比')
toc_section('8.4  与LLM智能体的对比')
toc_section('8.5  互补架构与混合路线设计')
toc_section('8.6  YLYW与Qwen-VLA的对比分析')

toc_chapter('', '第9章  通向通用智能：从元胞到生态', 12)
toc_section('9.1  智能元胞：YLYW作为通用推理基元')
toc_section('9.2  智能组织：元胞群体的自组织涌现')
toc_section('9.3  智能生态：从封闭系统到开放世界')
toc_section('9.4  通向AGI的路线图与时间线')
toc_section('9.5  道器合一：从符号先验到通用智能的工程哲学')

toc_chapter('', '第10章  局限与开放问题', 12)
toc_section('10.1  感知前端：视觉到13维特征的语义鸿沟')
toc_section('10.2  先验的边界：当八卦不够用的时候')
toc_section('10.3  学习的上限：六爻精化的参数天花板')
toc_section('10.4  语言理解的爻编码难题')
toc_section('10.5  十个开放问题——给后续探索者的导航')
toc_section('10.6  结语：本书的最后陈词')

doc.add_page_break()

# ============================================================
# 正文处理函数（先定义，后使用）
# ============================================================

images_base = BASE  # where png files live

def heading(text, level):
    if level == 1:
        # 章标题 → Heading 1（Word自动目录识别）
        p = doc.add_heading(text, level=1)
        pf = p.paragraph_format
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(24); pf.space_after = Pt(14)
        # 直接修改Heading 1样式的run
        for run in p.runs:
            run.bold = True; run.font.size = Pt(18); run.font.name = '黑体'
    elif level == 2:
        # 节标题 → Heading 2
        p = doc.add_heading(text, level=2)
        pf = p.paragraph_format
        pf.space_before = Pt(16); pf.space_after = Pt(8)
        for run in p.runs:
            run.bold = True; run.font.size = Pt(15); run.font.name = '黑体'
    else:
        # 子节标题 → Heading 3
        p = doc.add_heading(text, level=3)
        pf = p.paragraph_format
        pf.space_before = Pt(8); pf.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(12); run.font.name = '黑体'
            run.bold = False

def parse_bold(text):
    segs = []
    pos = 0
    for m in re.finditer(r'\*\*(.+?)\*\*', text):
        if m.start() > pos:
            segs.append((False, text[pos:m.start()]))
        segs.append((True, m.group(1)))
        pos = m.end()
    if pos < len(text):
        segs.append((False, text[pos:]))
    if not segs:
        segs = [(False, text)]
    return segs

def add_p(text, indent_first=True, indent_left=0):
    segs = parse_bold(text)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(5)
    if indent_first and indent_left == 0:
        pf.first_line_indent = Cm(0.74)
    if indent_left > 0:
        pf.left_indent = Cm(indent_left)
    for is_b, txt in segs:
        run = p.add_run(txt)
        run.font.size = Pt(11)
        run.font.name = '宋体'
        if is_b: run.bold = True

def add_table(header, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(header))
    table.style = 'Table Grid'
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]; pf = p.paragraph_format
        pf.space_before = Pt(2); pf.space_after = Pt(2)
        r = p.add_run(h.strip()); r.bold = True; r.font.size = Pt(9); r.font.name = '宋体'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            p = cell.paragraphs[0]; pf = p.paragraph_format
            pf.space_before = Pt(2); pf.space_after = Pt(2)
            r = p.add_run(val.strip()); r.font.size = Pt(9); r.font.name = '宋体'
    doc.add_paragraph()

def add_image(filename, alt=''):
    full = os.path.join(images_base, filename)
    if os.path.exists(full):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(); r.add_picture(full, width=Inches(5.5))
        if alt:
            p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(alt); r2.font.size = Pt(9); r2.font.name = '宋体'; r2.italic = True
    else:
        add_p(f'[图片缺失: {filename}]')

def divider():
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(6); pf.space_after = Pt(6); pf.first_line_indent = Cm(0)
    r = p.add_run('—'*30); r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(150,150,150)

def is_table_sep(line):
    return re.match(r'^\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?\s*$', line)

def parse_table_cells(line):
    parts = line.strip().split('|')
    if parts and not parts[0].strip(): parts = parts[1:]
    if parts and not parts[-1].strip(): parts = parts[:-1]
    return [p.strip() for p in parts]

def process_md(filepath, force_chapter=None):
    """处理md文件。force_chapter不为None时，在正文前注入一个# 第X章 的heading"""
    with open(filepath) as f:
        lines = f.readlines()
    
    # 如果需要注入章标题
    if force_chapter:
        # 确保文件开头没有空白行再写
        inject = '# ' + force_chapter + '\n\n'
        lines = [inject] + lines
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        s = line.strip()
        if not s:
            i += 1; continue

        # 表格
        if '|' in s and i+1 < len(lines) and is_table_sep(lines[i+1].rstrip('\n')):
            header = parse_table_cells(s)
            i += 2
            rows = []
            while i < len(lines) and '|' in lines[i].rstrip('\n'):
                rows.append(parse_table_cells(lines[i].rstrip('\n')))
                i += 1
            add_table(header, rows)
            continue

        # 图片
        img = re.match(r'^!\[([^\]]*)\]\(([^\)]+)\)', s)
        if img:
            add_image(img.group(2), img.group(1))
            i += 1; continue

        # 一级标题（章标题）——以#开头但不以##开头
        # 如果格式为# x.y（数字点数字），则为节标题→Heading 2
        # 否则为章标题→Heading 1
        if s.startswith('# ') and not s.startswith('## '):
            text = s[2:]
            # 匹配 # 1.1、# 2.3 等格式
            if re.match(r'^\d+\.\d', text):
                heading(text, 2)
            else:
                heading(text, 1)
            i += 1; continue

        # 二级标题
        if s.startswith('## '):
            text = s[3:]
            # 匹配 ## 1.1.1、## 3.2.4 等格式 → Heading 3
            if re.match(r'^\d+\.\d+\.\d+', text):
                heading(text, 3)
            else:
                heading(text, 2)
            i += 1; continue

        # 三级标题
        if s.startswith('### '):
            heading(s[4:], 3)
            i += 1; continue

        # 分隔线
        if s == '---':
            divider()
            i += 1; continue

        # 无序列表
        if s.startswith('- ') or s.startswith('* '):
            add_p('•  ' + s[2:], indent_first=False, indent_left=1.2)
            i += 1; continue

        # 编号列表
        m = re.match(r'^(\d+)[\.\、]\s(.+)', s)
        if m:
            add_p(f'{m.group(1)}.  {m.group(2)}', indent_first=False, indent_left=1.2)
            i += 1; continue

        # 普通段落
        add_p(s)
        i += 1

# ============================================================
# 合并各章节
# ============================================================
# ============================================================
# 序言
# ============================================================
preface_path = os.path.join(BASE, 'ch00_preface.md')
with open(preface_path) as f:
    preface_lines = f.readlines()

i = 0
while i < len(preface_lines):
    line = preface_lines[i].rstrip('\n')
    s = line.strip()
    if not s:
        i += 1; continue
    if s.startswith('## '):
        heading(s[3:], 2)
    elif s.startswith('### '):
        heading(s[4:], 3)
    elif s.startswith('---'):
        divider()
    elif s.startswith('- ') or s.startswith('* '):
        add_p('•  ' + s[2:], indent_first=False, indent_left=1.2)
    else:
        add_p(s)
    i += 1

doc.add_page_break()

# ============================================================
# 第1-7章
# ============================================================
# 格式：(md文件名, 强制插入的章标题, 是否分页)
files = [
    ('ch01_1_1.md', '第1章 易理研物的哲学基础', False),
    ('ch01_1_2.md', None, False),
    ('ch01_1_3.md', None, False),
    ('ch02.md', None, True),                                # 第2章
    ('ch03_1.md', '第3章 六十四卦联邦架构设计', True),
    ('ch03_2.md', None, False),
    ('ch03_3.md', None, False),
    ('ch03_4.md', None, False),
    ('ch03_5.md', None, False),
    ('ch03_6.md', None, False),
    ('ch04.md', None, True),                                # 第4章
    ('ch05_1.md', '第5章 天生安全：双八卦约束与反幻觉', True),
    ('ch05_2.md', None, False),
    ('ch05_3.md', None, False),
    ('ch05_4_5_6.md', None, False),
    ('ch06.md', None, True),
    ('ch06_2.md', None, False),
    ('ch06_3.md', None, False),
    ('ch06_4.md', None, False),
    ('ch06_5_6.md', None, False),
    ('ch07.md', None, True),                                # 第7章
    ('ch07_1.md', None, False),
    ('ch07_2.md', None, False),
    ('ch07_3.md', None, False),
    ('ch07_4.md', None, False),
    ('ch07_5_6_7.md', None, False),
    ('ch08.md', None, True),                                # 第8章
    ('ch08_2.md', None, False),
    ('ch08_3.md', None, False),
    ('ch08_4.md', None, False),
    ('ch08_5.md', None, False),
    ('ch08_6.md', None, False),
    ('ch09.md', None, True),                                # 第9章
    ('ch09_1.md', None, False),
    ('ch09_2.md', None, False),
    ('ch09_3.md', None, False),
    ('ch09_4.md', None, False),
    ('ch09_5.md', None, False),
    ('ch10.md', None, True),                               # 第10章
    ('ch10_1.md', None, False),
    ('ch10_2.md', None, False),
    ('ch10_3.md', None, False),
    ('ch10_4.md', None, False),
    ('ch10_5.md', None, False),
    ('ch10_6.md', None, False),
]

for fname, chapter, page_break in files:
    fpath = os.path.join(BASE, fname)
    print(f'Processing {fname} ...')
    process_md(fpath, force_chapter=chapter)
    if page_break:
        doc.add_page_break()

output = os.path.join(BASE, '易理研物_序言+第1-10章.docx')
doc.save(output)
print(f'\nDone: {output}')
