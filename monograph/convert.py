#!/usr/bin/env python3
"""通用 md→docx 转换器，支持表格、粗体、列表"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re, sys, os

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        from lxml import etree
        tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for edge in ('start', 'top', 'end', 'bottom'):
        if edge in kwargs:
            element = tcBorders.find(qn(f'w:{edge}'))
            if element is None:
                from lxml import etree
                element = etree.SubElement(tcBorders, qn(f'w:{edge}'))
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))

def border_single():
    return {'val': 'single', 'sz': '4', 'color': '000000'}

def add_table(doc, header, rows):
    """添加带边框的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Table Grid'
    # 表头
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h.strip())
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '宋体'
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val.strip())
            run.font.size = Pt(9)
            run.font.name = '宋体'
    doc.add_paragraph()  # 表后空行
    return table

def is_table_sep(line):
    return re.match(r'^\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?\s*$', line)

def parse_table_cells(line):
    """解析表格行，返回单元格列表"""
    parts = line.strip().split('|')
    # 去掉首尾空
    if parts and not parts[0].strip():
        parts = parts[1:]
    if parts and not parts[-1].strip():
        parts = parts[:-1]
    return [p.strip() for p in parts]

def make_doc(doc, md_path):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)

    def heading(text, level):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(18)
            pf.space_after = Pt(12)
            run = p.add_run(text)
            run.bold = True; run.font.size = Pt(16); run.font.name = '黑体'
        elif level == 2:
            pf.space_before = Pt(14); pf.space_after = Pt(8)
            run = p.add_run(text)
            run.bold = True; run.font.size = Pt(14); run.font.name = '黑体'
        else:
            pf.space_before = Pt(10); pf.space_after = Pt(6)
            run = p.add_run(text)
            run.bold = True; run.font.size = Pt(12); run.font.name = '黑体'

    def parse_bold(text):
        """返回 (clean_text, [(is_bold, text_seg), ...])"""
        segs = []
        pos = 0
        for m in re.finditer(r'\*\*(.+?)\*\*', text):
            if m.start() > pos:
                segs.append((False, text[pos:m.start()]))
            segs.append((True, m.group(1)))
            pos = m.end()
        if pos < len(text):
            segs.append((False, text[pos:]))
        if not segs:
            segs = [(False, text)]
        return segs

    def add_p(segs, indent_first=True, indent_left=0, heading_run=False):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(5)
        if indent_first and indent_left == 0:
            pf.first_line_indent = Cm(0.74)
        if indent_left > 0:
            pf.left_indent = Cm(indent_left)
        for is_b, txt in segs:
            run = p.add_run(txt)
            run.font.size = Pt(11)
            run.font.name = '宋体'
            if is_b: run.bold = True

    def divider():
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(6); pf.space_after = Pt(6)
        pf.first_line_indent = Cm(0)
        run = p.add_run('—' * 30)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150,150,150)

    with open(md_path) as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        s = line.strip()

        if not s:
            i += 1; continue

        # 表格检测
        if '|' in s and i+1 < len(lines) and is_table_sep(lines[i+1].rstrip('\n')):
            # 表头
            header = parse_table_cells(s)
            i += 2  # 跳过表头和分隔行
            rows = []
            while i < len(lines) and '|' in lines[i].rstrip('\n'):
                rows.append(parse_table_cells(lines[i].rstrip('\n')))
                i += 1
            # 表标题（上一行）
            add_table(doc, header, rows)
            continue

        # 图片 ![alt](path)
        img_match = re.match(r'^!\[([^\]]*)\]\(([^\)]+)\)', s)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            # 相对于 md 文件的路径
            md_dir = os.path.dirname(os.path.abspath(md_path))
            full_img = os.path.join(md_dir, img_path) if not os.path.isabs(img_path) else img_path
            if os.path.exists(full_img):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(full_img, width=Inches(5.5))
                if alt_text:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = cap.add_run(alt_text)
                    r.font.size = Pt(9)
                    r.font.name = '宋体'
                    r.italic = True
            else:
                print(f'WARNING: image not found: {full_img}')
            i += 1; continue

        # 标题
        if s.startswith('# ') and not s.startswith('## '):
            i += 1; continue
        if s.startswith('## '):
            heading(parse_bold(s[3:])[0][1] if parse_bold(s[3:]) else s[3:], 2)
            i += 1; continue
        if s.startswith('### '):
            txt = ''.join(t for is_b,t in parse_bold(s[4:]) for _ in [0])
            heading(txt, 3)
            i += 1; continue

        # 分隔线
        if s == '---':
            divider()
            i += 1; continue

        # 无序列表
        if s.startswith('- ') or s.startswith('* '):
            segs = parse_bold(s[2:])
            add_p([(False, '•  ')] + segs, indent_first=False, indent_left=1.2)
            i += 1; continue

        # 编号列表
        m = re.match(r'^(\d+)[\.\、]\s(.+)', s)
        if m:
            segs = parse_bold(m.group(2))
            add_p([(False, f'{m.group(1)}. ')] + segs, indent_first=False, indent_left=1.2)
            i += 1; continue

        # 普通段落
        add_p(parse_bold(s))
        i += 1

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python3 convert.py input.md output.docx')
        sys.exit(1)
    doc = Document()
    make_doc(doc, sys.argv[1])
    doc.save(sys.argv[2])
    print(f'已保存: {sys.argv[2]}')
