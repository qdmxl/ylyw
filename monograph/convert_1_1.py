#!/usr/bin/env python3
"""将 ch01_1_1.md 转换为 docx —— 彻底清理所有 markdown 语法"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)

def add_heading_custom(text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(18)
        pf.space_after = Pt(12)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = '黑体'
    elif level == 2:
        pf.space_before = Pt(14)
        pf.space_after = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = '黑体'
    else:
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = '黑体'

def parse_inline_format(text):
    """
    解析 **bold** 标记，返回 (cleaned_text, bold_ranges)
    例如 "这是**粗体**文字" → ("这是粗体文字", [(2,4)])
    """
    parts = []
    bold_ranges = []
    cleaned = ''
    pos = 0
    # 找所有 **...** 对
    pattern = re.compile(r'\*\*(.+?)\*\*')
    for m in pattern.finditer(text):
        # m.start() 之前的普通文字
        before = text[pos:m.start()]
        if before:
            parts.append(('normal', before))
            cleaned += before
        # ** 内的粗体文字
        bold_text = m.group(1)
        start_in_cleaned = len(cleaned)
        parts.append(('bold', bold_text))
        cleaned += bold_text
        bold_ranges.append((start_in_cleaned, len(cleaned)))
        pos = m.end()
    # 最后一段
    if pos < len(text):
        tail = text[pos:]
        parts.append(('normal', tail))
        cleaned += tail
    return cleaned, parts

def add_rich_para(parts, indent_first=True, indent_left=0):
    """添加富文本段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    if indent_first and indent_left == 0:
        pf.first_line_indent = Cm(0.74)
    if indent_left > 0:
        pf.left_indent = Cm(indent_left)
    for kind, text in parts:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = '宋体'
        if kind == 'bold':
            run.bold = True

def add_divider():
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    run = p.add_run('—' * 30)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)

def process_md(content):
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 一级标题跳过
        if stripped.startswith('# ') and not stripped.startswith('## '):
            i += 1
            continue

        # ## 二级标题
        if stripped.startswith('## '):
            text = stripped[3:]
            cleaned, _ = parse_inline_format(text)
            add_heading_custom(cleaned, 2)
            i += 1
            continue

        # ### 三级标题
        if stripped.startswith('### '):
            text = stripped[4:]
            cleaned, _ = parse_inline_format(text)
            add_heading_custom(cleaned, 3)
            i += 1
            continue

        # 分隔线
        if stripped == '---':
            add_divider()
            i += 1
            continue

        # 无序列表
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            cleaned, parts = parse_inline_format(text)
            # 添加 •  前缀
            bullet_parts = [('normal', '•  ')] + parts
            add_rich_para(bullet_parts, indent_first=False, indent_left=1.2)
            i += 1
            continue

        # 编号列表 (数字. 或 数字、)
        m = re.match(r'^(\d+)[\.\、]\s(.+)', stripped)
        if m:
            num, text = m.group(1), m.group(2)
            cleaned, parts = parse_inline_format(text)
            bullet_parts = [('normal', f'{num}. ')] + parts
            add_rich_para(bullet_parts, indent_first=False, indent_left=1.2)
            i += 1
            continue

        # 普通段落
        cleaned, parts = parse_inline_format(stripped)
        add_rich_para(parts)
        i += 1

with open('/home/lijinhan/MXL/科研/ylyw/monograph/ch01_1_1.md', 'r') as f:
    content = f.read()

process_md(content)

output_path = '/home/lijinhan/MXL/科研/ylyw/monograph/第1章_1.1_易经作为先验知识体系.docx'
doc.save(output_path)
print(f'已保存: {output_path}')
