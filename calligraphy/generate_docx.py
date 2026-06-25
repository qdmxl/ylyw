#!/usr/bin/env python3
"""生成 YLYW 书法学习技术论文 docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import re

OUTDIR = Path('/home/lijinhan/MXL/科研/ylyw/calligraphy/output')

doc = Document()

# ---- 页面设置 ----
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ---- 辅助函数 ----
def add_title(text, level=0):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return p
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_bold_para(bold_text, rest_text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    r1 = p.add_run(bold_text)
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = p.add_run(rest_text)
    r2.font.size = Pt(12)
    return p

def add_table(headers, rows, caption=""):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(10)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table

def add_figure(image_path, caption="", width=5.5):
    if Path(image_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width))
        if caption:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p2.add_run(caption)
            r.font.size = Pt(10)
            r.italic = True

# ================================================================
# 正文
# ================================================================

# === 标题 ===
add_title('基于易经符号先验的机器人书法学习系统：\n观物取象与以卦驭笔')

# === 作者 ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('马兴录  李金函  张国安  于敬涛  李望  马圣洁*')
r.font.size = Pt(14)
r = p.add_run('\n（青岛科技大学 信息科学技术学院，山东 青岛 266061）')
r.font.size = Pt(11)

doc.add_paragraph()

# === 摘要 ===
add_title('摘  要', level=1)
add_para(
    '机器人精细操作是具身智能的关键挑战。当前的深度强化学习方法依赖数十万至数百万次试错交互才能习得基本技能，且决策过程不可解释、物理约束缺乏显式保证。本文提出基于YLYW（易理研物）三层符号先验架构的机器人书法学习系统，实现从字帖图像到机械臂书写轨迹的零样本生成与少样本优化。系统将八卦基元映射为八种基本笔法原型（乾→中锋、坤→侧锋、震→提按等），六爻编码建模笔画方向、粗细、曲率、规整性和重心位置，六十四卦匹配输出结构化的书写策略。在此基础上，提出知几学习闭环——通过比较字帖卦象与书写结果卦象的爻位级差异，精确定位偏差来源并进行定向参数修正。MuJoCo仿真实验表明：在「大」字的零样本基线测试中，系统仅凭先验知识即可产生合理的书写策略（卦象距离0.324）；经过8轮知几迭代，卦象距离降至0.087（改善73.2%），评级从"中"提升至"优"，且第6轮卦象从错误的"坎"自发回归到正确的"离"。爻位收敛分析显示初爻（方向）偏差从+0.267降至+0.014（改善94.8%），验证了爻位级可解释学习的有效性。同一架构可直接迁移至狭小空间焊接场景，仅需替换领域语义映射，展现了YLYW作为跨域精细操作通用框架的潜力。'
)

p = doc.add_paragraph()
r = p.add_run('关键词：')
r.bold = True
r.font.size = Pt(12)
r = p.add_run('易经符号先验；机器人书法；知几学习；零样本学习；可解释具身智能；六爻编码')
r.font.size = Pt(12)

# === 1 引言 ===
add_title('1  引言', level=1)

add_title('1.1  机器人精细操作的挑战', level=2)
add_para(
    '机器人精细操作（Fine Manipulation）——如书法书写、精密焊接、微创手术——是具身智能最具挑战性的领域之一。与抓取和推动等粗粒度操作不同，精细操作要求亚毫米级的轨迹精度、毫秒级的力控制以及结构化的任务理解。当前主流方法可分为两条技术路线。'
)
add_para(
    '端到端学习的代表性工作如QT-Opt和RT-2，通过海量示范数据学习从视觉到动作的直接映射。然而，精细操作的示范数据极难采集——书法的"笔法"、焊接的"熔深控制"无法通过遥操作准确还原。更根本的是，这些方法将物理世界的因果结构埋藏在深度网络的权重矩阵中，丧失了可解释性和可审计性。'
)
add_para(
    '示教编程仍是工业焊接机器人的主流方案，但存在两大痛点：新工件需要重新示教（周期以天计），狭小空间内的避碰轨迹需要反复调试。一种能"看懂工件结构、自主规划操作顺序"的机器人系统，是工业界的迫切需求。'
)

add_title('1.2  书法作为精细操作的理想实验场', level=2)
add_para(
    '书法为精细操作研究提供了一个独特且理想的实验场景。首先，书法是二维平面的轨迹任务（加上笔尖压力控制作为第三维），变量可控且易于量化评估。其次，书法有明确的"字帖"作为目标图像，天然形成"目标→执行→对比→修正"的闭环。再次，书法的笔法系统（永字八法：侧、勒、弩、趯、策、掠、啄、磔）已提供了结构化的操作原语分类。最后，书法的文化价值赋予了研究独特的叙事维度。'
)
add_para(
    '更具深意的是，汉字与《易经》共享"观物取象"的认知源头——"象-数-理"三位一体的结构同构使YLYW模型的八卦基元在书法域具有天然的亲和性。'
)

add_title('1.3  书法与焊接的结构同构性', level=2)
add_para(
    '本文还揭示了一个值得关注的结构同构：书法书写与狭小空间焊接在任务结构上几乎完全对应。汉字由笔画组成，焊缝由焊道组成；书法的"横竖撇捺"对应焊接的"打底焊、填充焊、盖面焊"；书法的笔顺规则（从上到下、从左到右、先横后竖）对应焊接的层序规则（从内到外、从下到上）；书法的"提按顿挫"对应焊接的"电流-速度-摆幅"参数调节。这意味着一次架构设计可以实现两个场景的验证——书法作为低成本、低风险的先行验证场，焊接作为高价值的工业泛化场景。'
)

add_title('1.4  本文贡献', level=2)
add_para(
    '本文的主要贡献包括：（1）将YLYW三层符号先验架构从抓取域系统性地映射到书法域，建立了八卦-笔法、六爻-笔画特征、六十四卦-书写策略的完整对应；（2）提出知几学习闭环，在MuJoCo仿真环境中验证了从"观帖"到"临摹"到"自省"到"精进"的完整学习链路；（3）通过爻位级收敛分析，首次揭示了可解释学习中的"爻位分化"现象——部分爻位可通过参数修正改善，部分爻位需要轨迹重规划；（4）论证了书法与狭小空间焊接的结构同构性，提出了YLYW作为跨域精细操作通用框架的路线图。'
)

# === 2 方法 ===
add_title('2  方法', level=1)

add_title('2.1  YLYW三层架构的书法域映射', level=2)
add_para(
    'YLYW的核心架构由三层组成：L1八卦基元（模糊隶属度计算）、L2六爻编码（物理状态→符号编码）、L3六十四卦规则库（符号→策略映射）。在书法域，这三层的语义发生了系统性的迁移，但计算结构完全保留。'
)

add_title('2.1.1  L1：八卦基元——从物理到视觉到笔法', level=3)
add_para(
    '在原始的物理抓取域，八卦映射为8种物体物理原型（乾=刚性强力、坤=柔性包容等）。在视觉书法域，八卦映射为8种视觉原型（乾=高边缘+高规整的正方结构、坤=圆转流畅的曲线型、震=高对比的方向性笔画等）。更重要的是，在书写执行域，八卦映射为8种基本笔法，每种笔法包含特定的压力基准、速度范围和笔尖姿态。'
)

add_table(
    ['卦名', '符号', '物理原型', '笔法', '特征'],
    [
        ['乾', '☰', '刚健', '中锋', '力均笔正'],
        ['坤', '☷', '柔顺', '侧锋', '铺毫舒展'],
        ['震', '☳', '动态', '提按', '轻重交替'],
        ['艮', '☶', '静止', '顿笔', '驻留蓄力'],
        ['离', '☲', '附丽', '露锋', '出锋见尖'],
        ['坎', '☵', '险陷', '藏锋', '回锋内敛'],
        ['兑', '☱', '悦', '轻灵', '短小跳跃'],
        ['巽', '☴', '入', '绵密', '细长柔韧'],
    ],
    '表1  八卦→笔法映射'
)

add_title('2.1.2  L2：六爻编码——笔画结构特征', level=3)
add_para(
    '六爻在书法域的语义映射如表2所示。与物理域使用公式计算的六爻编码不同，书法视觉域通过12维图像特征提取器（含笔画方向性、粗细对比、曲率复杂度、间架规整度、重心位置等特征）直接产生[0,1]范围的爻值。'
)

add_table(
    ['爻位', '名称', '书法域语义', '说明'],
    [
        ['初爻', '基础', '笔画方向主导度', '笔画方向是否集中'],
        ['二爻', '展现', '粗细对比度', '笔画宽度变化幅度'],
        ['三爻', '变化', '曲直复杂度', '笔画曲线的丰富程度'],
        ['四爻', '稳固', '间架规整度', '结构布局的规则性'],
        ['五爻', '中心', '重心横向位置', '字心是否居中'],
        ['上爻', '完成', '重心纵向位置', '重心是否偏上/下'],
    ],
    '表2  六爻语义映射'
)

add_title('2.2  视觉YLYW：读帖与结构卦象生成', level=2)
add_para(
    '视觉YLYW实现了从字帖图像到结构卦象的完整分析链路。首先，通过Zhang-Suen细化算法提取笔画骨架，检测端点和分叉点来分割笔画段。然后，对每个笔画段提取12维结构特征（含基础特征、方向特征和结构特征三组）。最后，将6维主要特征编码为六爻向量，与八卦原型计算隶属度，匹配到最接近的六十四卦。'
)

add_title('2.3  书写YLYW：以卦驭笔的轨迹生成', level=2)
add_para(
    '书写YLYW将结构卦象转换为可执行的机械臂书写轨迹。系统维护了一个包含"永、大、人、中、心、山"等基础汉字的参考模板，定义每个笔画的起终点和基本形态。卦象驱动的书写策略根据以下参数生成具体的轨迹：笔法选择（根据卦象隶属度选择主导笔法）、几何修正（笔画角度修正对应初爻、宽度因子对应二爻、曲率因子对应三爻、抖动幅度对应四爻）、重心偏移（根据字帖重心调整整体位置，对应五爻和上爻）、压力曲线（起笔轻→行笔重→收笔轻的三段式压力剖面）。'
)

add_title('2.4  知几学习闭环：自省与精进', level=2)
add_para(
    '知几学习是本文最核心的方法论创新。与深度强化学习的事后奖励驱动不同，知几学习以"事先洞察变化的征兆"为哲学内核。在书法场景中，其闭环流程为：'
)
add_para('（1）观帖：视觉YLYW分析字帖，生成目标卦象 H_target；')
add_para('（2）临摹：书写YLYW根据目标卦象生成首次书写计划；')
add_para('（3）自省：视觉YLYW分析书写结果，生成实际卦象 H_actual，计算爻位差异 Δy = y_actual − y_target；')
add_para('（4）精进：根据爻位差异定向修正参数，重新书写。')
add_para(
    '参数修正公式为 p(t+1) = p(t) + η · D · Δy(t)，其中 p 为可调参数向量（约443维），η=0.5 为学习率，D 为爻位-参数映射矩阵。这一简洁的更新规则使系统能够在5~8次迭代内实现显著改善，而非深度强化学习的数千至百万次试错。'
)

add_title('2.5  卦象驱动的笔画自主分解与排序', level=2)
add_para(
    '为实现对任意新字的自主学习（无需预定义模板），系统包含了一个笔画自主分解模块。从字帖图像出发：二值化→Zhang-Suen细化→骨架修剪→端点/分叉点检测→笔画追踪分割→合并共线段→笔画类型分类（横/竖/撇/捺/点/钩/弧）。笔画书写顺序由卦象驱动的排序算法决定，核心规则为：（1）空间约束：从顶部到底部（y坐标升序）、从左到右（x坐标升序）；（2）卦象优先级：乾(横)→艮(竖)→震(撇)→坤(捺)→兑(点)→坎(钩)→巽(折)→离(弧)；（3）爻位关系：乘承比应→先承后乘、先比后应、当位优先。'
)

add_title('2.6  安全八卦：纸面保护与书写约束', level=2)
add_para(
    '仿照YLYW的双八卦架构，书法系统同样包含安全八卦层。在书法场景中，安全八卦建模纸面损坏风险（笔尖压力过大导致纸张破损）和墨量约束。当压力超过阈值时触发变卦机制，自动降力或提笔。这一设计同样可迁移至焊接场景的安全八卦——建模碰撞避免和热影响区控制。'
)

# === 3 实验 ===
add_title('3  实验', level=1)
add_figure(str(OUTDIR / 'figures' / 'fig4_architecture.png'), '图1  YLYW书法学习系统架构', width=5.5)

add_title('3.1  实验设置', level=2)
add_para(
    '实验在MuJoCo物理引擎中进行。仿真环境包含一个300×300 mm的虚拟纸面、一个由三个滑块关节控制的毛笔末端执行器（x位置、y位置、z压力），以及一个正上方的虚拟相机用于渲染256×256像素的书写结果图像。笔触渲染采用基于高斯核的墨色混合模型，笔触半径和墨色浓度与压力值正相关。所有实验在一台配备Ubuntu 26.04的VirtualBox虚拟机上完成，使用EGL离屏渲染后端。'
)

add_title('3.2  零样本基线', level=2)
add_para(
    '「大」字的零样本基线测试结果表明，系统仅凭卦象先验（视觉YLYW将字帖分析为"离"卦，即疏朗通透的结构），生成了第一次书写计划。书写结果与字帖的卦象距离为0.324，对应评级为"中"——即"基本合理但存在显著偏差"。零样本基线已具备合理的结构（三笔：横→撇→捺），但存在以下偏差：初爻（笔画方向）偏差+0.267，表明笔画方向过于统一僵硬；二爻（粗细对比）偏差+0.158，表明笔画宽度变化不足。这与端到端学习的"随机初始"形成鲜明对比——YLYW的第1次书写已是"先验驱动"的合理尝试，而非盲目探索。'
)

add_title('3.3  学习曲线与爻位收敛', level=2)
add_para('表3给出了「大」字8轮知几迭代的完整数据，图2展示了卦象距离的下降趋势和爻位收敛情况。')

add_table(
    ['迭代', '卦象距离', '评级', '目标卦', '实际卦', '最大偏差爻'],
    [
        ['1', '0.324', '中', '离', '坎', '初爻(+0.267)'],
        ['2', '0.300', '中', '离', '坎', '初爻(+0.233)'],
        ['3', '0.307', '中', '离', '坎', '二爻(+0.231)'],
        ['4', '0.199', '良', '离', '坎', '二爻(+0.133)'],
        ['5', '0.159', '良', '离', '坎', '二爻(+0.107)'],
        ['6', '0.134', '良', '离', '离', '二爻(+0.103)'],
        ['7', '0.124', '良', '离', '离', '二爻(+0.103)'],
        ['8', '0.087', '优', '离', '离', '二爻(+0.064)'],
    ],
    '表3  「大」字8轮迭代学习数据'
)

add_figure(str(OUTDIR / 'figures' / 'fig1_learning_curve.png'), '图2  「大」字8轮学习曲线（左）与改善量瀑布图（右）', width=5.5)

add_para(
    '总改善率：73.2%（从0.324降至0.087），评级从"中"提升至"优"。最引人注目的是第6轮发生了卦象回归——书写结果的主导卦象从错误的"坎"（对应藏锋内敛的笔法）自发回归到正确的"离"（疏朗通透），且此后保持稳定。'
)

add_table(
    ['爻位', '初始偏差', '最终偏差', '改善量', '改善率', '趋势'],
    [
        ['初爻(方向)', '+0.267', '+0.014', '0.253', '94.8%', '✓'],
        ['二爻(粗细)', '+0.158', '+0.064', '0.094', '59.5%', '✓'],
        ['三爻(曲直)', '+0.073', '+0.045', '0.028', '38.4%', '✓'],
        ['四爻(规整)', '0.000', '0.000', '0.000', '--', '--'],
        ['五爻(重心x)', '+0.056', '+0.031', '0.025', '44.6%', '✓'],
        ['上爻(重心y)', '+0.012', '+0.018', '-0.006', '-50.0%', '微退'],
    ],
    '表4  爻位收敛数据'
)

add_figure(str(OUTDIR / 'figures' / 'fig2_yao_convergence.png'), '图3  爻位收敛分析：初始偏差 vs 最终偏差', width=5.5)

add_title('3.4  卦象回归现象', level=2)
add_para(
    '第6轮出现的卦象回归是本次实验最核心的定性发现。前5轮书写结果均被视觉YLYW判定为"坎"卦，但从第6轮开始回归到目标"离"卦。从数据角度看，这是因为在前5轮迭代中，初爻（方向）偏差从+0.267逐步降至+0.054，二爻（粗细）偏差从+0.158波动至+0.103——当笔画方向和粗细接近字帖时，视觉YLYW的卦象分类器自然地将书写结果重新归类为"离"。这一可观测的符号层面变迁，是YLYW可解释性优势的生动体现。'
)

add_title('3.5  消融实验', level=2)
add_para(
    '为验证知几修正机制的有效性，我们进行了消融实验：在相同条件下执行5轮书写但不应用参数修正。结果显示卦象距离在0.335±0.005范围内随机波动，无下降趋势。这证明YLYW的学习改善来自于爻位定向修正，而非随机扰动或环境噪声。'
)

# === 4 讨论 ===
add_title('4  讨论', level=1)

add_title('4.1  爻位分化现象', level=2)
add_para(
    '实验揭示了爻位在可修正性上的系统性分化。初爻（方向）、二爻（粗细）、三爻（曲直）和五爻（重心）可以通过调整轨迹生成参数进行有效修正——这些爻位对应的"风格参数"（笔法选择、角度修正、压力曲线）在轨迹生成器中已有直接的调节接口。然而，四爻（间架规整度）在当前架构中无法通过参数修正来改善。这是因为间架结构取决于笔画的绝对位置关系（例如，"大"字的横与撇捺的位置关系），而当前轨迹生成基于预定义模板，无法重新规划笔画的空间布局。这一发现指明了系统下一步的升级方向——从模板轨迹生成升级到卦象驱动的自由轨迹规划。'
)

add_title('4.2  与深度强化学习的对比', level=2)
add_para(
    '表5从五个维度对比了YLYW知几学习与深度强化学习。YLYW的核心优势在于先验驱动的零样本合理性（而非随机初始）、极小样本需求（8次 vs 数千/百万次）、完全可解释的学习过程（爻位级诊断 vs 黑箱loss）、以及结构内建的安全约束。'
)

add_table(
    ['维度', 'YLYW知几学习', '深度强化学习(DRL)'],
    [
        ['初始性能', '先验驱动，零样本合理', '随机策略，无意义'],
        ['样本效率', '5~8次定向修正', '数千~百万次试错'],
        ['可解释性', '爻位级可追溯', '黑箱，仅输出标量loss'],
        ['学习迁移', '参数空间直接迁移', '需重新训练'],
        ['安全性', '安全八卦内建约束', '依赖试错中"学到"'],
    ],
    '表5  YLYW知几学习 vs 深度强化学习'
)

add_figure(str(OUTDIR / 'figures' / 'fig3_ylyw_vs_drl.png'), '图4  YLYW知几学习 vs 深度强化学习五维对比', width=4.5)

add_title('4.3  从书法到焊接的迁移路径', level=2)
add_para(
    '本文的核心论点之一是：书法和焊接共享同一套任务分解和执行的底层结构。具体而言，YLYW的三层架构在焊接域的迁移仅需替换领域语义映射（表6），而推理逻辑完全保持不变。焊缝图像→焊道骨架提取→焊道分段→卦象驱动焊接排序→焊枪轨迹生成的完整链路，与书法系统的处理链路结构完全相同。唯一的差异在于视觉特征提取器的具体算子和六爻编码公式的领域系数——这些可以通过少量焊接样本进行快速校准。'
)

add_table(
    ['YLYW组件', '书法域', '焊接域'],
    [
        ['L1 八卦基元', '笔法类型(中锋/侧锋等)', '焊接工艺(深熔/盖面/脉冲等)'],
        ['L2 六爻编码', '笔画结构特征', '焊缝结构特征'],
        ['L3 六十四卦', '书写策略', '焊接顺序策略'],
        ['安全八卦', '纸面压力保护', '碰撞避免+热影响控制'],
        ['知几学习', '看字对比修正', '焊缝质量检测修正'],
    ],
    '表6  书法→焊接：领域语义迁移'
)

add_figure(str(OUTDIR / 'figures' / 'fig5_welding_isomorphism.png'), '图5  书法→焊接结构同构对比', width=5.5)

# === 5 结论 ===
add_title('5  结论', level=1)
add_para(
    '本文在YLYW三层符号先验架构的基础上，系统性地构建了一个机器人书法学习系统，并在MuJoCo仿真环境中完成了完整的实验验证。系统实现了从字帖图像到机械臂书写轨迹的零样本生成（仅凭先验知识），以及通过知几学习闭环实现少样本定向优化（8轮迭代改善73.2%）。'
)
add_para(
    '在方法论层面，本文的核心贡献是论证了符号先验知识在机器人精细操作中的独特价值。与"越多数据越好"的主流叙事不同，YLYW的实证结果表明：一套结构良好的、可解释的先验知识体系（八卦、六爻、六十四卦），可以在零样本条件下产生合理的行为，并且该行为可以通过精确诊断和定向修正实现快速改善——而无需海量训练数据。'
)
add_para(
    '在技术层面，爻位分化现象的发现为系统升级提供了清晰路径：下一阶段的重点在于将笔画轨迹生成从模板升级为卦象驱动的自由规划，特别是解耦四爻（间架结构）的生成能力。同时，将系统从MuJoCo仿真迁移到真实机械臂（如UR系列），是验证实际可行性的关键步骤。'
)
add_para(
    '从更宏观的视角看，本研究暗示了一种值得关注的AGI路径：不是堆算力和数据，而是将人类文明中已经积累的、经过数千年验证的知识体系（如《易经》符号系统）形式化为可计算的先验结构。这种"道器合一"的路径——以哲学智慧提供架构层面的强归纳偏置，以工程实现完成具体的形式化和验证——为具身智能提供了当前主流方法之外的一种可能。在后续工作中，我们将进行实物机械臂的书法验证和焊接场景的跨域迁移实验，并在更丰富的汉字集合上评估笔画自主分解模块的泛化性能。'
)

# === 致谢 ===
add_title('致  谢', level=1)
add_para('本研究得到青岛科技大学信息科学技术学院的支持。')

# === 参考文献 ===
add_title('参考文献', level=1)

refs = [
    '[1] 马兴录, 张国安, 李金函, 等. YLYW：一种基于《易经》先验符号知识的联邦式神经符号具身决策框架[R]. 青岛科技大学, 2025.',
    '[2] 马兴录, 李金函, 张国安, 等. 知几学习：从"事后奖惩"到"事先征兆"的具身学习范式革新[R]. 青岛科技大学, 2025.',
    '[3] 马兴录, 李金函, 等. YLYW汉语言文字处理新范式：观物取象与乘承比应[R]. 青岛科技大学, 2025.',
    '[4] 周文王, 孔子, 等. 周易[M]. 公元前11世纪--前5世纪.',
    '[5] 孔子, 等. 系辞传[M]//周易. 公元前5世纪.',
    '[6] Driess D, Xia F, Sajjadi M S M, et al. PaLM-E: An embodied multimodal language model[C]. ICML, 2023.',
    '[7] Brohan A, Brown N, Carbajal J, et al. RT-2: Vision-language-action models transfer web knowledge to robotic control[C]. CoRL, 2023.',
    '[8] Kim M J, Pertsch K, Karamcheti S, et al. OpenVLA: An open-source vision-language-action model[J]. arXiv:2406.09246, 2024.',
    '[9] Team Octo, Ghosh D, Walke H, et al. Octo: An open-source generalist robot policy[C]. RSS, 2024.',
    '[10] Black K, Brown N, Driess D, et al. π₀: A vision-language-action flow model for general robot control[J]. arXiv:2410.24164, 2024.',
    '[11] Yao F, Shao G, Yi J. Extracting the trajectory of writing brush in Chinese character calligraphy[J]. EAAI, 2004.',
    '[12] Wu R, Zhou C, Chao F, et al. A developmental evolutionary learning framework for robotic Chinese calligraphy[J]. IEEE TCDS, 2023.',
    '[13] Chao F, Lin G, Zheng L, et al. A learning framework for robotic calligraphy with style transfer[J]. Complex & Intelligent Systems, 2023.',
    '[14] Chen S B, Lv N. Research evolution on intelligentized technologies for arc welding process[J]. JMP, 2014.',
    '[15] Xu Y, Lv N, Fang G, et al. Welding seam tracking in robotic gas metal arc welding[J]. JMP, 2022.',
    '[16] Wang X, Zhou Q, Chen B, et al. Deep learning-based welding seam recognition and tracking[J]. RCIM, 2023.',
    '[17] Garcez A, Lamb L C. Neurosymbolic AI: The 3rd wave[J]. AI Review, 2023.',
    '[18] Yu D, Yang B, Liu D, et al. A survey of neurosymbolic visual reasoning[J]. arXiv:2305.07625, 2023.',
    '[19] Levine S, Pastor P, Krizhevsky A, et al. Learning hand-eye coordination for robotic grasping with deep learning and large-scale data collection[J]. IJRR, 2018.',
    '[20] Kalashnikov D, Irpan A, Pastor P, et al. QT-Opt: Scalable deep reinforcement learning for vision-based robotic manipulation[C]. CoRL, 2018.',
    '[21] Zadeh L A. Fuzzy logic[J]. Computer, 1988.',
    '[22] Wang Y, et al. VLA模型在具身智能中的语义泛化与执行泛化瓶颈综述[J]. 中国科学：信息科学, 2025.',
    '[23] Todorov E, Erez T, Tassa Y. MuJoCo: A physics engine for model-based control[C]. IROS, 2012.',
    '[24] Coumans E, Bai Y. PyBullet, a Python module for physics simulation for games, robotics and machine learning[J]. http://pybullet.org, 2016--2021.',
    '[25] 魏宏森. 系统论的基本规律[J]. 自然辩证法研究, 1995.',
]

for ref in refs:
    p = doc.add_paragraph()
    r = p.add_run(ref)
    r.font.size = Pt(10)

# ---- 保存 ----
output_path = OUTDIR / 'YLYW书法学习技术论文.docx'
doc.save(str(output_path))
print(f"✅ 论文已生成: {output_path}")
print(f"   大小: {output_path.stat().st_size / 1024:.0f} KB")
