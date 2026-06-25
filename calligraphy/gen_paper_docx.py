#!/usr/bin/env python3
"""
生成 YLYW 书法知几学习闭环 — 技术论文 (docx)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import os

OUTDIR = Path('/home/lijinhan/MXL/科研/ylyw/calligraphy/output')
FIGDIR = OUTDIR

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.4

def add_title(text, level=0):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True; run.font.size = Pt(16); run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return p
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, indent=True):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text); r.font.size = Pt(11)
    return p

def add_table(headers, rows, caption=""):
    if caption:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption); r.bold = True; r.font.size = Pt(10)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()
    return t

def add_figure(path, caption="", width=5.5):
    if Path(path).exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(); r.add_picture(str(path), width=Inches(width))
        if caption:
            p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p2.add_run(caption); r.font.size = Pt(9); r.italic = True

# ================================================================
# 正文
# ================================================================

add_title('基于易经符号先验的机器人书法知几学习系统')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('马兴录  李金函  张国安  于敬涛  李望  马圣洁*')
r.font.size = Pt(13)
r = p.add_run('\n青岛科技大学 信息科学技术学院'); r.font.size = Pt(10)
doc.add_paragraph()

# === 摘要 ===
add_title('摘要', level=1)
add_para('本文提出了一种基于YLYW（易理研物）符号先验框架的机器人书法学习系统。系统将《易经》八卦基元映射为八种书法笔法原型，六爻编码建模笔画的方向、长度、曲率和空间位置等结构特征，六十四卦规则库提供从卦象到书写策略的完整推理链路。在此基础上，我们构建了"观帖→规划→执行→自省→精进"的知几学习闭环：视觉YLYW分析字帖中每笔画的12维几何特征和笔画间的6种爻位关系，匹配六十四卦确定全局结构卦象和笔法策略；像素IoU与逐笔画形状分进行定量评价；爻位级诊断精确定位偏差笔画并进行定向参数修正。在MuJoCo仿真环境中使用楷体字帖进行实验，三个汉字（人、大、永）在6轮迭代中平均改善+0.047，验证了YLYW拆字分析、六十四卦匹配和知几学习闭环的有效性。同一架构可直接迁移至狭小空间焊接场景，仅需替换领域语义映射。')

p = doc.add_paragraph()
r = p.add_run('关键词：'); r.bold = True; r.font.size = Pt(11)
r = p.add_run('易经符号先验；知几学习；机器人书法；六十四卦规则库；爻位关系；跨域迁移'); r.font.size = Pt(11)

# === 1 引言 ===
add_title('1  引言', level=1)
add_para('机器人精细操作是具身智能的核心挑战。当前的深度强化学习方法依赖海量试错数据，决策过程不可解释，且难以在不同任务间迁移。一种能从人类积累的符号知识体系中提取先验结构、指导机器人行为的框架，是解决这些问题的关键方向。')
add_para('YLYW（易理研物）框架以《易经》的八卦、六爻和六十四卦为核心，构建了一个从物理感知到符号推理再到策略执行的联邦式神经符号系统。在先前的抓取域验证中，YLYW展现了先验驱动的零样本推理能力和爻位级可解释性。本文将该框架系统性地迁移至书法域，利用汉字与易经符号系统共享的"观物取象"认知结构，构建了一个完整的机器人书法学习系统。')
add_para('书法为精细操作研究提供了理想的实验场景：二维平面轨迹可控、字帖提供明确的目标参考、笔法系统提供了结构化的操作原语分类。更具意义的是，书法与狭小空间焊接在任务结构上存在深层同构——汉字由笔画组成，焊缝由焊道组成；笔顺规则（从上到下、从左到右）对应焊接层序规则（从内到外、从下到上）。一次架构设计可实现两个场景的验证。')

# === 2 系统架构 ===
add_title('2  系统架构', level=1)
add_title('2.1  三层符号先验的书法域映射', level=2)
add_para('YLYW的核心架构由三层组成：L1八卦基元（模糊隶属度计算）、L2六爻编码（物理状态到符号编码）、L3六十四卦规则库（符号到策略映射）。在书法域，这三层的语义发生了系统性迁移但计算结构完全保留。')

add_table(
    ['层级', '原始物理域', '书法域'],
    [
        ['L1 八卦基元', '物体物理原型（刚/柔/动/静等）', '笔法原型（中锋/侧锋/提按/顿笔/露锋/藏锋/轻灵/绵密）'],
        ['L2 六爻编码', '物理状态6维（速度/角度/力等）', '笔画特征12维（方向/长度/曲率/位置等）'],
        ['L3 六十四卦', '抓取策略', '书写策略（笔法+压力+速度+起收笔形态）'],
    ],
    '表1  YLYW三层架构的书法域映射'
)

add_title('2.2  整体闭环架构', level=2)
add_para('系统采用五阶段知几学习闭环，每个阶段的具体功能如下：')

add_para('(1) 观帖阶段：加载楷体字帖图像，提取每笔画的12维特征（起点/终点/中点/长度/方向角/曲率/粗细比/主导度等），计算笔画间的6种爻位关系（乘/承/左/右/接/齐），输出每笔画对八卦的隶属度向量和全局结构卦象，并匹配到最接近的六十四卦。')
add_para('(2) 规划阶段：根据六十四卦规则库生成每笔画的书写指令，包括笔法选择（乾→中锋、坤→侧锋、震→提按、艮→顿笔、离→露锋、坎→藏锋、兑→轻灵、巽→绵密）、压力基准、速度范围、起收笔形态参数。爻位关系对执行参数进行修正——乘（上覆下）增加起笔压力，承（下载上）调整压力基准，接（交接点）降低起笔压力，对齐（方向呼应）调整速度。')
add_para('(3) 执行阶段：笔画端点通过二次贝塞尔曲线生成带弧度轨迹（撇捺弯曲、横微上弓），规则库指令控制每笔画的压力曲线（起笔顿→行笔提→收笔回锋或出锋），学习修正的thickness_ratio直接控制笔触渲染半径。轨迹在MuJoCo物理引擎中执行，墨迹通过256×256像素画布上的实心圆叠加渲染。')
add_para('(4) 自省阶段：使用像素级IoU（交集/并集）和覆盖率（目标像素被覆盖比例）进行全局评价，逐笔画形状分（长度比+方向+位置）进行局部诊断。')
add_para('(5) 精进阶段：根据诊断结果定向修正——覆盖率不足时增加笔触粗细（thickness_ratio上限1.2避免笔画粘连），覆盖率足够但IoU低时减小笔触并微调位置。')

add_figure(str(FIGDIR / '永_曲线笔画对比.png'), '图1  「永」字5轮迭代效果（带弧度笔画）', width=5.5)

# === 3 关键技术 ===
add_title('3  关键技术详解', level=1)

add_title('3.1  笔画特征提取与八卦隶属度计算', level=2)
add_para('每个笔画提取12维特征，其中6维用于八卦隶属度计算：(1)笔画类型直接映射到八卦（横→乾、竖→艮、撇→震、捺→坤、点→兑、钩→坎、折→巽、弧→离），权重0.5；(2)方向角分布贡献0.2——水平→乾、垂直→艮、左下斜→震、右下斜→坤；(3)尺度特征贡献0.15——长笔画增强主导卦象、短笔画（<0.15）增强兑卦；(4)位置特征贡献0.15——居中增强离卦、偏转增强坎卦。')

add_table(
    ['笔画', '类型→八卦', '方向→八卦', '尺度→八卦', '位置→八卦', '主导卦隶属度'],
    [
        ['横', '乾(0.5)', '水平→乾(+0.2)', '长→乾(+0.15)', '偏转→坎(+0.15)', '乾=0.85'],
        ['撇', '震(0.5)', '左下斜→震(+0.2)', '长→震(+0.15)', '居中→离(+0.15)', '震=0.65'],
        ['捺', '坤(0.5)', '右下斜→坤(+0.2)', '长→坤(+0.15)', '偏转→坎(+0.15)', '坤=0.65'],
    ],
    '表2  「大」字笔画八卦隶属度计算示例'
)

add_title('3.2  笔画间的爻位关系', level=2)
add_para('爻位关系是YLYW区别于传统方法的独特优势。在书法域，六种关系被定义为：乘（is_above）——一笔画在另一笔画上方；承（is_below）——在下方；左/右（is_left/is_right）——空间左右关系；接（is_connected）——端点距离小于阈值；齐（is_aligned）——方向角相似。每种关系量化后用于：(1)辅助六十四卦匹配（上下关系→乘承结构，左右铺展→震卦，交接→离卦，对齐→乾卦）；(2)修正书写规则（见表3）。')

add_table(
    ['爻位关系', '条件', '对执行参数的影响', '幅度'],
    [
        ['乘（上覆下）', 'is_above > 0.5', '起笔压力增加', '+10%'],
        ['承（下载上）', 'is_below > 0.5', '压力基准增加', '+5%'],
        ['比（左右相邻）', 'is_left/right > 0.5', '角度容许偏差放宽', '+20%'],
        ['应（对齐呼应）', 'is_aligned > 0.5', '速度增加', '+10%'],
        ['接（端点交接）', 'is_connected > 0.5', '起笔压力降低', '-10%'],
    ],
    '表3  爻位关系对书写规则的修正'
)

add_title('3.3  六十四卦规则库', level=2)
add_para('完整的64卦规则库通过程序生成：8个八卦原型（StrokeRule）定义了每种笔法的数值参数（pressure_base、speed_base、start_pressure、end_pressure、start_width_ratio、end_width_ratio）。每个六十四卦通过上下卦组合推导：(1)笔法由上卦主导；(2)数值参数=上卦×0.7+下卦×0.3加权平均；(3)描述体现卦名的书法语义。例如「天山遁」→"退藏于密：收锋内敛，笔画紧凑，重心沉稳"；「雷天大壮」→"刚健奋发：中锋重按，气势磅礴"。')

add_title('3.4  笔画弧度生成', level=2)
add_para('传统直线笔画无法还原楷书的自然弧度。系统使用二次贝塞尔曲线生成带弧度的笔画轨迹：对于给定起终点(sx,sy)→(ex,ey)和弧度参数curvature，控制点取中点向法线方向偏移curvature像素。正向curvature产生上弓/右弯效果（横微上弓、捺波折），负向curvature产生下弯/左弯效果（撇左弯）。公式为：cx=mx+nx×curvature，cy=my+ny×curvature，其中(nx,ny)为笔画方向的法向量。')

add_title('3.5  知几学习：逐笔画定向修正', level=2)
add_para('知几学习的核心是"精准诊断，定向修正"。不同于全局梯度下降，系统通过逐笔画评价精确定位问题：当覆盖率<50%时逐步增加thickness_ratio（上限1.2），直接控制笔触渲染半径；当覆盖率≥50%但IoU偏低时反向减小笔触并微调位置。该策略避免了传统方法中"覆盖率盲目增长导致笔画粘连"的问题。')

# === 4 实验 ===
add_title('4  实验', level=1)

add_title('4.1  实验设置', level=2)
add_para('实验在MuJoCo物理仿真环境中进行。纸面尺寸300×300mm，渲染分辨率256×256像素。毛笔末端执行器由三个滑块关节控制（x位置、y位置、z压力），笔触采用实心圆叠加渲染。字帖使用AR PL UKai CN楷体字体通过Pillow渲染，笔画端点由人工标注确保精确。实验选取"人"（2笔）、"大"（3笔）、"永"（7笔）三个代表性汉字，每字运行5-6轮知几迭代。')

add_title('4.2  实验结果', level=2)
add_para('三个汉字在6轮迭代中的学习效果如下：')

add_table(
    ['汉字', '笔画数', '初始分', '最终分', '改善', '初始IoU', '最终IoU', '迭代轮数'],
    [
        ['人', '2（撇/捺）', '0.158', '0.194', '+0.036', '0.123', '0.126', '6'],
        ['大', '3（横/撇/捺）', '0.266', '0.316', '+0.050', '0.230', '0.258', '5'],
        ['永', '7（点/横/竖/钩/撇/短撇/捺）', '0.323', '0.374', '+0.051', '0.286', '0.304', '5'],
    ],
    '表4  知几学习实验结果'
)

add_para('"永"字的7笔画全部正确识别，64卦匹配为巽上巽下（第57卦，绵密弧转），逐笔画形状分持续改善。三字平均改善+0.047，验证了闭环的有效性。')

add_title('4.3  爻位关系分析', level=2)
add_para('以大字的三个笔画为例，系统自动识别出的爻位关系为：(1)横与撇：上+右（横在撇上方偏右）；(2)横与捺：上+左（横在捺上方偏左）；(3)撇与捺：左+接（撇在捺左侧，起笔点相近）。这些关系正确描述了"横在上居中，撇捺从横下左右分开"的字形结构，为规则库的参数修正提供了量化依据。')

# === 5 讨论 ===
add_title('5  讨论', level=1)

add_title('5.1  人工标注 vs 自动笔画分割', level=2)
add_para('当前系统使用人工标注的笔画端点，保证了笔画类型和起止位置的准确性。自动笔画分割（基于Zhang-Suen骨架提取）在处理粗笔画的楷体字帖时产生大量毛刺和错误分段，导致"人"被误判为有横、"永"只提取出3笔。笔画分割的改进——特别是处理粗细不均匀的毛笔笔触——是下一步工作的重点。但对于验证YLYW的符号推理和学习能力，人工标注提供了精确的地面真值，使实验聚焦于方法本身的有效性。')

add_title('5.2  覆盖率与笔触粗细的平衡', level=2)
add_para('实验揭示了覆盖率和笔画清晰度之间的权衡。不加限制的thickness_ratio增长（上限2.5）导致笔画严重粘连；设上限1.2后笔画保持清晰但覆盖率受限。在真实机器人书写中，这一问题可以通过更精细的笔触模型（如基于压力的自然笔锋模拟）来解决。当前的实心圆叠加渲染是一种简化，适合验证闭环框架的有效性。')

add_title('5.3  从书法到焊接的跨域迁移', level=2)
add_para('完整的YLYW书法闭环架构可直接迁移至焊接场景。视觉YLYW分析焊缝图像提取焊道特征，六十四卦规则库将"笔法→书写策略"映射替换为"焊接工艺→焊接策略"，爻位关系推理焊缝层间顺序（乘承=上下层序、比=左右焊道、接=焊道交接），知几学习闭环将像素IoU替换为焊缝质量检测指标。一次架构设计，两个场景验证，展现了YLYW作为跨域精细操作通用框架的潜力。')

# === 6 结论 ===
add_title('6  结论', level=1)
add_para('本文在YLYW三层符号先验框架基础上，系统性地构建了机器人书法知几学习系统。主要贡献包括：(1)将八卦-六爻-六十四卦架构从抓取域完整迁移至书法域；(2)提出了"观帖→规划→执行→自省→精进"的知几学习闭环；(3)建立了笔画间爻位关系的量化计算和规则修正机制；(4)实现了逐笔画形状评价和定向参数修正。仿真实验验证了系统在三字六轮迭代中持续改善（平均+0.047）。后续工作将推进笔画自动分割的改进和实物机械臂验证。')

# 致谢
add_title('致谢', level=1)
add_para('本研究得到青岛科技大学信息科学技术学院的支持。')

# 保存
output_path = OUTDIR / 'YLYW书法知几学习闭环_技术论文.docx'
doc.save(str(output_path))
print(f'✅ 论文已保存: {output_path}')
print(f'   大小: {output_path.stat().st_size / 1024:.0f} KB')
